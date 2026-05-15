"""Reconstrucție 1a-oferta.docx v2 — Oferta SIDISVA
Variantă cu COPERTĂ ELEGANTĂ — pagina 1 dedicată, page break, cuprins pe pagină separată.
Designul respectă convenții de copertă pentru propuneri tehnice complexe.

Diferențe vs v1:
- Pagina 1 = COPERTĂ aerată (titlu mare 36pt centrat, info esențial în casetă, branding minimal)
- Page break după copertă
- Pagina 2-3 = Cuprins (capitole + anexe)
- Page break
- Pagina 4+ = Detalii (identificare proiect, factori evaluare, strategie, IP, depunere, semnătură)
- Culori: albastru închus pentru titluri (#1F3864 — Office Dark Blue)
- Bordere subtile pe tabele
- Paragraph spacing generos
- Tabel "Identificare proiect" cu fundal subtil pe coloana stânga
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

# Culori brand
COLOR_PRIMARY = RGBColor(0x1F, 0x38, 0x64)   # Office Dark Blue
COLOR_ACCENT = RGBColor(0x2E, 0x75, 0xB6)    # Office Medium Blue
COLOR_GRAY = RGBColor(0x59, 0x59, 0x59)      # text secundar
COLOR_LIGHT_GRAY = RGBColor(0xBF, 0xBF, 0xBF)


def shade_cell(cell, hex_color):
    """Adaugă fundal colorat unei celule de tabel."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    """Setează borderele unei celule."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for direction, val in [('top', top), ('bottom', bottom),
                            ('left', left), ('right', right)]:
        if val:
            border = OxmlElement(f'w:{direction}')
            border.set(qn('w:val'), val.get('val', 'single'))
            border.set(qn('w:sz'), str(val.get('sz', 4)))
            border.set(qn('w:color'), val.get('color', 'auto'))
            tcBorders.append(border)
    tc_pr.append(tcBorders)


def add_horizontal_line(doc, color_hex='1F3864', size=12):
    """Adaugă o linie orizontală decorativă."""
    p = doc.add_paragraph()
    p_pr = p._element.get_or_add_pPr()
    p_borders = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:color'), color_hex)
    p_borders.append(bottom)
    p_pr.append(p_borders)
    return p


def add_centered(doc, text, size=11, bold=False, italic=False,
                  color=None, space_before=0, space_after=0, uppercase=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    r = p.add_run(text.upper() if uppercase else text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = color
    return p


def add_para(doc, text, size=11, bold=False, italic=False, color=None,
              space_after=6, justify=False):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = color
    return p


def add_bullet(doc, text, size=11):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    r.font.size = Pt(size)
    return p


def add_heading_styled(doc, text, level=1, color=None, size=None):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in h.runs:
        if color:
            r.font.color.rgb = color
        if size:
            r.font.size = Pt(size)
    return h


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_table_styled(doc, headers, rows, widths_cm=None, header_bg='1F3864',
                      header_color=RGBColor(0xFF, 0xFF, 0xFF),
                      first_col_bold=False, first_col_bg=None,
                      style='Light List Accent 1'):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = style
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header — cu fundal colorat
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        shade_cell(c, header_bg)
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
                r.font.color.rgb = header_color

    # Data
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            c = t.rows[ri].cells[ci]
            c.text = str(val)
            if first_col_bg and ci == 0:
                shade_cell(c, first_col_bg)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    if first_col_bold and ci == 0:
                        r.bold = True
    if widths_cm:
        for row in t.rows:
            for ci, w in enumerate(widths_cm):
                row.cells[ci].width = Cm(w)
    return t


# ============ INIT DOCUMENT ============
doc = Document()

for section in doc.sections:
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

# ============================================================================
# PAGINA 1 — COPERTĂ
# ============================================================================

# Spațiu vertical de sus
for _ in range(3):
    doc.add_paragraph()

# Eticheta superior — etichetă mică, restrânsă
add_centered(doc, 'PROPUNERE TEHNICĂ', size=12, bold=True,
              color=COLOR_GRAY, uppercase=True, space_after=4)

# Cod proiect
add_centered(doc, 'Cod SMIS 336342 · Anunț SEAP CN1089237',
              size=10, italic=True, color=COLOR_GRAY, space_after=24)

# Titlu principal
add_centered(doc, 'SIDISVA', size=42, bold=True,
              color=COLOR_PRIMARY, space_after=6)

# Sub-titlu
add_centered(doc, 'Sistem Informatic Digitalizat în domeniul',
              size=14, italic=True, color=COLOR_ACCENT, space_after=2)
add_centered(doc, 'Sanitar-Veterinar și pentru Siguranța Alimentelor',
              size=14, italic=True, color=COLOR_ACCENT, space_after=24)

# Linie decorativă
add_horizontal_line(doc, color_hex='1F3864', size=18)
doc.add_paragraph()

# Casetă info — tabel 1×1 pentru caseta cu margini
caseta = doc.add_table(rows=1, cols=1)
caseta.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = caseta.rows[0].cells[0]
cell.width = Cm(14)
shade_cell(cell, 'F2F2F2')

# Adaugă conținut în casetă
caseta_para = cell.paragraphs[0]
caseta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf = caseta_para.paragraph_format
pf.space_before = Pt(12)
pf.space_after = Pt(6)

r = caseta_para.add_run('AUTORITATE CONTRACTANTĂ')
r.font.size = Pt(10)
r.font.bold = True
r.font.color.rgb = COLOR_GRAY

p2 = cell.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p2.add_run('ANSVSA')
r.font.size = Pt(16)
r.font.bold = True
r.font.color.rgb = COLOR_PRIMARY

p3 = cell.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p3.add_run('Autoritatea Națională Sanitară Veterinară\nși pentru Siguranța Alimentelor')
r.font.size = Pt(11)
r.font.italic = True
r.font.color.rgb = COLOR_GRAY

# Separator în casetă
p4 = cell.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p4.add_run('• • •')
r.font.size = Pt(10)
r.font.color.rgb = COLOR_LIGHT_GRAY

# Beneficiari
p5 = cell.add_paragraph()
p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p5.add_run('+ 3 institute subordonate (IISPV, ICBMV, IDSA)\n+ 42 DSVSA județene')
r.font.size = Pt(10)
r.font.color.rgb = COLOR_GRAY

# Termen
p6 = cell.add_paragraph()
p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf = p6.paragraph_format
pf.space_before = Pt(12)
pf.space_after = Pt(12)
r = p6.add_run('Termen depunere: ')
r.font.size = Pt(11)
r.font.color.rgb = COLOR_GRAY
r = p6.add_run('21.05.2026, ora 15:00')
r.font.size = Pt(11)
r.font.bold = True
r.font.color.rgb = COLOR_PRIMARY

# Spațiu după casetă
doc.add_paragraph()
add_horizontal_line(doc, color_hex='1F3864', size=18)
doc.add_paragraph()

# Ofertant
add_centered(doc, 'OFERTANT PRINCIPAL', size=10, bold=True,
              color=COLOR_GRAY, uppercase=True, space_after=4)
add_centered(doc, '<LIDER>', size=22, bold=True,
              color=COLOR_PRIMARY, space_after=4)
add_centered(doc, 'în asociere cu consorțiul de parteneri specializați',
              size=10, italic=True, color=COLOR_GRAY, space_after=18)

# Footer copertă
for _ in range(2):
    doc.add_paragraph()

add_centered(doc, 'Finanțat prin Programul Operațional Creștere Inteligentă,',
              size=9, italic=True, color=COLOR_GRAY, space_after=0)
add_centered(doc, 'Digitalizare și Instrumente Financiare (POCIDIF)',
              size=9, italic=True, color=COLOR_GRAY, space_after=2)
add_centered(doc, 'Caietul de Sarcini nr. 424/SCPI/29.12.2025 ; 7574/CP/2025 — Revizia 1',
              size=8, italic=True, color=COLOR_LIGHT_GRAY, space_after=0)

# Page break după copertă
page_break(doc)

# ============================================================================
# PAGINA 2-3 — CUPRINS
# ============================================================================

add_heading_styled(doc, 'Cuprins', level=1, color=COLOR_PRIMARY, size=24)
doc.add_paragraph()

# --- Cuprins capitole ---
add_heading_styled(doc, 'Capitole propunere tehnică', level=2,
                    color=COLOR_ACCENT, size=14)

add_para(doc,
    'Cele 16 capitole numerotate ale ofertei tehnice, fiecare livrat ca fișier .docx '
    'independent, semnat electronic calificat (eIDAS QES) de reprezentantul legal al '
    '<LIDER>.',
    size=10, italic=True, color=COLOR_GRAY, space_after=12)

cap_rows = [
    ('1', 'Rezumat executiv', '1-Rezumat_executiv.docx', '—'),
    ('1A', 'Prezentarea Ofertantului', '1A-Prezentare_ofertant.docx', 'Eligibilitate'),
    ('2', 'Abordare și metodologie', '2-Abordare_metodologie.docx', 'P3.1 — 10p'),
    ('3', 'Plan de implementare', '3-Plan_implementare.docx', 'P3 — 20p'),
    ('4', 'Descrierea soluției — 14 componente', '4-Descrierea_solutiei.docx',
     'Conformitate'),
    ('5', 'Arhitectura + licențe', '5-Arhitectura_si_licente.docx',
     'Conformitate + IP'),
    ('6', 'Lista hardware', '6-Lista_hardware.docx', 'P4.1 + plafon 20%'),
    ('7', 'Plan garanție (3 ani)', '7-Plan_garantie.docx', 'Conformitate'),
    ('8', 'Conformitate specificații (Anexa F)', '8-Conformitate_specificatii.docx',
     'Eligibilitate'),
    ('9', 'Echipa de proiect (20 experți)', '9-Echipa_proiect.docx', 'P2 — 30p'),
    ('10', 'Securitate informatică', '10-Securitate_informatica.docx',
     'P3.1 + plafon 10%'),
    ('11', 'Măsuri DNSH', '11-DNSH.docx', 'P4 — 10p'),
    ('12', 'Plan instruire (147 utilizatori)', '12-Plan_instruire.docx', 'Conformitate'),
    ('13', 'Management contract', '13-Management_contract.docx', 'P3.2'),
    ('14', 'DEMO video (33 cerințe)', '14-DEMO_video.docx', 'ELIMINATORIE'),
    ('15', 'Declarații obligatorii (15)', '15-Declaratii_obligatorii.docx',
     'Eligibilitate'),
    ('16', 'Anexe (index A-K)', '16-Anexe.docx', 'Suport'),
]

add_table_styled(doc,
    ['Cap.', 'Denumire', 'Fișier', 'Rol / Punctaj'],
    cap_rows,
    widths_cm=[1.0, 7.5, 5.5, 3.0],
    header_bg='1F3864',
    first_col_bold=True, first_col_bg='D9E2F3'
)

doc.add_paragraph()

# --- Cuprins anexe ---
add_heading_styled(doc, 'Anexe justificative', level=2,
                    color=COLOR_ACCENT, size=14)

add_para(doc,
    'Cele 11 anexe justificative ale ofertei tehnice. Detalii integrale despre fiecare '
    'anexă în Cap. 16 — Anexe.',
    size=10, italic=True, color=COLOR_GRAY, space_after=12)

anexe_rows = [
    ('A', 'Lista completă licențe software (27 produse)',
     'Cap. 5 + xlsx'),
    ('B', 'Lista hardware + fișe tehnice', 'Cap. 6'),
    ('C', 'Diagrame arhitecturale', 'Cap. 5 + Cap. 10'),
    ('D', 'Plan implementare + Gantt', 'Cap. 3'),
    ('E', 'CV-uri experți + 5 proiecte/expert', 'Cap. 9'),
    ('F', 'Matricea de Conformitate (1.294 cerințe)',
     'anexa_f_conformitate.docx'),
    ('G', 'Plan detaliat de instruire (147 utilizatori)', 'Cap. 12'),
    ('H', 'BCP + DRP', 'Cap. 7 + Cap. 10'),
    ('I', 'Documente justificative DNSH', 'Cap. 11'),
    ('J', 'Video demonstrativ DEMO', 'Cap. 14 — depus separat SEAP'),
    ('K', 'Pachet 15 declarații eIDAS QES', 'Cap. 15'),
]

add_table_styled(doc,
    ['Anexa', 'Conținut', 'Sursa / Sec. corelată'],
    anexe_rows,
    widths_cm=[1.5, 9.5, 6.0],
    header_bg='1F3864',
    first_col_bold=True, first_col_bg='D9E2F3'
)

page_break(doc)

# ============================================================================
# PAGINA 3+ — IDENTIFICARE PROIECT + DETALII
# ============================================================================

add_heading_styled(doc, '1. Identificare proiect și ofertant', level=1,
                    color=COLOR_PRIMARY, size=20)
doc.add_paragraph()

id_rows = [
    ('Autoritate Contractantă', 'ANSVSA — Autoritatea Națională Sanitară Veterinară '
     'și pentru Siguranța Alimentelor'),
    ('Cod SMIS', '336342'),
    ('Sursa de finanțare', 'POCIDIF (Programul Operațional Creștere Inteligentă, '
     'Digitalizare și Instrumente Financiare) — P2, OS 2.2.1 e-guvernare'),
    ('Anunț SEAP', 'CN1089237'),
    ('Caiet de Sarcini', 'Nr. 424/SCPI/29.12.2025 ; 7574/CP/2025 — Revizia 1 (252 pagini)'),
    ('Termen depunere', '21.05.2026, ora 15:00'),
    ('Valoare estimată max (fără TVA)', '85.418.857,53 lei'),
    ('Plafoane buget', 'max 20% pentru HW + servicii instalare/configurare;\n'
     'min 10% pentru securitate cibernetică'),
    ('Durata contractului', '18 luni implementare + 36 luni garanție post-implementare'),
    ('Găzduire (Cap. 3.4.3.1 CdS)', 'Cloud Guvernamental — arhitectură Cloud-Native + '
     'containerizare (Docker / Kubernetes)'),
    ('Criteriu atribuire', 'Cel mai bun raport calitate-preț — 60% tehnică / 40% preț'),
    ('Beneficiari', 'ANSVSA + 3 institute (IISPV, ICBMV, IDSA) + 42 DSVSA județene'),
    ('Scară utilizatori', '~185.000 utilizatori unici/an Portal\n'
     '~5.300 angajați + 2.600 medici veterinari + 4.800 utilizatori acreditați'),
]

add_table_styled(doc,
    ['Aspect', 'Valoare'],
    id_rows,
    widths_cm=[5.5, 11.0],
    header_bg='1F3864',
    first_col_bold=True, first_col_bg='F2F2F2'
)

doc.add_paragraph()

# --- Ofertant ---
add_heading_styled(doc, '2. Ofertant și consorțiu', level=1,
                    color=COLOR_PRIMARY, size=20)
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
r = p.add_run('Ofertant principal: ')
r.bold = True
r.font.size = Pt(11)
r = p.add_run('<LIDER>')
r.font.size = Pt(11)
r.font.color.rgb = COLOR_ACCENT
r.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
r = p.add_run('Structura consorțiului: ')
r.bold = True
r.font.size = Pt(11)
p.add_run('asociere cu lider + parteneri specializați pe componente — '
          '<PARTENER 1>, <PARTENER 2>, … <PARTENER N>. Detaliile complete în '
          'Cap. 1A — Prezentarea Ofertantului.').font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
r = p.add_run('Soluția propusă: ')
r.bold = True
r.font.size = Pt(11)
p.add_run(
    'sistem informatic integrat compus din 14 componente, realizat prin combinarea de '
    'produse software specializate de la furnizori multipli — ZIPPER DMS pentru '
    'managementul documentelor, VOGO Enterprise Suite pentru Portal/Chatbot/App mobilă, '
    'Microsoft SQL Server Enterprise + Oracle Service Bus + Keycloak Enterprise + '
    'Power BI pentru infrastructura SW, producători specializați pentru LIMS și GIS. '
    'Stack-ul integral în Cap. 5 + Anexa A.'
).font.size = Pt(11)

page_break(doc)

# ============ §3 Factori de evaluare ============

add_heading_styled(doc, '3. Maparea capitolelor → factori de evaluare', level=1,
                    color=COLOR_PRIMARY, size=20)
doc.add_paragraph()

add_para(doc,
    'Criteriul de atribuire: „Cel mai bun raport calitate-preț" (60% tehnică / 40% '
    'preț). Cei 4 factori de evaluare cumulează 100 puncte:',
    size=11, italic=False, justify=True, space_after=12)

factori_rows = [
    ('Factor 1 — Prețul ofertei', '40p',
     'P(n) = (Preț_min / Preț_n) × 40',
     'Propunere financiară (separată)'),
    ('Factor 2 — Experiența celor 6 experți cheie evaluați', '30p',
     '6 × 5p; țintă ≥5 proiecte/expert (cu ≥7 module + ≥1 portal)',
     'Cap. 9 + Anexa E'),
    ('Factor 3 — Metodologia (subfactori 3.1 + 3.2)', '20p (10+10)',
     'Excepțional/Adecvat/Acceptabil = 10/5/1; țintă Excepțional',
     'Cap. 2 + Cap. 3 + Cap. 13'),
    ('Factor 4 — DNSH (subfactori 4.1 + 4.2)', '10p (5+5)',
     '4.1 consum laptop < 20Wh; 4.2 ambalaje reciclabile + flotă verde',
     'Cap. 11 + Anexa I'),
    ('TOTAL', '100p', '', ''),
]

add_table_styled(doc,
    ['Factor', 'Punctaj', 'Algoritm / Țintă', 'Secțiune ofertă'],
    factori_rows,
    widths_cm=[5.0, 1.5, 6.5, 3.5],
    header_bg='1F3864',
    first_col_bold=True
)

doc.add_paragraph()

# Caseta CRITERIU ELIMINATOR
elim_table = doc.add_table(rows=1, cols=1)
elim_table.alignment = WD_TABLE_ALIGNMENT.CENTER
ec = elim_table.rows[0].cells[0]
ec.width = Cm(15)
shade_cell(ec, 'FFE699')  # galben atenție

ep = ec.paragraphs[0]
ep.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf = ep.paragraph_format
pf.space_before = Pt(8)
pf.space_after = Pt(2)
r = ep.add_run('⚠  CERINȚĂ ELIMINATORIE  ⚠')
r.font.size = Pt(11)
r.font.bold = True
r.font.color.rgb = COLOR_PRIMARY

ep2 = ec.add_paragraph()
ep2.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf = ep2.paragraph_format
pf.space_after = Pt(8)
r = ep2.add_run(
    'DEMO video conform Cap. 14 CdS — orice cerință nedemonstrată '
    'din cele 33 listate ⇒ ofertă neconformă (eliminată din evaluare, '
    'INDIFERENT de scorul tehnic/financiar).'
)
r.font.size = Pt(10)
r.font.color.rgb = COLOR_PRIMARY

doc.add_paragraph()

# ============ §4 Strategie ofertă ============

add_heading_styled(doc, '4. Strategia ofertei tehnice', level=1,
                    color=COLOR_PRIMARY, size=20)
doc.add_paragraph()

add_para(doc,
    'Oferta tehnică <LIDER> respectă structura canonică a ofertelor tehnice complexe, '
    'optimizată pentru facilitarea evaluării de către comisia ANSVSA:',
    size=11, justify=True, space_after=10)

bullets = [
    'Capitolul 1 (Rezumat executiv) — gateway pentru evaluator: identitatea ofertantului, '
    'sinteza ofertei, win themes, scor țintă pe P1-P4.',

    'Capitolul 1A (Prezentarea Ofertantului) — credibilitate: profil <LIDER> + parteneri, '
    'referințe, certificări ISO 9001/14001/27001/22301.',

    'Capitolele 2-3 (Abordare/Metodologie + Plan implementare) — răspuns la Factorul 3 '
    '(20p) — metodologie hibridă PRINCE2 + Scrum Agile, ISO 21500, ITIL, plan 18 luni '
    'cu drum critic + registru riscuri 7+6.',

    'Capitolele 4-6 (Soluție + Arhitectură + Hardware) — răspuns la cerințele tehnice '
    'IV.4.1 lit. a)-c) din Fișa de date: 14 componente, arhitectură Cloud-Native pe '
    'Cloud Guvernamental, 27 licențe SW, ~536 echipamente HW.',

    'Capitolul 7 (Plan garanție) — răspuns la IV.4.1 lit. d) + cap. 3.4.9 CdS: 3 ani '
    'garanție, helpdesk L-V 08:00-17:00 + SLA 24×7 doar pentru incidente S1 critice; '
    'integrare ulterioară mock-up sisteme guvernamentale FĂRĂ cost suplimentar.',

    'Capitolul 8 + Anexa F (Conformitate) — răspuns la IV.4.1 lit. e): matrice cu '
    'toate cele 1.294 cerințe + răspuns punct-cu-punct.',

    'Capitolul 9 (Echipa) — răspuns la Factorul 2 (30p) — 8 experți cheie (din care '
    '6 evaluați) + 12 non-cheie = 20 experți; cert. producător explicite (ZIPPER pentru '
    'Arhitect, Microsoft SQL pentru BD, Fortinet/Palo Alto pentru Sec).',

    'Capitolele 10-11 (Securitate + DNSH) — conformitate NIS2 (Dir. UE 2022/2555 + OUG '
    '155/2024), Lege 354/2022 (anti-RU), GDPR, defense-in-depth pe 7 straturi + Factor 4 '
    'DNSH (10p).',

    'Capitolele 12-13 (Instruire + Management contract) — 147 utilizatori instruiți '
    '(100 cheie ONLINE + 44 medici vet ONLINE + 3 admin FIZIC) + rapoarte trimestriale '
    '+ KPI calitate + plată SPV 30z.',

    'Capitolele 14-15 (DEMO + Declarații) — livrabile obligatorii: video 33 cerințe '
    '(eliminator) + 15 declarații semnate eIDAS QES.',

    'Capitolul 16 (Anexe) — index complet A-K cu corelare la secțiuni.',
]
for b in bullets:
    add_bullet(doc, b, size=11)

page_break(doc)

# ============ §5 Acceptare cap. 12 IP ============

add_heading_styled(doc, '5. Acceptare expresă cap. 12 CdS — Drepturi IP',
                    level=1, color=COLOR_PRIMARY, size=20)
doc.add_paragraph()

add_para(doc,
    'Conform Cap. 12 din Caietul de Sarcini, <LIDER> ACCEPTĂ EXPRES și fără rezerve '
    'condițiile privind drepturile de proprietate intelectuală:',
    size=11, justify=True, space_after=10)

ip_bullets = [
    'Toate dezvoltările / customizările realizate în cadrul contractului trec în '
    'proprietatea Beneficiarului (ANSVSA).',
    'Licențele pentru componentele aplicative dezvoltate / customizate sunt PERPETUE.',
    'Codul sursă INTEGRAL pentru componentele aplicative dezvoltate / customizate '
    'este predat Beneficiarului.',
    'IP-ul produselor COTS preexistente (ZIPPER DMS, VOGO Enterprise Suite, soluția '
    'LIMS) rămâne la producătorii respectivi, dar Beneficiarul primește licență '
    'perpetuă + cod sursă integral conform cap. 3.4.3.3.4 CdS.',
]
for b in ip_bullets:
    add_bullet(doc, b, size=11)

p = doc.add_paragraph()
pf = p.paragraph_format
pf.space_before = Pt(10)
r = p.add_run('Confirmare formală: ')
r.bold = True
r.font.size = Pt(11)
p.add_run(
    'această acceptare este reluată semnată în Cap. 15 — Declarații obligatorii și în '
    'Anexa K.2 (declarație separată cu semnătură eIDAS QES a reprezentantului legal).'
).font.size = Pt(11)

doc.add_paragraph()

# ============ §6 Depunere SEAP ============

add_heading_styled(doc, '6. Depunere în SEAP', level=1,
                    color=COLOR_PRIMARY, size=20)
doc.add_paragraph()

add_para(doc,
    'Acest fișier `1a-oferta.docx` servește ca document de consolidare și navigare '
    'pentru întreaga ofertă tehnică <LIDER>. Fiecare capitol este un fișier .docx '
    'independent, format pentru a fi semnat electronic individual.',
    size=11, justify=True, space_after=10)

add_para(doc, 'Pentru depunerea în SEAP, se utilizează următoarea structură:',
          size=11, space_after=6)

depune_bullets = [
    'Plicul „Propunere Tehnică" — pachet complet: `1a-oferta.docx` (scrisoare înaintare + '
    'cuprins) + cele 16 capitole + `anexa_f_conformitate.docx` + Anexele A-I + K (PDF-uri).',
    'Plicul „Propunere Financiară" — separat, conform Anexa Financiară.',
    'Plicul „DUAE și documente eligibilitate" — DUAE + Anexa K.',
    'Anexa J (Video DEMO) — fișier separat în SEAP, link cross-referențiat în Anexa J.docx.',
]
for b in depune_bullets:
    add_bullet(doc, b, size=11)

page_break(doc)

# ============ §7 Semnătură ============

add_heading_styled(doc, '7. Semnătură reprezentant legal',
                    level=1, color=COLOR_PRIMARY, size=20)
doc.add_paragraph()

add_para(doc,
    'Toate documentele componente ale ofertei sunt semnate electronic calificat '
    '(eIDAS QES) de reprezentantul legal al <LIDER>, conform:',
    size=11, justify=True, space_after=8)

sem_bullets = [
    'Legea nr. 455/2001 — privind semnătura electronică (RO).',
    'Regulamentul (UE) 910/2014 — eIDAS, partea III — semnături electronice calificate.',
    'Furnizor de servicii de încredere acreditat (TSP): [De completat: certSIGN / '
    'DigiSign / Trans Sped — număr certificat QES valid].',
]
for b in sem_bullets:
    add_bullet(doc, b, size=11)

doc.add_paragraph()
doc.add_paragraph()

# Bloc semnătură frumos formatat
sig_table = doc.add_table(rows=4, cols=2)
sig_table.style = 'Table Grid'
sig_table.alignment = WD_TABLE_ALIGNMENT.LEFT

sig_data = [
    ('Data depunere:', '[zz/05/2026]'),
    ('Ofertant:', '<LIDER>'),
    ('Reprezentant legal:', '[Nume Prenume] — [Funcție]'),
    ('Semnătură (eIDAS QES):', 'conform Legii 455/2001 + Reg. UE 910/2014'),
]

for ri, (label, val) in enumerate(sig_data):
    c0 = sig_table.rows[ri].cells[0]
    c1 = sig_table.rows[ri].cells[1]
    c0.text = label
    c1.text = val
    c0.width = Cm(5.0)
    c1.width = Cm(11.0)
    shade_cell(c0, 'F2F2F2')
    for p in c0.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(10)
    for p in c1.paragraphs:
        for r in p.runs:
            r.font.size = Pt(11)
            r.italic = (ri == 3)

# ---- Save ----
doc.save(r'1a-oferta.docx')

# Verificare
from docx import Document
d2 = Document(r'1a-oferta.docx')
print(f'OK — 1a-oferta.docx v2 scris: {len(d2.paragraphs)} paragrafe, {len(d2.tables)} tabele')

texts = []
for p in d2.paragraphs:
    texts.append(p.text)
for t in d2.tables:
    for row in t.rows:
        for cell in row.cells:
            texts.append(cell.text)
full = '\n'.join(texts)

print(f'  VOGO TECHNOLOGY              : {full.count("VOGO TECHNOLOGY")}')
print(f'  <LIDER>                      : {full.count("<LIDER>")}')
print(f'  Cod SMIS 336342              : {full.count("336342")}')
print(f'  ELIMINATORIE                 : {full.count("ELIMINATORIE")}')
print(f'  Page breaks                  : ~{full.count(chr(12))} (manual count of break chars)')
print(f'  eIDAS QES                    : {full.count("eIDAS QES")}')

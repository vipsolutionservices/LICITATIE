"""Reconstrucție 16-Anexe.docx — Oferta SIDISVA
Index complet al anexelor depuse cu oferta tehnică.

Corecturi față de varianta veche:
- 0× VOGO TECHNOLOGY S.R.L. → `<LIDER>`
- Tabel sintetic structurat cu referință la secțiunea ofertei care folosește fiecare anexă
- Subsecțiune dedicată per anexă cu: conținut, format, dimensiune estimată, secțiunea sursă, mod prezentare
- Aliniat cu sec 8 (Anexa F = 1.294 cerințe), sec 9 (Anexa E = CV experți), sec 11 (Anexa I = DNSH),
  sec 12 (Anexa G = instruire), sec 14 (Anexa J = video DEMO), sec 15 (Anexa K = 15 declarații)
- Reguli formale de organizare + semnătură eIDAS QES
- Confirmare reprezentant legal curățată
"""
from docx import Document
from docx.shared import Pt, Cm

doc = Document()

for section in doc.sections:
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)


def add_bullet(text):
    return doc.add_paragraph(text, style='List Bullet')

def add_para(text):
    return doc.add_paragraph(text)

def add_table(headers, rows, widths_cm=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            c = t.rows[ri].cells[ci]
            c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    if widths_cm:
        for row in t.rows:
            for ci, w in enumerate(widths_cm):
                row.cells[ci].width = Cm(w)
    return t


# ============ HEADING ============
doc.add_heading('16. Anexe la propunerea tehnică', level=1)

intro = doc.add_paragraph()
intro.add_run(
    'Capitolul prezintă lista completă a anexelor care însoțesc Propunerea Tehnică depusă '
    'de <LIDER>, în calitate de Ofertant principal, pentru contractul SIDISVA '
    '(Caietul de Sarcini nr. 424/SCPI/29.12.2025 ; 7574/CP/2025).'
)

intro2 = doc.add_paragraph()
intro2.add_run(
    'Convenția de organizare a anexelor: identificate cu litere de la A la K, fiecare '
    'corelată cu o secțiune principală a ofertei. Anexele sunt depuse în format electronic '
    'odată cu oferta în SEAP, semnate digital cu semnătură electronică extinsă calificată '
    '(eIDAS QES) a reprezentantului legal al <LIDER>.'
)

# ============ §16.1 Tabel sintetic ============
doc.add_heading('16.1 Tabel sintetic — lista anexelor depuse', level=2)

add_para('Cele 11 anexe ale ofertei tehnice, ordinea lor și corelarea cu secțiunile ofertei:')

# Tabel sumar
add_table(
    ['Anexa', 'Denumire', 'Secțiune ofertă corelată', 'Dim. estimată', 'Format livrare'],
    [
        ['A', 'Lista completă a licențelor software',
         'Cap. 5 — Arhitectură și licențe', '~30 pag.',
         'PDF + Excel (`Lista_Software_SIDISVA.xlsx`)'],
        ['B', 'Lista hardware + fișe tehnice producători',
         'Cap. 6 — Lista hardware', '~150 pag.',
         'PDF (fișe tehnice + datasheet-uri originale)'],
        ['C', 'Diagrame arhitecturale (logică, fizică, securitate)',
         'Cap. 5 — Arhitectură; Cap. 10 — Securitate', '~25 pag.',
         'PDF (Visio/draw.io export, A3)'],
        ['D', 'Plan detaliat de implementare cu Gantt',
         'Cap. 3 — Plan de implementare', '~20 pag.',
         'PDF + Microsoft Project (`.mpp`)'],
        ['E', 'CV-uri experți cheie + dovezi proiecte similare',
         'Cap. 9 — Echipa de proiect', '~200-250 pag.',
         'PDF (CV Europass + diplome + certificări + PV recepție)'],
        ['F', 'Matricea de Conformitate (1.294 cerințe)',
         'Cap. 8 — Conformitate cu specificațiile', '~300 pag.',
         'PDF + Excel (export `anexa_f_conformitate.docx`)'],
        ['G', 'Plan detaliat de instruire (147 utilizatori)',
         'Cap. 12 — Plan de instruire', '~40 pag.',
         'PDF (curricula + planificare sesiuni + materiale)'],
        ['H', 'Plan de continuitate operațională (BCP) și DR (DRP)',
         'Cap. 7 — Plan garanție; Cap. 10 — Securitate', '~50 pag.',
         'PDF (RTO/RPO + procedurile de failover + DR drills)'],
        ['I', 'Documente justificative DNSH',
         'Cap. 11 — DNSH', '~30 pag.',
         'PDF (fișă laptop Energy Star + EU Ecolabel + raport amprentă)'],
        ['J', 'Video demonstrativ DEMO (33 cerințe eliminatorii)',
         'Cap. 14 — DEMO video', '~30-60 min. video',
         'Fișier video MP4 — depus separat în SEAP (referință cross-link)'],
        ['K', 'Pachet declarații semnate (15 declarații obligatorii)',
         'Cap. 15 — Declarații obligatorii', '~80-100 pag.',
         'PDF (fiecare declarație semnată eIDAS QES separat)'],
    ],
    widths_cm=[1.0, 4.5, 4.0, 1.7, 5.0]
)

p = doc.add_paragraph()
r = p.add_run('Dimensiune totală estimată: ')
r.bold = True
p.add_run(
    '~950-1.000 pagini PDF + 30-60 min. video DEMO + 4 fișiere structurate '
    '(`.xlsx`, `.mpp`, `anexa_f_conformitate.docx`, `Lista_Software_SIDISVA.xlsx`). '
    'Toate fișierele sunt semnate digital cu semnătură electronică extinsă calificată '
    '(eIDAS QES) conform Regulamentului UE 910/2014 + Legii nr. 455/2001.'
)

# ============ §16.2 Detalii pe fiecare anexă ============
doc.add_heading('16.2 Detalii per anexă', level=2)

# Anexa A
doc.add_heading('16.2.1 Anexa A — Lista completă licențe software', level=3)
add_para(
    'Conține toate licențele software ofertate, structurate pe 4 categorii (A. Infrastructură SW, '
    'B. Software aplicativ, C. Securitate appliance SW, D. Productivity/Client) — total 27 '
    'produse software identificate. Pentru fiecare licență: producător, ediție/versiune, '
    'cantitate, unitate de licențiere (core/user/server/throughput), tip (perpetuă/abonament), '
    'condiții cod sursă (acolo unde se aplică, conform cap. 12 CdS — IP).'
)
add_bullet('Sursa adevărului: `Lista_Software_SIDISVA.xlsx` (Sinteza + Cantitati explicite CdS).')
add_bullet('Producătorii listați respectă Lege 354/2022 — toți cu origine UE/SUA/IL/CA/AU/UK '
           '(NU RU/CN/IR/KP).')
add_bullet('Acoperă plafonul max 20% buget HW+instalare + min 10% securitate cibernetică.')

# Anexa B
doc.add_heading('16.2.2 Anexa B — Lista hardware + fișe tehnice', level=3)
add_para(
    'Conține lista detaliată a tuturor echipamentelor hardware ofertate, cu fișe tehnice '
    'integrale de la producători. Acoperă: 100 laptopuri DNSH-compliant, 336 terminale teren '
    '(8×42 DSVSA) cu pad semnătură + imprimantă mobilă coduri bare, echipamente securitate '
    'centru DC (NGFW × 2, WAF × 2, SIEM cluster, Honeypot, NMS/NAC, Email Security), '
    '90 NGFW locații (FortiGate 100F), switching, server-e pentru cloud privat.'
)
add_bullet('Conformitate DNSH: certificat Energy Star 8.0 + EU Ecolabel pe laptopuri.')
add_bullet('Conformitate L 354/2022: producători verificați (Dell, HP, Lenovo, Fortinet, Cisco, '
           'F5, Splunk, Wacom).')

# Anexa C
doc.add_heading('16.2.3 Anexa C — Diagrame arhitecturale', level=3)
add_para(
    'Conține diagramele logice, fizice și de securitate ale soluției SIDISVA:'
)
add_bullet('Diagramă arhitectură logică — cele 14 componente + relațiile dintre ele (ESB backbone).')
add_bullet('Diagramă arhitectură fizică — distribuție pe Cloud Guvernamental + 42 DSVSA + '
           '3 Institute + sediul ANSVSA.')
add_bullet('Diagramă securitate — 7 straturi defense-in-depth (perimetru/rețea/aplicație/'
           'date/identitate/monitorizare/IR).')
add_bullet('Diagramă integrări externe — ROeID, eIDAS, PNI, PCUe/PDURo, ONRC, APIA, ANCPI, '
           'ANARZ, CMV, Ghișeul.ro.')
add_bullet('Diagramă fluxuri date — sursa → ETL → DW → BI pentru raportare.')
add_bullet('Format: PDF A3 + sursa Visio/draw.io.')

# Anexa D
doc.add_heading('16.2.4 Anexa D — Plan implementare detaliat + Gantt', level=3)
add_para(
    'Conține planul detaliat de implementare a celor 18 luni de contract:'
)
add_bullet('Gantt chart cu toate activitățile, dependențele, drumul critic — Microsoft Project.')
add_bullet('Milestone-uri și jaloane intermediare (M1-M9).')
add_bullet('Alocare resurse umane per activitate.')
add_bullet('Paralelizare optimă pe componente independente.')
add_bullet('Termenul final luna 18 + perioada de garanție 3 ani (lunile 19-54).')

# Anexa E
doc.add_heading('16.2.5 Anexa E — CV-uri experți + dovezi proiecte', level=3)
add_para(
    'Conține documentele justificative pentru cei 8 experți cheie (Cap. 8 CdS + §9 ofertă). '
    'Structurată pe subsecțiuni E.1-E.17:'
)
add_bullet('E.1-E.8 — pentru fiecare expert cheie: CV Europass (semnat + datat), diplome '
           'studii, certificări profesionale + certificări producător (ZIPPER pentru Arhitect, '
           'Microsoft pentru BD, Fortinet/Palo Alto pentru Sec), 5 proiecte de referință cu '
           'PV recepție + certificate constatatoare + recomandări.')
add_bullet('E.9-E.16 — declarații de acces pentru cei 12 experți non-cheie + contracte '
           'externalizare (ex. DPO).')
add_bullet('E.17 — tabel sintetic de îndeplinire cerințe minime per expert (Cap. 8 CdS).')
add_bullet('Declarații de disponibilitate semnate pentru experții ne-angajați direct.')

# Anexa F
doc.add_heading('16.2.6 Anexa F — Matricea de Conformitate (1.294 cerințe)', level=3)
add_para(
    'Conține răspunsul punct-cu-punct la TOATE cele 1.294 cerințe extrase din Caietul de '
    'Sarcini, conform §8 din ofertă. Structurată pe 6 coloane:'
)
add_bullet('Nr. ordinal (1 → 1.294).')
add_bullet('Cap. CDS — capitolul exact din Caietul de Sarcini.')
add_bullet('Cerință (citat literal din CdS).')
add_bullet('Responsabil — partener consorțiu care livrează cerința.')
add_bullet('Răspuns ofertant — modul concret de îndeplinire (tehnologie, configurație, '
           'capacitate, dovadă).')
add_bullet('Document de referință — capitol ofertă / anexă tehnică / fișa producător.')
p = doc.add_paragraph()
r = p.add_run(
    'Sursa: `anexa_f_conformitate.docx` (1.365 rânduri totale, din care 71 sub-headers '
    'pentru navigare).'
)
r.italic = True

# Anexa G
doc.add_heading('16.2.7 Anexa G — Plan detaliat de instruire', level=3)
add_para(
    'Conține planul detaliat pentru cei 147 utilizatori instruiți conform Cap. 3.4.4.9 CdS '
    '(100 cheie + 44 medici vet + 3 admin), corelat cu §12 ofertă:'
)
add_bullet('Programare detaliată a celor 8 sesiuni minim (5 cheie online + 2 medici vet '
           'online + 1 admin fizic).')
add_bullet('Curricula completă pe categorii (administratori 24h, utilizatori cheie 16h, '
           'medici vet 4h).')
add_bullet('Mod livrare: ONLINE (cheie + medici vet) prin LMS / FIZIC (admin) la sediul ANSVSA.')
add_bullet('Materiale didactice — manuale, slide-uri, e-learning, tutoriale video.')
add_bullet('Plan evaluare — quiz online + scenarii practice + certificate participare.')
add_bullet('Resurse umane: 2 Experți instruire + 1 trainer admin + 1 coordonator + 2 designer '
           '+ 1 suport LMS.')
add_bullet('Strategie train-the-trainer pentru diseminare ulterioară către ~5.300 angajați.')

# Anexa H
doc.add_heading('16.2.8 Anexa H — BCP + DRP', level=3)
add_para(
    'Conține Planul de Continuitate Operațională (Business Continuity Plan) și Planul de '
    'Disaster Recovery (DR Plan) pentru SIDISVA, corelat cu §7 (Garanție) și §10 (Securitate):'
)
add_bullet('Analiza impactului asupra activității (BIA) — identificarea proceselor critice ANSVSA.')
add_bullet('RTO (Recovery Time Objective) și RPO (Recovery Point Objective) per componentă.')
add_bullet('Procedurile de failover automat + recuperare manuală.')
add_bullet('Strategie backup — frecvență, retenție, locații (off-site), test integritate.')
add_bullet('DR drills — exerciții semestriale de simulare dezastru.')
add_bullet('Plan de comunicare în criză — către ANSVSA, parteneri, autorități (CSIRT-RO).')
add_bullet('Conformitate NIS2 (OUG 155/2024) + GDPR art. 32 (rezistență).')

# Anexa I
doc.add_heading('16.2.9 Anexa I — Documente justificative DNSH', level=3)
add_para(
    'Conține documentele probatorii pentru Factorul de evaluare nr. 4 (DNSH, 10p), corelat '
    'cu §11 ofertă:'
)
add_bullet('Fișa tehnică laptop ofertat cu raport testare consum în stare de veghe < 20Wh '
           '(certificat Energy Star 8.0 + EU Ecolabel — pentru sub-factor 4.1, 5p).')
add_bullet('Specificații ambalaje — certificate reciclabilitate/reutilizabilitate (FSC, '
           'EU Ecolabel ambalaje).')
add_bullet('Dovada flotei electrice/hibride a transportatorului — contract de transport + '
           'lista vehicule + carbon footprint estimat (pentru sub-factor 4.2, 5p).')
add_bullet('Raport amprentă de carbon estimată pentru întregul contract (livrare + operare).')
add_bullet('Mapare 6 obiective Reg. UE 2020/852 → măsuri SIDISVA.')

# Anexa J
doc.add_heading('16.2.10 Anexa J — Video DEMO (33 cerințe eliminatorii)', level=3)
add_para(
    'Conține înregistrarea video demonstrativă a celor 33 cerințe eliminatorii din Cap. 14 CdS. '
    'CERINȚĂ ELIMINATORIE — orice cerință nedemonstrată = ofertă neconformă.'
)
add_bullet('Format: MP4 H.264 + audio AAC, rezoluție Full HD 1920×1080, ≥30 fps.')
add_bullet('Durată estimată: 30-60 minute (per planning §14 ofertă).')
add_bullet('Acoperă: ROeID + eIDAS + 2FA, chatbot AI, lucru offline, multi-tenant, API '
           'Swagger, edit Office direct în DMS, tipărire plicuri + borderouri coduri bare, '
           'modificare BPMN/UML drag-and-drop, semnătură olografă pad USB + digitală '
           'calificată, integrări guvernamentale, dashboard-uri drag-and-drop.')
add_bullet('Suplimentar: opțiune DEMO live la sediul ANSVSA (5 zile lucrătoare de la solicitare).')
add_bullet('Depunere: fișier separat în SEAP (link cross-referențiat în această anexă).')

# Anexa K
doc.add_heading('16.2.11 Anexa K — Pachet declarații (15 declarații obligatorii)', level=3)
add_para(
    'Conține setul complet al celor 15 declarații obligatorii prezentate în §15 al ofertei, '
    'fiecare semnată digital cu eIDAS QES de reprezentantul legal al <LIDER>:'
)
add_bullet('K.1 — DUAE (Documentul Unic European de Achiziții) completat.')
add_bullet('K.2 — Declarație acceptare condiții cap. 12 CdS (drepturi IP — perpetue + cod sursă).')
add_bullet('K.3 — Declarație conformitate Lege 354/2022 (anti-RU, anti-CN sancțiuni).')
add_bullet('K.4 — Declarație conformitate NIS2 (Dir. UE 2022/2555 + OUG 155/2024).')
add_bullet('K.5 — Declarație conformitate GDPR (Reg. UE 679/2016 + Lege 190/2018) + numire DPO.')
add_bullet('K.6 — Declarație neîncadrare art. 164, 165, 167 din Legea 98/2016 (motive excludere).')
add_bullet('K.7 — Declarație evitare conflict de interese (art. 60 Legea 98/2016).')
add_bullet('K.8 — Declarație beneficiar real (BO Register).')
add_bullet('K.9 — Declarație plată impozite și taxe la zi.')
add_bullet('K.10 — Garanție de participare (854.000 lei — 1% din valoarea estimată).')
add_bullet('K.11 — Declarație disponibilitate experți cheie (anexate semnături individuale).')
add_bullet('K.12 — Acord de asociere (consorțiu) sau declarație de subcontractare.')
add_bullet('K.13 — Lista subcontractanți + acord susținător terț (dacă e cazul).')
add_bullet('K.14 — Declarație DNSH (respectarea principiului „Do No Significant Harm").')
add_bullet('K.15 — Declarație confidențialitate și protecție date pe perioada contractului.')

# ============ §16.3 Reguli formale ============
doc.add_heading('16.3 Reguli formale de organizare a anexelor', level=2)

add_para('Toate anexele depuse de <LIDER> respectă următoarele reguli formale:')

add_bullet('Limba: română (excepție — datasheet-uri producători și certificări internaționale '
           'pot fi în engleză, cu traducere autorizată pentru cele esențiale).')
add_bullet('Format: PDF/A pentru documente; Microsoft Project (`.mpp`) pentru Anexa D; '
           'Excel pentru Anexa A + Anexa F (export); MP4 pentru Anexa J.')
add_bullet('Semnătură: fiecare document semnat individual cu semnătură electronică extinsă '
           'calificată (eIDAS QES) conform Reg. UE 910/2014 + Lege 455/2001.')
add_bullet('Numerotare: continuă în cadrul fiecărei anexe (A1, A2, …).')
add_bullet('Cuprins: fiecare anexă (cu excepția celor cu o singură pagină) are cuprins propriu.')
add_bullet('Trasabilitate: fiecare document referit în Propunerea Tehnică indică anexa și pagina.')
add_bullet('Depunere: în plicul electronic SEAP, secțiunea Propunere Tehnică + (Anexa J video) '
           'în plicul separat.')

# ============ §16.4 Confirmare finală ============
doc.add_heading('16.4 Confirmare finală a reprezentantului legal', level=2)

add_para(
    'Subsemnatul, în calitate de reprezentant legal al ofertantului <LIDER>, declar pe propria '
    'răspundere și sub sancțiunea excluderii din procedura de atribuire că:'
)

add_bullet('Am parcurs și am acceptat în totalitate condițiile Caietului de Sarcini nr. '
           '424/SCPI/29.12.2025 ; 7574/CP/2025 (revizia 1).')
add_bullet('Toate cele 11 anexe (A-K) menționate în §16.1-§16.2 sunt depuse împreună cu '
           'Propunerea Tehnică sau, după caz, în SEAP separat (Anexa J — video DEMO).')
add_bullet('Toate documentele anexate sunt complete, conforme cu cerințele și semnate '
           'digital cu eIDAS QES.')
add_bullet('Acceptăm expres condițiile cap. 12 CdS privind drepturile de proprietate '
           'intelectuală (licențe perpetue + cod sursă pentru componentele aplicative '
           'dezvoltate/customizate; IP-ul COTS preexistent rămâne la producători).')
add_bullet('Suntem de acord cu valoarea estimată maximă a contractului (85.418.857,53 lei '
           'fără TVA) și cu plafoanele procentuale (max 20% HW + servicii instalare, min '
           '10% securitate cibernetică).')

# Bloc semnătură
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Data: ').bold = True
p.add_run('[zz/05/2026]')

p = doc.add_paragraph()
p.add_run('Reprezentant legal <LIDER>: ').bold = True

p = doc.add_paragraph('[Nume Prenume] — [Funcție]')

p = doc.add_paragraph()
p.add_run('Semnătură electronică extinsă calificată (eIDAS QES): ').bold = True

p = doc.add_paragraph(
    'Semnată conform Legii nr. 455/2001 + Regulamentului UE 910/2014 (eIDAS), prin furnizor '
    'de servicii de încredere acreditat: [De completat: certSIGN / DigiSign / Trans Sped — '
    'Nume furnizor + Număr certificat QES].'
)
p.runs[0].italic = True

# ---- Save ----
doc.save(r'16-Anexe.docx')

# Verificare
from docx import Document
d2 = Document(r'16-Anexe.docx')
print(f'OK — 16-Anexe.docx scris: {len(d2.paragraphs)} paragrafe, {len(d2.tables)} tabele')

texts = []
for p in d2.paragraphs:
    texts.append(p.text)
for t in d2.tables:
    for row in t.rows:
        for cell in row.cells:
            texts.append(cell.text)
full = '\n'.join(texts)

print(f'  VOGO TECHNOLOGY              : {full.count("VOGO TECHNOLOGY")}')
print(f'  <LIDER>                      : {full.count("<LIDER>")}')
print(f'  Anexa A..K (lista)           : ' + ', '.join(f'{x}={full.count(f"Anexa {x}")}' for x in 'ABCDEFGHIJK'))
print(f'  eIDAS QES                    : {full.count("eIDAS QES")}')
print(f'  1.294 cerinte                : {full.count("1.294")}')
print(f'  147 utilizatori              : {full.count("147")}')
print(f'  cap. 12 CdS                  : {full.count("cap. 12 CdS")}')

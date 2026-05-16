"""Extensii pe 2-Abordare_metodologie.docx (post-matrice):
1. Adaugă §2.10.5 — Trasabilitate drum critic → cap. 3 (tabel 5 componente).
2. Adaugă sinteză „Impact cuantificat" la final de §2.13.2, §2.13.3, §2.13.5.
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from copy import deepcopy
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PATH = r'C:\Users\adria\Documents\Claude\Projects\Licitatie 1\doc-output\oferta\2-Abordare_metodologie.docx'

doc = Document(PATH)
body = doc.element.body

def find_paragraph_index(prefix):
    """Returnează indexul paragrafului care începe cu prefix (sau None)."""
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith(prefix):
            return i
    return None

def insert_paragraph_before(target_p, text, style='Normal', bold=False):
    """Inserează un paragraf nou înaintea target_p."""
    new_p = OxmlElement('w:p')
    target_p._element.addprevious(new_p)
    new_para = doc.paragraphs[[i for i, p in enumerate(doc.paragraphs) if p._element is new_p][0]]
    new_para.style = doc.styles[style]
    run = new_para.add_run(text)
    if bold:
        run.bold = True
    return new_para

def insert_table_before(target_p, headers, rows, widths_cm=None):
    """Inserează un tabel înaintea target_p."""
    # Creăm tabelul la finalul documentului, apoi îl mutăm
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    for style_name in ('Light Grid Accent 1', 'Table Grid', 'Tabel grilă'):
        try:
            t.style = style_name
            break
        except KeyError:
            continue
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, txt in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = txt
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            cell = t.rows[ri].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    if widths_cm:
        for i, w in enumerate(widths_cm):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    # Mutăm elementul XML al tabelului înaintea target_p
    tbl_el = t._element
    tbl_el.getparent().remove(tbl_el)
    target_p._element.addprevious(tbl_el)
    return t

# ============================================================
# STEP 1 — Adaug §2.10.5 Trasabilitate drum critic → cap. 3
# Inserăm înainte de §2.11 (paragraful care începe cu "2.11 Registrul...")
# ============================================================
print("STEP 1: Adăugare §2.10.5 — Trasabilitate drum critic → cap. 3")

# Găsim paragraful §2.11
target_p = None
for p in doc.paragraphs:
    if p.style.name.startswith('Heading') and p.text.strip().startswith('2.11 '):
        target_p = p
        break
if target_p is None:
    print("[EROARE] Nu găsesc §2.11 ca punct de inserare")
    sys.exit(1)

# Heading §2.10.5
heading_2_10_5 = insert_paragraph_before(target_p, '2.10.5 Trasabilitate drum critic → activități din Planul de implementare (cap. 3)', style='Heading 3')

# Paragraf introductiv
para_intro = insert_paragraph_before(
    target_p,
    'Tabelul de mai jos arată maparea componentelor de pe drumul critic identificate în §2.10.3 către '
    'activitățile concrete din Planul de implementare (cap. 3 al ofertei), resursele umane alocate (cap. 9 — Echipa proiect), '
    'resursele materiale (cap. 5 — Arhitectură și licențe; cap. 6 — Lista hardware) și rezultatele intermediare / finale '
    'corespunzătoare. Această trasabilitate susține direct calificativul „Excepțional" la subfactorul 3.2 — '
    'corelația dintre datele de intrare și rezultatele intermediare / finale (vezi §2.14.2).'
)

# Tabel trasabilitate
insert_table_before(
    target_p,
    ['Componentă / cap. CdS', 'Activitate din Planul de implementare (cap. 3)', 'Resursă umană alocată (cap. 9)',
     'Resursă materială (cap. 5 / 6)', 'Rezultat intermediar / final'],
    [
        [
            'BND-SNIIA — cap. 3.4.2.7',
            'Analiză (L1-L6, sub-echipă dedicată) → Proiectare (L7) → Dezvoltare integrată (L8-L13) → Testare sistem (L14) → Migrare date din SNIIA actual + APIA + ANARZ (L15-L16) → UAT (L16) → Punere în Producție (L17-L18)',
            'Team leader software <LIDER> + 2 Experți dezvoltare software + Expert integrare + Expert migrare + Expert testare',
            'VOGO Enterprise Suite (modul SNIIA + aplicații mobile native iOS + Android cu offline sync); Oracle Service Bus pentru integrări APIA / ANARZ / Colegiul Medicilor Vet; MS SQL Server Enterprise; FURNIZOR_GIS pentru cartografia exploatațiilor',
            'Aplicații mobile publicate în AppStore + GooglePlay (intermediar — L13); Bază de date SNIIA migrată cu validare 100% (intermediar — L16); Sistem funcțional în producție cu 4 categorii animale + integrările APIA/ANARZ active (final — L18)'
        ],
        [
            'Integrări guvernamentale — cap. 3.4.2.14',
            'Analiză inventar integrări (L1-L3) → Dezvoltare adaptoare ESB + mock-uri OpenAPI (L4-L10) → Testare integrări reale acolo unde sistemele țintă sunt disponibile (L11-L14) → Comutare mock → real prin feature flags (L15-L17, sau ulterior în garanție pentru sisteme indisponibile)',
            'Expert integrare + Expert comunicații și securitate + Team leader software <LIDER>',
            'Oracle Service Bus (cluster activ-pasiv, 2×16 cores); Mirth Connect pentru HL7; Anti-Corruption Layer pe ESB; mock-uri OpenAPI generate automat din specificații; feature flags configurabile prin Keycloak Enterprise (IAM)',
            'Catalog API publicat (intermediar — L10); Integrări active cu ROeID + eIDAS + ONRC + APIA + ANCPI + Ghișeul.ro (intermediar — L14); Mock-uri funcționale pentru PNI / PCUe / PDURo / PJN dacă acestea nu sunt disponibile la PIP (intermediar — L17); Comutare automată mock → real în garanție, fără cost suplimentar (final post-PIP)'
        ],
        [
            'LIMS + migrare date laborator — cap. 3.4.2.1 + 3.4.3.3.3',
            'Analiză fluxuri 3 institute + 41 laboratoare DSVSA (L1-L6) → Proiectare integrare echipamente IoT (L7) → Configurare LIMS COTS (L8-L11) → Integrare HL7 prin Mirth Connect (L12-L13) → Migrare istorice analize (L14-L15) → Acreditare RENAR a fluxurilor (L16) → PIP (L17-L18)',
            'Arhitect sistem (cu certificare producător LIMS) + Expert administrare BD + Expert integrare + Expert migrare + Expert testare',
            'FURNIZOR_LIMS_COTS (cu licență nelimitată + cod sursă); Mirth Connect HL7 FHIR gateway; 126 dispozitive IoT (laboratoare); Microsoft SQL Server Enterprise pentru baza de date LIMS; conformitate ISO/IEC 17025',
            'Sistem LIMS pilotat în 1 institut (intermediar — L11); Integrări HL7 active cu 126 dispozitive IoT (intermediar — L13); 280.000 cereri analiză migrate cu validare RENAR (intermediar — L16); Sistem LIMS în producție pentru toate 3 institute + 41 lab DSVSA (final — L18)'
        ],
        [
            'DMS național — cap. 3.4.2.2 + 3.4.3.3.1',
            'Analiză fluxuri documentare 46 instituții (L1-L6) → Proiectare BPMN configurabilă (L7) → Configurare cele 7 sub-module DMS (L8-L11) → Configurare workflow-uri pilot (L12-L13) → Migrare documente istorice (L14-L15) → Roll-out instituție-cu-instituție (L16-L17) → PIP (L18)',
            'Arhitect sistem (cu certificare producător DMS ZIPPER) + 2 Experți dezvoltare software + Expert migrare + Expert testare + Coordonator instruire',
            'ZIPPER DMS (licență nelimitată + cod sursă); Microsoft SQL Server Enterprise; Elasticsearch pentru căutare full-text + OCR; Keycloak Enterprise pentru SSO; integrare nativă Office pentru editare in-place; pad USB pentru semnătura olografă',
            'Pilot DMS funcțional în ANSVSA centru (intermediar — L11); Migrare arhivă documentară (intermediar — L15); Roll-out pentru cele 46 instituții cu 5.300 utilizatori (intermediar — L17); Sistem DMS în producție națională cu BPMN configurabilă + semnătură eIDAS + olografă (final — L18)'
        ],
        [
            'Portal Servicii Publice + Chatbot AI + App mobilă cetățeni — cap. 3.4.2.6 + 3.4.2.11 + 3.4.3.3.2',
            'Analiză cele 56 servicii electronice (L1-L6) → Proiectare UX/UI cu accesibilitate WCAG 2.1 AA (L7) → Dezvoltare iterativă Suite VOGO (L8-L13) → Integrare ROeID + eIDAS + Ghișeul.ro (L11-L14) → Antrenare Chatbot NLP RO pe corpus ANSVSA (L10-L14) → UAT + publicare în AppStore/GooglePlay (L15-L16) → Acreditare ADR + PIP (L17-L18)',
            'Expert portal + Expert BI + Team leader software <LIDER> + 2 Experți dezvoltare software + Coordonator instruire',
            'VOGO Enterprise Suite (Portal + Chatbot + App mobilă, suită integrată — vezi §2.13.1); NGINX Plus pentru reverse proxy; Microsoft IIS pentru server aplicație; Keycloak Enterprise pentru SSO + ROeID + eIDAS + 2FA SMS; Microsoft Power BI pentru telemetrie RCR11',
            'MVP Portal cu 10 servicii (intermediar — L11); Aplicații mobile publicate (intermediar — L13); 56 servicii electronice active (intermediar — L16); Indicator RCR11 ≥185.000 utilizatori unici/an raportat automat către ADR (final — post-PIP, în perioada de sustenabilitate POCIDIF)'
        ],
    ],
    widths_cm=[3.0, 4.5, 3.2, 3.7, 3.6]
)

# Paragraf de încheiere
insert_paragraph_before(
    target_p,
    'Această trasabilitate este sincronizată cu Diagrama Gantt din cap. 3.2 al Planului de implementare, '
    'cu Registrul de Riscuri din §2.11 al prezentului capitol și cu Lista Software (Lista_Software_SIDISVA.xlsx, '
    'sheet „Sinteza") care leagă fiecare cerință CdS de produsul software ofertat și de cantitatea / sizing-ul aferent.'
)
print("  [OK] §2.10.5 adăugat cu tabel 5 componente × 5 coloane")

# ============================================================
# STEP 2 — Adaug „Impact cuantificat" la §2.13.2, §2.13.3, §2.13.5
# ============================================================
print("STEP 2: Adăugare paragrafe Impact cuantificat în §2.13.x")

# Definim impactele cuantificate pe fiecare sub-secțiune
IMPACT = {
    '2.13.3': (  # heading înaintea căruia inserăm
        '2.13.3',
        '__pre_2.13.3__',  # marker
        'Impact cuantificat: reducere estimată cu 60-70% a efortului manual de revizuire date BND-SNIIA '
        '(cca. 8.000 ore/an în baseline operațional ANSVSA → 2.400-3.200 ore/an după implementare); '
        'acuratețe țintă Anomaly Detection ≥85% pe cazurile reale (validare pe seturi istorice); '
        'recall Fraud Detection ≥75% pe pattern-urile cunoscute (crotalii reutilizate, exploatații fantomă); '
        'Data Quality Score crescut la ≥80 / 100 per înregistrare BND (baseline estimat 55-60); '
        'amortizare costurilor de implementare AI/ML estimată la ≤18 luni de la Punerea în Producție prin reducerea efortului operațional.'
    ),
    '2.13.4': (
        '2.13.4',
        '__pre_2.13.4__',
        'Impact cuantificat: timp de comutare mock → API real per integrare guvernamentală ≤4 ore lucrătoare '
        '(prin feature flags, fără modificare de cod; comparativ cu cca. 2-4 săptămâni necesare pentru re-dezvoltare clasică); '
        'acoperire 100% a cerinței cap. 3.4.2.14 CdS privind integrarea ulterioară a sistemelor guvernamentale neimplementate; '
        'zero downtime la activarea integrărilor reale (rollout canary + capacitate rollback rapid); '
        'cost integrare ulterioară: 0 lei pentru ANSVSA (acoperit prin garanția de 3 ani conform cap. 3.4.2.14 CdS).'
    ),
    '2.13.5': (
        '2.13.5',
        '__pre_2.13.5__',
        'Impact cuantificat: reducere efort raportare manuală indicator RCR11 cu cca. 95% '
        '(de la cca. 40-60 ore/lună efort manual estimat → 1-2 ore/lună validare exporturi automate); '
        'latență raportare reduse la maxim 24 ore între acțiunea utilizator și actualizarea indicatorului '
        '(de la minim 1 lună în raportarea manuală bazată pe Excel); acuratețe raportare 100% '
        '(date direct din baza de date operațională, fără riscul erorilor de transcriere); '
        'retenție audit trail 5 ani conform cerinței de sustenabilitate POCIDIF; '
        'amortizare cost implementare telemetrie estimată ≤6 luni de la Punerea în Producție.'
    ),
}
# Pentru §2.13.5 nu putem insera „înainte de 2.13.6" (nu există) — așa că inserăm înainte de paragraful
# „Aceste 5 elemente de inovatie au impact direct verificabil asupra:" (P209).

# §2.13.2 → inserăm înainte de heading §2.13.3
# §2.13.3 → inserăm înainte de heading §2.13.4
# §2.13.5 → inserăm înainte de paragraful "Aceste 5 elemente..."

INSERT_POINTS = [
    ('2.13.3', 'Impact cuantificat: reducere estimată cu 60-70% a efortului manual de revizuire date BND-SNIIA '
        '(cca. 8.000 ore/an în baseline operațional ANSVSA → 2.400-3.200 ore/an după implementare); '
        'acuratețe țintă Anomaly Detection ≥85% pe cazurile reale (validare pe seturi istorice); '
        'recall Fraud Detection ≥75% pe pattern-urile cunoscute (crotalii reutilizate, exploatații fantomă); '
        'Data Quality Score crescut la ≥80 / 100 per înregistrare BND (baseline estimat 55-60); '
        'amortizare cost implementare AI/ML estimată la ≤18 luni de la Punerea în Producție prin reducerea efortului operațional.'),
    ('2.13.4', 'Impact cuantificat: timp de comutare mock → API real per integrare guvernamentală ≤4 ore lucrătoare '
        '(prin feature flags, fără modificare de cod; comparativ cu cca. 2-4 săptămâni necesare pentru re-dezvoltare clasică); '
        'acoperire 100% a cerinței cap. 3.4.2.14 CdS privind integrarea ulterioară a sistemelor guvernamentale neimplementate; '
        'zero downtime la activarea integrărilor reale (rollout canary + capacitate rollback rapid); '
        'cost integrare ulterioară: 0 lei pentru ANSVSA (acoperit prin garanția de 3 ani conform cap. 3.4.2.14 CdS).'),
    ('Aceste 5 elemente', 'Impact cuantificat: reducere efort raportare manuală indicator RCR11 cu cca. 95% '
        '(de la cca. 40-60 ore/lună efort manual estimat → 1-2 ore/lună validare exporturi automate); '
        'latență raportare redusă la maxim 24 ore între acțiunea utilizatorului și actualizarea indicatorului '
        '(de la minim 1 lună în raportarea manuală bazată pe Excel); acuratețe raportare 100% '
        '(date direct din baza de date operațională, fără riscul erorilor de transcriere); '
        'retenție audit trail 5 ani conform cerinței de sustenabilitate POCIDIF; '
        'amortizare cost implementare telemetrie estimată ≤6 luni de la Punerea în Producție.'),
]

# Pentru fiecare punct de inserare găsim paragraful țintă și inserăm înaintea lui un paragraf cu "Impact cuantificat: ..."
for prefix, impact_text in INSERT_POINTS:
    target = None
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            target = p
            break
    if target is None:
        print(f"  [WARN] Nu găsesc paragraful '{prefix}' — sar peste.")
        continue
    new_p = insert_paragraph_before(target, impact_text)
    # Bold pentru "Impact cuantificat:"
    for run in new_p.runs:
        # toate runs în noul paragraf — bold doar primele 19 caractere (lungimea "Impact cuantificat:")
        if 'Impact cuantificat:' in run.text:
            run.bold = True
    print(f"  [OK] Impact cuantificat adăugat înainte de '{prefix}'")

doc.save(PATH)
print(f"\n[OK] Salvat: {PATH}")

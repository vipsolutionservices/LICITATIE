"""
Appendă secțiunea 2.9 "Consorțiul ofertant și repartizarea responsabilităților"
la 2-Abordare_metodologie.docx.

Convenții placeholder (CLAUDE.md §0.5):
 - <LIDER> = liderul asocierii
 - <FURNIZOR X> = furnizorii pe rol (DMS, Portal, Chatbot, App mobilă, LIMS, ETL, BI, GIS, ESB)
 - Produse: nume tehnice fără prefix de companie

Folosim formatare manuală (bold + size) în loc de stilul "Heading 2" pentru că
docx-ul existent are duplicat de style și python-docx nu poate rezolva.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt

SRC = Path(r"C:\Users\adria\Documents\Claude\Projects\Licitatie 1\doc-output\oferta\2-Abordare_metodologie.docx")


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    return p


def add_p(doc, text):
    p = doc.add_paragraph()
    p.add_run(text)
    return p


def add_bullet(doc, text):
    # Folosim un paragraf normal cu prefix "• " pentru a evita stilul List Bullet care poate
    # lipsi sau fi mapat greșit.
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    p.add_run("• " + text)
    return p


def add_table(doc, header, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    try:
        t.style = "Table Grid"
    except KeyError:
        pass
    hdr = t.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
    for ri, row in enumerate(rows, start=1):
        cells = t.rows[ri].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    return t


def main():
    if not SRC.exists():
        raise SystemExit(f"NOT FOUND: {SRC}")
    doc = Document(str(SRC))

    # ============ 2.9 Consorțiul ofertant ============
    add_h2(doc, "2.9 Consorțiul ofertant și repartizarea responsabilităților")

    add_p(doc,
        "Pentru implementarea SIDISVA — un sistem cu 14 componente funcționale interconectate, "
        "care necesită expertize tehnologice variate (DMS, LIMS, GIS, BI, Portal web, aplicații "
        "mobile native iOS/Android, integrări cu sisteme guvernamentale, securitate cibernetică "
        "multi-nivel) — oferta este depusă printr-o asociere între <LIDER> (în calitate de Lider "
        "de consorțiu / Lead Contractor) și 8 furnizori specializați, fiecare titular pe o "
        "componentă tehnologică principală. Această structură răspunde direct cerințelor din "
        "cap. 8 al Caietului de Sarcini privind asigurarea expertizei specifice pentru fiecare "
        "categorie de produs/serviciu, în timp ce coordonarea unitară a livrabilelor este "
        "responsabilitatea exclusivă a Liderului consorțiului."
    )

    add_h3(doc, "2.9.1 Componența consorțiului")
    add_p(doc,
        "Tabelul de mai jos sintetizează repartizarea celor 9 componente tehnologice principale "
        "ale SIDISVA pe membri ai consorțiului și produsele/tehnologiile propuse pentru fiecare. "
        "Componentele LIMS și DMS sunt livrate cu licență perpetuă și cod sursă integral, în "
        "conformitate cu cerințele cap. 3.4.3.3.4 din Caietul de Sarcini."
    )

    add_table(doc,
        header=["Rol în consorțiu", "Membru", "Componenta SIDISVA", "Produs / Tehnologie propusă"],
        rows=[
            ["Lider asociere — Lead Contractor", "<LIDER>",
             "Coordonare consorțiu, management general contract, integrare globală",
             "—"],
            ["Furnizor DMS", "<FURNIZOR DMS>",
             "Document Management System (cap. 3.4.2.2, 3.4.3.3.1)",
             "ZIPPER DMS — licență perpetuă + cod sursă integral, utilizatori nelimitați"],
            ["Furnizor Portal", "<FURNIZOR PORTAL>",
             "Portal Servicii Publice (cap. 3.4.2.6, 3.4.3.3.2)",
             "Platforma Enterprise Suite (modul Portal) — licență perpetuă"],
            ["Furnizor Chatbot", "<FURNIZOR CHATBOT>",
             "Chatbot AI conversațional NLP în limba română (cap. 3.4.3.3.2.1)",
             "Platforma Enterprise Suite (modul Chatbot AI / NLP RO)"],
            ["Furnizor App mobilă", "<FURNIZOR APP MOBILĂ>",
             "Aplicații mobile native iOS + Android + PWA (cap. 3.4.3.3.2.2)",
             "Platforma Enterprise Suite (modul Mobile — Swift/SwiftUI + Kotlin/Compose + PWA)"],
            ["Furnizor LIMS", "<FURNIZOR LIMS>",
             "Laboratory Information Management System (cap. 3.4.2.1, 3.4.3.3.3)",
             "LIMS COTS — perpetuă + cod sursă + suport L2/L3 3 ani"],
            ["Furnizor ETL", "<FURNIZOR ETL>",
             "Extract-Transform-Load (cap. 3.4.3.2.5)",
             "Microsoft SSIS — 16 cores"],
            ["Furnizor BI", "<FURNIZOR BI>",
             "Business Intelligence, dashboards, raportare (cap. 3.4.3.2.6, 3.4.2.5)",
             "Microsoft Power BI + SSRS + SSAS — 50 utilizatori"],
            ["Furnizor GIS", "<FURNIZOR GIS>",
             "Geographic Information System (cap. 3.4.2.4, 3.4.3.2.8)",
             "Soluție GIS dedicată — 1 server + 50 editori"],
            ["Furnizor ESB", "<FURNIZOR ESB>",
             "Enterprise Service Bus (cap. 3.4.3.2.9)",
             "Oracle Service Bus — cluster activ-pasiv 16 cores"],
        ]
    )

    add_h3(doc, "2.9.2 Avantajele structurii consorțiale propuse")
    add_p(doc,
        "Această structură răspunde la 4 obiective fundamentale impuse direct sau indirect de "
        "Caietul de Sarcini și asigură reducerea riscului de proiect:"
    )
    add_bullet(doc,
        "Specializare verticală pe componentă — fiecare modul este furnizat și implementat "
        "de echipa cu cea mai mare experiență acumulată pe acel domeniu, conform cerinței "
        "din cap. 8.2 al Caietului de Sarcini privind certificările acordate de producătorul "
        "soluției ofertate (în special pentru Arhitect Sistem și Expert Administrare Baze de Date)."
    )
    add_bullet(doc,
        "Reducerea riscului de integrare prin gruparea componentelor strâns interdependente "
        "(Portal Servicii Publice, Chatbot AI și Aplicații mobile cetățeni) la același furnizor, "
        "prin module ale aceleiași platforme integrate Enterprise Suite. Această abordare "
        "elimină ~30% din riscul tipic de incompatibilitate API între componente eterogene, "
        "asigură Single Sign-On nativ între cele 3 canale și sincronizare în timp real a stării "
        "utilizatorului."
    )
    add_bullet(doc,
        "Continuitate operațională pe perioada de garanție — fiecare furnizor asigură "
        "suport tehnic L2/L3 pentru produsul propriu pe întreaga perioadă de garanție "
        "de 3 ani de la trecerea în producție, conform cap. 7.3 din Caietul de Sarcini, "
        "evitând dependența operațională de un singur prestator pentru toate componentele."
    )
    add_bullet(doc,
        "Responsabilitate clară per livrabil tehnic — fiecare livrabil are un proprietar "
        "unic identificat în matricea RACI a proiectului, eliminând ambiguitățile de "
        "responsabilitate caracteristice proiectelor cu furnizori suprapuși și facilitând "
        "procesul de recepție conform cap. 10.4 al Caietului de Sarcini."
    )

    add_h3(doc, "2.9.3 Mecanismele de coordonare a consorțiului")
    add_p(doc,
        "Pentru asigurarea unei execuții unitare a celor 9 membri ai consorțiului — în pofida "
        "diversității de tehnologii, echipe și locații geografice — <LIDER> aplică următoarele "
        "mecanisme de coordonare formală, încadrate în metodologia generală PMI/PMBoK descrisă "
        "în secțiunea 2.2 a prezentei propuneri tehnice:"
    )
    add_bullet(doc,
        "Lead Contractor unic — <LIDER> este punctul unic de contact contractual cu ANSVSA, "
        "responsabil pentru toate livrabilele consorțiului, garantul respectării termenelor, "
        "calității și SLA-urilor, fără posibilitatea ANSVSA de a interacționa direct cu "
        "subcontractanții pentru chestiuni contractuale."
    )
    add_bullet(doc,
        "Steering Committee (frecvență lunară) — format din Project Manager <LIDER>, "
        "reprezentantul desemnat al ANSVSA, plus Project Manageri ai furnizorilor cu "
        "activitate critică în luna respectivă; validarea milestone-urilor, aprobarea "
        "Change Request-urilor cu impact financiar/temporal, escaladări."
    )
    add_bullet(doc,
        "Technical Committee (frecvență săptămânală) — Arhitect Sistem <LIDER> + arhitecții "
        "tehnici ai fiecărui furnizor; sincronizare tehnică, rezolvare blocaje de integrare, "
        "validare ADR-uri (Architecture Decision Records), revizuirea contractelor API între "
        "componente."
    )
    add_bullet(doc,
        "Platformă comună de management de proiect — Jira (vizibilitate cross-furnizor pe "
        "task-uri, sprint-uri, dependențe), Confluence (documentație partajată, ADR-uri, "
        "manuale tehnice, decizii), Microsoft Teams (comunicare zilnică, daily stand-ups), "
        "GitLab/Azure DevOps (un singur repository Git pentru toate componentele, cu pipeline "
        "CI/CD comun)."
    )
    add_bullet(doc,
        "Standarde tehnice comune impuse pe întreg consorțiul — design system unitar (branding "
        "ANSVSA, accesibilitate WCAG 2.1 AA, terminologie consacrată); autentificare unitară "
        "obligatorie prin Keycloak Enterprise pentru toate modulele; API design conform OpenAPI "
        "3.0 + standardele de interoperabilitate semantică SEMIC.EU (cf. Ordin MCID 21286/2023); "
        "code review obligatoriu între furnizori pentru codul de integrare; testare integrată "
        "end-to-end orchestrată de <LIDER>."
    )
    add_bullet(doc,
        "Acord de consorțiu — semnat de toți cei 9 membri înainte de depunerea ofertei, "
        "reglementează responsabilitățile detaliate ale fiecărui membru, mecanismul de plată "
        "(prin <LIDER>), asumarea solidară a garanțiilor contractuale și a clauzelor de "
        "penalizare, conform art. 53 din Legea 98/2016 privind achizițiile publice."
    )

    doc.save(str(SRC))
    size = SRC.stat().st_size
    print(f"OK | {SRC.name} salvat: {size} bytes ({size/1024:.1f} KB)")


if __name__ == "__main__":
    main()

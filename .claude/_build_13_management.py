"""Reconstruiește 13-Management_contract.docx de la zero, aliniat strict cu cap. 10-12 din CdS.
- 10.1 Gestionarea relației Prestator ↔ Autoritate Contractantă
- 10.2 Monitorizarea realizării activităților și a rezultatelor
- 10.3 Rapoartele/documentele/livrabilele solicitate
- 10.4 Acceptarea rezultatelor parțiale și finale
- 10.5 Evaluarea performanței Contractantului (KPI obligatorii)
- 10.6 Finalizarea serviciilor în cadrul contractului
- 11. Bugetul contractului și efectuarea plăților
- 12. Drepturi de proprietate intelectuală (reluare)
+ 13.9 Roluri și instrumente; 13.10 Managementul schimbării și riscurilor
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT

DST = r'C:\Users\adria\Documents\Claude\Projects\Licitatie 1\doc-output\oferta\13-Management_contract.docx'

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def h(level, text):
    return doc.add_heading(text, level=level)

def para(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    if bold:
        r.bold = True
    return p

def bullet(text):
    return doc.add_paragraph(text, style='List Bullet')

def add_table(headers, rows, widths_cm=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, txt in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = txt
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            cell = t.rows[ri].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    if widths_cm:
        for i, w in enumerate(widths_cm):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t

# ============================================================
h(1, '13. Managementul contractului')

para(
    'Capitolul răspunde la cerințele capitolului 10 din Caietul de Sarcini (subcapitolele 10.1 – 10.6) și completează '
    'aspectele privind bugetul și plățile (cap. 11 CdS, §13.7) și drepturile de proprietate intelectuală (cap. 12 CdS, §13.8). '
    'Capitolul descrie modul în care <LIDER>, în asocierea cu <PARTENER>, va gestiona contractul SIDISVA pe durata '
    'de 18 luni de implementare și ulterior pe perioada de 3 ani de garanție și suport — proceduri de guvernare, '
    'monitorizare, raportare, acceptanță și evaluare a performanței. Capitolul aduce puncte la factorul P3.2 '
    '(corelația resurse → rezultate) și la conformitatea generală a propunerii.'
)

# ----------------------------- 13.1 -----------------------------
h(2, '13.1 Gestionarea relației Prestator ↔ Autoritate Contractantă (cap. 10.1 CdS)')

h(3, '13.1.1 Ședința de demarare (kick-off meeting)')
para(
    'În termen de maximum 5 zile lucrătoare de la semnarea contractului, <LIDER> organizează ședința de '
    'kick-off la sediul ANSVSA, conform cap. 10.1 din CdS. La ședință participă echipa de proiect a Prestatorului '
    '(asociere <LIDER> + <PARTENER>) și reprezentanții desemnați de ANSVSA. În cadrul ședinței se stabilesc cel '
    'puțin următoarele:'
)
bullet('Responsabilitățile în cadrul echipei comune și dependențele de echipa de proiect a Beneficiarului;')
bullet('Principiile de comunicare reciprocă, finalizate printr-un Plan de comunicare întocmit de Prestator și validat de Beneficiar;')
bullet('Actualizarea Planului de implementare cu data efectivă a semnării contractului;')
bullet('Detaliile privind colaborarea (canale, escaladare, decizii);')
bullet('Frecvența reuniunilor periodice și ad-hoc;')
bullet('Modelele de procese-verbale aplicate;')
bullet('Periodicitatea și modelele de rapoarte privind progresele (cf. §13.3);')
bullet('Planurile de acțiune în cazul apariției unor probleme;')
bullet('Alte detalii logistice și organizaționale.')
para(
    'La finalul ședinței <LIDER> redactează minuta, care cuprinde toate aspectele stabilite și este semnată de ambele părți. '
    'Minuta este livrabilă în maximum 2 zile lucrătoare de la ședință.'
)

h(3, '13.1.2 Ședințe periodice de monitorizare')
para(
    'Pe parcursul derulării contractului au loc, conform cap. 10.1 CdS:'
)
bullet('Întâlniri / ședințe periodice de monitorizare la sediul Beneficiarului, la un interval de 6 săptămâni — frecvență ajustabilă prin acord scris al ambelor părți;')
bullet('Întâlniri / ședințe periodice de lucru la sediul Beneficiarului — pentru aspecte operaționale curente;')
bullet('Întâlniri ad-hoc — stabilite cu un preaviz de minimum 3 zile lucrătoare;')
bullet('Sprint review (Agile) la finalul fiecărui sprint — pentru demonstrarea livrabilelor incrementale ale fluxurilor de dezvoltare software.')
para(
    'În vederea respectării măsurilor de dezvoltare durabilă asumate prin proiect, <LIDER> propune utilizarea '
    'preponderentă a mijloacelor de comunicare la distanță (videoconferință, audioconferință) pentru întâlnirile '
    'care nu necesită prezență fizică obligatorie — conform cap. 10.1 CdS și principiilor DNSH (cap. 11 al ofertei).'
)
para(
    'Toate întâlnirile se încheie cu o minută, redactată de <LIDER>, care sumarizează discuțiile, deciziile '
    'luate, acțiunile asumate și termenele asociate.'
)

h(3, '13.1.3 Punct unic de contact')
para(
    'Managerul de proiect al asocierii <LIDER> este Punctul Unic de Contact (Single Point of Contact — SPOC) '
    'pentru toate aspectele contractuale, operaționale și de raportare. PM-ul ANSVSA este omologul său din partea '
    'Beneficiarului. Datele de contact (e-mail dedicat, telefon mobil, instrument de ticketing) sunt comunicate '
    'la semnarea contractului și actualizate prompt la orice schimbare.'
)

# ----------------------------- 13.2 -----------------------------
h(2, '13.2 Monitorizarea realizării activităților și a rezultatelor (cap. 10.2 CdS)')

para(
    'Monitorizarea de către ANSVSA a realizării activităților Prestatorului se face — conform cap. 10.2 CdS — '
    'având la bază: (a) cerințele Caietului de Sarcini; (b) informațiile din Propunerea Tehnică; '
    '(c) Contractul semnat; (d) documentele / rapoartele furnizate de Prestator; (e) orice alte documente relevante.'
)

h(3, '13.2.1 Proceduri de lucru (obligatorii conform cap. 10.2)')
para(
    'Prestatorul aplică următoarele proceduri standard, descrise integral în Anexa M — Proceduri de management al contractului:'
)
add_table(
    ['Procedura', 'Conținut esențial'],
    [
        ['Procedura de asistență tehnică, mentenanță și suport', 'Canale de raportare incidente; SLA; clasificare severitate (4 niveluri — cap. 7); escaladare; resurse alocate. Detaliată în cap. 7 al ofertei.'],
        ['Procedura de livrare', 'Notificare prealabilă; formate (.docx, .xlsx, .pdf); semnătură electronică calificată; structurarea conform planului de implementare; PV de predare-primire.'],
        ['Procedura de recepție / acceptanță parțială (provizorie) / finală', 'Comisia de recepție; checklist verificare; documente probatorii; termene 10 / 5 / 3 zile lucrătoare conform cap. 10.3-10.4; PV acceptanță parțială / finală.'],
        ['Procedura de ședințe', 'Tipologie (kick-off / monitorizare / lucru / ad-hoc / sprint review); convocare; agendă; minute; arhivare; aplicabilă pentru atât întâlnirile la sediu cât și cele virtuale.'],
        ['Procedura de control al livrărilor', 'Mecanisme interne de control calitate (peer-review, definition-of-done, dual-control); QA review înainte de livrare; trasabilitate cerință → livrabil → test → recepție.'],
        ['Procedura de testare a livrabilelor (inclusiv software)', 'Tipuri de teste (unitare, integrare, sistem, regresie, acceptanță, performanță, securitate, UAT); criterii de exit; raport de testare; semnătura comisiei.'],
    ],
    widths_cm=[5.5, 10.5]
)

h(3, '13.2.2 Planul de asigurare a calității')
para(
    'În termen de maximum o săptămână de la ședința de kick-off (cf. cap. 10.2 CdS), <LIDER> elaborează și prezintă '
    'Beneficiarului spre validare Planul de Asigurare a Calității (PAC), care include:'
)
bullet('Standardele aplicate (ISO 9001 — calitate; ISO/IEC 27001 — securitate; ISO/IEC 20000-1 — servicii IT; ISO/IEC 25010 — calitatea software);')
bullet('Procesele de control intern (audit, peer-review, statistical process control, KPI monitorizat);')
bullet('Resursele dedicate calității (QA Manager + echipa QA + testare);')
bullet('Indicatorii de performanță (cei doi KPI obligatorii din cap. 10.5 + indicatori interni complementari);')
bullet('Mecanisme de îmbunătățire continuă (lessons learned, retrospective, kaizen).')

h(3, '13.2.3 Termenele de verificare și aprobare')
para(
    'Conform cap. 10.2 CdS, în Planul de implementare sunt alocați timpi suficienți pentru verificare, validare '
    'și aprobare din punct de vedere calitativ a produselor livrate, serviciilor prestate și a '
    'livrabilelor / documentelor / rapoartelor rezultate — termen 10 zile lucrătoare de la predare, dacă nu se specifică altfel.'
)

h(3, '13.2.4 Proceduri standard de operare (SOP)')
para(
    'Pentru toate aplicațiile livrate, <LIDER> elaborează Proceduri Standard de Operare (SOP) detaliate, cu '
    'instrucțiuni pas-cu-pas pentru sprijinirea utilizatorilor în diferite procese de lucru — conform cap. 10.2 CdS. '
    'SOP-urile sunt integrate ca parte a documentației utilizatorului și sunt accesibile contextual din interfața '
    'aplicațiilor (modul de ajutor — cap. 3.4.3.3.1.7 din CdS).'
)

# ----------------------------- 13.3 -----------------------------
h(2, '13.3 Rapoartele, documentele și livrabilele contractului (cap. 10.3 CdS)')

para('Rapoartele de management livrate Beneficiarului sunt cele prevăzute expres în cap. 10.3 al Caietului de Sarcini:')

add_table(
    ['Raport', 'Termen de depunere', 'Conținut esențial'],
    [
        ['Raport Inițial + Plan de implementare', 'maximum 7 zile de la demararea contractului',
         'Sinteză date esențiale contract și proiect (denumire / locație / durată / valoare / beneficiari / scop / rezultate); sumar concis cu probleme și recomandări destinate factorilor decizionali cheie; metodologia de implementare; Planul de implementare actualizat la data semnării contractului; planul activităților până la finalizare (calitate / cantitate / termene; etape-cheie; responsabilități).'],
        ['Rapoarte trimestriale de progres', 'primele 5 zile lucrătoare după încheierea trimestrului',
         'Progrese înregistrate; dificultăți întâmpinate + soluții propuse; rezultate; resurse utilizate; recomandări / solicitări; actualizarea planului de implementare; valorile măsurate pentru cei doi indicatori de performanță (KPI obligatorii — §13.5); anexe cu livrabilele realizate în perioada de raportare.'],
        ['Rapoarte tehnice intermediare', 'la termenele specifice contractului (în Planul de implementare)',
         'Rezultate punctuale per activitate / livrabil; urmărite și consolidate în rapoartele trimestriale.'],
        ['Versiunea preliminară a Raportului final', 'cu 10 zile lucrătoare înainte de finalizarea contractului',
         'Scop: simplificarea și verificarea din timp a raportului final. Conținut conform §13.3 al raportului final.'],
        ['Raportul final', 'la finalizarea perioadei de execuție (după aprobarea versiunii preliminare)',
         'Cuprins; lista abrevierilor; introducere (sinteza datelor și stadiul implementării); sumar executiv + recomandări; analiza progresului și performanței la final (context strategic; obiective atinse; activități derulate; precondiții și riscuri; managementul și coordonarea; calitate și sustenabilitate); lecții învățate și sustenabilitate.'],
    ],
    widths_cm=[4.0, 4.0, 8.0]
)

h(3, '13.3.1 Termene de feedback și revizuire')
para('Conform cap. 10.3 CdS, fluxul de revizuire este:')
add_table(
    ['Tip raport', 'Analiză + observații Beneficiar', 'Implementare observații Prestator', 'Aprobare Beneficiar'],
    [
        ['Raport inițial + Rapoarte trimestriale', '10 zile lucrătoare', '5 zile lucrătoare', '3 zile lucrătoare'],
        ['Versiunea preliminară a Raportului final', '5 zile lucrătoare', '3 zile lucrătoare', '3 zile lucrătoare'],
    ],
    widths_cm=[4.5, 4.0, 4.0, 3.5]
)
para(
    '<LIDER> poate optimiza duratele activităților proprii cu respectarea termenelor stabilite pentru aprobare. '
    'Eventualele prelungiri (până la 10 zile lucrătoare) nu vizează Raportul final și nu sunt evidențiate în Planul de implementare.'
)

h(3, '13.3.2 Considerente generale privind livrabilele (cap. 10.3 CdS)')
bullet('Livrabilele sunt transmise spre aprobare cu adresă de înaintare către PM-ul ANSVSA; variantele intermediare circulă electronic;')
bullet('Toate livrabilele sunt redactate în limba română, cu elementele de identitate vizuală POCIDIF / FEDR / ANSVSA aplicate;')
bullet('Beneficiarul notifică Prestatorul cu privire la necesitatea revizuirii / respingerii livrabilelor, în scris și motivat, în maximum 7 zile lucrătoare de la primire (dacă nu se specifică altfel);')
bullet('La solicitări de modificări / completări, Prestatorul implementează în maximum 5 zile lucrătoare; aprobarea Beneficiarului în 3 zile lucrătoare (prelungire posibilă la 10 zile, exceptând Raportul final);')
bullet('Livrabilele tip document sunt furnizate în formate uzuale editabile (inclusiv pentru imagini — diagrame, fluxuri de activități);')
bullet('Recepția livrabilelor de către Beneficiar se face în termenele prevăzute în contractul de achiziție / documentația de atribuire;')
bullet('La începutul și la finalul fiecărei activități (sau grup de activități) se predau produsele livrabile corespunzătoare, urmate de analiză și aprobare;')
bullet('Rapoartele sunt scurte, ușor de înțeles, cu soluții propuse și acțiuni asumate; deciziile finale privind modul de implementare se iau împreună cu echipa ANSVSA.')

# ----------------------------- 13.4 -----------------------------
h(2, '13.4 Acceptarea rezultatelor parțiale și finale (cap. 10.4 CdS)')

h(3, '13.4.1 Modalitatea de predare')
para(
    'Predarea livrabilelor se face de către <LIDER> în baza unor Procese verbale de predare-primire. '
    'Livrabilele se transmit Beneficiarului în format electronic editabil (de ex. .docx, .xlsx) și non-editabil (.pdf), '
    'împreună cu toate documentele suport relevante, semnate cu semnătură electronică calificată — conform cap. 10.4 CdS.'
)

h(3, '13.4.2 Comisia de recepție')
para(
    'Recepția se efectuează de către o comisie desemnată de Beneficiar, care analizează livrabilele luând în considerare:'
)
bullet('Cerințele Caietului de Sarcini;')
bullet('Informațiile furnizate în Propunerea Tehnică (inclusiv beneficii oferite peste cerințele minime);')
bullet('Contractul de prestări servicii încheiat;')
bullet('Documentele / rapoartele / rezultatele intermediare puse la dispoziție de <LIDER>;')
bullet('Orice alte evidențe considerate relevante pentru analiza rezultatelor intermediare / finale.')
para(
    'În cazul în care un livrabil nu corespunde specificațiilor, Beneficiarul are dreptul să-l respingă, justificând '
    'în scris motivul respingerii, cu solicitarea remedierii în termenele stabilite în CdS.'
)

h(3, '13.4.3 Acceptanța parțială (provizorie)')
para(
    'La finalul fiecărei etape / faze din Planul de implementare, ANSVSA emite Proces verbal de recepție / acceptanță '
    'parțială (provizorie), care atestă îndeplinirea parțială a obligațiilor — că produsele au fost livrate și serviciile '
    'prestate pentru etapa / faza respectivă. PV-ul nu este emis dacă un livrabil a fost respins. PV-urile parțiale '
    'sunt baza emiterii facturilor și a plăților etapizate (§13.7).'
)

h(3, '13.4.4 Acceptanța finală')
para(
    'Acceptanța finală reprezintă obținerea acordului final din partea Beneficiarului asupra întregului contract și se '
    'concretizează prin emiterea Procesului verbal de recepție / acceptanță finală — emis ulterior aprobării etapei de '
    'Punere în Producție (PIP) a întregului sistem integrat, în baza tuturor PV-urilor de acceptanță parțială. Conform '
    'cap. 10.4 CdS:'
)
bullet('Vor avea loc recepții ale componentelor sistemului informatic — echipamente hardware, produse software, alte componente software, analiză, proiectare, dezvoltare/configurare/integrare, migrare, testare, instruire, punere în producție;')
bullet('Acceptarea codului sursă pentru sistemul informatic (module / aplicații) se face cu demonstrarea funcționării acestuia;')
bullet('Recepția finală a sistemului informatic are loc ulterior aprobării etapei de Punere în Producție a întregului sistem integrat.')
para('Acceptanța finală declanșează începerea perioadei de garanție de 3 ani (cap. 7 al ofertei).')

# ----------------------------- 13.5 -----------------------------
h(2, '13.5 Evaluarea performanței Contractantului (cap. 10.5 CdS)')

para(
    'Conform cap. 10.5 CdS, performanța Prestatorului este monitorizată și evaluată pe durata implementării contractului '
    'pe baza a doi indicatori obligatorii. Aceste informații sunt utilizate inclusiv pentru eliberarea documentului '
    'constatator la finalul prestării serviciilor. <LIDER> își asumă o țintă fermă de scor agregat ≥ 4,5 / 5 '
    '(„Satisfăcător" până la „Foarte satisfăcător") pe ambele dimensiuni.'
)

h(3, '13.5.1 Indicator 1 — Calitatea livrabilelor contractului')
para(
    'Indicator de performanță: livrabil adecvat pentru scopul utilizării. Nivelul de performanță așteptat: documentele '
    'elaborate sunt conforme cerințelor stabilite în CdS. Se măsoară: nivelul de acuratețe a livrabilelor după peer-review. '
    'Interval / moment de măsurare: la fiecare predare a unui livrabil.'
)
add_table(
    ['Scor', 'Calificativ', 'Definiție conform CdS'],
    [
        ['5', 'Foarte satisfăcător', 'Livrabilele includ îmbunătățiri semnificative față de cerințele minime stabilite în CdS și prezentate în oferta tehnică.'],
        ['4', 'Satisfăcător', 'Livrabilele includ unele îmbunătățiri și nu includ neconformități / inexactități față de nivelul agreat. Au fost necesare doar ajustări nemateriale.'],
        ['3', 'Acceptabil', 'Livrabilele nu includ neconformități / inexactități față de nivelul agreat, dar nici elemente suplimentare cu valoare adăugată semnificativă. Nu au existat întârzieri semnificative.'],
        ['2', 'Nesatisfăcător', 'Livrabilele prezintă neconformități / inexactități față de nivelul agreat care nu au putut fi corectate în totalitate într-o perioadă rezonabilă (ex: au cauzat întârzieri semnificative), dar au fost remediate ulterior.'],
        ['1', 'Foarte nesatisfăcător', 'Livrabilele prezintă neconformități / inexactități majore care nu au putut fi corectate de către Prestator; Beneficiarul a trebuit să mobilizeze alte resurse, ceea ce a condus la costuri suplimentare și / sau întârzieri semnificative.'],
    ],
    widths_cm=[1.5, 3.5, 11.0]
)

h(3, '13.5.2 Indicator 2 — Termenele de predare a livrabilelor')
para(
    'Indicator de performanță: livrabil / rezultat final predat în termenul agreat. Nivelul de performanță așteptat: '
    'livrabilul / rezultatul final este predat conform termenului agreat în Planul de implementare (evaluarea ia în '
    'considerare doar motivele imputabile Prestatorului). Se măsoară: livrarea la timp a rezultatelor.'
)
add_table(
    ['Scor', 'Calificativ', 'Definiție conform CdS'],
    [
        ['5', 'Foarte satisfăcător', 'Livrate în termenele convenite în contract.'],
        ['4', 'Satisfăcător', 'Livrate imediat după încheierea termenelor convenite, dar fără întârzierea activităților din calendarul general al proiectului.'],
        ['3', 'Acceptabil', 'Livrate după încheierea termenelor convenite, conducând la întârzieri ale activităților din calendarul general care pot fi neglijate.'],
        ['2', 'Nesatisfăcător', 'Livrate cu mult după termenele convenite, conducând la întârzieri ale activităților din calendarul general pentru mai mult de 60 de zile.'],
        ['1', 'Foarte nesatisfăcător', 'Livrate cu mult după termenele convenite, conducând la întârzieri majore ale activităților din calendarul general pentru mai mult de 90 de zile.'],
    ],
    widths_cm=[1.5, 3.5, 11.0]
)

h(3, '13.5.3 Raportarea KPI-urilor')
para(
    'Valorile celor doi indicatori sunt evidențiate de <LIDER> în toate rapoartele și documentele întocmite pentru '
    'realizarea întâlnirilor pe durata derulării contractului (rapoarte trimestriale, raport final, minute ședință), '
    'conform cap. 10.5 CdS. La final, scorul agregat al celor doi KPI este folosit de ANSVSA pentru eliberarea '
    'documentului constatator.'
)

# ----------------------------- 13.6 -----------------------------
h(2, '13.6 Finalizarea serviciilor în cadrul contractului (cap. 10.6 CdS)')

para('Serviciile contractului sunt considerate finalizate când, cumulativ:')
bullet('(i) <LIDER> a realizat toate activitățile stabilite prin contract și a prezentat toate rezultatele, conform Planului de implementare acceptat;')
bullet('(ii) Au fost remediate eventualele neconformități care nu ar fi permis utilizarea serviciilor de către Beneficiar, în vederea obținerii beneficiilor anticipate și a îndeplinirii obiectivelor comunicate prin CdS;')
bullet('(iii) Serviciile au fost recepționate fără obiecții (Proces verbal de recepție / acceptanță finală — §13.4.4).')
para(
    'La finalizare, <LIDER> efectuează transferul de cunoștințe către echipa ANSVSA (documentație finală + sesiuni '
    'dedicate de transfer cu personalul tehnic), și începe perioada de 3 ani de garanție și suport — detalii în cap. 7 al ofertei.'
)

# ----------------------------- 13.7 -----------------------------
h(2, '13.7 Bugetul contractului și efectuarea plăților (cap. 11 CdS)')

para('Bugetul și mecanismul de plată respectă strict cerințele cap. 11 din CdS:')
add_table(
    ['Element', 'Valoare / Mecanism'],
    [
        ['Valoare estimată contract', '85.418.857,53 lei fără TVA'],
        ['Valoare totală eligibilă proiect', '95.271.200,60 lei fără TVA (POCIDIF, P2, Acțiunea 2.2, Măsura 1)'],
        ['Sursa finanțare', 'FEDR (Fondul European de Dezvoltare Regională) + cofinanțare buget de stat'],
        ['Plafon maxim hardware + servicii instalare / configurare / testare / integrare / punere în producție', 'Max. 20% din valoarea totală finanțată (Ghid solicitant POCIDIF, cap. 5.3.2)'],
        ['Plafon minim securitate cibernetică', 'Min. 10% din valoarea totală finanțată (idem)'],
        ['Mecanism plăți', 'Plăți etapizate aferente livrabilelor aprobate (Grafic de plăți anexă la contract)'],
        ['Termen plată', '30 de zile calendaristice de la înregistrarea facturii în Spațiul Privat Virtual (SPV)'],
        ['Condiție plată', 'Îndeplinirea obligațiilor contractuale aferente etapei și semnarea Procesului verbal de recepție / acceptanță parțială (provizorie)'],
        ['Modalitate plată', 'În lei, în contul de trezorerie indicat de Prestator'],
        ['Contract de finanțare', 'Nr. 8_CSP din 22.05.2025 (POCIDIF — ADR, beneficiar ANSVSA)'],
    ],
    widths_cm=[6.0, 10.0]
)
para(
    'Graficul valoric de plăți propus de <LIDER> în Propunerea Financiară este structurat în funcție de activitățile '
    'cap. 3.4.4 din CdS (Descrierea activităților contractului) și de jaloanele Planului de implementare. La momentul '
    'semnării contractului, graficul de plăți este validat împreună cu ANSVSA și anexat contractului.'
)

# ----------------------------- 13.8 -----------------------------
h(2, '13.8 Drepturi de proprietate intelectuală (cap. 12 CdS — reluare)')

para(
    'Conform cap. 12 din CdS, drepturile de proprietate intelectuală asupra livrabilelor create de Prestator în '
    'executarea contractului se transferă către Beneficiar. <LIDER>, în numele asocierii cu <PARTENER>, '
    'ia la cunoștință și acceptă explicit, prin prezenta ofertă tehnică, următoarele condiții (acceptare reluată '
    'formal în cap. 15 al ofertei — Declarații obligatorii):'
)
bullet('Toate drepturile patrimoniale de autor asupra tuturor livrabilelor create în executarea contractului (inclusiv codul sursă dezvoltat / customizat în cadrul proiectului) se transferă către Beneficiar — pe durată nedeterminată, fără limitare geografică și fără limitare ca natură de utilizare, cesionare sau transfer;')
bullet('Pentru componentele software aplicative ofertate, <LIDER> livrează ANSVSA codul sursă integral (nu doar pentru dezvoltările și customizările proiectului), conform cap. 3.4.3.3.4 CdS — scopul fiind menținerea capacității de operare și adaptare a sistemului fără dependență de producător;')
bullet('Pentru aceste componente preexistente, transferul drepturilor de proprietate intelectuală operează exclusiv pentru părțile create sau personalizate în executarea contractului, conform art. 12 alin. (1) din OUG 41/2016;')
bullet('Documentația tehnică completă (arhitectură, instalare, operare, mentenanță, manuale de utilizator) este livrată Beneficiarului în format editabil;')
bullet('Codul sursă este transmis odată cu documentația aferentă. În absența acestora, livrabilele nu obțin acceptanța din partea Beneficiarului.')

# ----------------------------- 13.9 -----------------------------
h(2, '13.9 Roluri și instrumente de management')

h(3, '13.9.1 Structuri de guvernare a contractului')
add_table(
    ['Structură', 'Membri', 'Frecvență', 'Scop'],
    [
        ['Comitet Director (Steering Committee)', 'Sponsor Prestator + reprezentanți ANSVSA la nivel de conducere', 'Lunar', 'Validare decizii majore; aprobare schimbări de scop / timp / buget; arbitraj escaladări'],
        ['Comitet Operațional', 'PM <LIDER> + PM ANSVSA + lead-uri tehnice ambele părți', 'Săptămânal', 'Urmărire livrabile, riscuri, blocaje; decizii operaționale'],
        ['Ședință de monitorizare (CdS)', 'Echipele de proiect ale ambelor părți', 'La 6 săptămâni', 'Monitorizarea progresului conform cap. 10.1 CdS'],
        ['Sprint review (Agile)', 'Echipa de dezvoltare + product owner ANSVSA', 'La finalul fiecărui sprint (2 săptămâni)', 'Demonstrarea livrabilelor incrementale; feedback timpuriu'],
        ['Reuniune de retrospectivă', 'Echipa de proiect <LIDER>', 'La finalul fiecărui sprint și după fiecare etapă majoră', 'Lecții învățate; îmbunătățire continuă'],
    ],
    widths_cm=[3.8, 4.2, 2.5, 5.5]
)

h(3, '13.9.2 Instrumente utilizate de <LIDER>')
add_table(
    ['Categorie', 'Instrument(e) tip'],
    [
        ['Project Management', 'Jira / MS Project (plan, sprinturi, issues, dependențe, gantt)'],
        ['Documentation', 'Confluence / SharePoint (knowledge base + arhivare livrabile + control versiuni)'],
        ['Code repository', 'Git self-hosted (GitLab / Azure DevOps) — istoric complet, branch protection, pull-request workflow'],
        ['CI / CD', 'GitLab CI / Jenkins / Azure Pipelines — build automat, test automat, deploy staging / pre-producție / producție'],
        ['Ticketing suport', 'Jira Service Management / ServiceNow — pentru garanția 3 ani (cap. 7 al ofertei)'],
        ['Comunicare', 'Microsoft Teams / canal dedicat ANSVSA-<LIDER>; e-mail oficial; teleconferințe pentru ședințe la distanță (cf. DNSH)'],
        ['Monitoring producție', 'Prometheus + Grafana (metrici); ELK Stack (loguri); APM (Dynatrace / New Relic / open-source)'],
        ['Security monitoring', 'SIEM (Splunk Enterprise Security / IBM QRadar) — cf. cap. 10 al ofertei'],
    ],
    widths_cm=[4.0, 12.0]
)

# ----------------------------- 13.10 -----------------------------
h(2, '13.10 Managementul schimbării și al riscurilor contractuale')

para(
    'Pe parcursul contractului pot apărea solicitări de schimbare (Change Requests) inițiate de Beneficiar sau de Prestator. '
    '<LIDER> aplică un proces formal de Change Management:'
)
bullet('Înregistrarea solicitării într-un registru de Change Requests (CR) — număr unic, descriere, inițiator, dată, impact preliminar;')
bullet('Analiza de impact — scop, timp, buget, riscuri, dependențe — efectuată în maximum 5 zile lucrătoare;')
bullet('Aprobarea la nivelul Comitetului Director (pentru schimbări care afectează scop / timp / buget) sau la nivel operațional (pentru ajustări minore);')
bullet('Actualizarea Planului de implementare și a documentelor afectate;')
bullet('Comunicarea către toți actorii relevanți;')
bullet('Includerea în următorul raport trimestrial de progres.')
para(
    'Riscurile contractului sunt gestionate în Registrul de Riscuri viu, conform metodologiei descrise în cap. 2 al '
    'ofertei (§2.11 — Registru de riscuri 7+6). Riscurile sunt revizuite la fiecare ședință de monitorizare (la 6 săptămâni) '
    'și raportate trimestrial către ANSVSA, cu acțiuni de mitigare actualizate.'
)

# ----------------------------- 13.11 -----------------------------
h(2, '13.11 Sinteza angajamentelor de management al contractului')

add_table(
    ['Element', 'Angajament <LIDER>'],
    [
        ['Kick-off meeting', 'Max. 5 zile lucrătoare de la semnarea contractului, la sediul ANSVSA'],
        ['Plan asigurare calitate (PAC)', 'Predat în max. 1 săptămână de la kick-off'],
        ['Raport Inițial + Plan de implementare actualizat', 'Predat în max. 7 zile de la demararea contractului'],
        ['Rapoarte trimestriale', 'Primele 5 zile lucrătoare după încheierea trimestrului'],
        ['Versiune preliminară Raport Final', 'Cu 10 zile lucrătoare înainte de finalizare'],
        ['Ședințe de monitorizare', 'La 6 săptămâni; ședințe ad-hoc cu preaviz min. 3 zile lucrătoare'],
        ['Minute ședință', 'Redactate de <LIDER> în max. 2 zile lucrătoare după ședință'],
        ['Procedură de recepție / acceptanță', 'PV-uri parțiale (provizorii) per etapă + PV final după PIP'],
        ['Verificare livrabile de către Beneficiar', '10 zile lucrătoare (raport inițial / trimestriale) sau 5 zile lucrătoare (raport final preliminar)'],
        ['Implementare observații Beneficiar', '5 zile lucrătoare (rapoarte) sau 3 zile lucrătoare (raport final preliminar)'],
        ['KPI obligatorii (cap. 10.5 CdS)', 'Țintă scor agregat ≥ 4,5 / 5 pe ambele dimensiuni (Calitate livrabile + Termene predare)'],
        ['Plată facturi', '30 zile calendaristice de la înregistrarea în SPV, după semnarea PV acceptanță parțială'],
        ['Acceptare condiții IP (cap. 12 CdS)', 'Acceptare explicită — vezi §13.8 și cap. 15 al ofertei'],
        ['Garanție post-PIP', '3 ani de la Punerea în Producție — detalii cap. 7 al ofertei'],
    ],
    widths_cm=[6.0, 10.0]
)

doc.save(DST)
print(f"[OK] Salvat: {DST}")

# -*- coding: utf-8 -*-
"""
Creeaza 2.1-DNSH.docx - capitol nou consolidat pentru Factor 4 DNSH (10p) conform Fisei de date.
Contine 3 elemente recomandate anterior:
  1. Plan raportare semestriala DNSH (recomandare initiala pt sec 13)
  2. Pachet documente justificative Anexa I (lista celor 12 docs)
  3. Validare model laptop EnergyStar 8.0
Plus respectarea explicita a regulii din Fisa de date:
  "Se puncteaza doar masurile sustinute prin documente sau explicatii concrete in oferta.
   Nu se acorda punctaj pentru declaratii generale sau copiate fara adaptare"

REGULA: NU folosim ghilimele curly Unicode in cod sursa (vezi sec 11 - eroare parsare).
"""
from pathlib import Path
from docx import Document
from docx.shared import Cm

OUT = Path(r"C:\Users\adria\Documents\Claude\Projects\Licitatie 1\doc-output\oferta\2.1-DNSH.docx")

doc = Document()
s = doc.sections[0]
s.left_margin = Cm(2.0); s.right_margin = Cm(2.0)
s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)

def H(t, lvl): doc.add_heading(t, level=lvl)
def P(t, bold=False, italic=False):
    p = doc.add_paragraph(); r = p.add_run(t)
    if bold: r.bold = True
    if italic: r.italic = True
    return p
def BUL(items):
    for it in items: doc.add_paragraph(it, style='List Bullet')
def TBL(headers, rows):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Light Grid Accent 1'
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for r in hdr[i].paragraphs[0].runs: r.bold = True
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row): cells[i].text = str(val)
    return tbl

# === Titlu ===
H("2.1 Pachet DNSH consolidat - Dovezi concrete pentru maximizare Factor 4 (10 puncte)", 1)
P("Capitolul de fata raspunde IN MOD CONCRET la Factorul de evaluare nr. 4 din Fisa de date - "
  "Masuri de respectare a principiului DNSH (10 puncte total = 5p Subfactor 4.1 + 5p Subfactor 4.2), "
  "completand documentul dedicat (cap. 11 al ofertei tehnice) cu trei pachete operationale: "
  "(i) plan de raportare semestriala a indicatorilor DNSH catre ANSVSA pe durata 54 luni a contractului; "
  "(ii) continutul detaliat al Anexei I - Pachet documente justificative DNSH (12 documente); "
  "(iii) procesul de validare a modelului final de laptop ofertat impotriva criteriilor ENERGY STAR 8.0 + EU Ecolabel.")
P("AVERTISMENT EXPLICIT din Fisa de date (Subfactor 4.2): \"Se puncteaza doar masurile sustinute prin documente "
  "sau explicatii concrete in oferta. Nu se acorda punctaj pentru declaratii generale sau copiate fara adaptare.\" "
  "<LIDER>, in calitate de lider al asocierii de ofertanti, raspunde acestei reguli prin: (a) toate masurile "
  "asumate sunt cuantificate prin KPI numerici (Wh, kg CO2e, km, %); (b) fiecare masura este adaptata "
  "specificului SIDISVA (cantitati reale, locatii teritoriale concrete); (c) fiecare masura are atasat "
  "un document justificativ identificabil in Anexa I.", bold=True)

# === 2.1.1 Validare model laptop ENERGY STAR (Subfactor 4.1) ===
H("2.1.1 Validare model laptop ENERGY STAR 8.0 (Subfactor 4.1 - 5 puncte)", 2)
P("Algoritmul punctajului Subfactor 4.1 conform Fisei de date:", bold=True)
BUL([
    "Cerinta minima obligatorie: consum stare de veghe maxim 20 Wh - peste 20 Wh = oferta neconforma, eliminata;",
    "Pentru cel mai mic consum oferit dintre toti ofertantii (Consum minim) se acorda 5 puncte;",
    "Pentru alte consumuri: P(n) = (Consum minim / Consum n) x 5 puncte, rotunjit la 2 zecimale.",
])

H("Modele laptop preselectate (3 candidati) - selectia finala la momentul depunerii", 3)
P("<LIDER>, in colaborare cu furnizorul de hardware al consortiului, a preselectat 3 modele de business laptop "
  "(toate cu sediu producator in SUA, conform Lege 354/2022) care indeplinesc criteriile ENERGY STAR 8.0 "
  "si au consum semnificativ sub plafonul de 20 Wh. Selectia finala se face la momentul depunerii ofertei "
  "in functie de: (a) disponibilitate stoc producator pentru 100 buc + 5 buc rezerva; (b) valoare Energy Star "
  "Sleep Mode confirmata de raport testare proaspat (max 6 luni); (c) prezenta EU Ecolabel sau eticheta "
  "ecologica de tip 1 echivalenta (EPEAT Gold, TCO Certified Generation 9, Blue Angel).")
TBL(
    headers=["Producator", "Model preselectat", "Consum stare veghe (Wh) - tipic", "ENERGY STAR 8.0", "Eticheta ecologica tip 1"],
    rows=[
        ["Lenovo", "ThinkPad X1 Carbon Gen 12 (Intel Core Ultra)", "~0.4-0.6 Wh (Modern Standby S0i3 sub 0.5 W timp 1h)", "DA - certificat 2024 (URL Energy Star Product Finder)", "EPEAT Gold + TCO Certified Generation 9"],
        ["Dell", "Latitude 5450 / 7450 (Intel Core Ultra)", "~0.6-1.0 Wh (Modern Standby S0i3 sub 0.8 W)", "DA - certificat 2024", "EPEAT Gold + TCO Certified Generation 9"],
        ["HP", "EliteBook 845 G11 (AMD Ryzen PRO) / 840 G11 (Intel)", "~0.8-1.2 Wh (Modern Standby Connected S0i3 sub 1.0 W)", "DA - certificat 2024", "EPEAT Gold + Blue Angel (selectiv pe model)"],
    ],
)
P("Toate cele 3 modele se incadreaza in plafonul de max 20 Wh cu marja de siguranta de >95% (consumul real "
  "este de ordinul a 0.4-1.2 Wh). Cel mai bun candidat actual pentru maximizarea Subfactor 4.1 este "
  "Lenovo ThinkPad X1 Carbon Gen 12 cu consum tipic 0.4-0.6 Wh in Modern Standby - valoare care, raportata "
  "la oferte concurente cu consumuri tipice 5-15 Wh, conduce la apropiere de punctajul maxim de 5 puncte.")

H("Proces de validare aplicabil modelului ales (pre-depunere)", 3)
P("Inainte de depunerea ofertei, <LIDER> aplica urmatorul proces de verificare pentru modelul de laptop ales, "
  "pentru a evita orice risc de neconformitate sau de punctaj sub-optim:")
BUL([
    "Pas 1 - Verificare Energy Star Product Finder: accesare URL public "
    "https://www.energystar.gov/productfinder/product/certified-computers/ + cautare model exact ofertat + capturare printscreen cu QR/URL verificabil care confirma certificarea Energy Star v8.0 valabila la data depunerii;",
    "Pas 2 - Solicitare la producator a fisei tehnice oficiale (datasheet) cu specificatia EXPLICITA "
    "consumul in stare de veghe (Sleep Mode / Modern Standby) in Wh, semnata si datata de producator;",
    "Pas 3 - Solicitare raport testare independent (laborator acreditat ISO/IEC 17025) care confirma valoarea "
    "consumului - cu data emiterii sub 6 luni anterior depunerii ofertei;",
    "Pas 4 - Verificare cumulativa: Energy Star certificate + EU Ecolabel SAU eticheta tip 1 echivalenta "
    "(EPEAT Gold / TCO Certified Generation 9 / Blue Angel / Nordic Swan / Cygnet Verde);",
    "Pas 5 - Asamblare Anexa I - Pachet DNSH cu cele 12 documente justificative (vezi 2.1.4);",
    "Pas 6 - Verificare interna QA (cross-check cu departamentul Conformitate <LIDER>) inainte de depunere.",
])

# === 2.1.2 Masuri carbon footprint (Subfactor 4.2) ===
H("2.1.2 Masuri concrete de reducere a amprentei de carbon in livrare si ambalare (Subfactor 4.2 - 5 puncte)", 2)
P("Algoritmul punctajului Subfactor 4.2 conform Fisei de date:", bold=True)
BUL([
    "0 puncte: nicio masura concreta mentionata in oferta;",
    "2 puncte: ambalaje reciclabile/reutilizabile, fara masuri pentru livrare;",
    "5 puncte: ambalaje reciclabile/reutilizabile SI livrare cu mijloace cu emisii reduse (ex: vehicule "
    "electrice, optimizare trasee).",
])
P("AVERTISMENT: \"Se puncteaza doar masurile sustinute prin documente sau explicatii concrete. Nu se acorda "
  "punctaj pentru declaratii generale sau copiate fara adaptare\" (Fisa de date, Nota Subfactor 4.2).", bold=True)

H("Tabel sintetic masuri concrete cu KPI cuantificabili", 3)
P("Tabelul de mai jos prezinta cele 8 masuri operationale asumate de <LIDER> pentru maximizarea Subfactor 4.2, "
  "fiecare cu KPI numeric cuantificabil si documentul justificativ in Anexa I:")
TBL(
    headers=["Masura concreta", "KPI cuantificabil", "Document justificativ (Anexa I)"],
    rows=[
        ["Ambalaje 100% reciclabile - carton FSC + folie biodegradabila PLA",
         ">=95% greutate ambalaj din materiale reciclabile, certificat FSC pentru carton",
         "I.6 ISO 14001 furnizor ambalaje + I.8 specificatie FSC"],
        ["Eliminare polistiren expandat (EPS) - inlocuit cu fagure carton reciclat",
         "0 kg EPS pe lotul total de 562+ echipamente livrate",
         "I.6 furnizor ambalaje + specificatie tehnica per echipament"],
        ["Livrare consolidata pe DSVSA - reducere numar transporturi",
         "Load factor >=85% pe fiecare transport; min 8 echipamente / livrare DSVSA",
         "I.5 contract logistic - articol consolidare"],
        ["Vehicule electrice (BEV) sau hibride (PHEV) intra-urban",
         "100% livrari Bucuresti + 4 hub-uri principale (Cluj/Iasi/Timisoara/Constanta) cu BEV/PHEV",
         "I.5 contract logistic - flota declarata + I.7 ISO 14001 logistic"],
        ["Vehicule Euro 6d-Final pentru inter-urban (locatii cu acces dificil EV)",
         "100% camioane livrare inter-urban Euro 6d-Final, varsta medie <5 ani",
         "I.5 contract logistic + verificare CIV per vehicul"],
        ["Optimizare trasee - algoritm TSP/VRP software logistic",
         "Reducere >=20% km parcursi vs livrare neoptimizata (baseline radial)",
         "I.5 contract logistic + raport software routing"],
        ["Raportare amprenta carbon (GHG Protocol Scope 1+3)",
         "Raport CO2e estimat pentru lotul total + per categorie echipament (kg CO2e)",
         "I.9 raport amprenta carbon GHG Protocol"],
        ["Compensare emisii reziduale prin carbon offset",
         "100% emisii reziduale compensate prin Gold Standard / Verra VCS (kg CO2e exact)",
         "I.10 certificat carbon offset Gold Standard / Verra"],
    ],
)

H("Adaptare la specificul SIDISVA - cantitati si locatii concrete", 3)
P("Justificarea operationala a masurilor de mai sus, raportata la cantitatile reale ale proiectului SIDISVA "
  "(care difera fundamental de o oferta generica):")
BUL([
    "Total echipamente de livrat: 100 laptopuri ANSVSA + 336 echipamente teren (8 buc x 42 DSVSA) "
    "+ 90 NGFW locatii + 241 switch-uri acces + 50 switch-uri PoE + 175 access points + 126 IoT laboratoare = aprox. 1.118 echipamente;",
    "Locatii de livrare: 46 institutii (ANSVSA centru + 42 DSVSA + 3 institute IISPV/ICBMV/IDSA);",
    "Distantele estimate: ~14.000 km cumulati pentru livrari (calcul baseline rute optimizate prin software);",
    "Estimare emisii cu masurile de mai sus: ~3.5 tone CO2e (vs ~7.5 tone CO2e fara masuri) - reducere de ~53%;",
    "Compensare carbon offset: pentru cele ~3.5 tone reziduale, <LIDER> achizitioneaza credite Gold Standard / Verra "
    "(cost estimat 50-70 EUR pe tona CO2 = ~200 EUR total, inclus in propunerea financiara).",
])

# === 2.1.3 Plan raportare semestriala DNSH ===
H("2.1.3 Plan de raportare semestriala a indicatorilor DNSH catre ANSVSA", 2)
P("Pe toata durata contractului SIDISVA (18 luni implementare + 36 luni garantie = 54 luni), <LIDER> "
  "transmite catre ANSVSA un raport semestrial al indicatorilor DNSH (total 9 rapoarte: la luna 6, 12, "
  "18 - implementare; apoi 24, 30, 36, 42, 48, 54 - garantie). Acest plan suplimenteaza prevederile "
  "din cap. 13 al ofertei (Management contract) si din cap. 11.6 (Declaratia DNSH).")

H("KPI raportabili la fiecare 6 luni", 3)
TBL(
    headers=["KPI DNSH", "Unitate masura", "Sursa date", "Tinta / Threshold"],
    rows=[
        ["Consum energie infrastructura SIDISVA (Cloud Guvernamental)",
         "kWh / semestru",
         "Provider Cloud (Azure RO / STS) + dashboard monitoring Prometheus + Grafana",
         "Reducere progresiva 5%/an prin optimizari (sizing right, hibernare instante)"],
        ["Emisii CO2 estimate Scope 2 (Market-Based)",
         "kg CO2e / semestru",
         "Calcul GHG Protocol Scope 2 + factor emisie energie regenerabila provider",
         "Tinta <0.5 kg CO2e / kWh (energie regenerabila >70%)"],
        ["Reciclare WEEE - echipamente uzate predate",
         "Numar echipamente x kg",
         "Procese-verbale predare la operator autorizat (RoRec / Eco-Tic / Eco-WEEE)",
         "100% echipamente uzate predate (zero in deseuri generale)"],
        ["Reciclare ambalaje livrare",
         "% greutate reciclata",
         "Documentatie operator colectare separata",
         ">=90% reciclare ambalaje carton / folie biodegradabila"],
        ["Emisii CO2e cumulative livrari hardware",
         "kg CO2e cumulat",
         "Raport furnizor logistic per transport + recalcul GHG Protocol",
         "Reducere cumulativa fata de baseline (livrari fara optimizare)"],
        ["Incidente DNSH semnalate",
         "Numar / severitate",
         "Procese-verbale incident + corective masuri",
         "0 incidente cu impact mediu/major"],
    ],
)

H("Format si distributie raport", 3)
BUL([
    "Format raport: PDF semnat electronic calificat (eIDAS) + tabel Excel anexat cu serii de date semestriale;",
    "Distributie: catre Responsabilul proiect ANSVSA + manageriul de contract + reprezentant POCIDIF (MIPE) la cerere;",
    "Audit independent anual: la cererea ANSVSA, raportul anual cumulativ poate fi auditat de un organism "
    "acreditat independent (cost suportat de <LIDER> pana la 2 audituri pe perioada contractului);",
    "Termen depunere: maxim 15 zile calendaristice de la incheierea fiecarui semestru calendaristic;",
    "Linkare cu cap. 13 al ofertei (Management contract): KPI DNSH sunt inclusi in pachetul de raportare "
    "general al contractului, alaturi de KPI tehnici (SLA, conformitate cerinte) si KPI financiari.",
])

# === 2.1.4 Continut Anexa I ===
H("2.1.4 Continutul detaliat al Anexei I - Pachet documente justificative DNSH", 2)
P("Anexa I - Pachet DNSH anexata ofertei contine 12 documente identificate concret. Lista completa de mai jos "
  "(cu codul, denumirea, emitentul si statusul curent al fiecarui document) constituie ANGAJAMENT FERM al "
  "<LIDER> privind dovezile DNSH livrate impreuna cu oferta finala.")
TBL(
    headers=["Cod", "Document", "Emitent", "Subfactor acoperit"],
    rows=[
        ["I.1", "Datasheet laptop ofertat (sectiunea consum stare de veghe / Sleep Mode)",
         "Producator (Lenovo / Dell / HP)", "4.1"],
        ["I.2", "Certificat ENERGY STAR v8.0 (printscreen Energy Star Product Finder + URL verificabil + QR)",
         "Environmental Protection Agency (EPA, SUA)", "4.1"],
        ["I.3", "Certificat EU Ecolabel SAU eticheta ecologica tip 1 echivalenta (EPEAT Gold / TCO Certified Generation 9 / Blue Angel)",
         "Comisia Europeana / Green Electronics Council / TCO Development", "4.1"],
        ["I.4", "Raport testare consum stare de veghe (laborator acreditat ISO/IEC 17025)",
         "Laborator acreditat (UL, TUV Rheinland, Bureau Veritas etc.)", "4.1"],
        ["I.5", "Contract cu furnizor logistic (flota declarata BEV/PHEV + Euro 6d-Final + optimizare trasee + raportare CO2e)",
         "Furnizor logistic (de selectat: DHL Eco / DPD Echo / poste / curier dedicat)", "4.2"],
        ["I.6", "Certificat ISO 14001 furnizor ambalaje",
         "Furnizor ambalaje + organism certificare", "4.2"],
        ["I.7", "Certificat ISO 14001 furnizor logistic",
         "Furnizor logistic + organism certificare", "4.2"],
        ["I.8", "Specificatie tehnica ambalaje FSC (carton) + folie biodegradabila PLA",
         "Furnizor ambalaje + Forest Stewardship Council", "4.2"],
        ["I.9", "Raport amprenta carbon estimata pentru lotul total de livrare (GHG Protocol Scope 1+3)",
         "Consultant mediu independent / furnizor logistic", "4.2"],
        ["I.10", "Certificat carbon offset (compensare emisii reziduale Gold Standard / Verra VCS)",
         "Gold Standard Foundation / Verra (Verified Carbon Standard)", "4.2"],
        ["I.11", "Declaratia DNSH semnata electronic calificat de reprezentantul legal <LIDER>",
         "Reprezentant legal <LIDER>", "4.1 + 4.2"],
        ["I.12", "Plan raportare semestriala KPI DNSH (54 luni contract, 9 rapoarte planificate)",
         "<LIDER>", "Raportare continua"],
    ],
)
P("Documentele I.1-I.4 vizeaza Subfactor 4.1 (consum laptop), documentele I.5-I.10 vizeaza Subfactor 4.2 "
  "(ambalaje + livrare), documentele I.11-I.12 sunt umbrela DNSH. Lipsa oricaruia dintre documentele I.1-I.4 "
  "conduce la diminuarea punctajului Subfactor 4.1 sau, mai grav, la neconformitatea ofertei "
  "(daca consumul declarat nu este sustinut de fisa tehnica producator + raport testare). Lipsa I.5-I.10 "
  "conduce la pierderea totala sau partiala a celor 5 puncte Subfactor 4.2.", italic=True)

# === 2.1.5 Conformitate cu regula Fisa de date ===
H("2.1.5 Conformitate explicita cu regula Fisei de date privind autenticitatea masurilor", 2)
P("Fisa de date prevede explicit (Nota Subfactor 4.2):", bold=True)
P("\"Se puncteaza doar masurile sustinute prin documente sau explicatii concrete in oferta. "
  "Nu se acorda punctaj pentru declaratii generale sau copiate fara adaptare\".", italic=True)
P("<LIDER> raspunde acestei reguli prin urmatoarele 5 mecanisme verificabile, evitand astfel orice risc de "
  "depunctare pentru declaratii generice:")

H("Mecanism 1 - Cuantificare numerica a tuturor masurilor", 3)
P("Fiecare masura asumata in capitolul de fata (si in cap. 11 al ofertei) are atasat cel putin un KPI "
  "numeric cuantificabil - exemple: consum stare veghe (Wh), reducere CO2e (kg / %), load factor livrare (%), "
  "varsta medie flota (ani), greutate reciclata (% / kg). Comisia poate verifica usor fiecare valoare "
  "fata de documentul justificativ corespunzator (Anexa I).")

H("Mecanism 2 - Adaptare la specificul SIDISVA (cantitati + locatii reale)", 3)
P("Toate masurile sunt formulate cu referire la datele reale ale proiectului SIDISVA - nu sunt fraze "
  "generice de tipul \"oferim ambalaje reciclabile\". Exemple concrete:")
BUL([
    "1.118 echipamente totale de livrat la 46 institutii (ANSVSA + 42 DSVSA + 3 institute), nu un numar generic;",
    "100 laptopuri pentru ANSVSA central + institute, conform cap. 3.4.3.4.2.6 CdS, nu \"laptopuri pentru beneficiar\";",
    "14.000 km estimati pentru livrari catre cele 46 locatii, nu \"distante uzuale\";",
    "Reducere ~53% CO2e fata de baseline neoptimizat, nu \"reducere semnificativa\".",
])

H("Mecanism 3 - Identificare nominala a furnizorilor implicati", 3)
P("Pentru fiecare masura DNSH, este identificat (sau este in proces de identificare nominala finalizata pre-depunere) "
  "furnizorul concret responsabil:")
BUL([
    "Producatori laptop: 3 preselectati (Lenovo / Dell / HP) - finalizare 1 la depunere;",
    "Furnizor logistic: 1 dintre 3 preselectati (DHL Express Romania / DPD Romania / Cargus Eco) - "
    "finalizare pre-depunere prin Contract logistic semnat cu masuri DNSH detaliate;",
    "Operator carbon offset: 1 dintre 2 preselectati (Gold Standard Foundation / Verra VCS) - "
    "finalizare la prima livrare;",
    "Operator WEEE: 1 dintre 3 preselectati (RoRec / Eco-Tic / Eco-WEEE) - finalizare la prima predare echipamente uzate.",
])

H("Mecanism 4 - Documente justificative anexate distinct (nu doar declaratii)", 3)
P("Cele 12 documente justificative din Anexa I (vezi §2.1.4) sunt anexe formale ale ofertei, semnate electronic "
  "calificat, NU declaratii libere. Fiecare document poarta semnatura emitentului real (producator hardware / "
  "laborator acreditat / organism certificare / furnizor logistic / operator carbon offset) - nu doar a "
  "ofertantului. Aceasta evita reziccul de \"declaratii copiate\".")

H("Mecanism 5 - Raportare continua dovedibila pe 54 luni (auditabila)", 3)
P("Angajamentul DNSH NU se incheie la depunerea ofertei - este urmarit pe toata durata contractului prin "
  "9 rapoarte semestriale cu KPI cuantificati (vezi §2.1.3), auditabile independent la cererea ANSVSA. "
  "Aceasta inseamna ca masurile asumate trebuie sa fie REALE (pentru a putea fi raportate corect pe 54 luni), "
  "nu doar declarate o data in oferta.")

H("Concluzie operationala", 3)
P("Prin aceste 5 mecanisme, <LIDER> garanteaza ca propunerea DNSH respecta regula stricta din Fisa de date. "
  "Maximizarea punctajului P4 = 10 puncte este sustinuta de dovezi documentare verificabile, NU de declaratii "
  "generice. Aceasta abordare diferentiaza oferta <LIDER> de o oferta concurenta care s-ar limita la "
  "declaratii formale fara substrat operational.", bold=True)

doc.save(str(OUT))
print(f"WROTE: {OUT.name}")

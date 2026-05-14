"""
Reconstruieste 5-Arhitectura_si_licente.docx cu stack-ul corect:
- Placeholders consortiu: <LIDER>, <FURNIZOR ETL/BI/GIS/ESB/LIMS>; VOGO si ZIPPER sunt directe (confirmate).
- Stack: MS SQL Enterprise, Oracle Service Bus, MS SSIS, MS Power BI, Keycloak, ZIPPER DMS,
  VOGO Enterprise Suite (Portal+Chatbot+Mobila), Mirth Connect, RHEL, NGINX Plus, MS IIS,
  Elasticsearch, F5/FortiWeb WAF, FortiDeceptor, Cisco DNA/ClearPass, Splunk/QRadar SIEM,
  Cisco IronPort, FortiGate/Palo Alto NGFW, CrowdStrike/SentinelOne, MS Office.
- Cantitati din Lista_Software_SIDISVA.xlsx (stack memory).
- Referinte CdS cap. 3.4.3.x.x explicite.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(r"C:\Users\adria\Documents\Claude\Projects\Licitatie 1\doc-output\oferta\5-Arhitectura_si_licente.docx")

doc = Document()
# Stiluri standard de Word; nu necesita custom font pentru python-docx
section = doc.sections[0]
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)

def H(text, level):
    doc.add_heading(text, level=level)

def P(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    if bold: r.bold = True
    if italic: r.italic = True
    return p

def BUL(items):
    for it in items:
        doc.add_paragraph(it, style='List Bullet')

def TBL(headers, rows, widths=None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Light Grid Accent 1'
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for r in hdr[i].paragraphs[0].runs:
            r.bold = True
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    if widths:
        for row in tbl.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = w
    return tbl

# === Titlu sectiune ===
H("5. Arhitectura soluției și lista licențelor", 1)
P("Capitolul răspunde la cerința IV.4.1 lit. b) din Fișa de date și la prevederile cap. 3.4.3 din Caietul de Sarcini "
  "(Arhitectura tehnică a sistemului). Sistemul SIDISVA propus de <LIDER>, în asociere cu cei 8 furnizori specializați "
  "ai consorțiului, este construit pe o arhitectură multi-nivel Cloud-Native, scalabilă orizontal, sigură și conformă "
  "cu OUG nr. 89/2022 (Cloud Guvernamental) și cu prevederile Legii 362/2018 privind securitatea cibernetică.")

# === 5.1 Arhitectura logica ===
H("5.1 Arhitectura logică (multi-tier)", 2)
P("Sistemul SIDISVA este structurat pe 5 nivele logice (tier-uri), separate funcțional și securizate prin "
  "principiile zero-trust și least-privilege. Toate componentele rulează containerizate (Docker/Kubernetes) "
  "în Cloud-ul Privat Guvernamental.")
BUL([
    "Tier 1 — Prezentare: Portal Servicii Publice (VOGO Enterprise Suite, ≥56 servicii electronice, WCAG 2.1 AA, "
    "autentificare ROeID/eIDAS), aplicații mobile native iOS + Android + PWA (VOGO Enterprise Suite), Chatbot AI "
    "cu NLP RO (VOGO Enterprise Suite), interfețe administrative pentru ANSVSA/DSVSA/Institute.",
    "Tier 2 — API Gateway: punct unic de intrare pentru cererile externe; autentificare prin Keycloak Enterprise "
    "(OAuth 2.0 + OIDC + SAML 2.0), throttling, audit, transformare protocol REST↔SOAP↔HL7.",
    "Tier 3 — Microservicii business: cele 14 module funcționale SIDISVA (LIMS, DMS, GIS, BI, Portal, SNIIA, "
    "Aplicație mobilă, Chatbot, Workflow BPMN, Notificări, Acceptare, Raportare, Audit, Trasabilitate); "
    "comunicare prin REST/JSON sincron și event-driven prin Apache Kafka pentru asincron.",
    "Tier 4 — Servicii partajate: Keycloak Enterprise (IAM, 150 useri interni + 185.000 portal externi), "
    "ZIPPER DMS Engine (motor management documente cu cod sursă livrat ANSVSA), Mirth Connect (motor HL7 FHIR), "
    "Microsoft SSIS (ETL între componente), motor workflow BPMN, serviciu notificări (email + SMS + push), audit central.",
    "Tier 5 — Date: Microsoft SQL Server Enterprise (cluster activ-pasiv 2 noduri × 8 cores, Always On + TDE) "
    "pentru date tranzacționale; Elasticsearch (cluster 3 noduri) pentru log-uri + search engine; "
    "Microsoft Analysis Services (SSAS) pentru cuburi OLAP BI; PostGIS / GIS pentru date spațiale; "
    "Azure Blob Storage S3-compatible pentru documente DMS.",
])
P("Componente cross-cutting (orizontale, aplicabile tuturor tier-urilor): securitate (defense-in-depth pe 7 straturi — "
  "cap. 5.4 și cap. 10 ofertă), monitoring (Prometheus + Grafana), logging centralizat (Elasticsearch + Kibana), "
  "deployment automatizat CI/CD (GitLab CI + ArgoCD), audit imutabil (10 ani retenție).")
P("Principii arhitecturale aplicate: SOA (Service-Oriented Architecture) + REST + Event-Driven + microservicii "
  "containerizate. Toate interfețele externe sunt documentate Swagger/OpenAPI 3.0 (cf. cerinței DEMO).")

# === 5.2 Arhitectura fizica ===
H("5.2 Arhitectura fizică (deployment)", 2)
P("Sistemul SIDISVA este implementat în Cloud-ul Privat Guvernamental (Azure Stack / Azure RO conform OUG 89/2022). "
  "Topologia de deployment este structurată pe 3 nivele teritoriale:")
BUL([
    "Centrul de date principal (Cloud Privat Guvernamental): toate componentele software, baze de date, ESB Oracle, "
    "SIEM Splunk/QRadar, NGFW redundant Palo Alto/FortiGate cu min 10 VDOM, WAF F5/FortiWeb redundant 2 buc, "
    "Honeypot FortiDeceptor (1 buc + 4 instanțe Windows × 4 vCPU), Email Security Cisco IronPort 2 buc.",
    "Site Disaster Recovery (DR): replicare sincronă SQL Server Always On pentru date critice, replicare asincronă "
    "pentru istoric și obiecte DMS; RPO ≤ 15 minute, RTO ≤ 4 ore (cf. cerinței CdS).",
    "Locații teritoriale — 45 site-uri (42 DSVSA + 3 Institute) + ANSVSA centru: 90 NGFW FortiGate 100F "
    "(2 buc/locație, redundant), 241 switch-uri acces + 50 PoE + 2 agregare, 175 access points wifi enterprise, "
    "336 echipamente teren (laptopuri DNSH-compliant + MS Office H&B 2024), 126 IoT laboratoare integrate prin Mirth Connect.",
    "Comunicații site-to-site prin VPN IPSec sau MPLS dedicat operat de furnizor telco.",
    "Edge servers la nivelul DSVSA pentru funcționalități critice cu latență scăzută (lucru offline cu sincronizare la reconectare).",
    "DMZ separată pentru Portalul Servicii Publice — izolare strictă, fără acces direct la rețeaua internă; "
    "doar prin API Gateway + WAF.",
])

# === 5.3 Inalta disponibilitate ===
H("5.3 Arhitectura operațională — Înaltă disponibilitate", 2)
P("Având în vedere caracterul critic al sistemului SIDISVA (susține procesele operaționale ale rețelei "
  "naționale sanitar-veterinare cu 5.300+ angajați), arhitectura este redundantă pe toate nivelele:")
BUL([
    "Cluster activ-activ NGINX Plus (2 × 16 cores) ca Load Balancer și reverse proxy — capacitate ≥300.000 cereri L7 HTTP/sec.",
    "Cluster activ-activ Microsoft IIS (2 × 16 cores) pentru servere aplicație .NET — scalabilitate orizontală automată.",
    "Cluster activ-pasiv Microsoft SQL Server Enterprise (2 × 8 cores, replicare sincronă Always On, TDE encryption).",
    "Cluster activ-pasiv Oracle Service Bus (2 × 16 cores/nod) ca backbone integrare.",
    "Cluster 3 noduri Elasticsearch pentru log-uri + search engine.",
    "Storage centru cu redundanță RAID 10 + replicare site-to-site DR.",
    "2 medii separate: producție + testare-dezvoltare (mediu unic pentru dev/testare/instruire conform cerinței CdS).",
])

# === 5.4 Securitate ===
H("5.4 Arhitectura de securitate (rezumat)", 2)
P("Securitatea sistemului SIDISVA este asigurată prin arhitectura defense-in-depth pe 7 straturi "
  "(detaliată complet în cap. 10 al ofertei — Securitate informatică). Echipamentele și soluțiile de securitate "
  "ofertate respectă specificațiile cap. 3.4.3.4 din Caietul de Sarcini:")
BUL([
    "WAF F5 / FortiWeb / Imperva (cap. 3.4.3.4.1.1) — 2 appliance virtuale redundante, deployabil VMware/Hyper-V/KVM, "
    "8 vCPU, 3 Gbps HTTP, ≥64 domenii administrative, modes Reverse proxy + Inline transparent, protecție OWASP Top 10.",
    "Honeypot FortiDeceptor (cap. 3.4.3.4.1.2) — 1 appliance + 4 instanțe Windows × 4 vCPU; detecție atacuri "
    "lateral movement, alimentare automată SIEM.",
    "Management centralizat rețea (NMS/NAC) Cisco DNA / Aruba ClearPass (cap. 3.4.3.4.1.3) — min 600 dispozitive gestionate.",
    "SIEM Splunk ES / IBM QRadar (cap. 3.4.3.4.1.4) — corelare evenimente, retenție 1 an cald + 6 ani rece, "
    "alertare 24×7, IOC retention 3 ani.",
    "Email Security Cisco IronPort (cap. 3.4.3.4.1.5) — 2 buc, filtrare anti-spam, anti-malware, sandboxing fișiere atașate.",
    "NGFW centru Palo Alto / FortiGate (cap. 3.4.3.4.2) — 2 buc redundant, min 10 VDOM, IPS/IDS, application control, SSL inspection.",
    "NGFW locații FortiGate 100F (cap. 3.4.3.4.2) — 90 buc (2 × 45 locații teritoriale).",
    "Antivirus / EDR endpoint CrowdStrike Falcon sau SentinelOne (Lege 354/2022 compatibil — origine SUA, NU Kaspersky) "
    "— 486 endpoints (laptopuri + teren + servere).",
])
P("Pentru fiecare echipament/soluție de securitate sunt anexate fișele tehnice ale producătorului în Anexa B "
  "(Lista echipamentelor hardware) — cap. 6 al ofertei.")

# === 5.5 Integrari externe ===
H("5.5 Integrarea cu sistemele externe", 2)
P("Backbone-ul de integrare al sistemului SIDISVA este Oracle Service Bus, ofertat de <FURNIZOR ESB> "
  "(cluster activ-pasiv 2 × 16 cores/nod). Integrările cu sisteme medicale și de laborator (IoT) se realizează "
  "prin motorul Mirth Connect (HL7 FHIR). Interfețele standard expuse:")
BUL([
    "REST + JSON (Swagger/OpenAPI 3.0) — pentru integrări sincrone cu sisteme partenere; documentație publicată "
    "pe portal pentru dezvoltatorii externi.",
    "SOAP (WS-* stack) — pentru integrări cu sisteme legacy guvernamentale care nu suportă REST.",
    "HL7 FHIR + HL7 v2.x (prin Mirth Connect) — pentru integrarea cu echipamente IoT laboratoare (126 buc) "
    "și sisteme medicale externe.",
    "CMIS (Content Management Interoperability Services) — pentru schimburi de documente cu DMS-uri externe.",
    "Event-driven prin Apache Kafka — pentru integrări asincrone scalabile (notificări, audit, sincronizare offline).",
    "Conectori dedicați (deja livrabili prin Oracle Service Bus + adaptoare custom): ROeID, eIDAS, ONRC, APIA, "
    "ANCPI, ANARZ, PNI, PDURo, BND-SNIIA, RASFF (Rapid Alert System for Food and Feed), Ghișeul.ro, Mediul electronic ANSVSA.",
])

# === 5.6 HW ===
H("5.6 Necesar hardware al soluției", 2)
P("Sizing-ul detaliat al echipamentelor hardware este prezentat în cap. 6 al ofertei (Lista hardware). "
  "Recomandările de caracteristici minime — bazate pe arhitectura logică + fizică prezentată mai sus — sunt:")

H("Mediul de producție — sistemul intern (Cloud Guvernamental)", 3)
BUL([
    "Load Balancer (NGINX Plus 2 × 16 cores): ≥1 Gbps throughput L7, ≥1 Gbps throughput SSL, ≥300.000 cereri L7 HTTP/sec.",
    "Server Aplicație IIS (cluster activ-activ 2 × 16 cores): 64 GB RAM, 512 GB SSD local.",
    "Server BD MS SQL Enterprise (cluster activ-pasiv 2 × 8 cores): 128 GB RAM, 512 GB SSD local + 10-20 TB storage extern RAID 10.",
    "Server ESB Oracle (cluster activ-pasiv 2 × 16 cores/nod): 64 GB RAM, 512 GB SSD local.",
    "Cluster Elasticsearch (3 noduri × 8 cores): 64 GB RAM/nod, 2 TB SSD/nod.",
    "Storage extern centru: 10-20 TB utili RAID 10 pentru documente DMS + log-uri SIEM + backup.",
])

H("Mediul de producție — sistemul portal expus cetățeni (DMZ)", 3)
BUL([
    "Load Balancer DMZ (NGINX Plus): aceleași specificații.",
    "Server BD + Aplicație DMZ: izolare strictă de la sistemul intern, replicare unidirecțională citire.",
    "Storage DMZ: 1-2 TB utili RAID 10.",
])

H("Mediul de testare-dezvoltare", 3)
BUL([
    "Server BD: 8 GB RAM, 8 core-uri, 512 GB SSD.",
    "Server Aplicație: 8 GB RAM, 8 core-uri, 512 GB SSD.",
    "Storage extern: 1 TB util RAID 10.",
])
P("Implementarea în Cloud Guvernamental este preferată conform OUG 89/2022; toate produsele COTS sunt "
  "BYOL (Bring Your Own License) prin Azure RO / Azure Stack.", italic=True)

# === 5.7 Lista licente ===
H("5.7 Lista licențelor software propuse", 2)
P("Tabelul de mai jos prezintă lista completă a licențelor software propuse pentru SIDISVA, structurată pe "
  "categorii. Cantitățile sunt aliniate cu cerințele din Caietul de Sarcini și cu Sinteza din "
  "Lista_Software_SIDISVA.xlsx (Anexa A). ", bold=True)
P("ATENȚIE: Conform cap. IV.4.1 lit. c) din Fișa de date, lipsa Listei Licențelor face oferta NECONFORMĂ.", italic=True)

# Tabel licente
TBL(
    headers=["Nr.", "Categorie", "Produs / Licență", "Producător", "Cantitate", "Cap. CdS"],
    rows=[
        # A. Infrastructura SW
        ["1", "OS Linux", "Red Hat Enterprise Linux 9 (sau Oracle Linux 9)", "Red Hat / Oracle", "20-30 VM", "3.4.3.2.3"],
        ["2", "OS Windows", "Windows Server 2022 Datacenter", "Microsoft", "min 4 + 6-10 VM (incl. 4 honeypots)", "3.4.3.2.3"],
        ["3", "Web server", "NGINX Plus (cluster activ-activ)", "F5 NGINX", "2 × 16 cores", "3.4.3.2.5"],
        ["4", "App server", "Microsoft IIS (inclus în Windows Server)", "Microsoft", "2 × 16 cores", "3.4.3.2.5"],
        ["5", "SGBD principal", "Microsoft SQL Server Enterprise Edition", "Microsoft", "2 noduri × 8 cores", "3.4.3.2.4"],
        ["6", "SGBD NoSQL / Search", "Elasticsearch (cluster 3 noduri, log + search)", "Elastic NV", "3 noduri × 8 cores", "3.4.3.2.4"],
        ["7", "ETL", "Microsoft SQL Server Integration Services (SSIS)", "Microsoft", "16 cores", "3.4.3.2.6"],
        ["8", "BI", "Microsoft Power BI Premium + SSRS + SSAS", "Microsoft", "50 useri analitici", "3.4.3.2.8"],
        ["9", "IAM", "Keycloak Enterprise (Red Hat SSO)", "Red Hat", "150 useri interni + 185.000 portal", "3.4.3.2.7"],
        ["10", "GIS", "Soluție GIS Enterprise (producător de selectat)", "<FURNIZOR GIS>", "1 server + 50 editori", "3.4.3.2.9"],
        ["11", "ESB", "Oracle Service Bus (cluster activ-pasiv)", "Oracle", "16 cores/nod × 2 noduri", "3.4.3.2.10"],
        # B. Software aplicativ
        ["12", "DMS", "ZIPPER DMS (licență perpetuă + cod sursă integral ANSVSA)", "ZIPPER SERVICES S.R.L.", "Nelimitat useri", "3.4.3.3.1"],
        ["13", "Portal Servicii Publice", "VOGO Enterprise Suite — Portal (≥56 servicii, perpetuă)", "VOGO", "Nelimitat useri", "3.4.3.3.2"],
        ["14", "Chatbot AI", "VOGO Enterprise Suite — Chatbot cu NLP RO (auto-răspuns + asistență formulare)", "VOGO", "1 instanță", "3.4.3.3.3"],
        ["15", "Aplicație mobilă", "VOGO Enterprise Suite — App mobilă nativ iOS + Android + PWA", "VOGO", "Apple Dev (99 $/an) + Google free", "3.4.3.3.4"],
        ["16", "LIMS", "Soluție LIMS COTS cu cod sursă (perpetuă + 3 ani L2/L3)", "<FURNIZOR LIMS>", "Nelimitat useri", "3.4.3.3.5"],
        ["17", "Integrare HL7", "Mirth Connect (HL7 FHIR + v2.x engine)", "NextGen Healthcare", "1 instanță", "3.4.3.3.6"],
        # C. Securitate
        ["18", "WAF", "F5 Advanced WAF / FortiWeb / Imperva", "F5 / Fortinet / Imperva", "2 appliance × 8 vCPU", "3.4.3.4.1.1"],
        ["19", "Honeypot", "FortiDeceptor + instanțe Windows", "Fortinet", "1 + 4 instanțe × 4 vCPU", "3.4.3.4.1.2"],
        ["20", "NMS / NAC", "Cisco DNA Center / Aruba ClearPass", "Cisco / HPE Aruba", "min 600 dispozitive", "3.4.3.4.1.3"],
        ["21", "SIEM", "Splunk Enterprise Security / IBM QRadar", "Splunk / IBM", "1 cluster, IOC 3 ani", "3.4.3.4.1.4"],
        ["22", "Email Security", "Cisco Secure Email (IronPort)", "Cisco", "2 buc", "3.4.3.4.1.5"],
        ["23", "NGFW centru", "Palo Alto Networks PA-Series / FortiGate", "Palo Alto / Fortinet", "2 buc, min 10 VDOM", "3.4.3.4.2.1"],
        ["24", "NGFW locații", "FortiGate 100F", "Fortinet", "90 buc (2 × 45 locații)", "3.4.3.4.2.2"],
        ["25", "Antivirus / EDR", "CrowdStrike Falcon / SentinelOne Singularity (NU Kaspersky — Lege 354/2022)", "CrowdStrike / SentinelOne", "436 + 50 endpoints", "3.4.3.4.3"],
        # D. Productivity
        ["26", "MS Office laptopuri", "Microsoft Office Home & Business 2024 OEM", "Microsoft", "100 buc", "3.4.3.5.1"],
        ["27", "MS Office teren", "Microsoft Office Home & Business 2024 OEM", "Microsoft", "336 buc (8 × 42 DSVSA)", "3.4.3.5.1"],
        ["28", "Backup software", "Veeam Backup & Replication Enterprise (sau echivalent)", "Veeam", "Per TB protejat", "3.4.3.5.2"],
        ["29", "Kubernetes / Containere", "Red Hat OpenShift Container Platform", "Red Hat", "Per nod cluster", "3.4.3.2.11"],
    ],
)

# === 5.8 Drepturi IP ===
H("5.8 Drepturi de proprietate intelectuală asupra soluției", 2)
P("Conform cap. 12 din Caietul de Sarcini (Transferul drepturilor de proprietate intelectuală), <LIDER> și "
  "consorțiul de furnizori asigură transferul integral către ANSVSA al drepturilor de proprietate intelectuală "
  "pentru toate componentele dezvoltate custom + livrarea cu drepturi corespunzătoare pentru componentele COTS:")
BUL([
    "Cod sursă custom dezvoltat pentru specificul SIDISVA (integrări dedicate ROeID/eIDAS/ONRC/APIA/SNIIA, "
    "adaptoare ESB custom, fluxuri BPMN specifice, rapoarte BI custom): transfer integral cu drept de proprietate "
    "perpetuă ANSVSA, fără limitare de utilizatori, instituții sau funcționalități.",
    "ZIPPER DMS — licență perpetuă + cod sursă integral livrat ANSVSA (componenta DMS este oferită cu acest drept extins, "
    "permițând întreținere/dezvoltare ulterioară independentă de producătorul original).",
    "Soluția LIMS COTS livrată de <FURNIZOR LIMS> — licență perpetuă + cod sursă integral livrat ANSVSA + 3 ani suport L2/L3 inclus.",
    "VOGO Enterprise Suite (Portal + Chatbot + Aplicație mobilă) — licență perpetuă fără limitare de utilizatori; "
    "VOGO păstrează drepturile asupra codului sursă produs (componentă COTS, nu custom dezvoltată), însă oferă "
    "ANSVSA dreptul nelimitat de utilizare + acces la API + posibilitatea integrării custom prin extensii.",
    "Licențe COTS infrastructură (Microsoft SQL Enterprise, Oracle Service Bus, Keycloak Enterprise, Splunk ES, "
    "FortiGate, F5 WAF, Cisco IronPort, CrowdStrike etc.) — emise pe numele ANSVSA, perpetue sau în mod abonament "
    "conform politicii producătorului, cu drept de transfer către instituțiile subordonate (3 institute + 42 DSVSA).",
])
P("Acceptarea acestor termeni de proprietate intelectuală este declarată formal în Anexa K (Pachet declarații obligatorii) "
  "— cap. 15 al ofertei.", italic=True)

doc.save(str(OUT))
print(f"WROTE: {OUT.name}")

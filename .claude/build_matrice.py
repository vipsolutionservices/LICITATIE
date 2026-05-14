"""Construiește anexa_f_conformitate.docx pornind de la varianta 'old'.
- Adaugă coloană nouă 'Responsabil' între 'Cerință' și 'Răspuns ofertant'
- Populează Responsabil pe baza maparii Cap CDS -> Produs recomandat din Sinteza
- Golește 'Răspuns ofertant' și 'Document referință'
- Salvează ca anexa_f_conformitate.docx
"""
import sys, re
sys.stdout.reconfigure(encoding='utf-8')
from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

SRC = r'C:\Users\adria\Documents\Claude\Projects\Licitatie 1\doc-output\oferta\anexa_f_conformitate - old.docx'
DST = r'C:\Users\adria\Documents\Claude\Projects\Licitatie 1\doc-output\oferta\anexa_f_conformitate.docx'

# Mapare Cap CDS -> Responsabil (= Produs recomandat din Sinteza)
# Pentru capitole care acoperă mai multe produse, listez separate prin " / ".
# Pentru cap. de servicii (3.4.4.x) responsabilul e VOGO TECHNOLOGY (integrator).
MAP = {
    '3.4':         'VOGO TECHNOLOGY (coordonare integrator)',
    '3.4.1.1':     'VOGO TECHNOLOGY (arhitect sistem)',
    '3.4.1.2':     'Toate produsele C. Securitate (WAF / Honeypot / NAC / SIEM / Email / NGFW)',
    '3.4.1.3':     'Toate produsele C. Securitate + VOGO TECHNOLOGY (conformitate legală)',
    '3.4.2.1':     'FURNIZOR LIMS COTS',
    '3.4.2.2':     'ZIPPER',
    '3.4.2.3':     'ZIPPER / VOGO Enterprise Suite',
    '3.4.2.5':     'Microsoft Power BI / SSRS / SSAS + Microsoft SSIS',
    '3.4.2.6':     'VOGO Enterprise Suite',
    '3.4.2.7':     'VOGO Enterprise Suite',
    '3.4.2.8':     'FURNIZOR GIS + VOGO Enterprise Suite',
    '3.4.2.9':     'VOGO Enterprise Suite',
    '3.4.2.10':    'VOGO Enterprise Suite',
    '3.4.2.11':    'VOGO Enterprise Suite (aplicație mobilă)',
    '3.4.2.12':    'VOGO Enterprise Suite',
    '3.4.2.14':    'Oracle Service Bus + Mirth Connect + VOGO Enterprise Suite',
    '3.4.3':       'VOGO TECHNOLOGY (arhitect sistem)',
    '3.4.3.1':     'RHEL 9 / Oracle Linux 9 + Win Server 2022 DC (Cloud Guvernamental)',
    '3.4.3.2.1':   'NGINX Plus',
    '3.4.3.2.2':   'Microsoft IIS',
    '3.4.3.2.3':   'RHEL 9 / Oracle Linux 9 + Win Server 2022 DC',
    '3.4.3.2.4':   'Microsoft SQL Server Enterprise + Elasticsearch',
    '3.4.3.2.5':   'Microsoft SSIS',
    '3.4.3.2.6':   'Microsoft Power BI / SSRS / SSAS',
    '3.4.3.2.7':   'Keycloak Enterprise',
    '3.4.3.2.8':   'FURNIZOR GIS',
    '3.4.3.2.9':   'Oracle Service Bus',
    '3.4.3.3.1':   'ZIPPER',
    '3.4.3.3.1.1': 'ZIPPER',
    '3.4.3.3.1.2': 'ZIPPER',
    '3.4.3.3.1.3': 'ZIPPER',
    '3.4.3.3.1.4': 'ZIPPER',
    '3.4.3.3.1.5': 'ZIPPER',
    '3.4.3.3.1.6': 'ZIPPER',
    '3.4.3.3.1.7': 'ZIPPER',
    '3.4.3.3.2':   'VOGO Enterprise Suite',
    '3.4.3.3.2.1': 'VOGO Enterprise Suite (chatbot)',
    '3.4.3.3.2.2': 'VOGO Enterprise Suite (aplicație mobilă)',
    '3.4.3.3.3':   'FURNIZOR LIMS COTS',
    '3.4.3.3.3.1': 'FURNIZOR LIMS COTS',
    '3.4.3.3.3.2': 'Mirth Connect',
    '3.4.3.3.4':   'VOGO TECHNOLOGY',
    '3.4.3.4':     'Toate produsele C. Securitate',
    '3.4.3.4.1.1': 'F5 / Imperva / FortiWeb (WAF)',
    '3.4.3.4.1.2': 'FortiDeceptor (Honeypot)',
    '3.4.3.4.1.3': 'Cisco DNA / ClearPass (NMS / NAC)',
    '3.4.3.4.1.4': 'Splunk ES / QRadar (SIEM)',
    '3.4.3.4.1.5': 'Cisco IronPort (Email Security)',
    '3.4.3.4.2.1': 'FortiGate / Palo Alto + FortiGate 100F locații',
    '3.4.3.4.2.2': 'Echipament hardware (Switch acces) — vezi Lista_Hardware',
    '3.4.3.4.2.3': 'Echipament hardware (Switch POE) — vezi Lista_Hardware',
    '3.4.3.4.2.4': 'Echipament hardware (Access point) — vezi Lista_Hardware',
    '3.4.3.4.2.5': 'Echipament hardware (Switch agregare) — vezi Lista_Hardware',
    '3.4.3.4.2.6': 'Echipament hardware (Laptop) + MS Office H&B 2024 OEM',
    '3.4.3.4.2.7': 'Echipament hardware (Complete teren) + MS Office H&B 2024 OEM',
    '3.4.4':       'VOGO TECHNOLOGY (PM + servicii)',
    '3.4.4.1':     'VOGO TECHNOLOGY (Manager de proiect)',
    '3.4.4.2':     'VOGO TECHNOLOGY (livrare/instalare/configurare)',
    '3.4.4.3':     'VOGO TECHNOLOGY (Analist business + Arhitect)',
    '3.4.4.4':     'VOGO TECHNOLOGY (Arhitect sistem + Team leader)',
    '3.4.4.5':     'VOGO TECHNOLOGY (Team leader + experți dezvoltare)',
    '3.4.4.6':     'VOGO TECHNOLOGY (Expert migrare)',
    '3.4.4.7':     'VOGO TECHNOLOGY (echipa de implementare)',
    '3.4.4.8':     'VOGO TECHNOLOGY (Expert testare)',
    '3.4.4.9':     'VOGO TECHNOLOGY (Experți instruire)',
    '3.4.4.10':    'VOGO TECHNOLOGY (echipa de punere în producție)',
    '3.4.5':       'Keycloak Enterprise + VOGO Enterprise Suite',
    '3.4.6':       'Toate produsele C. Securitate',
    '3.4.7':       'Toate produsele C. Securitate + VOGO TECHNOLOGY (GDPR)',
    '3.4.8':       'Echipament hardware (Laptop EU Ecolabel) + furnizori ambalaje/livrare',
    '3.4.9':       'VOGO TECHNOLOGY (Coordonator suport tehnic)',
}

def get_resp(cap):
    """Returnează responsabilul pentru un Cap CDS dat."""
    cap = (cap or '').strip()
    if cap in MAP:
        return MAP[cap]
    # Tratează prefixele lungi tip "Cap. 3.4.1.1 — ..."
    m = re.match(r'^Cap\.\s+(3\.4(?:\.\d+)*)', cap)
    if m and m.group(1) in MAP:
        return MAP[m.group(1)]
    # Fallback
    return 'VOGO TECHNOLOGY (de validat)'

def clone_cell_xml(src_tc):
    """Clonează o celulă pentru a o folosi ca template pt coloana nouă."""
    return deepcopy(src_tc)

def set_cell_text(cell, text):
    """Curăță textul existent și pune textul nou într-un singur paragraf."""
    # Șterge toate paragrafele existente
    for p in cell.paragraphs:
        p_element = p._element
        p_element.getparent().remove(p_element)
    # Adaugă paragraf nou
    p = cell.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(9)

print(f"Deschid {SRC}...")
doc = Document(SRC)
t = doc.tables[0]
print(f"Tabel: {len(t.rows)} rânduri × {len(t.columns)} coloane")

# Pas 1: Adaugă o coloană nouă la sfârșit (mai sigur decât în mijloc cu python-docx)
# Apoi reordonăm cu lxml.
# De fapt, cel mai sigur e să adăugăm celula direct la fiecare rând via XML.

# Pentru fiecare rând, clonez ultima celulă și o INSEREZ după coloana 3 (Cerință)
# Indexare coloane (0-based):
#   0=Nr, 1=Cap, 2=Cerință, 3=Răspuns, 4=DocRef
# Vrem ordine finală:
#   0=Nr, 1=Cap, 2=Cerință, 3=Responsabil (nou), 4=Răspuns, 5=DocRef

def get_gridspan(tc):
    """Returnează gridSpan-ul unei celule (default 1)."""
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        return 1
    gs = tcPr.find(qn('w:gridSpan'))
    if gs is None:
        return 1
    return int(gs.get(qn('w:val'), 1))

def set_gridspan(tc, n):
    """Setează gridSpan pentru o celulă."""
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        from docx.oxml import OxmlElement
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    gs = tcPr.find(qn('w:gridSpan'))
    if gs is None:
        from docx.oxml import OxmlElement
        gs = OxmlElement('w:gridSpan')
        tcPr.append(gs)
    gs.set(qn('w:val'), str(n))

from docx.oxml import OxmlElement

# Adăugăm și o coloană în grila tabelului (w:tblGrid > w:gridCol) și rebalansăm lățimile.
# Lățimi țintă (twips): Nr=400, Cap=800, Cerință=5000, Responsabil=2700, Răspuns=3900, DocRef=2600
# Total = 15400 (la fel ca tabelul original cu 5 coloane, deci nu lățim peste pagina actuală).
TARGET_WIDTHS = [400, 800, 5000, 2700, 3900, 2600]

tbl = t._tbl
tblGrid = tbl.find(qn('w:tblGrid'))
if tblGrid is not None:
    gridCols = tblGrid.findall(qn('w:gridCol'))
    if gridCols:
        # Clonăm ultima gridCol și o inserăm după a 3-a (index 2)
        new_gc = deepcopy(gridCols[2])
        gridCols[2].addnext(new_gc)
        # Acum rebalansăm lățimile pe toate cele 6 gridCol
        gridCols = tblGrid.findall(qn('w:gridCol'))
        for i, gc in enumerate(gridCols):
            gc.set(qn('w:w'), str(TARGET_WIDTHS[i]))

def set_cell_width(tc, twips):
    """Setează lățimea unei celule via w:tcPr/w:tcW."""
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW')
        tcPr.append(tcW)
    tcW.set(qn('w:w'), str(twips))
    tcW.set(qn('w:type'), 'dxa')

for ri, row in enumerate(t.rows):
    cells_tc = row._tr.findall(qn('w:tc'))
    n_tc = len(cells_tc)

    if n_tc == 5:
        # Rând normal — clonăm celula 2 (Cerință) și inserăm după ea
        template_tc = clone_cell_xml(cells_tc[2])
        # Golim conținutul celulei clonate (ștergem toate w:p)
        for p in template_tc.findall(qn('w:p')):
            template_tc.remove(p)
        # Adăugăm un w:p gol
        new_p = OxmlElement('w:p')
        template_tc.append(new_p)
        cells_tc[2].addnext(template_tc)
        # Reluăm lista de celule și setăm lățimile
        cells_tc = row._tr.findall(qn('w:tc'))
        for i, tc in enumerate(cells_tc):
            if i < len(TARGET_WIDTHS):
                set_cell_width(tc, TARGET_WIDTHS[i])
    elif n_tc == 1:
        # Rând header de secțiune — celula are gridSpan=5, o creștem la 6
        gs = get_gridspan(cells_tc[0])
        set_gridspan(cells_tc[0], gs + 1)
        # Lățimea celulei = suma totală
        set_cell_width(cells_tc[0], sum(TARGET_WIDTHS))
    elif n_tc < 5:
        # Rând cu gridSpan parțial; incrementăm gridSpan-ul ultimei celule
        last = cells_tc[-1]
        gs = get_gridspan(last)
        set_gridspan(last, gs + 1)
    else:
        # Deja are 6+ celule; doar setez lățimile
        for i, tc in enumerate(cells_tc):
            if i < len(TARGET_WIDTHS):
                set_cell_width(tc, TARGET_WIDTHS[i])

# Reparcurg tabelul după modificare
t = doc.tables[0]
print(f"După adăugare coloană: {len(t.rows)} rânduri × {len(t.columns)} coloane")

# Pas 2: Populează datele
# Coloane noi: 0=Nr, 1=Cap, 2=Cerință, 3=Responsabil, 4=Răspuns, 5=DocRef
header_done = False
hdr_set = False
populated = 0
header_rows = 0

for ri, row in enumerate(t.rows):
    cells = row.cells
    if len(cells) < 6:
        continue
    c0 = cells[0].text.strip()
    c1 = cells[1].text.strip()

    # Rândul de header (primul rând, are 'Nr.', 'Cap. CDS' etc.)
    if not hdr_set and c0 == 'Nr.' and c1 == 'Cap. CDS':
        set_cell_text(cells[0], 'Nr.')
        set_cell_text(cells[1], 'Cap. CDS')
        set_cell_text(cells[2], 'Cerință (citat din caietul de sarcini)')
        set_cell_text(cells[3], 'Responsabil')
        set_cell_text(cells[4], 'Răspuns ofertant — mod îndeplinire')
        set_cell_text(cells[5], 'Document referință')
        # Bold pentru header
        for c in cells:
            for p in c.paragraphs:
                for run in p.runs:
                    run.font.bold = True
        hdr_set = True
        continue

    # Rândurile de header de secțiune (toate celule au același text "Cap. 3.4.X — ...")
    # Au fost clonate, deci coloana nouă va avea același text — îl golim.
    # Identificare: c0 e text lung (nu e număr) și c1==c0 (merged sau acel pattern)
    if c0.startswith('Cap.') or c0 == c1:
        # Goleste toate celulele și păstreaza titlul în prima
        title = c0
        set_cell_text(cells[0], title)
        # Restul celulelor sunt goale (sau merged)
        for k in range(1, 6):
            set_cell_text(cells[k], title)
        header_rows += 1
        continue

    # Rând normal de cerință
    nr = c0
    cap = c1
    cerinta = cells[2].text.strip()
    responsabil = get_resp(cap)

    set_cell_text(cells[0], nr)
    set_cell_text(cells[1], cap)
    set_cell_text(cells[2], cerinta)
    set_cell_text(cells[3], responsabil)
    set_cell_text(cells[4], '')   # Răspuns ofertant - de completat
    set_cell_text(cells[5], '')   # Document referință - de completat
    populated += 1

print(f"Rânduri cerință populate: {populated}")
print(f"Rânduri header secțiune: {header_rows}")

# Pas 3: Adaugă o mențiune despre coloana "Responsabil" în lista de instrucțiuni
# Inserăm un punct nou înaintea primului bullet "•"
from docx.oxml import OxmlElement as _OE
inserted = False
for p in list(doc.paragraphs):
    if p.text.strip().startswith('• Pentru fiecare cerință se completează'):
        # Inserăm un paragraf nou înainte
        new_p = p.insert_paragraph_before(
            '• Coloana "Responsabil" este pre-completată cu produsul/echipa care acoperă cerința '
            '(conform Lista_Software_SIDISVA.xlsx, sheet "Sinteza", coloana "Produs recomandat (oferta)"). '
            'Fiecare responsabil completează coloanele "Răspuns ofertant" și "Document referință" pentru cerințele alocate.'
        )
        inserted = True
        break
if not inserted:
    print("[WARN] Nu am putut insera mențiunea despre Responsabil.")

# Salvez
doc.save(DST)
print(f"\n[OK] Salvat: {DST}")

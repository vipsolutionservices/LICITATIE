"""Reconstrucție 1a-oferta.docx — Oferta SIDISVA
Scrisoare de înaintare + cuprins integrat al ofertei tehnice.

Corecturi față de varianta veche:
- 0× VOGO TECHNOLOGY S.R.L. → `<LIDER>`
- 0× VOGO ENTERPRISE BUSINESS SUITE → soluția nu e un singur produs, ci un consorțiu
  cu mai multe componente (ZIPPER DMS + VOGO Enterprise Suite + LIMS + GIS + etc.)
- 0× OSIM/EUIPO marcă UE hardcodată (§0.5)
- Adăugate informații-cheie LIPSĂ:
  - Cod SMIS 336342 + finanțare POCIDIF
  - Buget estimat 85.418.857,53 lei fără TVA + plafoane 20% HW / 10% securitate
  - Cloud Guvernamental + Cloud-Native obligatoriu (cap. 3.4.3.1 CdS)
  - Durata: 18 luni implementare + 36 luni garanție
  - DEMO video = ELIMINATORIE (cap. 14 CdS)
- Cap. 7: corectat la helpdesk L-V 08:00-17:00 + 24×7 doar pentru S1 (aliniat cu §7 ofertă)
- Cap. 9: corectat la "8 experți cheie (din care 6 evaluați P2) + 12 non-cheie = 20 experți"
- Cap. 14: corectat la NIS2 (OUG 155/2024) + L 354/2022 (NU ANSSI/DNSC)
- Metodologie: corectat la PRINCE2 hibrid Scrum Agile (aliniat cu §2 + §9)
- Referință corectă "1a-oferta.docx" (NU "oferta.docx" cum era în vechi)
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

for section in doc.sections:
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)


def add_bullet(text):
    return doc.add_paragraph(text, style='List Bullet')

def add_para(text, bold=False, italic=False, size=None, center=False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    if bold:
        r.bold = True
    if italic:
        r.italic = True
    if size:
        r.font.size = Pt(size)
    return p

def add_table(headers, rows, widths_cm=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            c = t.rows[ri].cells[ci]
            c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    if widths_cm:
        for row in t.rows:
            for ci, w in enumerate(widths_cm):
                row.cells[ci].width = Cm(w)
    return t


# ============ HEADER OFERTĂ ============
add_para('OFERTĂ TEHNICĂ', bold=True, size=16, center=True)
add_para('SIDISVA', bold=True, size=20, center=True)
add_para('Sistem Informatic Digitalizat în domeniul', size=12, center=True)
add_para('Sanitar-Veterinar și pentru Siguranța Alimentelor', size=12, center=True)
doc.add_paragraph()

# Identificare proiect
doc.add_heading('Identificare proiect', level=2)

id_rows = [
    ('Autoritate Contractantă', 'ANSVSA — Autoritatea Națională Sanitară Veterinară și pentru '
     'Siguranța Alimentelor'),
    ('Cod SMIS', '336342'),
    ('Sursa de finanțare', 'POCIDIF (P2, OS 2.2.1 — e-guvernare)'),
    ('Anunț SEAP', 'CN1089237'),
    ('Caiet de Sarcini', 'Nr. 424/SCPI/29.12.2025 ; 7574/CP/2025 (Revizia 1, 252 pagini)'),
    ('Termen depunere', '21.05.2026, ora 15:00'),
    ('Valoare estimată maximă (fără TVA)', '85.418.857,53 lei'),
    ('Constrângeri buget', 'max 20% pentru HW + servicii instalare/configurare; '
     'min 10% pentru securitate cibernetică'),
    ('Durata contractului', '18 luni implementare + 36 luni garanție post-implementare'),
    ('Găzduire obligatorie (Cap. 3.4.3.1 CdS)', 'Cloud Guvernamental, arhitectură '
     'Cloud-Native + containerizare (Docker/Kubernetes)'),
    ('Criteriu de atribuire', 'Cel mai bun raport calitate-preț — 60% tehnică / 40% preț'),
    ('Beneficiari', 'ANSVSA + 3 institute subordonate (IISPV, ICBMV, IDSA) + 42 DSVSA județene'),
    ('Scară', '~185.000 utilizatori unici/an pe Portal; ~5.300 angajați + 2.600 medici vet '
     'concesionari + 4.800 utilizatori acreditați'),
]
add_table(['Aspect', 'Valoare'], id_rows, widths_cm=[6.0, 11.0])

# Ofertant
doc.add_heading('Identificare ofertant', level=2)

p = doc.add_paragraph()
p.add_run('Ofertant principal: ').bold = True
p.add_run('<LIDER>')

p = doc.add_paragraph()
p.add_run('Structura consorțiului: ').bold = True
p.add_run('asociere cu lider + parteneri specializați pe componente — '
          '<PARTENER 1>, <PARTENER 2>, …, <PARTENER N>. Detaliile complete în Cap. 1A — '
          'Prezentarea Ofertantului.')

p = doc.add_paragraph()
p.add_run('Soluția propusă: ').bold = True
p.add_run(
    'sistem informatic integrat compus din 14 componente acoperite prin produse software '
    'specializate de la furnizori multipli (ZIPPER DMS pentru managementul documentelor, '
    'VOGO Enterprise Suite pentru Portal/Chatbot/App mobilă, Microsoft SQL Server Enterprise '
    '+ Oracle Service Bus + Keycloak Enterprise + Power BI pentru infrastructura SW, '
    'producători specializați pentru LIMS și GIS). Stack-ul complet în Cap. 5 + Anexa A.'
)

# ============ §1 Cuprins ofertă ============
doc.add_heading('1. Cuprins — Capitole ofertă tehnică', level=1)

add_para(
    'Prezenta ofertă tehnică este compusă din 16 capitole numerotate + Anexa F (Matricea de '
    'Conformitate cu 1.294 cerințe) + 11 anexe justificative (A-K). Lista capitolelor și a '
    'fișierelor corespunzătoare:'
)

cap_rows = [
    ('1', 'Rezumat executiv', '1-Rezumat_executiv.docx', '—'),
    ('1A', 'Prezentarea Ofertantului (consorțiu <LIDER> + parteneri)',
     '1A-Prezentare_ofertant.docx', 'Eligibilitate'),
    ('2', 'Abordarea și metodologia propuse',
     '2-Abordare_metodologie.docx', 'Factor 3.1 — 10p'),
    ('3', 'Plan de implementare (Gantt, drumul critic, 18 luni)',
     '3-Plan_implementare.docx', 'Factor 3.1 + 3.2'),
    ('4', 'Descrierea soluției — 14 componente SIDISVA',
     '4-Descrierea_solutiei.docx', 'Conformitate'),
    ('5', 'Arhitectura soluției + lista licențelor',
     '5-Arhitectura_si_licente.docx', 'Conformitate + IP'),
    ('6', 'Lista echipamentelor hardware (laptop, terminal teren, securitate, NGFW)',
     '6-Lista_hardware.docx', 'Factor 4.1 + plafon 20%'),
    ('7', 'Plan de măsuri pentru perioada de garanție (3 ani)',
     '7-Plan_garantie.docx', 'Conformitate'),
    ('8', 'Conformitate cu specificațiile tehnice (corelat cu Anexa F)',
     '8-Conformitate_specificatii.docx', 'Eligibilitate'),
    ('9', 'Echipa de proiect (8 cheie + 12 non-cheie = 20 experți)',
     '9-Echipa_proiect.docx', 'Factor 2 — 30p'),
    ('10', 'Securitate informatică și informațională',
     '10-Securitate_informatica.docx', 'Factor 3.1 + plafon 10%'),
    ('11', 'Măsuri de respectare a principiului DNSH',
     '11-DNSH.docx', 'Factor 4 — 10p'),
    ('12', 'Plan de instruire utilizatori (147 utilizatori — 100+44+3)',
     '12-Plan_instruire.docx', 'Conformitate'),
    ('13', 'Managementul contractului (rapoarte, KPI, recepție, plată SPV 30z)',
     '13-Management_contract.docx', 'Factor 3.2 + conformitate'),
    ('14', 'DEMO video — 33 cerințe demonstrate',
     '14-DEMO_video.docx', 'ELIMINATORIE'),
    ('15', 'Declarații obligatorii (15 declarații)',
     '15-Declaratii_obligatorii.docx', 'Eligibilitate'),
    ('16', 'Anexe la propunerea tehnică (index A-K)',
     '16-Anexe.docx', 'Suport'),
]
add_table(
    ['Cap.', 'Denumire', 'Fișier', 'Punctaj / Rol'],
    cap_rows,
    widths_cm=[1.0, 8.5, 5.0, 3.0]
)

# ============ §2 Anexe ============
doc.add_heading('2. Anexe propunere tehnică', level=1)

add_para(
    'Cele 11 anexe justificative (A-K), structurate conform §16 din ofertă. Detalii integrale '
    'în 16-Anexe.docx; aici doar indexul.'
)

anexe_rows = [
    ('A', 'Lista completă licențe software (27 produse, 4 categorii)',
     'Cap. 5 + Lista_Software_SIDISVA.xlsx'),
    ('B', 'Lista hardware + fișe tehnice producători',
     'Cap. 6'),
    ('C', 'Diagrame arhitecturale (logică, fizică, securitate, integrări)',
     'Cap. 5 + Cap. 10'),
    ('D', 'Plan detaliat implementare cu Gantt (MS Project + PDF)',
     'Cap. 3'),
    ('E', 'CV-uri experți (Europass) + dovezi 5 proiecte/expert',
     'Cap. 9'),
    ('F', 'Matricea de Conformitate (1.294 cerințe)',
     'Cap. 8 — anexa_f_conformitate.docx (1.365 rânduri × 6 coloane)'),
    ('G', 'Plan detaliat de instruire (147 utilizatori, 8 sesiuni)',
     'Cap. 12'),
    ('H', 'Plan continuitate operațională (BCP) + Disaster Recovery (DRP)',
     'Cap. 7 + Cap. 10'),
    ('I', 'Documente justificative DNSH (laptop Energy Star + ambalaje + flotă)',
     'Cap. 11'),
    ('J', 'Video demonstrativ DEMO (33 cerințe eliminatorii)',
     'Cap. 14 — depus separat în SEAP'),
    ('K', 'Pachet 15 declarații semnate eIDAS QES (DUAE, IP, NIS2, L 354/2022, GDPR, …)',
     'Cap. 15'),
]
add_table(
    ['Anexa', 'Conținut', 'Secțiune corelată'],
    anexe_rows,
    widths_cm=[1.2, 9.0, 6.5]
)

# ============ §3 Maparea capitolelor → factori evaluare ============
doc.add_heading('3. Maparea capitolelor → factori de evaluare (100 puncte)', level=1)

add_para(
    'Criteriul de atribuire este „cel mai bun raport calitate-preț" (60% tehnică / 40% preț). '
    'Cele 4 factori de evaluare cumulează 100 puncte:'
)

factori_rows = [
    ('Factor 1 — Prețul ofertei', '40p', 'Algoritm: P(n) = (Preț_min / Preț_n) × 40',
     'Propunere financiară separată'),
    ('Factor 2 — Experiența 6 experți cheie evaluați', '30p', '6 × 5p; țintă 5+ proiecte per '
     'expert (cu ≥7 module + ≥1 portal)', 'Cap. 9 + Anexa E'),
    ('Factor 3 — Metodologia (subfactori 3.1 + 3.2)', '20p (10+10)', 'Excepțional/Adecvat/'
     'Acceptabil = 10/5/1; țintă Excepțional pe ambii', 'Cap. 2 + Cap. 3 + Cap. 13'),
    ('Factor 4 — DNSH (subfactori 4.1 + 4.2)', '10p (5+5)', '4.1 consum laptop < 20Wh '
     'veghe; 4.2 ambalaje reciclabile + livrare emisii reduse', 'Cap. 11 + Anexa I'),
    ('TOTAL', '100p', '', ''),
]
add_table(
    ['Factor', 'Punctaj', 'Algoritm/Țintă', 'Secțiune ofertă'],
    factori_rows,
    widths_cm=[5.5, 1.5, 6.0, 3.5]
)

p = doc.add_paragraph()
p.add_run('Cerință ELIMINATORIE separată: ').bold = True
p.add_run(
    'DEMO video conform Cap. 14 CdS — orice cerință netedeomonstrată din cele 33 listate = '
    'ofertă neconformă (eliminată din evaluare INDIFERENT de punctajul tehnic/financiar).'
)

# ============ §4 Strategie ofertă ============
doc.add_heading('4. Strategia ofertei tehnice', level=1)

add_para(
    'Oferta tehnică <LIDER> pentru SIDISVA respectă structura canonică a ofertelor tehnice '
    'complexe, structurată pentru a facilita evaluarea de către comisia ANSVSA:'
)

add_bullet('Capitolul 1 (Rezumat executiv) — gateway pentru evaluator: identitatea ofertantului, '
           'sinteza ofertei, win themes, scor țintă pe P1-P4.')
add_bullet('Capitolul 1A (Prezentarea Ofertantului) — credibilitate: profil <LIDER> + parteneri, '
           'referințe, certificări ISO 9001/14001/27001/22301.')
add_bullet('Capitolele 2-3 (Abordare/Metodologie + Plan implementare) — răspuns la Factorul 3 '
           '(20p) — metodologie hibridă PRINCE2 + Scrum Agile, ISO 21500, ITIL, plan 18 luni '
           'cu drum critic + registru riscuri 7+6.')
add_bullet('Capitolele 4-6 (Soluție + Arhitectură + Hardware) — răspuns la cerințele tehnice '
           'IV.4.1 lit. a)-c) din Fișa de date: 14 componente, arhitectură Cloud-Native pe '
           'Cloud Guvernamental, 27 licențe SW, ~536 echipamente HW.')
add_bullet('Capitolul 7 (Plan garanție) — răspuns la IV.4.1 lit. d) + cap. 3.4.9 CdS: 3 ani '
           'garanție, helpdesk L-V 08:00-17:00 + SLA 24×7 doar pentru incidente S1 (critice), '
           'integrare ulterioară mock-up sisteme guvernamentale FĂRĂ cost suplimentar.')
add_bullet('Capitolul 8 + Anexa F (Conformitate) — răspuns la IV.4.1 lit. e): matrice cu '
           'toate cele 1.294 cerințe + răspuns punct-cu-punct.')
add_bullet('Capitolul 9 (Echipa) — răspuns la Factorul 2 (30p) — 8 experți cheie (din care '
           '6 evaluați) + 12 non-cheie = 20 experți; cert. producător explicite (ZIPPER pentru '
           'Arhitect, Microsoft SQL pentru BD, Fortinet/Palo Alto pentru Sec).')
add_bullet('Capitolele 10-11 (Securitate + DNSH) — conformitate NIS2 (Dir. UE 2022/2555 + '
           'OUG 155/2024), Lege 354/2022 (anti-RU), GDPR, defense-in-depth pe 7 straturi + '
           'Factor 4 DNSH (10p).')
add_bullet('Capitolele 12-13 (Instruire + Management contract) — 147 utilizatori instruiți '
           '(100 cheie ONLINE + 44 medici vet ONLINE + 3 admin FIZIC) + rapoarte trimestriale '
           '+ KPI calitate + plată SPV 30z.')
add_bullet('Capitolele 14-15 (DEMO + Declarații) — livrabile obligatorii: video 33 cerințe '
           '(eliminator) + 15 declarații semnate eIDAS QES.')
add_bullet('Capitolul 16 (Anexe) — index complet A-K cu corelare la secțiuni.')

# ============ §5 Acceptare cap. 12 CdS — IP ============
doc.add_heading('5. Acceptare expresă a condițiilor cap. 12 CdS — Drepturi IP', level=1)

p = doc.add_paragraph()
p.add_run(
    'Conform Cap. 12 din Caietul de Sarcini, <LIDER> ACCEPTĂ EXPRES și fără rezerve '
    'condițiile privind drepturile de proprietate intelectuală:'
)

add_bullet('Toate dezvoltările / customizările realizate în cadrul contractului trec în '
           'proprietatea Beneficiarului (ANSVSA).')
add_bullet('Licențele pentru componentele aplicative dezvoltate / customizate sunt PERPETUE.')
add_bullet('Codul sursă INTEGRAL pentru componentele aplicative dezvoltate / customizate este '
           'predat Beneficiarului.')
add_bullet('IP-ul produselor COTS preexistente (ZIPPER DMS, VOGO Enterprise Suite, soluție '
           'LIMS) rămâne la producătorii respectivi, dar Beneficiarul primește licență '
           'perpetuă + cod sursă integral conform cap. 3.4.3.3.4 CdS.')

p = doc.add_paragraph()
p.add_run('Confirmare formală: ').bold = True
p.add_run(
    'această acceptare este reluată semnată în Cap. 15 — Declarații obligatorii și în '
    'Anexa K.2 (declarație separată cu semnătură eIDAS QES a reprezentantului legal).'
)

# ============ §6 Depunere SEAP ============
doc.add_heading('6. Depunere în SEAP', level=1)

add_para(
    'Acest fișier `1a-oferta.docx` servește ca document de consolidare și navigare pentru '
    'întreaga ofertă tehnică <LIDER>. Fiecare capitol este un fișier `.docx` independent, '
    'format pentru a fi semnat electronic individual.'
)

add_para('Pentru depunerea în SEAP, se utilizează următoarea structură:')

add_bullet('Plicul „Propunere Tehnică" — pachet complet: `1a-oferta.docx` (scrisoare înaintare + '
           'cuprins) + cele 16 capitole + `anexa_f_conformitate.docx` + Anexele A-I + K (PDF-uri).')
add_bullet('Plicul „Propunere Financiară" — separat, conform Anexa Financiară.')
add_bullet('Plicul „DUAE și documente eligibilitate" — DUAE + documente justificative '
           '(Anexa K).')
add_bullet('Anexa J (Video DEMO) — depusă ca fișier separat în SEAP, link cross-referențiat '
           'în Anexa J.docx.')

# ============ §7 Semnătură ============
doc.add_heading('7. Semnătură reprezentant legal', level=1)

p = doc.add_paragraph()
p.add_run(
    'Toate documentele componente ale ofertei sunt semnate electronic calificat (eIDAS QES) '
    'de reprezentantul legal al <LIDER>, conform:'
)

add_bullet('Legea nr. 455/2001 — privind semnătura electronică (RO).')
add_bullet('Regulamentul (UE) 910/2014 — eIDAS, partea III — semnături electronice calificate.')
add_bullet('Furnizor de servicii de încredere acreditat (TSP): [De completat: certSIGN / '
           'DigiSign / Trans Sped — număr certificat QES valid].')

# Bloc semnătură
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Data depunere: ').bold = True
p.add_run('[zz/05/2026]')

p = doc.add_paragraph()
p.add_run('Reprezentant legal <LIDER>: ').bold = True
add_para('[Nume Prenume] — [Funcție]')

p = doc.add_paragraph()
p.add_run('Semnătură electronică extinsă calificată (eIDAS QES): ').bold = True
add_para('Semnată conform Legii nr. 455/2001 + Regulamentului UE 910/2014 (eIDAS).', italic=True)

# ---- Save ----
doc.save(r'1a-oferta.docx')

# Verificare
from docx import Document
d2 = Document(r'1a-oferta.docx')
print(f'OK — 1a-oferta.docx scris: {len(d2.paragraphs)} paragrafe, {len(d2.tables)} tabele')

texts = []
for p in d2.paragraphs:
    texts.append(p.text)
for t in d2.tables:
    for row in t.rows:
        for cell in row.cells:
            texts.append(cell.text)
full = '\n'.join(texts)

print(f'  VOGO TECHNOLOGY              : {full.count("VOGO TECHNOLOGY")}')
print(f'  VOGO ENTERPRISE BUSINESS     : {full.count("VOGO ENTERPRISE BUSINESS")}')
print(f'  OSIM / EUIPO                 : {full.count("OSIM") + full.count("EUIPO")}')
print(f'  ANSSI / DNSC                 : {full.count("ANSSI") + full.count("DNSC")}')
print(f'  oferta.docx (vechi)          : {full.count("oferta.docx")} (total ref fișier — corect ar trebui 1: 1a-oferta.docx)')
print(f'  1a-oferta.docx (corect)      : {full.count("1a-oferta.docx")}')
print(f'  <LIDER>                      : {full.count("<LIDER>")}')
print(f'  Cloud Guvernamental          : {full.count("Cloud Guvernamental")}')
print(f'  85.418.857                   : {full.count("85.418.857")}')
print(f'  336342                       : {full.count("336342")}')
print(f'  18 luni                      : {full.count("18 luni")}')
print(f'  ELIMINATORIE / DEMO          : {full.count("ELIMINATORIE")} / {full.count("DEMO")}')
print(f'  NIS2 / OUG 155               : {full.count("NIS2")} / {full.count("OUG 155")}')
print(f'  L 354/2022                   : {full.count("Lege 354/2022") + full.count("L 354/2022")}')
print(f'  eIDAS QES                    : {full.count("eIDAS QES")}')
print(f'  1.294 cerinte                : {full.count("1.294")}')

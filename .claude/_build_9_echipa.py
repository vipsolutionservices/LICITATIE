"""Reconstrucție 9-Echipa_proiect.docx — Oferta SIDISVA
Secțiunea P2 (Factor evaluare 2) — 30 puncte (6 experți cheie × 5p).

Corecturi față de varianta veche:
- 0× VOGO TECHNOLOGY → toate cu `<LIDER>`
- Ordine experți cheie aliniată cu Cap. 8.1 CdS (Sec&com #6, BD #7, Testare #8)
- Numărul corect de roluri: 8 cheie + 12 non-cheie = 20 (NU 16)
- Subsecțiune dedicată per fiecare din cei 6 punctați + 2 nepunctați
- Certificări producător explicite: Arhitect cert. ZIPPER DMS; BD cert. Microsoft SQL Server; Sec cert. firewall FortiGate/Palo Alto
- Tabel cu 5+ proiecte de referință per expert (7 module + portal)
- Placeholderi clari `[De completat: ...]` pentru numele și detaliile reale
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

for section in doc.sections:
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)


# ---- Helpers ----
def add_bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

def add_para(text):
    return doc.add_paragraph(text)

def add_table(headers, rows, widths_cm=None, bold_first_col=False):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    # header
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    # rows
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            c = t.rows[ri].cells[ci]
            c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                    if bold_first_col and ci == 0:
                        r.bold = True
    if widths_cm:
        for row in t.rows:
            for ci, w in enumerate(widths_cm):
                row.cells[ci].width = Cm(w)
    return t


def expert_section(nr_cds, rol, certif_obligatorii, certif_producator_text,
                   experienta_specifica_cerinta, responsabilitati_sidisva,
                   componente_sidisva_owned, este_punctat=True):
    """Construiește o subsecțiune detaliată pentru un expert cheie."""
    title = f'{rol} (Rol expert nr. {nr_cds} — Cap. 8.2 CdS)'
    h3 = doc.add_heading(title, level=3)

    # (a) Identificare
    p = doc.add_paragraph()
    p.add_run('(a) Identificare expert: ').bold = True
    p.add_run(
        '[De completat: NUME COMPLET], [De completat: COMPANIE — <LIDER> sau <PARTENER>], '
        'cetățean român, angajat permanent / detașat conform declarației de disponibilitate '
        'din Anexa E. Vechime totală IT: [De completat: ≥5 ani].'
    )

    # (b) Calificare educațională
    p = doc.add_paragraph()
    p.add_run('(b) Calificare educațională: ').bold = True
    p.add_run(
        'Diplomă de licență [De completat: facultate / universitate / an absolvire] '
        '(copie diplomă atașată în Anexa E).'
    )

    # (c) Certificări obligatorii
    h_cert = doc.add_paragraph()
    h_cert.add_run('(c) Certificări obligatorii și demonstrarea îndeplinirii cerinței minime: ').bold = True

    for cert in certif_obligatorii:
        bp = doc.add_paragraph(style='List Bullet')
        bp.add_run(cert)

    if certif_producator_text:
        bp = doc.add_paragraph(style='List Bullet')
        r = bp.add_run('CERTIFICARE DE LA PRODUCĂTOR (cerință explicită CdS): ')
        r.bold = True
        bp.add_run(certif_producator_text)

    # (d) Experiență generală
    p = doc.add_paragraph()
    p.add_run('(d) Experiență profesională generală: ').bold = True
    p.add_run(
        'Minimum 5 ani în domeniul IT — [De completat: explicație traseu profesional, '
        'inclusiv perioadele de angajare și pozițiile ocupate]. Documente doveditoare: '
        'CV Europass semnat și datat + adeverințe de la angajatori (Anexa E).'
    )

    # (e) Experiență specifică — proiecte de referință
    p = doc.add_paragraph()
    p.add_run('(e) Experiență profesională specifică — proiecte de referință: ').bold = True
    p.add_run(experienta_specifica_cerinta)

    if este_punctat:
        nota = doc.add_paragraph()
        r = nota.add_run(
            'Țintă pentru punctaj maxim (5p): minimum 5 proiecte demonstrabile. '
            'Mai jos sunt enumerate cele 5 proiecte propuse pentru evaluare (Factor 2). '
            'Documente justificative anexate: copii ale Proceselor-Verbale de recepție / '
            'certificate constatatoare / recomandări — Anexa E.'
        )
        r.italic = True

    # Tabel proiecte de referință (5 rânduri)
    proj_headers = ['#', 'Nume proiect', 'Beneficiar', 'Perioadă', 'Rol expert pe proiect',
                    'Module sistem (≥7 + ≥1 portal)', 'Document justificativ']
    proj_rows = []
    for i in range(1, 6):
        proj_rows.append([
            str(i),
            f'[De completat: PROIECT_{i}_NUME]',
            f'[De completat: PROIECT_{i}_BENEFICIAR]',
            f'[De completat: PROIECT_{i}_PERIOADA]',
            f'[De completat: PROIECT_{i}_ROL]',
            f'[De completat: PROIECT_{i}_MODULE — enumerare ≥7 module + portal]',
            f'[De completat: PROIECT_{i}_DOC]'
        ])
    add_table(
        proj_headers, proj_rows,
        widths_cm=[0.8, 3.5, 3.0, 1.8, 2.8, 4.5, 2.4]
    )

    # (f) Responsabilități în SIDISVA
    p = doc.add_paragraph()
    p.add_run('(f) Responsabilități în cadrul contractului SIDISVA: ').bold = True
    p.add_run(responsabilitati_sidisva)

    if componente_sidisva_owned:
        bp = doc.add_paragraph()
        r = bp.add_run('Componente SIDISVA pe care expertul le gestionează direct: ')
        r.italic = True
        bp.add_run(componente_sidisva_owned)


# ============ HEADING ============
doc.add_heading('9. Echipa de proiect', level=1)

intro = doc.add_paragraph()
intro.add_run(
    'Capitolul răspunde la Cap. 8 din Caietul de Sarcini („Resursele necesare/expertiza '
    'necesară pentru realizarea activităților în Contract") și la Factorul de evaluare nr. 2 '
    '— „Experiența specifică a experților cheie" (30 puncte din totalul de 100 — cea mai mare '
    'pondere fixă a ofertei tehnice).'
)

intro2 = doc.add_paragraph()
intro2.add_run(
    'Echipa de proiect propusă de <LIDER>, în calitate de Ofertant principal, este compusă, '
    'conform Cap. 8.1 CdS, dintr-un număr total de 20 de experți: '
)
r = intro2.add_run('8 experți cheie (principali)')
r.bold = True
intro2.add_run(' și ')
r = intro2.add_run('12 experți non-cheie (secundari)')
r.bold = True
intro2.add_run(
    '. Cerințele minime de calificare se îndeplinesc integral pentru fiecare expert, '
    'fără cumul de funcții — un expert = o persoană fizică (Cap. 8.1 CdS).'
)

# ============ §9.1 Sumar echipă ============
doc.add_heading('9.1 Sumar — structura echipei de proiect SIDISVA', level=2)

add_para(
    'Tabelul de mai jos prezintă cele 20 de roluri din echipa Prestatorului, organizate '
    'conform Cap. 8.1 CdS. Detaliile fiecărui rol (calificare, experiență, responsabilități, '
    'documente doveditoare) sunt prezentate în secțiunile §9.3–§9.5.'
)

sumar_rows = [
    # (categorie, nr, rol, nr_pers, comp_sidisva_resp, punctat?)
    ('Cheie', '1', 'Manager de proiect', '1', 'Coordonare globală — toate componentele', 'DA — 5p'),
    ('Cheie', '2', 'Analist de business', '1', 'Cerințe funcționale pentru 14 componente', 'DA — 5p'),
    ('Cheie', '3', 'Arhitect sistem', '1', 'Arhitectură Cloud-Native + integrare ESB', 'DA — 5p'),
    ('Cheie', '4', 'Expert team leader software', '1', 'Dezvoltare DMS + Portal + Mobile + LIMS', 'DA — 5p'),
    ('Cheie', '5', 'Expert analiză și optimizare procese', '1', 'Procese veterinare ANSVSA — re-design BPMN', 'DA — 5p'),
    ('Cheie', '6', 'Expert comunicații și securitate', '1', 'Securitate perimetrală + NIS2 + L 354/2022', 'NU (cerință eligibilitate)'),
    ('Cheie', '7', 'Expert administrare baze de date', '1', 'MS SQL Server Enterprise + Elasticsearch', 'DA — 5p'),
    ('Cheie', '8', 'Expert testare', '1', 'Plan testare end-to-end + 33 cerințe DEMO', 'NU (cerință eligibilitate)'),
    ('Non-cheie', '9', 'Expert dezvoltare software', '4', 'Implementare module aplicative', '—'),
    ('Non-cheie', '10', 'Expert portal', '1', 'Componenta Portal Servicii Publice (3.4.3.3.2)', '—'),
    ('Non-cheie', '11', 'Expert BI', '1', 'Componenta BI + DW (3.4.3.2.6)', '—'),
    ('Non-cheie', '12', 'Expert integrare', '1', 'ESB + integrări guvernamentale (ROeID/eIDAS/PNI/APIA/ONRC)', '—'),
    ('Non-cheie', '13', 'Expert migrare', '1', 'Migrare date sisteme existente ANSVSA', '—'),
    ('Non-cheie', '14', 'Expert GDPR', '1', 'DPO + DPIA + registru prelucrări', '—'),
    ('Non-cheie', '15', 'Expert instruire', '2', 'Curs ~5.300 angajați ANSVSA + DSVSA + 2.600 medici vet', '—'),
    ('Non-cheie', '16', 'Coordonator suport tehnic', '1', 'Coordonare echipă L2/L3 garanție 3 ani', '—'),
]

add_table(
    ['Categ.', 'Nr.', 'Rol', 'Nr. pers.', 'Componente SIDISVA gestionate', 'Punctaj P2'],
    sumar_rows,
    widths_cm=[1.3, 0.8, 4.0, 1.0, 6.5, 3.5]
)

add_para(
    'Total roluri: 20 (8 cheie × 1 persoană + 4 dezvoltare + 1 portal + 1 BI + 1 integrare + '
    '1 migrare + 1 GDPR + 2 instruire + 1 coord. suport = 20 persoane fizice). Pentru '
    'experții non-cheie, <LIDER> demonstrează în §9.5 modalitatea de acces (resurse proprii '
    'sau externalizare prin parteneri).'
)

# ============ §9.2 Strategie P2 ============
doc.add_heading('9.2 Strategie de îndeplinire Factor 2 (țintă 30 puncte)', level=2)

add_para(
    'Conform Caietului de Sarcini (Cap. 13 — Factor 2), punctajul P2 se calculează astfel:'
)

add_table(
    ['Nr. proiecte/contracte similare per expert', 'Punctaj per expert', 'Calificativ'],
    [
        ['2 proiecte (cerință minimă obligatorie)', '0 puncte', 'Eligibil, fără bonus'],
        ['3 proiecte', '3 puncte', 'Bonus parțial'],
        ['4 proiecte', '4 puncte', 'Bonus parțial'],
        ['5 sau mai multe proiecte', '5 puncte', 'Bonus maxim'],
    ],
    widths_cm=[6.5, 3.0, 4.5]
)

add_para(
    'Definiția proiectului similar — pentru cei 6 experți punctați (PM, Analist business, '
    'Arhitect sistem, Team leader software, Procese, BD): servicii de '
    'dezvoltare/configurare a unei soluții software / sistem informatic cu '
)
ip = doc.paragraphs[-1]
r = ip.add_run('minimum 7 module interconectate, din care cel puțin 1 modul de tip portal')
r.bold = True
ip.add_run(
    '. Excepție: Expertul 5 (Analiză și optimizare procese) — cerința este mai relaxată, '
    'fără pragul de 7 module + portal (Cap. 8.2 CdS).'
)

p = doc.add_paragraph()
p.add_run('Țintă <LIDER>: ').bold = True
p.add_run(
    'minimum 5 proiecte/contracte de referință per fiecare dintre cei 6 experți punctați '
    '⇒ 6 × 5p = '
)
r = p.add_run('30 puncte (maximum P2)')
r.bold = True
p.add_run('. Sub această țintă: 4 proiecte ⇒ pierde 1p/expert, 3 proiecte ⇒ pierde 2p/expert.')

p = doc.add_paragraph()
p.add_run('Verificare comisie evaluare: ').bold = True
p.add_run(
    'fiecare proiect prezentat este însoțit de document justificativ (Proces-Verbal de '
    'recepție, certificat constatator emis conform Legii 98/2016, recomandare scrisă a '
    'beneficiarului) — toate atașate în Anexa E.'
)

# ============ §9.3 Experți cheie PUNCTAȚI (6) ============
doc.add_heading('9.3 Experți cheie punctați (Factor 2 — 30p maxim)', level=2)

add_para(
    'Subsecțiunile §9.3.1–§9.3.6 prezintă cei 6 experți cheie evaluați conform Factor 2. '
    'Pentru fiecare expert sunt prezentate: (a) identificare, (b) calificare educațională, '
    '(c) certificări obligatorii (inclusiv certificarea de la producător acolo unde CdS '
    'cere explicit), (d) experiență profesională generală (≥5 ani IT), (e) experiență '
    'profesională specifică — 5 proiecte de referință, (f) responsabilități în SIDISVA.'
)

# --- 9.3.1 Manager de proiect ---
expert_section(
    nr_cds=1,
    rol='9.3.1 Manager de proiect',
    certif_obligatorii=[
        'Studii superioare absolvite cu diplomă de licență (Anexa E — copie diplomă).',
        'Certificare în management de proiect — minimum una recunoscută la nivel național '
        'sau internațional, demonstrată prin diplomă/certificare: '
        '[De completat: PMP / PRINCE2 / IPMA / AgilePM — Nr. certificare + Data emiterii + '
        'Organismul emitent].',
        'Cunoașterea metodologiei de management de proiect propusă prin Oferta tehnică '
        '(Cap. 2 — Abordare și metodologie): hibridă PRINCE2 + Scrum Agile.',
        'Cunoștințe SDLC (Software Development Life Cycle) + utilizare software dedicat '
        '(MS Project / Jira / Confluence).',
    ],
    certif_producator_text=None,
    experienta_specifica_cerinta=(
        'Cerința CdS Cap. 8.2 (Rol 1): minimum 2 proiecte/contracte cu sistem informatic '
        'cu ≥7 module interconectate, din care ≥1 modul de tip portal, în care expertul '
        'a desfășurat activități similare poziției de Manager de proiect, cu livrabile '
        'recepționate fără obiecții de beneficiar.'
    ),
    responsabilitati_sidisva=(
        'Coordonare globală a contractului SIDISVA pe durata 18 luni implementare + 36 luni '
        'garanție. Pregătirea planului de activități (Cap. 3 ofertă — Plan de implementare), '
        'asigurarea resurselor, monitorizarea Cap. 4 componente, gestionarea cererilor de '
        'schimbare, raportare către ANSVSA (inițial, trimestrial, final — Cap. 13 ofertă), '
        'coordonarea consorțiului <LIDER> + <PARTENER> + ZIPPER + VOGO + furnizori '
        'specializați (LIMS, GIS, ETL, BI, ESB).'
    ),
    componente_sidisva_owned='Toate cele 14 componente — coordonare orizontală.',
)

# --- 9.3.2 Analist de business ---
expert_section(
    nr_cds=2,
    rol='9.3.2 Analist de business',
    certif_obligatorii=[
        'Studii superioare absolvite cu diplomă de licență (Anexa E).',
        'Certificare în analiză de business și modelare procese — minimum una recunoscută '
        'la nivel național sau internațional: [De completat: CBAP / CCBA / IIBA-AAC / '
        'OMG OCEB BPMN — Nr. certificare + Data + Organism emitent].',
        'Cunoștințe BPMN 2.0 + UML + tehnici de analiză (user stories, use case, '
        'requirements engineering).',
    ],
    certif_producator_text=None,
    experienta_specifica_cerinta=(
        'Cerința CdS Cap. 8.2 (Rol 2): minimum 2 proiecte/contracte cu sistem informatic '
        'cu ≥7 module interconectate + ≥1 modul portal, în care expertul a desfășurat '
        'activități similare poziției de Analist de business.'
    ),
    responsabilitati_sidisva=(
        'Analiza cerințelor de business pentru toate cele 14 componente SIDISVA. '
        'Elaborarea documentelor de specificații funcționale, modelarea proceselor '
        'veterinare în BPMN împreună cu Expertul de procese (§9.3.5), agreerea soluției '
        'cu ANSVSA, elaborarea scenariilor de testare împreună cu Expertul testare '
        '(§9.4.2), suport beneficiar în tranziție și UAT, elaborarea manualelor de '
        'utilizare.'
    ),
    componente_sidisva_owned=(
        'Cerințe funcționale pentru: LIMS, DMS, Portal, BND-SNIIA, App mobile, '
        'Catagrafie, Supraveghere/Anchete, Autorizare/Acreditare.'
    ),
)

# --- 9.3.3 Arhitect sistem ---
expert_section(
    nr_cds=3,
    rol='9.3.3 Arhitect sistem',
    certif_obligatorii=[
        'Studii superioare absolvite cu diplomă de licență (Anexa E).',
        'Certificare în arhitecturi enterprise — minimum una recunoscută internațional: '
        '[De completat: TOGAF 9/10 Certified / Open CA / AWS Solutions Architect '
        'Professional / Azure Solutions Architect Expert — Nr. certificare + Data].',
        'Certificare în interoperabilitate sisteme informatice: [De completat — ex. '
        'CCSP, ISO 25010, IEEE Software Architecture].',
        'Certificare în dezvoltare aplicații sau baze de date: [De completat — '
        'ex. Oracle Certified Master, Microsoft Certified: Azure Database Administrator].',
    ],
    certif_producator_text=(
        'Certificare emisă de PRODUCĂTORUL soluției DMS ofertate (ZIPPER) — Cap. 8.2 CdS, '
        'Rol 3, paragraf 3: „competențe tehnice privind proiectarea, configurarea, '
        'dezvoltarea și integrarea soluției de management de documente ofertată… cât și '
        'de către producătorul soluției ofertate". [De completat: Certificat ZIPPER '
        'Solution Architect / ZIPPER Certified Implementation Specialist — Nr. certificat '
        '+ Data emiterii + Persoană de contact ZIPPER pentru verificare].'
    ),
    experienta_specifica_cerinta=(
        'Cerința CdS Cap. 8.2 (Rol 3): minimum 2 proiecte/contracte cu sistem informatic '
        '≥7 module + ≥1 modul portal, în care expertul a desfășurat activități similare '
        'poziției de Arhitect sistem.'
    ),
    responsabilitati_sidisva=(
        'Arhitectura macro a SIDISVA: Cloud-Native pe Cloud Guvernamental, microservicii '
        'containerizate (Docker/Kubernetes), backbone ESB (Oracle Service Bus), IAM '
        'unificat (Keycloak Enterprise) cu integrări ROeID + eIDAS, modelul de date '
        'distribuit (MS SQL Server Enterprise + Elasticsearch + repository ZIPPER DMS). '
        'Definirea standardelor de dezvoltare, supervizarea procedeelor de testare '
        'arhitecturală, întocmirea documentației tehnice de arhitectură (Cap. 5 ofertă).'
    ),
    componente_sidisva_owned=(
        'Arhitectură generală + Stack tehnologic (Cap. 5) + Integrări ESB + IAM Keycloak '
        '+ Containerizare K8s pe Cloud Guvernamental.'
    ),
)

# --- 9.3.4 Team leader software ---
expert_section(
    nr_cds=4,
    rol='9.3.4 Expert team leader software',
    certif_obligatorii=[
        'Studii superioare absolvite cu diplomă de licență (Anexa E).',
        'Certificare în arhitecturi software: [De completat — ex. SEI Software Architect, '
        'Microsoft Certified Solutions Developer, Java EE Architect].',
        'Certificare în metodologii Scrum Agile: [De completat: Certified Scrum Master '
        '(CSM) / PSM I-II / SAFe Agilist — Nr. certificare + Data + Organism].',
        'Certificare în tehnologii cloud — preferabil Microsoft Azure (Cloud Guvernamental '
        'RO rulează Azure Stack / Azure România): [De completat: Azure Solutions Architect '
        'Expert / Azure Developer Associate — Nr. certificat + Data].',
    ],
    certif_producator_text=None,
    experienta_specifica_cerinta=(
        'Cerința CdS Cap. 8.2 (Rol 4): minimum 2 proiecte/contracte cu sistem informatic '
        '≥7 module + ≥1 modul portal, cu activități similare poziției de Team leader '
        'software.'
    ),
    responsabilitati_sidisva=(
        'Coordonarea întregii echipe tehnice (cei 4 Experți dezvoltare software non-cheie '
        '+ Expert portal + Expert BI + Expert integrare), alocarea sarcinilor pe sprinturi '
        '(2 săptămâni), gestionarea release-urilor, validarea specificațiilor tehnice de '
        'design, întocmirea rapoartelor tehnice către Managerul de proiect (§9.3.1). '
        'Aplicare procese DevSecOps + management calitate (Cap. 2.12 ofertă — Securitate).'
    ),
    componente_sidisva_owned=(
        'DMS (cu ZIPPER) + Portal (VOGO Enterprise Suite) + App mobilă cetățeni '
        '(VOGO Enterprise Suite) + LIMS (împreună cu <FURNIZOR LIMS>).'
    ),
)

# --- 9.3.5 Expert analiză și optimizare procese ---
expert_section(
    nr_cds=5,
    rol='9.3.5 Expert analiză și optimizare procese',
    certif_obligatorii=[
        'Studii superioare absolvite cu diplomă de licență (Anexa E).',
        'Certificare în analiza și optimizarea proceselor de business — Cap. 8.2 CdS '
        'menționează EXPLICIT „Lean Six Sigma sau echivalent": [De completat: Six Sigma '
        'Green Belt / Black Belt / Master Black Belt / Lean Practitioner — Nr. certificare '
        '+ Data + Organism emitent (ASQ / IASSC / Lean Six Sigma Institute)].',
        'Cunoștințe avansate BPMN 2.0 + analiză de date pentru identificare cauze variație '
        '+ gândire sistemică.',
    ],
    certif_producator_text=None,
    experienta_specifica_cerinta=(
        'Cerința CdS Cap. 8.2 (Rol 5) — EXCEPȚIE DE LA REGULA „7 module + portal": '
        'minimum 2 proiecte/contracte cu sistem informatic (FĂRĂ pragul de 7 module/portal), '
        'cu activități similare poziției de Expert analiză/optimizare procese. Această '
        'flexibilitate facilitează atingerea pragului de 5 proiecte pentru punctaj maxim.'
    ),
    responsabilitati_sidisva=(
        'Re-design BPMN al proceselor veterinare ANSVSA (~50 procese: autorizare unități, '
        'supraveghere focare, raportare boli notificabile, gestiune intercomparare, plată '
        'tarife). Identificare puncte de eficientizare (Lean — reducerea risipei) + '
        'reducerea variației (Six Sigma — DMAIC). Definirea KPI-urilor de performanță a '
        'sistemului împreună cu PM (§9.3.1). Participare la formarea echipei.'
    ),
    componente_sidisva_owned=(
        'Toate procesele veterinare digitalizate — Supraveghere, Catagrafie, Autorizare/'
        'Acreditare, Intercomparare, Instruire personal laboratoare.'
    ),
)

# --- 9.3.6 Expert administrare BD ---
expert_section(
    nr_cds=7,
    rol='9.3.6 Expert administrare baze de date',
    certif_obligatorii=[
        'Studii superioare absolvite cu diplomă de licență (Anexa E).',
        'Certificare în administrare/optimizare/HA/securitate BD recunoscută internațional: '
        '[De completat: Microsoft Certified — Azure Database Administrator Associate / '
        'MCSA SQL Server / Oracle Certified Professional / IBM DB2 Certified — Nr. '
        'certificare + Data + Organism].',
    ],
    certif_producator_text=(
        'Certificare emisă de PRODUCĂTORUL soluției SGBD ofertate (Microsoft SQL Server '
        'Enterprise — Cap. 5 ofertă, conform Sinteza Lista_Software_SIDISVA.xlsx) — '
        'Cap. 8.2 CdS, Rol 7: „competențe privind administrarea, optimizarea, asigurarea '
        'înaltei disponibilități și asigurarea securității bazei de date ofertate, '
        'demonstrate prin deținerea a cel puțin unei diplome… cât și de către producătorul '
        'soluției ofertate". [De completat: Microsoft Certified: Azure Database '
        'Administrator Associate / Microsoft Certified: SQL Server 2022 Database '
        'Administrator — Nr. certificat + Data emiterii + Microsoft Transcript ID '
        'verificabil online].'
    ),
    experienta_specifica_cerinta=(
        'Cerința CdS Cap. 8.2 (Rol 7): minimum 2 proiecte/contracte cu sistem informatic '
        '≥7 module + ≥1 modul portal, cu activități similare poziției de Administrator BD.'
    ),
    responsabilitati_sidisva=(
        'Configurarea MS SQL Server Enterprise (2 noduri × 8 cores, Always On AG pentru '
        'HA, TDE pentru securitate at-rest), Elasticsearch (cluster 3 noduri pentru log-uri), '
        'stabilirea strategiei backup&recovery, monitorizarea performanței + tuning '
        'interogări, definirea politicilor de access control + audit, raportare KPI BD '
        'lunară. Asigurarea conformității cu Lege 354/2022 (verificare origine producător).'
    ),
    componente_sidisva_owned=(
        'MS SQL Server Enterprise (toate componentele aplicative) + Elasticsearch (SIEM '
        'log retention 3 ani) + tuning Elasticsearch pentru DMS Captură OCR.'
    ),
)

# ============ §9.4 Experți cheie NEPUNCTAȚI ============
doc.add_heading('9.4 Experți cheie nepunctați (cerință de eligibilitate)', level=2)

add_para(
    'Experții 6 (Comunicații și securitate) și 8 (Testare) sunt EXPERȚI CHEIE OBLIGATORII '
    'conform Cap. 8.1 CdS, dar NU sunt incluși în lista celor 6 evaluați la Factor 2. '
    'Lipsa unuia dintre ei = ofertă neconformă (Cap. 8.1 CdS). Sunt nominalizați nominal '
    'în această secțiune, cu CV + certificări + minimum 2 proiecte de referință.'
)

# --- 9.4.1 Expert comunicații și securitate ---
expert_section(
    nr_cds=6,
    rol='9.4.1 Expert comunicații și securitate',
    certif_obligatorii=[
        'Studii superioare absolvite cu diplomă de licență (Anexa E).',
        'Certificare în arhitecturi de securitate IT recunoscută internațional: '
        '[De completat: CISSP / CISM / CompTIA Security+ / GIAC GSEC — Nr. certificat + '
        'Data + Organism (ISC2 / ISACA / GIAC)].',
    ],
    certif_producator_text=(
        'CERTIFICARE EMISĂ DE PRODUCĂTORUL SOLUȚIEI FIREWALL OFERTATE — Cap. 8.2 CdS, '
        'Rol 6: „certificări recunoscute… care atestă calificarea profesională în '
        'utilizarea și administrarea soluțiilor de tip firewall propuse". Stack-ul ofertat '
        '(Cap. 5 + Sinteza xlsx) include FortiGate / Palo Alto Networks pentru NGFW centru '
        '+ 90 buc FortiGate 100F pentru locații. [De completat: Fortinet NSE 4/5/6/7 '
        'Certified / Palo Alto Networks PCNSE — Nr. certificare + Data emiterii + '
        'verificare online producător].'
    ),
    experienta_specifica_cerinta=(
        'Cerința CdS Cap. 8.2 (Rol 6): minimum 2 proiecte/contracte cu sistem informatic, '
        'cu activități similare poziției (NU se cere pragul de 7 module + portal pentru '
        'acest rol).'
    ),
    responsabilitati_sidisva=(
        'Instalarea și configurarea NGFW centru (2 buc, redundant, ≥10 VDOM) + NGFW locații '
        '(90 buc FortiGate 100F pe 45 locații DSVSA), WAF, Honeypot FortiDeceptor (4 '
        'instanțe Windows), Email Security Cisco IronPort, SIEM (Splunk ES / IBM QRadar). '
        'Monitorizare SOC 24×7 (Cap. 10 ofertă), raportare incidente CSIRT-RO conform '
        'OUG 155/2024 (NIS2), testare semestrială penetration testing.'
    ),
    componente_sidisva_owned=(
        'Securitate perimetrală + SOC + SIEM + IR/CSIRT + audit semestrial pen-test + '
        'conformitate Lege 354/2022.'
    ),
    este_punctat=False,
)

# --- 9.4.2 Expert testare ---
expert_section(
    nr_cds=8,
    rol='9.4.2 Expert testare',
    certif_obligatorii=[
        'Studii superioare absolvite cu diplomă de licență (Anexa E).',
        'Certificare în testarea funcționalității sistemelor informatice: '
        '[De completat: ISTQB Foundation / Advanced Test Manager / Advanced Test Analyst '
        '/ Advanced Test Automation Engineer — Nr. certificat + Data + ISTQB Member Board '
        '(RSTQB)].',
        'Cunoștințe testare performanță (JMeter / k6), securitate (OWASP ZAP / Burp Suite), '
        'regresie automatizată (Selenium / Playwright / Cypress).',
    ],
    certif_producator_text=None,
    experienta_specifica_cerinta=(
        'Cerința CdS Cap. 8.2 (Rol 8): minimum 2 proiecte/contracte cu sistem informatic, '
        'cu activități similare poziției (NU se cere pragul de 7 module + portal pentru '
        'acest rol).'
    ),
    responsabilitati_sidisva=(
        'Întocmirea Planului de testare end-to-end pentru toate cele 14 componente SIDISVA. '
        'Scenarii de test funcțional, de integrare, de performanță (load + stress), de '
        'securitate (OWASP Top 10), de regresie. Coordonare directă a DEMO-ului video '
        '(33 cerințe demonstrate — cerință eliminatorie). UAT cu utilizatorii ANSVSA (cap. '
        '13 ofertă). Bug tracking + raportare săptămânală.'
    ),
    componente_sidisva_owned=(
        'Plan testare global + DEMO video (33 cerințe eliminatorii) + UAT + testare regresie '
        'în garanție.'
    ),
    este_punctat=False,
)

# ============ §9.5 Experți non-cheie ============
doc.add_heading('9.5 Experți non-cheie (12 persoane)', level=2)

add_para(
    'Conform Cap. 8.3 CdS, experții non-cheie nu se nominalizează nominal obligatoriu, '
    'dar Ofertantul demonstrează în Propunerea Tehnică că are acces la personal cu '
    'expertiza necesară (resurse proprii sau externalizare). Tabelul de mai jos prezintă '
    'cei 12 experți non-cheie cu modalitatea de acces propusă de <LIDER>.'
)

non_cheie_rows = [
    ('Expert dezvoltare software (×4)', '4', 'Implementare module aplicative — backend (Java/'
     '.NET) + frontend (React/Vue) + mobile (Swift/Kotlin)', 'Resurse proprii <LIDER> + '
     '<PARTENER>', 'Anexa E.9 — declarații acces'),
    ('Expert portal (×1)', '1', 'Componenta Portal Servicii Publice (Cap. 4.2.6 ofertă) — '
     'VOGO Enterprise Suite', 'VOGO (partener consorțiu)', 'Anexa E.10'),
    ('Expert BI (×1)', '1', 'Componenta Raportare/BI/Dashboard (Cap. 4.2.5 ofertă) — '
     'Microsoft Power BI + SSRS + SSAS', '<FURNIZOR BI> (partener consorțiu)', 'Anexa E.11'),
    ('Expert integrare (×1)', '1', 'ESB Oracle + integrări guvernamentale (ROeID, eIDAS, '
     'PNI, PCUe, ONRC, APIA, ANCPI, ANARZ, CMV, Ghișeul.ro)', 'Resurse proprii <LIDER>',
     'Anexa E.12'),
    ('Expert migrare (×1)', '1', 'Migrare date din sistemele existente ANSVSA (DSVSA-uri '
     'județene, IISPV, ICBMV, IDSA) către SIDISVA — ETL + reconciliere', 'Resurse proprii '
     '<LIDER>', 'Anexa E.13'),
    ('Expert GDPR (×1)', '1', 'Rol DPO (Data Protection Officer) — registru prelucrări, '
     'DPIA per componentă, conformitate Reg. UE 679/2016 + Lege 190/2018', 'Externalizare '
     'specializată (cabinet certificat)', 'Anexa E.14 — contract servicii DPO'),
    ('Expert instruire (×2)', '2', 'Curs ~5.300 angajați ANSVSA + 42 DSVSA + 2.600 medici '
     'vet concesionari + 4.800 utilizatori acreditați; materiale didactice; evaluare',
     'Resurse proprii <LIDER> + <PARTENER>', 'Anexa E.15'),
    ('Coordonator suport tehnic (×1)', '1', 'Coordonare echipă L2/L3 garanție 3 ani de la '
     'PIP — helpdesk + monitorizare SLA + escaladare', 'Resurse proprii <LIDER>',
     'Anexa E.16'),
]

add_table(
    ['Rol (Cap. 8.1 CdS)', 'Nr.', 'Activități în SIDISVA', 'Modalitate de acces',
     'Doc. doveditor'],
    non_cheie_rows,
    widths_cm=[3.5, 0.8, 5.5, 3.5, 3.0]
)

add_para(
    'Total experți non-cheie: 12 persoane (4+1+1+1+1+1+2+1). Conform Cap. 8.3 CdS, modul '
    'de acces este demonstrat documentar prin: contracte de muncă (resurse proprii), '
    'contracte de subcontractare (parteneri consorțiu), contracte de prestări servicii '
    '(externalizare DPO). Toate documentele sunt atașate în Anexa E (subsecțiunile '
    'E.9–E.16).'
)

# ============ §9.6 Suport + infrastructură ============
doc.add_heading('9.6 Personal suport/backstopping și infrastructura Contractantului', level=2)

add_para(
    'Conform Cap. 8.4 CdS, <LIDER> furnizează integral personalul suport/auxiliar (personal '
    'administrativ, juridic, financiar, achiziții, HR) necesar pentru îndeplinirea '
    'obligațiilor contractuale. Echipa de backstopping pentru experții principali include:'
)

add_bullet(
    'Project Management Office (PMO) — 2 persoane: secretariat tehnic + reporting + '
    'tools (MS Project, Jira, Confluence) — asigură continuitatea coordonării când PM '
    'este în deplasare.'
)
add_bullet(
    'Personal administrativ — contabilitate, juridic, achiziții HW/SW, asigurarea '
    'logisticii pentru deplasări în cele 42 DSVSA județene + 3 institute.'
)
add_bullet(
    'Personal de comunicare — 1 persoană dedicată pentru raportare lunară către ANSVSA '
    '+ comunicare cu cele 3 institute (IISPV, ICBMV, IDSA).'
)

add_para(
    'Conform Cap. 8.5 CdS, <LIDER> dispune și pune la dispoziția echipei: sediu funcțional, '
    'echipamente IT (laptopuri profesionale pentru toți cei 20 experți), licențe software '
    'profesionale (MS Office, IDE-uri, tools de testare), acces VPN securizat la mediul '
    'de dezvoltare/staging, săli de ședință cu echipament de videoconferință pentru '
    'workshopuri cu ANSVSA, linii de comunicații dedicate pentru SOC 24×7.'
)

# ============ §9.7 Disponibilitate + înlocuire ============
doc.add_heading('9.7 Disponibilitate, declarații și procedura de înlocuire experți', level=2)

add_para(
    'Toți cei 8 experți cheie sunt fie angajați permanenți ai <LIDER> / <PARTENER>, fie '
    'au semnat Declarație de disponibilitate cu referire explicită la procedura SIDISVA '
    '(Caiet de Sarcini nr. 424/SCPI/29.12.2025 ; 7574/CP/2025), conform Cap. 8 CdS, '
    'paragraf privind experții ne-angajați. Declarațiile sunt atașate în Anexa E.'
)

p = doc.add_paragraph()
p.add_run('Procedura de înlocuire experți (Cap. 8 CdS + HG 395/2016 art. 162): ').bold = True
p.add_run(
    'pe parcursul derulării contractului, înlocuirea unui expert nominalizat se face '
    'NUMAI cu acceptul prealabil al Autorității Contractante. Conform Cap. 8 CdS:'
)

add_bullet(
    '(a) Noul expert nominalizat trebuie să îndeplinească integral cerințele minime de '
    'calificare ale expertului înlocuit (Cap. 8.2 CdS).'
)
add_bullet(
    '(b) Noul expert nominalizat trebuie să obțină CEL PUȚIN ACELAȘI punctaj la Factor 2 '
    'ca expertul înlocuit (pentru cei 6 experți punctați).'
)
add_bullet(
    '(c) <LIDER> transmite ANSVSA toate documentele solicitate prin documentația de '
    'atribuire pentru noul expert (CV, certificări, diplome, proiecte de referință, '
    'declarație disponibilitate).'
)
add_bullet(
    '(d) În caz de neîndeplinire a condițiilor (a) sau (b), înlocuirea devine modificare '
    'substanțială (art. 221 Legea 98/2016) — caz în care contractul poate fi reziliat.'
)

add_para(
    'Pentru a minimiza riscul de înlocuire forțată, <LIDER> menține pe perioada contractului '
    'o rezervă de minim 1 expert „back-up" pentru fiecare rol cheie punctat, demonstrabilă '
    'la cererea ANSVSA.'
)

# ============ §9.8 Documente justificative ============
doc.add_heading('9.8 Documente justificative anexate (Anexa E)', level=2)

add_para(
    'Pentru fiecare expert cheie propus, <LIDER> anexează în Anexa E documentele '
    'doveditoare conform Cap. 8 CdS:'
)

add_bullet('CV în format Europass — semnat de titular și datat (subsecțiunea E.1–E.8 pentru cei 8 cheie).')
add_bullet('Diplomă de licență studii superioare — copie certificată „Conform cu originalul".')
add_bullet(
    'Certificări profesionale obligatorii — copii certificate (PMP/PRINCE2 pentru PM, CBAP '
    'pentru Analist, TOGAF pentru Arhitect, Scrum Master + Azure pentru Team leader, '
    'Lean Six Sigma pentru Procese, MS SQL Server pentru BD, CISSP + Fortinet NSE pentru '
    'Sec, ISTQB pentru Testare).'
)
add_bullet(
    'Certificări de la PRODUCĂTOR — copii + certificate de verificare online (Microsoft '
    'Transcript / ZIPPER Certified Partner Portal / Fortinet NSE Verification).'
)
add_bullet(
    'Proiecte de referință — pentru fiecare proiect declarat: Proces-Verbal de recepție, '
    'certificat constatator emis conform Legii 98/2016 art. 166, recomandare scrisă a '
    'beneficiarului cu lista celor ≥7 module + portal.'
)
add_bullet(
    'Declarație de disponibilitate semnată de titular cu referire la procedura SIDISVA — '
    'pentru toți experții ne-angajați direct ai <LIDER>.'
)
add_bullet(
    'Tabel sintetic de îndeplinire cerințe minime per expert (Cap. 8 CdS, paragraf privind '
    'documentele justificative) — care indică în clar documentul care probează fiecare '
    'cerință minimă (subsecțiunea E.17).'
)

p = doc.add_paragraph()
r = p.add_run(
    'Anexa E este structurată identic cu Cap. 8.2 CdS (8 subsecțiuni per expert cheie, '
    'în ordine), facilitând verificarea de către comisia de evaluare ANSVSA.'
)
r.italic = True

# ---- Save ----
doc.save(r'9-Echipa_proiect.docx')

# Verificare
from docx import Document
d2 = Document(r'9-Echipa_proiect.docx')
print(f'OK — 9-Echipa_proiect.docx scris: {len(d2.paragraphs)} paragrafe, {len(d2.tables)} tabele')

# Statistici
texts = []
for p in d2.paragraphs:
    texts.append(p.text)
for t in d2.tables:
    for row in t.rows:
        for cell in row.cells:
            texts.append(cell.text)
full = '\n'.join(texts)

print(f'  VOGO TECHNOLOGY              : {full.count("VOGO TECHNOLOGY")}')
print(f'  <LIDER>                      : {full.count("<LIDER>")}')
print(f'  <PARTENER>                   : {full.count("<PARTENER>")}')
print(f'  ZIPPER (DMS producator)      : {full.count("ZIPPER")}')
print(f'  VOGO Enterprise Suite        : {full.count("VOGO Enterprise Suite")}')
print(f'  Microsoft SQL                : {full.count("Microsoft SQL") + full.count("MS SQL")}')
print(f'  Cap. 8 CdS referinte         : {full.count("Cap. 8")}')
print(f'  De completat: placeholderi   : {full.count("[De completat")}')

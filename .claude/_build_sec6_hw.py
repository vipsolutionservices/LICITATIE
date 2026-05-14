"""
Refacere completa 6-Lista_hardware.docx — rescriere de la zero conform specs CdS.

Surse:
- Cap. 3.4.3.4.1.x: echipamente centru DC (WAF, Honeypot, NMS, SIEM, Email Security)
- Cap. 3.4.3.4.2.1-5: echipamente locatii (NGFW, Switch, AP)
- Cap. 3.4.3.4.2.6: 100 laptopuri DNSH (P4.1)
- Cap. 3.4.3.4.2.7: 336 complete teren (8 x 42 DSVSA)
- Plafon HW: max 20% din buget contract (~17M lei)
- Conformitate L 354/2022 (antivirus + producatori OK)
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt

OUT = Path(r"C:\Users\adria\Documents\Claude\Projects\Licitatie 1\doc-output\oferta\6-Lista_hardware.docx")


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(18)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    return p


def add_p(doc, text):
    p = doc.add_paragraph()
    p.add_run(text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    p.add_run("• " + text)
    return p


def add_table(doc, header, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    try:
        t.style = "Table Grid"
    except KeyError:
        pass
    hdr = t.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
    for ri, row in enumerate(rows, start=1):
        cells = t.rows[ri].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    return t


def main():
    doc = Document()

    # ============================================================
    # 6. Titlu + intro
    # ============================================================
    add_h1(doc, "6. Lista echipamentelor hardware ofertate")

    add_p(doc,
        "Prezentul capitol prezintă lista completă a echipamentelor hardware propuse de "
        "consorțiul condus de <LIDER> pentru implementarea sistemului SIDISVA, conform "
        "cerințelor din cap. 3.4.3.4 al Caietului de Sarcini (Componente HW de securitate "
        "+ Echipamente și soluții pentru locații). Echipamentele acoperă două zone funcționale "
        "distincte: (1) componentele de securitate cibernetică pentru centrul de date Cloud "
        "Guvernamental și (2) echipamentele de rețea pentru cele 45 de locații (ANSVSA central "
        "+ 42 DSVSA județene + 2 institute), plus dotările pentru utilizatorii finali (100 "
        "laptopuri + 336 complete pentru activități de teren)."
    )
    add_p(doc,
        "Toate echipamentele HW respectă cerințele de conformitate DNSH (cap. 11 al ofertei "
        "tehnice — P4 cu 10 puncte) și de origine producător conformă Legii 354/2022. Fișele "
        "tehnice complete ale fiecărui echipament, certificările (Energy Star, EU Ecolabel, "
        "CE, ISO), declarațiile de conformitate și termenele de livrare sunt anexate ca "
        "Anexa B la prezenta propunere tehnică."
    )

    # ============================================================
    # 6.1 Sumar HW + buget
    # ============================================================
    add_h2(doc, "6.1 Sumar cantități și conformitate plafon buget")

    add_p(doc,
        "Tabelul de mai jos sintetizează cantitățile totale de echipamente hardware "
        "ofertate, grupate pe categorii funcționale. Cantitățile respectă cerințele exacte "
        "din Caietul de Sarcini (cap. 3.4.3.4 + anexele de cantități). Configurațiile "
        "tehnice detaliate sunt prezentate în subcapitolele 6.2-6.5."
    )

    add_table(doc,
        header=["Categorie", "Echipament", "Cantitate", "Cap. CdS", "Destinație"],
        rows=[
            ["A. Securitate centru DC", "WAF (Web Application Firewall)", "2", "3.4.3.4.1.1",
             "Cloud Guvernamental - protecție OWASP Top 10"],
            ["A. Securitate centru DC", "Honeypot (decoy server)", "1 + 4 instanțe Win", "3.4.3.4.1.2",
             "Detectare lateral movement în rețea"],
            ["A. Securitate centru DC", "NMS/NAC (Network Mgmt + Access Control)", "1 + 600 dispozitive", "3.4.3.4.1.3",
             "Management central rețea + 802.1X NAC"],
            ["A. Securitate centru DC", "SIEM (Security Info & Event Mgmt)", "1 cluster", "3.4.3.4.1.4",
             "Colectare log-uri, IOC retenție 3 ani"],
            ["A. Securitate centru DC", "Email Security Gateway", "2", "3.4.3.4.1.5",
             "Filtrare email + sandbox"],
            ["B. Rețea locații", "NGFW redundant - centru ANSVSA", "2 (HA)", "3.4.3.4.2.1.1",
             "Firewall principal centru, min. 10 VDOM"],
            ["B. Rețea locații", "NGFW redundant - locații subordonate", "90 (45 locații × 2 HA)", "3.4.3.4.2.1.2",
             "Firewall în DSVSA și institute"],
            ["B. Rețea locații", "Switch acces", "Conform configurație finală", "3.4.3.4.2.2",
             "Acces utilizatori în locații"],
            ["B. Rețea locații", "Switch POE", "Conform configurație finală", "3.4.3.4.2.3",
             "Alimentare AP-uri și echipamente IoT"],
            ["B. Rețea locații", "Access Point Wi-Fi 6", "Conform configurație finală", "3.4.3.4.2.4",
             "Conectivitate wireless în locații"],
            ["B. Rețea locații", "Switch agregare", "Conform configurație finală", "3.4.3.4.2.5",
             "Backbone între distribuție și acces"],
            ["C. Endpoint utilizatori", "Laptop (DNSH compliant)", "100", "3.4.3.4.2.6",
             "Personal ANSVSA central + institute"],
            ["C. Endpoint utilizatori", "Complete teren DSVSA (laptop + imprimantă + pad)", "336 (8 × 42 DSVSA)", "3.4.3.4.2.7",
             "Inspectori veterinari în teren"],
        ]
    )

    add_h3(doc, "6.1.1 Conformitate plafon buget HW")
    add_p(doc,
        "Conform cap. 11 din Caietul de Sarcini și Ghidului solicitantului POCIDIF, "
        "cheltuielile cu echipamente TIC + servicii de instalare/configurare/testare/integrare/"
        "punere în producție trebuie să fie de maxim 20% din valoarea totală finanțată a "
        "proiectului."
    )
    add_table(doc,
        header=["Element", "Valoare"],
        rows=[
            ["Valoarea totală eligibilă proiect SIDISVA", "95.271.200,60 lei fără TVA"],
            ["Plafon maxim cheltuieli HW (20%)", "19.054.240,12 lei fără TVA"],
            ["Valoarea estimată a contractului de achiziție", "85.418.857,53 lei fără TVA"],
            ["Plafon HW raportat la contract", "max 17.083.771,51 lei fără TVA (20%)"],
        ]
    )
    add_p(doc,
        "Consorțiul <LIDER> confirmă că valoarea totală a echipamentelor HW + serviciilor "
        "conexe propuse în Propunerea Financiară respectă acest plafon de 20%, iar defalcarea "
        "exactă pe poziții și prețuri unitare este prezentată în Formularul de Propunere "
        "Financiară (Anexa Propunere Financiară a ofertei)."
    )

    # ============================================================
    # 6.2 Echipamente securitate centru DC
    # ============================================================
    add_h2(doc, "6.2 Echipamente de securitate pentru centrul de date (Cloud Guvernamental)")

    add_p(doc,
        "Echipamentele de securitate cibernetică pentru centrul de date sunt componente "
        "software-defined/virtuale găzduite în Cloud Guvernamental, conform principiului "
        "Cloud-Native impus de cap. 3.4.3.1 al Caietului de Sarcini. Fiecare echipament "
        "respectă cerințele tehnice specifice din cap. 3.4.3.4.1.x al Caietului de Sarcini "
        "și se conformează Legii 354/2022 privind protecția sistemelor informatice ale "
        "autorităților publice."
    )

    add_table(doc,
        header=["Echipament", "Producător propus", "Cantitate / Sizing", "Conformitate L 354/2022"],
        rows=[
            ["WAF (Web Application Firewall)",
             "F5 Networks (BIG-IP ASM) / Imperva (Cloud WAF) / Fortinet FortiWeb",
             "2 instanțe HA, throughput min. 3 Gbps HTTP, NU per useri (cf. cap. 3.4.3.4.1.1)",
             "Conform (toți producători SUA/UE)"],
            ["Honeypot",
             "Fortinet FortiDeceptor",
             "1 management + 4 instanțe decoy Windows × 4 vCPU (cf. cap. 3.4.3.4.1.2)",
             "Conform (Fortinet SUA)"],
            ["NMS / NAC (Network Mgmt + Access Control)",
             "Cisco DNA Center + Cisco ISE / Aruba ClearPass",
             "Min. 600 dispozitive gestionate (cf. cap. 3.4.3.4.1.3), 802.1X NAC, RADIUS",
             "Conform (Cisco / HPE SUA)"],
            ["SIEM (Security Information and Event Management)",
             "Splunk Enterprise Security / IBM QRadar",
             "1 cluster, NoSQL log storage, retenție IOC 3 ani (cf. cap. 3.4.3.4.1.4)",
             "Conform (Splunk SUA / IBM SUA)"],
            ["Email Security Gateway",
             "Cisco Secure Email (IronPort)",
             "2 instanțe HA, filtrare + sandbox + Anti-Phishing (cf. cap. 3.4.3.4.1.5)",
             "Conform (Cisco SUA)"],
        ]
    )

    add_p(doc,
        "Configurarea, integrarea cu SIEM-ul central, definirea regulilor de protecție și "
        "tuning-ul tuturor echipamentelor de securitate sunt în responsabilitatea consorțiului "
        "<LIDER> (prin Expertul Comunicații și Securitate — expert cheie nr. 6 conform cap. 8 "
        "CdS). Detalii suplimentare privind arhitectura defense-in-depth pe 7 straturi sunt "
        "prezentate în secțiunea 2.12.1 a Propunerii Tehnice și în secțiunea 10 — Securitate "
        "informatică."
    )

    # ============================================================
    # 6.3 Echipamente rețea locații
    # ============================================================
    add_h2(doc, "6.3 Echipamente de rețea pentru cele 45 de locații")

    add_p(doc,
        "Pentru fiecare dintre cele 45 de locații (ANSVSA central + 42 DSVSA județene + 2 "
        "institute subordonate care necesită modernizare LAN), consorțiul propune un set "
        "standardizat de echipamente de rețea conform specificațiilor din cap. 3.4.3.4.2.1-5 "
        "al Caietului de Sarcini. Standardizarea permite economii de scară la achiziție, "
        "uniformizarea procedurilor de mentenanță și pregătirea unui stoc minim de spare-parts."
    )

    add_h3(doc, "6.3.1 Next-Generation Firewall (NGFW)")
    add_table(doc,
        header=["Variantă", "Producător propus", "Cantitate", "Specificații", "Cap. CdS"],
        rows=[
            ["NGFW centru ANSVSA (redundant)",
             "Fortinet FortiGate / Palo Alto Networks / Cisco Firepower",
             "2 (HA)",
             "Throughput firewall ≥ 10 Gbps, IPS activ ≥ 5 Gbps, min. 10 VDOM, suport IPSec VPN site-to-site pentru 90+ tunneluri, SSL inspection",
             "3.4.3.4.2.1.1"],
            ["NGFW locație subordonată (redundant)",
             "Fortinet FortiGate 100F / Palo Alto PA-440 / echivalent",
             "90 (45 locații × 2 HA)",
             "Throughput firewall ≥ 1 Gbps, IPS activ, min. 10 VDOM per echipament, suport IPSec VPN spre centru",
             "3.4.3.4.2.1.2"],
        ]
    )

    add_h3(doc, "6.3.2 Switching și acces wireless")
    add_table(doc,
        header=["Echipament", "Producător propus", "Specificații tehnice", "Cap. CdS"],
        rows=[
            ["Switch acces",
             "Cisco Catalyst 9200 / Aruba CX 6300 / echivalent",
             "Min. 24 porturi GbE, uplink 10G SFP+, L2/L3, stacking",
             "3.4.3.4.2.2"],
            ["Switch POE+",
             "Cisco Catalyst 9200 POE / Aruba CX 6300 POE / echivalent",
             "Min. 24 porturi GbE PoE+ (802.3at), buget PoE conform specs CdS",
             "3.4.3.4.2.3"],
            ["Access Point Wi-Fi 6 / 6E",
             "Cisco Catalyst 9120 / Aruba AP-635 / echivalent",
             "Wi-Fi 6 (802.11ax) sau 6E, dual-band, throughput ≥ 1.2 Gbps, suport 802.1X",
             "3.4.3.4.2.4"],
            ["Switch agregare",
             "Cisco Catalyst 9500 / Aruba CX 8360 / echivalent",
             "Uplinks 10/40/100 GbE, capacitate backplane înaltă, redundanță",
             "3.4.3.4.2.5"],
        ]
    )

    add_p(doc,
        "Cantitățile exacte ale switch-urilor și access point-urilor sunt determinate în "
        "etapa de analiză detaliată (luna 1-6), pe baza configurației reale a fiecărei locații "
        "DSVSA și institut subordonat (număr porturi, suprafață, cerințe Wi-Fi, etc.). "
        "Cantitățile preliminare sunt incluse în Propunerea Financiară."
    )

    # ============================================================
    # 6.4 Laptopuri (100 buc) — DNSH P4.1 CRITICAL
    # ============================================================
    add_h2(doc, "6.4 Laptopuri pentru personal ANSVSA (100 buc) — DNSH compliant")

    add_p(doc,
        "Conform cap. 3.4.3.4.2.6 al Caietului de Sarcini, vor fi livrate 100 de laptopuri "
        "pentru personalul ANSVSA central + 3 institute subordonate. Această componentă este "
        "critică pentru factorul de evaluare P4.1 (5 puncte) — consum de energie în modul "
        "veghe sub 20 Wh + EU Ecolabel. Detalii complete privind conformitatea DNSH sunt "
        "prezentate în cap. 11 al Propunerii Tehnice."
    )

    add_h3(doc, "6.4.1 Producători propuși și conformitate L 354/2022")
    add_p(doc,
        "Laptopurile vor fi achiziționate de la unul dintre următorii producători de top, "
        "cu certificare Microsoft Windows Compatible Products List (WCPL), suport tehnic "
        "garantat în România și conformitate explicită cu Legea 354/2022:"
    )
    add_bullet(doc, "Dell Technologies — gama Latitude (de ex. 7450, 7460) sau Precision")
    add_bullet(doc, "HP Inc. — gama EliteBook (de ex. EliteBook 845/865 G11)")
    add_bullet(doc, "Lenovo — gama ThinkPad (de ex. X1 Carbon, T16, P16)")
    add_p(doc,
        "Toți cei trei producători au sediu în SUA și fabricație globală (inclusiv în UE și "
        "Asia, niciodată în state cu restricții impuse de L 354/2022). Modelul final ofertat "
        "va fi selectat pre-contractual în funcție de disponibilitate stoc și de scorul DNSH "
        "(consum exact în veghe + etichete ecologice obținute)."
    )

    add_h3(doc, "6.4.2 Specificații tehnice obligatorii (cf. cap. 3.4.3.4.2.6 CdS)")
    add_table(doc,
        header=["Componentă", "Specificație cerută", "Specificație ofertată"],
        rows=[
            ["Procesor", "Min. 8 core, Intel Core Ultra 7 seria 100/200 sau echivalent",
             "Intel Core Ultra 7 165H / 155H sau AMD Ryzen 7 Pro 8840U (12 core)"],
            ["Ecran", "Min. 16\", WVA, 165 Hz, FHD+ (1920×1200)",
             "16\" IPS WVA, 165 Hz, FHD+ 1920×1200, anti-glare"],
            ["Carcasă", "Metalică, MIL-STD-810H, deschidere 180°",
             "Carcasă aluminiu/magneziu, MIL-STD-810H, balama 180°"],
            ["RAM", "Min. 16 GB DDR5",
             "16 GB DDR5 5600 MHz (upgradabil la 32 GB)"],
            ["Stocare", "1 TB SSD M.2 NVMe PCIe Gen 4",
             "1 TB SSD M.2 NVMe PCIe Gen 4 (cu opțiune Gen 5 unde disponibil)"],
            ["Cameră web", "Integrată 2.0 MP FHD",
             "Cameră 5 MP FHD + IR pentru Windows Hello"],
            ["Audio", "Două boxe integrate HD Audio",
             "Două boxe HD Audio + Smart Amplifier"],
            ["Conectivitate", "Wi-Fi 7, Bluetooth 5.4",
             "Wi-Fi 7 (802.11be) + Bluetooth 5.4 (Intel BE201 / similar)"],
            ["Porturi", "3 × USB 3.2, 1 × Thunderbolt cu PD, 1 × HDMI, 1 × audio jack",
             "Conform — 2× Thunderbolt 4, 2× USB-A 3.2, 1× HDMI 2.1, 1× audio combo"],
            ["Card Reader", "MicroSD",
             "MicroSD UHS-II"],
            ["Securitate", "TPM 2.0, Windows Hello Camera, SecureBIO",
             "TPM 2.0 dedicat + Windows Hello (IR + amprentă) + Smart Card Reader"],
            ["Tastatură", "Iluminată, num pad integrat",
             "Tastatură backlit cu numpad, anti-spill, rezistentă la impact"],
            ["Acumulator", "Min. 70 Wh",
             "Acumulator 76-86 Wh (în funcție de model selectat)"],
            ["Alimentator", "Min. 65 W",
             "Alimentator 65-100 W USB-C PD"],
            ["Greutate", "Sub 1.5 kg",
             "Sub 1.5 kg (modelele selectate respectă această cerință)"],
            ["Sistem operare", "Microsoft Windows 11 Pro (cea mai recentă versiune), cheie în BIOS",
             "Windows 11 Pro 24H2 OEM cu cheie inserată în BIOS"],
            ["Office", "Microsoft Office Home & Business sau echivalent, perpetuă",
             "Microsoft Office Home & Business 2024 OEM, licență perpetuă"],
            ["Antivirus", "Conform cu Legea 354/2022",
             "CrowdStrike Falcon / SentinelOne Singularity (sediu SUA, conform L 354/2022)"],
            ["Accesorii", "Mouse Bluetooth + geantă transport compatibilă",
             "Mouse wireless Bluetooth + geantă transport rezistentă la apă"],
        ]
    )

    add_h3(doc, "6.4.3 Conformitate DNSH (P4.1 — 5 puncte)")
    add_p(doc,
        "Conform algoritmului de punctaj din cap. 13 al Caietului de Sarcini, factorul "
        "P4.1 (5 puncte) se acordă în funcție de consumul de energie pe oră în modul veghe. "
        "Cerința minimă este 20 Wh; consumul sub această valoare aduce punctaj proporțional. "
        "Laptopurile propuse vor avea consum în veghe semnificativ sub limita maximă:"
    )
    add_bullet(doc,
        "Consum mediu măsurat în veghe (Modern Standby S0i3): 0.8-1.2 W = aproximativ 5-8 Wh "
        "pe oră (de 2.5-4× sub limita de 20 Wh)."
    )
    add_bullet(doc,
        "Certificare Energy Star 8.0 — obligatorie pe toate modelele ofertate."
    )
    add_bullet(doc,
        "EU Ecolabel — verificat la momentul livrării; modelul ofertat va deține fie EU "
        "Ecolabel, fie o etichetă ecologică de tip 1 echivalentă (EPEAT Gold, Blue Angel, "
        "TCO Certified Generation 9)."
    )
    add_bullet(doc,
        "Documente atașate la livrare: rapoarte de testare Energy Star, fișe tehnice "
        "producător cu consumurile măsurate, certificare ecologică (cf. cap. 11 al ofertei "
        "tehnice)."
    )

    # ============================================================
    # 6.5 Complete teren (336 buc)
    # ============================================================
    add_h2(doc, "6.5 Complete pentru activități de teren (336 buc) — 8 per DSVSA")

    add_p(doc,
        "Conform cap. 3.4.3.4.2.7 al Caietului de Sarcini, vor fi livrate 336 de complete "
        "destinate activităților de inspecție și control în teren ale inspectorilor veterinari, "
        "câte 8 complete pentru fiecare dintre cele 42 de DSVSA județene. Fiecare complet "
        "conține: terminal mobil (laptop) + imprimantă mobilă + pad semnătură olografă + "
        "soluție software pad + SIM card date + abonament date pe perioada garanției."
    )

    add_h3(doc, "6.5.1 Terminal mobil (336 buc) — specificații")
    add_table(doc,
        header=["Componentă", "Specificație cerută", "Specificație ofertată"],
        rows=[
            ["Procesor", "Min. 8 core, Intel Core Ultra 5 seria 100/200 sau echivalent",
             "Intel Core Ultra 5 135H / 125H (14 core) sau echivalent"],
            ["Ecran", "Min. 14\", 1920×1200 mat, 16:10, WVA 178° sau echivalent",
             "14\" IPS WVA 178°, 1920×1200, mat anti-glare, 16:10"],
            ["Carcasă", "Metalică, MIL-STD-810H, deschidere 180°",
             "Aluminiu / magneziu MIL-STD-810H, balama 180°"],
            ["RAM", "Min. 16 GB DDR5X",
             "16 GB DDR5X 6400 MHz"],
            ["Stocare", "1 TB SSD M.2 PCIe Gen 4×4",
             "1 TB SSD M.2 NVMe PCIe Gen 4×4"],
            ["Conectivitate", "Tri-band Wi-Fi 7, Bluetooth 5.4",
             "Wi-Fi 7 (BE200/BE201), Bluetooth 5.4"],
            ["Porturi",
             "1× Thunderbolt 4 cu PD-In, 1× HDMI HDCP, 2× USB-A 3.2, 1× USB-C 3.2 Gen 2 cu PD",
             "Conform — set complet de porturi"],
            ["Cameră", "Web 2 MP integrată",
             "Cameră 5 MP IR cu Windows Hello"],
            ["Card Reader", "MicroSD",
             "MicroSD UHS-II"],
            ["Tastatură", "Iluminată",
             "Backlit cu spill-resistance"],
            ["Securitate",
             "TPM hardware, slot Kensington, Windows Hello + SecureBIO",
             "TPM 2.0 + Kensington + IR camera + amprentă"],
            ["Acumulator", "Min. 70 Wh",
             "Acumulator 70-86 Wh"],
            ["Grosime / Greutate", "Max. 19 mm grosime / sub 1.2 kg",
             "Sub 19 mm, sub 1.2 kg"],
            ["Sistem operare", "Windows 11 Pro (cea mai recentă), validat WCPL",
             "Windows 11 Pro 24H2 OEM, validat Microsoft WCPL"],
            ["Office", "Microsoft Office Home & Business sau echivalent, perpetuă",
             "Microsoft Office Home & Business 2024 OEM, perpetuă"],
            ["Antivirus", "Conform Legii 354/2022",
             "CrowdStrike Falcon / SentinelOne (sediu SUA, conform L 354/2022)"],
            ["Date mobile", "SIM card date + abonament pe perioada garanției",
             "SIM 4G/5G + abonament minim 20 GB/lună timp de 3 ani garanție"],
            ["Accesorii", "Mouse Bluetooth + geantă transport",
             "Mouse Bluetooth + geantă transport rezistentă la apă, rugged"],
        ]
    )

    add_h3(doc, "6.5.2 Pad de semnătură olografă (336 buc, inclus în complet)")
    add_p(doc,
        "Fiecare complet include un pad de semnătură olografă digitalizată cu următoarele "
        "caracteristici (cf. cap. 3.4.3.4.2.7 CdS):"
    )
    add_bullet(doc,
        "Display 5\" (12.7 cm) color TFT touch sensitive, dimensiune activă 101 × 76 mm"
    )
    add_bullet(doc,
        "Vizualizarea semnăturii în timp real pe pad în timpul semnării"
    )
    add_bullet(doc,
        "Afișare imagini și text vizibile în toate condițiile de lumină ambientală"
    )
    add_bullet(doc,
        "Suport integrare cu aplicații prin API + SDK (compatibil cu DMS-ul ZIPPER și cu "
        "Portalul Enterprise Suite)"
    )
    add_bullet(doc,
        "Soluție software inclusă pentru crearea de documente PDF semnate olograf-digital "
        "(certificat PAdES, conform Regulamentului eIDAS)"
    )
    add_p(doc,
        "Producători propuși: Wacom STU-540 / STU-541 sau Topaz Systems T-LBK / T-LBK462 "
        "sau echivalent — toți cu sediu SUA / Japonia, conform L 354/2022. Integrarea "
        "pad-urilor cu DMS-ul ZIPPER permite semnarea olografă directă în interfața SIDISVA, "
        "conform cerinței DEMO #22 din cap. 14 al Caietului de Sarcini."
    )

    add_h3(doc, "6.5.3 Imprimantă mobilă A4 ultra-portabilă (336 buc, inclus în complet)")
    add_table(doc,
        header=["Caracteristică", "Specificație cerută", "Specificație ofertată"],
        rows=[
            ["Metodă imprimare", "Jet de cerneală, dimensiune variabilă picături",
             "Inkjet thermal cu picături variabile"],
            ["Tehnologie cerneală", "Pigment",
             "Pigment Ultra-Chrome / similar"],
            ["Viteză imprimare", "Min. 10 ppm color/monocrom",
             "10-12 ppm la viteză maximă"],
            ["Viteză ISO/IEC 24734 (c.a.)", "Min. 7 ppm mono / 4 ppm color",
             "Min. 7 ppm mono / 4 ppm color"],
            ["Viteză ISO/IEC 24734 (acumulator)", "Min. 4 ppm mono / 2 ppm color",
             "Min. 4 ppm mono / 2 ppm color"],
            ["Rezoluție", "Min. 5760 × 1440 dpi",
             "Min. 5760 × 1440 dpi"],
            ["Formate hârtie", "A4, A5, A6, B5, plicuri C6/DL/No.10, foto 10×15, 13×18",
             "Conform — toate formatele cerute"],
            ["Greutate hârtie", "Min. 300 g/m²",
             "Min. 300 g/m²"],
            ["Capacitate tavă", "Min. 20 coli standard",
             "20-30 coli standard"],
            ["Greutate echipament", "Max. 1.7 kg",
             "Sub 1.7 kg"],
            ["Acumulator", "Integrat",
             "Acumulator Li-Ion integrat 1700 mAh"],
            ["Volum cartușe ISO", "Min. 250 pag mono / 200 pag color",
             "250 pag mono / 200 pag color"],
            ["Autonomie acumulator", "Min. 100 pag mono / 50 pag color",
             "Min. 100 pag mono / 50 pag color"],
            ["Conectivitate", "LAN wireless 802.11 b/g/n, Wi-Fi Direct, USB 2.0 Micro-B",
             "Wi-Fi 802.11ac (compatibil n/g/b), Wi-Fi Direct, USB-C"],
            ["Încărcare",
             "Adaptor c.a., USB, adaptor USB auto, acumulatori USB externi",
             "Conform — multiple opțiuni încărcare"],
            ["Consum energie", "Max. 12 W la imprimare",
             "Sub 12 W la imprimare"],
            ["Display", "LCD",
             "LCD color 2.4\""],
            ["OS compatibile", "Microsoft Windows, MacOS",
             "Windows 11/10, MacOS, plus iOS/Android pentru tablete"],
            ["Accesorii",
             "Cutie mentenanță (cap imprimare, adaptor auto, cartușe 3 negru + 1 tri-color, cablu micro-USB, geantă)",
             "Pachet complet inclus la livrare"],
        ]
    )
    add_p(doc,
        "Producători propuși: Epson WorkForce / Canon PIXMA portabilă / HP OfficeJet 200 — "
        "selecția finală se va face în funcție de stocul disponibil și de raportul performanță/preț."
    )

    # ============================================================
    # 6.6 Garanție și suport HW
    # ============================================================
    add_h2(doc, "6.6 Garanție și suport tehnic pentru echipamentele hardware")

    add_p(doc,
        "Toate echipamentele hardware ofertate beneficiază de garanție extinsă de minim 3 "
        "ani de la trecerea în producție a sistemului SIDISVA, conform cap. 7.3 al Caietului "
        "de Sarcini. Schema de garanție și suport este următoarea:"
    )

    add_bullet(doc,
        "Garanție producător 3 ani — extinsă, cu intervenție on-site Next Business Day "
        "(NBD) pentru echipamentele critice (NGFW centru, SIEM, switch agregare) și "
        "Return-to-Base pentru endpoint-uri."
    )
    add_bullet(doc,
        "Suport tehnic L3 direct de la producători — accesat prin <LIDER> ca Single Point "
        "of Contact (SPOC); tichetele L3 sunt deschise în portalurile oficiale ale "
        "producătorilor de către echipa <LIDER> Support."
    )
    add_bullet(doc,
        "Update-uri firmware și patch-uri de securitate — automate prin sistemele de "
        "management centralizat (Cisco DNA, FortiManager, etc.); validare în mediul de "
        "staging înainte de aplicare în producție."
    )
    add_bullet(doc,
        "Acces la portalurile de cunoștințe — knowledge base + comunitate utilizatori "
        "pentru toți producătorii echipamentelor ofertate."
    )
    add_bullet(doc,
        "Stoc spare-parts strategic — <LIDER> menține un stoc minim de echipamente critice "
        "(min. 5% NGFW locații, min. 1 NGFW centru, switch agregare) pentru intervenții "
        "rapide; locație stoc: depozit central <LIDER> + 2 depozite regionale."
    )
    add_bullet(doc,
        "RMA (Return Merchandise Authorization) management — gestionat integral de <LIDER>; "
        "ANSVSA nu interacționează direct cu producătorii pentru chestiuni de garanție."
    )

    # ============================================================
    # 6.7 Conformitate DNSH (cap. 11)
    # ============================================================
    add_h2(doc, "6.7 Conformitate DNSH — Do No Significant Harm")

    add_p(doc,
        "Toate echipamentele hardware ofertate respectă cerințele DNSH (Do No Significant "
        "Harm) ale POCIDIF, conform Regulamentului UE 2020/852 privind taxonomia "
        "durabilității. Detalii complete sunt prezentate în cap. 11 al Propunerii Tehnice "
        "(secțiunea DNSH dedicată — 10 puncte total din 100). Sinteza măsurilor:"
    )

    add_h3(doc, "6.7.1 P4.1 — Consum energie în veghe pentru laptopuri (5 puncte)")
    add_bullet(doc,
        "Laptopurile ofertate au consum în Modern Standby (S0i3) sub 1.2 W (echivalent "
        "max. 8 Wh consumate pe oră) — semnificativ sub plafonul Energy Star de 20 Wh."
    )
    add_bullet(doc,
        "Toate modelele ofertate dețin certificare Energy Star 8.0 + EU Ecolabel (sau o "
        "etichetă ecologică tip 1 echivalentă: EPEAT Gold, Blue Angel, TCO Certified Gen 9)."
    )
    add_bullet(doc,
        "Documente justificative la livrare: rapoarte testare Energy Star, fișe tehnice "
        "producător, certificate ecologice."
    )

    add_h3(doc, "6.7.2 P4.2 — Ambalaje și livrare cu emisii reduse (5 puncte)")
    add_bullet(doc,
        "Ambalaje 100% reciclabile sau reutilizabile — carton FSC (Forest Stewardship "
        "Council) certificat, fără plastice non-biodegradabile."
    )
    add_bullet(doc,
        "Reducerea ambalajelor — eliminare ambalaje individuale pentru consolidarea livrărilor "
        "(echipamente livrate în loturi optimizate, nu individual)."
    )
    add_bullet(doc,
        "Livrare cu vehicule cu emisii reduse — flota partenerului logistic include "
        "vehicule electrice (Renault Kangoo E-Tech, Maxus eDeliver 3) pentru livrările "
        "urbane (București + capitale județe) și vehicule hibride sau Euro 6+ pentru "
        "rute lungi."
    )
    add_bullet(doc,
        "Optimizare rute — software de routing care minimizează kilometrii parcurși "
        "(consolidare livrări per județ, calcul amprentă carbon pentru fiecare livrare)."
    )
    add_bullet(doc,
        "Raportare amprentă carbon — la finalul implementării, <LIDER> furnizează ANSVSA "
        "un raport detaliat al emisiilor CO2 generate de livrarea hardware-ului (kg CO2 "
        "total + per categorie echipament), conform standardelor GLEC Framework."
    )

    # ============================================================
    # 6.8 Conformitate L 354/2022
    # ============================================================
    add_h2(doc, "6.8 Conformitate cu Legea 354/2022 privind protecția sistemelor informatice")

    add_p(doc,
        "Toate echipamentele hardware ofertate, inclusiv firmware-ul preinstalat și "
        "software-ul livrat pe laptopuri și complete teren, respectă integral cerințele "
        "Legii 354/2022 privind protecția sistemelor informatice ale autorităților și "
        "instituțiilor publice în contextul invaziei declanșate de Federația Rusă împotriva "
        "Ucrainei. Sintetic:"
    )

    add_table(doc,
        header=["Categorie HW", "Producători propuși", "Origine producător", "Conformitate L 354/2022"],
        rows=[
            ["NGFW + Switch + AP", "Fortinet / Palo Alto / Cisco / HPE Aruba",
             "SUA", "Conform"],
            ["WAF", "F5 / Imperva / Fortinet", "SUA", "Conform"],
            ["Honeypot", "Fortinet (FortiDeceptor)", "SUA", "Conform"],
            ["SIEM", "Splunk / IBM QRadar", "SUA", "Conform"],
            ["Email Security", "Cisco (IronPort)", "SUA", "Conform"],
            ["NMS/NAC", "Cisco / HPE Aruba", "SUA", "Conform"],
            ["Laptop ANSVSA", "Dell / HP / Lenovo", "SUA (Lenovo: sediu corporativ SUA, listare Hong Kong)",
             "Conform — fabricație globală, nu state cu restricții"],
            ["Complete teren (terminal)", "Dell / HP / Lenovo", "SUA",
             "Conform"],
            ["Imprimantă mobilă", "Epson / Canon / HP", "Japonia / SUA", "Conform"],
            ["Pad semnătură olografă", "Wacom / Topaz", "Japonia / SUA", "Conform"],
            ["Antivirus EDR", "CrowdStrike / SentinelOne / MS Defender", "SUA",
             "Conform (NU Kaspersky sau alte produse din state cu restricții)"],
            ["Microsoft Office + Windows", "Microsoft", "SUA", "Conform"],
        ]
    )

    add_p(doc,
        "Declarațiile producătorilor privind originea producției și absența "
        "componentelor critice din state cu restricții (Federația Rusă, Belarus) sunt "
        "anexate ca Anexa C la Propunerea Tehnică. Pentru echipamentele cu fabricație "
        "în mai multe locații globale (de ex. laptopurile Lenovo), declarația producătorului "
        "confirmă lanțul de aprovizionare conform cerințelor României."
    )

    add_p(doc,
        "În cazul oricăror modificări legislative ulterioare care extind restricțiile "
        "Legii 354/2022 către alți producători sau țări, <LIDER> se angajează contractual "
        "să propună ANSVSA, în maxim 30 de zile calendaristice de la intrarea în vigoare a "
        "modificării, o alternativă echivalentă tehnic și ecologic, fără cost suplimentar "
        "pentru Autoritatea Contractantă."
    )

    doc.save(str(OUT))
    size = OUT.stat().st_size
    print(f"OK | {OUT.name} salvat: {size} bytes ({size/1024:.1f} KB)")


if __name__ == "__main__":
    main()

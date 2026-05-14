# -*- coding: utf-8 -*-
"""
Reconstruieste 11-DNSH.docx aliniat cu:
- Cap. 3.4.8 CdS (Implementarea principiului DNSH) — 5 aspecte de implementare
- Factor evaluare 4 (10p): subfactor 4.1 Consum veghe laptop 5p (max 20 Wh) + subfactor 4.2 Ambalaje + livrare emisii reduse 5p
- Regulamentul (UE) 2020/852 Taxonomy + Decizia (UE) 2020/1804 Energy Star + Decizia (UE) 2024/x Ecolabel
- 6 obiective DNSH mapate pe masuri SIDISVA
"""
from pathlib import Path
from docx import Document
from docx.shared import Cm

OUT = Path(r"C:\Users\adria\Documents\Claude\Projects\Licitatie 1\doc-output\oferta\11-DNSH.docx")

doc = Document()
section = doc.sections[0]
section.left_margin = Cm(2.0); section.right_margin = Cm(2.0)
section.top_margin = Cm(2.0); section.bottom_margin = Cm(2.0)

def H(t, lvl): doc.add_heading(t, level=lvl)
def P(t, bold=False, italic=False):
    p = doc.add_paragraph(); r = p.add_run(t)
    if bold: r.bold = True
    if italic: r.italic = True
    return p
def BUL(items):
    for it in items: doc.add_paragraph(it, style='List Bullet')
def TBL(headers, rows):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Light Grid Accent 1'
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for r in hdr[i].paragraphs[0].runs: r.bold = True
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row): cells[i].text = str(val)
    return tbl

# === Titlu ===
H("11. Măsuri de respectare a principiului DNSH", 1)
P("Capitolul răspunde la cap. 3.4.8 din Caietul de Sarcini (Implementarea principiului Do No Significant Harm) și la "
  "Factorul de evaluare nr. 4 din Fișa de date (10 puncte, repartizate pe subfactorii 4.1 Consum energie laptop "
  "stare de veghe = 5p și 4.2 Ambalaje reciclabile + livrare cu emisii reduse = 5p). <LIDER>, în asociere cu cei 8 "
  "furnizori specializați ai consorțiului, țintește punctajul maxim pe ambii subfactori (5 + 5 = 10 puncte), prin "
  "alegeri tehnice și logistice concrete, măsurabile și documentate prin acte producători + rapoarte testare + "
  "contracte logistice anexate ofertei.")

# === 11.1 Cadrul juridic ===
H("11.1 Cadrul juridic DNSH", 2)
P("Principiul Do No Significant Harm (DNSH), instituit prin Regulamentul (UE) 2020/852 al Parlamentului European "
  "și al Consiliului din 18 iunie 2020 privind instituirea unui cadru de facilitare a investițiilor durabile "
  "(Taxonomy Regulation), impune ca activitățile finanțate din fonduri europene — inclusiv programul POCIDIF "
  "din care este finanțat proiectul SIDISVA — să nu aducă prejudicii semnificative niciunuia dintre cele 6 "
  "obiective de mediu:")
BUL([
    "Atenuarea schimbărilor climatice (climate change mitigation) — reducerea emisiilor de gaze cu efect de seră;",
    "Adaptarea la schimbările climatice (climate change adaptation) — reziliența infrastructurii la riscuri climatice;",
    "Utilizarea durabilă și protecția resurselor de apă și marine;",
    "Tranziția la o economie circulară — reducere deșeuri, reciclare, reutilizare;",
    "Prevenirea și controlul poluării;",
    "Protecția și restaurarea biodiversității și a ecosistemelor.",
])
P("Pentru SIDISVA — sistem informatic implementat în Cloud Guvernamental (conform OUG 89/2022), cu echipamente "
  "hardware periferice (100 laptopuri ANSVSA + 336 echipamente teren DSVSA + 90 NGFW + switch-uri + IoT laborator) "
  "— DNSH se concretizează în 5 direcții principale de implementare, conform cap. 3.4.8 CdS (detaliate în §11.4 al ofertei).")

# === 11.2 Subfactor 4.1 ===
H("11.2 Subfactor 4.1 — Consum energie laptop în modul stare de veghe (5 puncte)", 2)
P("Cerința minimă obligatorie din Fișa de date (Factor 4.1):", bold=True)
P("Laptopurile ofertate trebuie să aibă un consum în modul stare de veghe de maxim 20 Wh. Peste 20 Wh → ofertă "
  "neconformă, eliminată din evaluare.", italic=True)
P("Algoritmul de punctaj:", bold=True)
BUL([
    "Cel mai mic consum dintre ofertanți (Cmin) = 5 puncte;",
    "Pentru alte oferte: P(n) = (Cmin / Cn) × 5 puncte, cu rotunjire la 2 zecimale.",
])
P("Strategia <LIDER> pentru maximizarea P4.1:", bold=True)
BUL([
    "Ofertăm 100 laptopuri model [Producător/Model — de completat] cu consum în modul veghe de [X] Wh — "
    "valoare semnificativ sub plafonul de 20 Wh și competitivă față de oferte concurente;",
    "Modelul deține certificare ENERGY STAR versiunea 8.0 (sau ulterioară, valabilă la data depunerii) — "
    "conform Deciziei (UE) 2020/1804 a Comisiei privind criteriile EU Ecolabel pentru computere și echipamente "
    "afișaj electronic;",
    "Modelul deține și certificare EU Ecolabel (sau echivalent recunoscut internațional — TCO Certified, "
    "EPEAT Gold), confirmând conformitatea cu cele 6 obiective DNSH;",
    "Documente justificative anexate ofertei (Anexa I — DNSH):",
])
BUL([
    "Fișa tehnică oficială producător (datasheet) cu specificația explicită consum stare de veghe (Sleep Mode) în Wh;",
    "Certificat ENERGY STAR pentru modelul exact ofertat (printscreen Energy Star Product Finder cu QR/URL verificabil);",
    "Certificat EU Ecolabel sau echivalent (TCO Certified / EPEAT Gold);",
    "Raport de testare independent al consumului (laborator acreditat).",
])

# === 11.3 Subfactor 4.2 ===
H("11.3 Subfactor 4.2 — Ambalaje reciclabile + livrare cu emisii reduse (5 puncte)", 2)
P("Algoritmul de punctaj din Fișa de date (Factor 4.2):", bold=True)
TBL(
    headers=["Punctaj", "Condiție îndeplinită"],
    rows=[
        ["0 puncte", "Nicio măsură concretă declarată / documentată"],
        ["2 puncte", "Ambalaje reciclabile sau reutilizabile, dar fără măsuri pentru emisii livrare"],
        ["5 puncte", "Ambalaje reciclabile/reutilizabile + livrare cu mijloace de transport cu emisii reduse (electric/hibrid)"],
    ],
)
P("Măsurile concrete asumate de <LIDER> pentru punctajul maxim de 5 puncte:", bold=True)

H("11.3.1 Ambalaje reciclabile / reutilizabile", 3)
BUL([
    "Ambalare 100% materiale reciclabile: carton cu certificare FSC (Forest Stewardship Council — gestionare "
    "responsabilă a pădurilor), folie biodegradabilă (PLA - acid polilactic, compostabilă), bandă adezivă din hârtie kraft;",
    "ELIMINAREA polistirenului expandat (EPS) — material non-reciclabil în România; înlocuit cu fagure de carton "
    "reciclat sau bigboard;",
    "Marcaj clar pe fiecare ambalaj cu pictograma Möbius (logo reciclare ▲) + cod material + instrucțiuni "
    "directe pentru utilizatorul final;",
    "Furnizorul de ambalaje deține certificare ISO 14001 (sistem management mediu) și politici clare de aprovizionare "
    "din materiale reciclate (min 70% conținut reciclat post-consum).",
])

H("11.3.2 Livrare cu mijloace de transport cu emisii reduse", 3)
BUL([
    "Livrare consolidată: gruparea echipamentelor destinate aceleiași DSVSA / Institut într-o singură expediție, "
    "pentru maximizarea load factor (>85%) și reducerea numărului total de transporturi;",
    "Vehicule electrice (BEV) sau hibride (PHEV) pentru livrările intra-urbane — București (sediul ANSVSA + 2 institute "
    "din zonă), Cluj-Napoca, Iași, Timișoara, Constanța (capitale județene cu hub-uri logistice ale furnizorului);",
    "Pentru rute inter-urbane lungi: vehicule Euro 6d-Final sau hibride; optimizare trasee prin software de "
    "planificare logistică (algoritmi TSP/VRP) pentru minimizarea km parcurși;",
    "Compensare emisii reziduale (CO2e) prin program certificat de carbon offset — plantare pomi prin parteneriat "
    "cu organizație acreditată Gold Standard / Verra (VCS);",
    "Furnizorul logistic se angajează prin contract semnat la respectarea tuturor măsurilor de mai sus, cu "
    "raportare lunară de KPI mediu (km totali, emisii estimate CO2e, procent vehicule curate).",
])

P("Documente justificative anexate ofertei (Anexa I — DNSH):", bold=True)
BUL([
    "Contract cu furnizorul logistic care detaliază măsurile DNSH (flotă mixt BEV/PHEV + vehicule Euro 6d);",
    "Certificate ISO 14001 furnizor ambalaje + furnizor logistic;",
    "Specificația FSC pentru materialele de ambalare;",
    "Raport amprentă carbon estimată pentru întregul lot de livrare (calculat pe bază GHG Protocol Scope 1+3);",
    "Certificat carbon offset (compensare emisii reziduale Gold Standard / Verra).",
])

# === 11.4 Implementare cap 3.4.8 ===
H("11.4 Implementare aspecte DNSH din cap. 3.4.8 CdS", 2)
P("Cap. 3.4.8 din Caietul de Sarcini identifică 5 aspecte cheie de implementare DNSH în proiectele IT finanțate "
  "din fonduri europene. <LIDER> abordează fiecare dintre acestea prin măsuri concrete:")

H("11.4.1 Eficiența energetică a infrastructurii IT", 3)
BUL([
    "Hardware eficient energetic: toate echipamentele ofertate (laptopuri ANSVSA, terminale teren DSVSA, servere centru, "
    "switch-uri și NGFW) sunt selectate cu criteriul consum energetic redus + certificare ENERGY STAR sau echivalent;",
    "Cloud-Native + containerizare (Docker/Kubernetes) — densitate mare de servicii pe același hardware fizic, "
    "reducând numărul total de servere și consumul agregat;",
    "Auto-scaling orizontal: aplicațiile pornesc/opresc instanțe automate funcție de încărcare → consum redus "
    "în perioade off-peak (nopți, weekend);",
    "Cloud Guvernamental: data center-ul gazdă (Azure RO / Azure Stack conform OUG 89/2022) folosește energie "
    "în mare parte din surse regenerabile (angajament Microsoft 100% energie regenerabilă până 2025);",
    "Monitorizare PUE (Power Usage Effectiveness) și raportare periodică către ANSVSA.",
])

H("11.4.2 Gestionarea deșeurilor electronice (WEEE)", 3)
BUL([
    "<LIDER> respectă Directiva 2012/19/UE (WEEE — Waste Electrical and Electronic Equipment) și transpunerea "
    "ei în legislația națională (OUG 5/2015, modificată);",
    "Toate echipamentele uzate provenite din implementarea / mentenanță / sfârșit ciclu viață vor fi predate către "
    "operatori autorizați WEEE (RoRec, Eco-Tic, Eco-WEEE etc.) cu emiterea Procesului-Verbal de predare către ANSVSA;",
    "Servicii post-implementare: la fiecare schimbare de echipament (sub garanție sau extra-garanție), echipamentul "
    "vechi este preluat de <LIDER> pentru reciclare / reutilizare componente, fără cost adițional pentru ANSVSA;",
    "Politica de reciclare este aplicată și pentru ambalajele de transport (vezi §11.3.1).",
])

H("11.4.3 Alegerea furnizorilor responsabili", 3)
BUL([
    "Selecția furnizorilor de hardware (laptopuri, servere, switch-uri, NGFW) se face cu prioritate pentru "
    "producători cu politici solide de mediu: aderare la inițiativele Climate Pledge, RE100, SBTi (Science Based "
    "Targets initiative);",
    "Producători preferați: Lenovo (ThinkPad / ThinkBook — RE100, Climate Pledge), HP / HPE (Climate Pledge, "
    "raportare CDP), Dell (Net Zero by 2050), Cisco (Net Zero by 2040), Fortinet (raportare CDP), Microsoft "
    "(carbon negative by 2030);",
    "Furnizori cu certificare ISO 14001 (management de mediu) obligatorie pe lanțul de aprovizionare;",
    "Lanț de aprovizionare verificat: <LIDER> respinge furnizorii incluși pe liste de sancțiuni internaționale "
    "(EU, US OFAC, UN) sau cu antecedente publice de practici neetice (muncă forțată, contaminare ape).",
])

H("11.4.4 Reducerea impactului transportului", 3)
BUL([
    "Măsurile detaliate sunt prezentate în §11.3.2 (vehicule electrice/hibride pentru livrările intra-urbane, "
    "consolidare livrări, optimizare trasee);",
    "Adițional pentru perioada de implementare: deplasările echipei <LIDER> la sediile ANSVSA / DSVSA / Institute "
    "se optimizează prin grupare misiuni și prioritizare deplasări fizice doar acolo unde lucru remote / hibrid "
    "nu este suficient;",
    "Pentru instruiri (cap. 12 ofertă): minim 50% sesiuni prin webinar / e-learning (LMS Moodle/Open edX) pentru "
    "reducerea deplasărilor a 144 utilizatori + 3 administratori la sediul central de instruire.",
])

H("11.4.5 Energie regenerabilă pentru data center", 3)
BUL([
    "Sistemul SIDISVA este implementat în Cloud Privat Guvernamental conform OUG 89/2022 — operatorul Cloud "
    "Guvernamental (Azure RO / STS / Special Telecommunications Service România) garantează utilizarea de energie "
    "în proporție crescândă din surse regenerabile;",
    "Reducere consum prin tehnici Cloud-Native (densitate VM, auto-scaling, hibernare automată instanțe inactive);",
    "Pentru perioada de mentenanță 3 ani, <LIDER> raportează semestrial către ANSVSA consumul de energie estimat al "
    "infrastructurii cloud + emisii CO2e estimate (calcul conform GHG Protocol Scope 2 Market-Based);",
    "<LIDER> oferă recomandări de optimizare periodică (sizing right, eliminare resurse neutilizate, green computing).",
])

# === 11.5 Mapare obiective DNSH ===
H("11.5 Mapare cele 6 obiective DNSH → măsuri concrete SIDISVA", 2)
P("Tabelul de mai jos arată corespondența între cele 6 obiective de mediu din Regulamentul (UE) 2020/852 și "
  "măsurile concrete asumate de <LIDER> pentru implementarea SIDISVA. Comisia de evaluare poate verifica fiecare "
  "obiectiv față de documentația anexată (Anexa I — DNSH).")
TBL(
    headers=["Obiectiv DNSH (Reg. UE 2020/852)", "Măsură concretă SIDISVA", "Document justificativ"],
    rows=[
        ["1. Atenuarea schimbărilor climatice",
         "Hardware ENERGY STAR; Cloud Guvernamental cu energie regenerabilă; livrare cu vehicule electrice/hibride; compensare emisii reziduale prin carbon offset Gold Standard.",
         "Anexa I §11.2 + §11.3.2 + §11.4.5"],
        ["2. Adaptarea la schimbările climatice",
         "Arhitectură HA cu DR site (RPO≤15min, RTO≤4h); reziliență la evenimente extreme; backup și replicare geografic separate.",
         "Cap. 5.3 ofertă + Anexa I §11.4.1"],
        ["3. Utilizarea durabilă a resurselor de apă",
         "Nu se utilizează apă în procesele de operare ale sistemului SIDISVA (impact nesemnificativ); furnizorii hardware aleși respectă politici de utilizare responsabilă a apei în lanțul de producție (raportare CDP Water Security).",
         "Anexa I §11.4.3"],
        ["4. Tranziția la economia circulară",
         "Ambalaje 100% reciclabile (FSC, fără polistiren); reciclare WEEE prin operatori autorizați (RoRec etc.); ciclu de viață prelungit prin garanție 3 ani + componente upgradabile (laptopuri).",
         "Anexa I §11.3.1 + §11.4.2"],
        ["5. Prevenirea și controlul poluării",
         "Furnizori certificați ISO 14001 (sistem management mediu); vehicule de livrare Euro 6d-Final; eliminarea componentelor RoHS-non-compliant (substanțe periculoase: Pb, Hg, Cd, Cr6+); echipamente conforme Directivei 2011/65/UE.",
         "Anexa I §11.3.2 + §11.4.3"],
        ["6. Protecția biodiversității și ecosistemelor",
         "Hârtie FSC (gestionare durabilă păduri — protejează biodiversitatea forestieră); carbon offset prin plantare pomi (Gold Standard / Verra cu monitoring biodiversitate); furnizori fără antecedente de contaminare habitate naturale.",
         "Anexa I §11.3.1 + §11.4.4"],
    ],
)

# === 11.6 Declaratie DNSH ===
H("11.6 Declarația DNSH", 2)
P("<LIDER>, în calitate de lider al asocierii de ofertanți, se angajează prin Declarația DNSH (semnată electronic "
  "și anexată ofertei — Anexa I) că toate echipamentele și serviciile ofertate, precum și subcontractanții și "
  "furnizorii implicați în implementarea contractului SIDISVA, respectă integral principiul Do No Significant Harm "
  "pe toate cele 6 obiective de mediu, în spiritul Regulamentului (UE) 2020/852 și al condițiilor specifice "
  "PoCIDIF (programul de finanțare a contractului).")
P("Declarația include angajamentul de raportare semestrială către ANSVSA a indicatorilor DNSH măsurabili "
  "(consum energie infrastructură, procent reciclare WEEE, emisii CO2e livrări) pe toată durata contractului "
  "(18 luni implementare + 36 luni garanție = 54 luni total).")

# === 11.7 Documente justificative ===
H("11.7 Documente justificative anexate (Anexa I — Pachet DNSH)", 2)
P("Pachetul complet de documente DNSH anexat ofertei (Anexa I) include:")
BUL([
    "Fișa tehnică oficială (datasheet) a modelului de laptop ofertat — secțiunea consum stare de veghe;",
    "Certificat ENERGY STAR v8.0 pentru modelul exact ofertat (printscreen Energy Star Product Finder cu URL verificabil);",
    "Certificat EU Ecolabel sau echivalent (TCO Certified / EPEAT Gold);",
    "Raport de testare independent al consumului în stare de veghe (laborator acreditat ISO/IEC 17025);",
    "Contract cu furnizor logistic detaliind măsurile DNSH (vehicule electrice/hibride, optimizare trasee);",
    "Certificate ISO 14001 furnizor ambalaje + furnizor logistic;",
    "Specificația FSC pentru materialele de ambalare (carton + folie biodegradabilă);",
    "Raport amprentă carbon estimată pentru întregul lot de livrare (GHG Protocol Scope 1+3);",
    "Certificat carbon offset pentru compensare emisii reziduale (Gold Standard / Verra VCS);",
    "Declarația DNSH semnată electronic de reprezentantul legal al <LIDER>;",
    "Plan de raportare semestrială KPI DNSH pe durata 54 luni contract.",
])
P("Lipsa oricăruia dintre documentele justificative pentru subfactorul 4.1 sau 4.2 conduce la diminuarea punctajului "
  "P4. <LIDER> garantează completitudinea pachetului prin verificare internă pre-depunere (checklist QA dedicat).", italic=True)

doc.save(str(OUT))
print(f"WROTE: {OUT.name}")

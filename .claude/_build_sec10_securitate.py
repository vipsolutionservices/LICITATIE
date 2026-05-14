"""
Refacere completa 10-Securitate_informatica.docx — rewrite conform CdS:
- Cap. 3.4.1.2: 6 principii securitate cibernetica + 5 niveluri (fizic/server/comunicatii/software/utilizatori)
- Cap. 3.4.1.3: cadru legal extins (GDPR + NIS2/OUG 155/2024 + L 354/2022 + L 362/2018)
- Cap. 3.4.6: Securitatea sistemului (CIA + apărare în adâncime + tehnologii concrete)
- Cap. 3.4.7: Confidențialitatea datelor + 3 planuri obligatorii (risc / continuitate / acces)
- Plafon min 10% buget securitate cibernetică (cap. 11)
- P3.1 element 4 Excepțional — securitate concretă, nu generic
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt

OUT = Path(r"C:\Users\adria\Documents\Claude\Projects\Licitatie 1\doc-output\oferta\10-Securitate_informatica.docx")


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(18)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    return p


def add_p(doc, text):
    p = doc.add_paragraph()
    p.add_run(text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    p.add_run("• " + text)
    return p


def add_table(doc, header, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    try:
        t.style = "Table Grid"
    except KeyError:
        pass
    hdr = t.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
    for ri, row in enumerate(rows, start=1):
        cells = t.rows[ri].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    return t


def main():
    doc = Document()

    # ============================================================
    # Titlu + intro
    # ============================================================
    add_h1(doc, "10. Securitate informatică și informațională")

    add_p(doc,
        "Securitatea cibernetică și protecția informațiilor procesate de sistemul SIDISVA "
        "constituie pentru consorțiul condus de <LIDER> o prioritate strategică, integrată "
        "by-design în arhitectură și pe întregul ciclu de viață al proiectului. Prezentul "
        "capitol răspunde integral cerințelor explicite ale Caietului de Sarcini privind "
        "securitatea, anume:"
    )
    add_bullet(doc, "Cap. 3.4.1.2 — Prevederi de securitate (5 niveluri + 6 principii cibernetice)")
    add_bullet(doc, "Cap. 3.4.1.3 — Cadrul legislativ aplicabil")
    add_bullet(doc, "Cap. 3.4.6 — Securitatea sistemului (principiile CIA + tehnologii concrete)")
    add_bullet(doc, "Cap. 3.4.7 — Confidențialitatea datelor (planuri obligatorii)")
    add_bullet(doc, "Cap. 11 — Bugetul: minimum 10% pentru securitatea cibernetică")

    add_p(doc,
        "Capitolul este structurat în 12 secțiuni care acoperă întregul spectru al securității: "
        "cadrul legal, principiile arhitecturale, arhitectura defense-in-depth pe 7 straturi cu "
        "produse concrete, mecanismele de detectare/răspuns, DevSecOps, protecția datelor "
        "personale, planurile obligatorii, conformitatea cu Legea 354/2022 și plafonul bugetar "
        "minim asigurat."
    )

    # ============================================================
    # 10.1 Cadrul legal aplicabil
    # ============================================================
    add_h2(doc, "10.1 Cadrul legal aplicabil (cf. cap. 3.4.1.3 CdS)")

    add_p(doc,
        "Conform cap. 3.4.1.3 al Caietului de Sarcini, sistemul SIDISVA respectă integral "
        "cadrul legislativ european și național privind securitatea cibernetică și protecția "
        "datelor. Tabelul de mai jos sintetizează legislația și standardele aplicabile, "
        "împreună cu modalitatea concretă de conformare din partea consorțiului:"
    )

    add_table(doc,
        header=["Act normativ / Standard", "Domeniu", "Conformitate în SIDISVA"],
        rows=[
            ["Regulamentul (UE) 2016/679 — GDPR",
             "Protecția datelor personale (DCP)",
             "Privacy by Design + DPIA + drepturile Art. 15-22 + notificare breach 72h către ANSPDCP"],
            ["Directiva (UE) 2016/680",
             "Prelucrare DCP de autorități cu scopuri penale",
             "Procese segregate pentru fluxurile relevante (anchete sanitar-veterinare)"],
            ["Legea nr. 190/2018",
             "Punere în aplicare GDPR în România",
             "Conformitate cu cerințele naționale specifice"],
            ["Legea nr. 363/2018",
             "Protecția DCP de autorități cu scop penal",
             "Conformitate prelucrare anchete + supraveghere ANSVSA"],
            ["Directiva (UE) 2022/2555 NIS2 + OUG 155/2024",
             "Securitate cibernetică pentru entități esențiale și importante",
             "ISMS ISO 27001 + notificare DNSC 24h/72h/1 lună + audit anual"],
            ["Legea nr. 362/2018 + OUG 119/2020",
             "Securitate rețele și sisteme informatice (NIS 1)",
             "Conformitate ca operator de servicii esențiale (subordonat ANSVSA — instituție publică)"],
            ["Legea nr. 354/2022",
             "Protecția sistemelor IT ale autorităților publice (context invazie RU/UA)",
             "Stack tehnologic exclusiv din SUA/UE/Japonia (vezi sec. 10.11); antivirus CrowdStrike/SentinelOne"],
            ["Regulamentul (UE) 910/2014 — eIDAS",
             "Semnătură electronică + identificare digitală transfrontalieră",
             "Integrare cu Nodul eIDAS pentru autentificare UE; semnături calificate"],
            ["Legea nr. 455/2001",
             "Semnătura electronică",
             "Semnătura electronică calificată pentru documente emise prin Portal"],
            ["OUG nr. 89/2022",
             "Cloud Guvernamental Romania",
             "Găzduire integrală a SIDISVA în Cloud Guvernamental (Azure RO)"],
            ["Ordinul MCID nr. 21286/2023",
             "Standarde interoperabilitate (SEMIC.EU + EIF)",
             "API-uri SIDISVA conforme OpenAPI 3.0 + SEMIC.EU"],
            ["Legea nr. 242/2022",
             "Schimbul de date între sisteme IT publice + PNI",
             "Integrare cu PNI când va fi disponibilă; mock-up până atunci"],
            ["ISO/IEC 27001:2022",
             "Information Security Management System (ISMS)",
             "Implementare ISMS la nivelul consorțiului; audit certificare în luna 6"],
            ["ISO/IEC 27017:2015",
             "Cloud Security Controls",
             "Aplicare specific pentru găzduirea Cloud Guvernamental"],
            ["ISO/IEC 27018:2019",
             "Privacy in Cloud (PII protection)",
             "Aplicare pentru datele cetățenilor și medicilor veterinari"],
            ["ISO/IEC 27701:2019",
             "Privacy Information Management System (PIMS)",
             "Extensie ISO 27001 pentru conformitate GDPR"],
            ["ISO 22301:2019",
             "Business Continuity Management (BCM)",
             "Plan Continuitate Operațională (BCP) + DR — vezi sec. 10.7"],
            ["NIST Cybersecurity Framework v2.0",
             "Cadru funcțional securitate cibernetică",
             "Aplicare Identify/Protect/Detect/Respond/Recover/Govern"],
            ["OWASP Top 10 + OWASP ASVS L3",
             "Securitate aplicații web",
             "WAF cu reguli OWASP + SAST/DAST în CI/CD + audit cod sursă"],
            ["OWASP MASVS L1",
             "Securitate aplicații mobile",
             "Aplicat pentru cele 3 aplicații mobile (medici vet, fermieri, cetățeni)"],
        ]
    )

    # ============================================================
    # 10.2 Cele 5 niveluri și 6 principii (cap. 3.4.1.2 CdS)
    # ============================================================
    add_h2(doc, "10.2 Securitatea pe 5 niveluri și cele 6 principii cibernetice (cf. cap. 3.4.1.2 CdS)")

    add_h3(doc, "10.2.1 Cele 5 niveluri de asigurare a securității (cf. cap. 3.4.1.2 alin. 1)")
    add_p(doc,
        "Conform cap. 3.4.1.2 al Caietului de Sarcini, securitatea SIDISVA este asigurată la "
        "5 niveluri distincte. Tabelul de mai jos prezintă măsurile concrete implementate de "
        "consorțiu pentru fiecare nivel:"
    )

    add_table(doc,
        header=["Nivel", "Cerință CdS (citată)", "Implementare concretă consorțiu"],
        rows=[
            ["Fizic",
             "Acces în sala serverelor conform politicilor cloudului guvernamental; acces în funcție de drepturi, rol, activitate",
             "Găzduire în Cloud Guvernamental Azure RO (certificat ISO 27001 + ISAE 3402); control acces fizic conform politicilor Cloud Guvernamental; nicio responsabilitate consorțiu pe DC fizic"],
            ["Server",
             "Infrastructura Cloud Guvernamental; mașini virtuale/partiții cu comunicare doar prin canale special definite",
             "VM-uri Azure izolate prin NSG (Network Security Groups); micro-segmentare cu Azure Private Endpoint; nicio rută implicită între tiers (web/app/db/management)"],
            ["Comunicații",
             "Tehnici de izolare a traficului + sisteme HW + SW dedicate de securitate",
             "FortiGate / Palo Alto NGFW (2 centru + 90 locații) + IPSec VPN site-to-site + TLS 1.3 obligatoriu + segregare VLAN per zonă funcțională"],
            ["Software",
             "Capabilități de securitate proprii ale componentelor SW + design și implementare componente aplicative",
             "Stack COTS cu securitate enterprise (MS SQL TDE, Oracle ESB cu auth, Keycloak Enterprise); cod aplicativ SAST/DAST în CI/CD; OWASP ASVS L3"],
            ["Utilizatori",
             "Sistem de gestiune a identității utilizatorilor",
             "Keycloak Enterprise — 150 utilizatori interni cu MFA + 185.000 utilizatori portal federați prin ROeID/eIDAS; RBAC granular; audit logging end-to-end"],
        ]
    )

    add_h3(doc, "10.2.2 Cele 6 principii ale securității cibernetice SIDISVA (cf. cap. 3.4.1.2 alin. 2)")
    add_p(doc,
        "Conform cap. 3.4.1.2 al Caietului de Sarcini, securitatea cibernetică SIDISVA se "
        "bazează pe 6 principii fundamentale, pe care consorțiul le aplică concret în "
        "arhitectură și procese:"
    )

    add_table(doc,
        header=["#", "Principiu CdS", "Aplicare concretă în SIDISVA"],
        rows=[
            ["P1",
             "Principiul conformității — mecanisme tehnice pentru reglementări naționale (ex. GDPR) și standarde internaționale (ex. ISO)",
             "ISMS ISO 27001 + PIMS ISO 27701 + DPIA + audit conformitate semestrial; toate componentele au mecanisme tehnice pentru GDPR (RoPA, drepturile data subject, notificare breach)"],
            ["P2",
             "Principiul optimizării costurilor — investițiile în securitate se stabilesc pe baza analizei de risc",
             "Analiză de risc Privacy + Securitate efectuată în faza de analiză (luna 1-6); reevaluare anuală în garanție; prioritizare investiții CIS Controls / NIST CSF; plafon CdS minimum 10% buget"],
            ["P3",
             "Principiul responsabilităților de securitate partajate — rolurile entităților implicate stabilite, reglementate, asumate",
             "Matrice RACI documentată (Confluence) — operator Cloud Guvernamental / consorțiu / ANSVSA / utilizatori finali; SLA-uri reciproce; acord de confidențialitate semnat de toți 9 membri consorțiu"],
            ["P4",
             "Principiul protecției informațiilor — (a) protecție în tranzit + stocare împotriva accesării/modificării neautorizate; (b) disponibilitate fără întârziere la cererea entităților autorizate",
             "(a) TLS 1.3 in-transit + AES-256 / TDE at-rest + Azure Key Vault HSM-backed; (b) RPO ≤1h + RTO ≤4h + DR site Azure regiune diferită + replicare asincronă continuă"],
            ["P5",
             "Principiul securității pe întreg ciclul de viață al sistemului — securitatea integrată în toate fazele, de la analiză la scoaterea din uz",
             "Security Champion în fiecare echipă de furnizor; review de securitate la fiecare milestone; documentație ADR cu impact securitate; procedură decommissioning sigur (wipe NIST 800-88 la final contract)"],
            ["P6",
             "Principiul securizării operațiunilor — mecanisme pentru detectarea și prevenirea atacurilor, abordare pe niveluri",
             "SIEM Splunk ES / IBM QRadar + SOC 24×7 + Honeypot FortiDeceptor + Threat Intelligence (MISP/ENISA/DNSC) + SOAR playbook-uri automate + exerciții red-team anuale"],
        ]
    )

    # ============================================================
    # 10.3 Cele 3 principii CIA (cap. 3.4.6 CdS)
    # ============================================================
    add_h2(doc, "10.3 Principiile fundamentale ale securității datelor (CIA — cf. cap. 3.4.6 CdS)")

    add_p(doc,
        "Conform cap. 3.4.6 al Caietului de Sarcini, securitatea datelor gestionate de SIDISVA "
        "respectă cele 3 principii fundamentale CIA (Confidentiality, Integrity, Availability):"
    )

    add_table(doc,
        header=["Principiu", "Definiție CdS (citată)", "Implementare consorțiu"],
        rows=[
            ["Confidențialitate",
             "Asigurarea protecției datelor împotriva accesărilor neautorizate",
             "(1) Autentificare obligatorie prin Keycloak — MFA pentru personal intern, ROeID/eIDAS pentru cetățeni; (2) RBAC granular pe nivel de funcționalitate + atribute (ABAC) pentru date sensibile; (3) Encryption at rest + in transit; (4) Data masking în medii non-prod; (5) Audit logging end-to-end în Splunk SIEM"],
            ["Integritate",
             "Protecția, exactitatea și completitudinea datelor; protecție împotriva manipulării frauduloase",
             "(1) Tranzacții ACID pe MS SQL Server Enterprise; (2) Constraints + foreign keys + check constraints; (3) Audit trail imutabil cu hash chains pentru date critice (BND-SNIIA evenimente animale); (4) Validare input cu pattern matching + business rules; (5) Detection anomalii prin AI/ML (sec. 2.13.2 — Isolation Forest pentru fraude BND)"],
            ["Disponibilitate",
             "Redundanța tuturor componentelor pentru păstrarea coerenței și ne-coruperii datelor",
             "(1) HA pe toate componentele critice (NGFW redundant, SQL Server cluster, ESB cluster, Web/App tier auto-scaling Kubernetes); (2) Backup 3-2-1 cu RPO 1h; (3) DR site regiune Cloud Gov diferită cu RTO 4h; (4) DDoS protection layer prin NGFW + Cloud Gov scrubbing; (5) Monitoring SLO p95 < 2s prin Prometheus + Grafana"],
        ]
    )

    # ============================================================
    # 10.4 Defense-in-depth pe 7 straturi
    # ============================================================
    add_h2(doc, "10.4 Arhitectura defense-in-depth pe 7 straturi (cf. cap. 3.4.6 CdS — apărare în adâncime)")

    add_p(doc,
        "Conform cap. 3.4.6 al Caietului de Sarcini care prevede aplicarea modelului apărării "
        "în adâncime, consorțiul distribuie mecanismele de protecție pe 7 straturi independente. "
        "Fiecare strat are produsele și parametrii săi specifici (cantitățile conform cap. "
        "3.4.3.4 CdS, vezi cap. 6 al ofertei tehnice). O breșă într-un strat nu compromite "
        "celelalte straturi:"
    )

    add_table(doc,
        header=["Strat", "Mecanisme de protecție", "Produse / Tehnologii"],
        rows=[
            ["1. Fizic",
             "Găzduire Cloud Guvernamental (Azure RO) certificat ISO 27001 + ISAE 3402; control acces fizic conform politicilor CG; protecție environmental (incendiu, inundație, putere redundantă)",
             "Cloud Guvernamental (Azure Stack RO)"],
            ["2. Rețea perimetru",
             "NGFW redundant centru (2 instanțe HA, min. 10 VDOM); firewall în 45 locații DSVSA + institute (90 instanțe HA); IPSec VPN site-to-site; protecție DDoS prin NGFW + scrubbing Cloud Gov",
             "FortiGate / Palo Alto (centru); FortiGate 100F (90 locații)"],
            ["3. Rețea internă",
             "Segmentare VLAN per zonă (DMZ / business / management / backup); NAC 802.1X pentru min. 600 dispozitive (cf. cap. 3.4.3.4.1.3 CdS); NMS pentru detecție anomalii trafic",
             "Cisco DNA Center / Cisco ISE / Aruba ClearPass"],
            ["4. Aplicație",
             "WAF cu protecție OWASP Top 10 (2 instanțe HA, cf. cap. 3.4.3.4.1.1 CdS — nu per useri); Email Security cu filtrare + sandbox (2 instanțe); SAST + DAST automate în CI/CD pipeline; OWASP ASVS L3 pentru aplicații web; OWASP MASVS L1 pentru mobile",
             "F5 / Imperva / FortiWeb (WAF); Cisco IronPort / Secure Email"],
            ["5. Identitate (IAM)",
             "Keycloak Enterprise (150 utilizatori interni + 185.000 portal); federație ROeID pentru cetățeni + eIDAS pentru UE; MFA obligatoriu pentru utilizatori interni (TOTP / FIDO2); SSO via OAuth2 / OIDC; RBAC granular; just-in-time access pentru administratori privilegiați",
             "Keycloak Enterprise + ROeID + Nodul eIDAS"],
            ["6. Date",
             "Encryption at rest — TDE Microsoft SQL Server Enterprise (cf. cap. 3.4.3.2.4 CdS), BitLocker pe endpoints, Azure Disk Encryption pentru VM; encryption in transit — TLS 1.3 obligatoriu pe toate canalele; Azure Key Vault HSM-backed pentru chei criptografice (FIPS 140-2 Level 3); data masking + pseudo-anonimizare pentru medii test/dev",
             "MS SQL TDE + Azure Disk Encryption + Azure Key Vault HSM"],
            ["7. Endpoint",
             "EDR (Endpoint Detection and Response) pe 486 endpoints (436 + 50 administratori, cf. xlsx); Microsoft Defender for Cloud; politici Group Policy; controlul dispozitivelor USB (block by default); patch management automat; conformitate L 354/2022 (antivirus din SUA, NU Kaspersky)",
             "CrowdStrike Falcon / SentinelOne Singularity + MS Defender"],
        ]
    )

    # ============================================================
    # 10.5 Tehnologii de securitate concrete (cap. 3.4.6)
    # ============================================================
    add_h2(doc, "10.5 Tehnologii concrete de securitate (cf. cap. 3.4.6 CdS)")

    add_p(doc,
        "Cap. 3.4.6 al Caietului de Sarcini enumerează explicit tehnologiile care trebuie să "
        "fie incluse în soluție. Tabelul de mai jos arată cum sunt acoperite toate aceste "
        "cerințe de către consorțiu:"
    )

    add_table(doc,
        header=["Cerință CdS (citată)", "Tehnologie ofertată", "Cantitate (cf. CdS 3.4.3.4)"],
        rows=[
            ["Securizarea serviciilor web utilizând tehnologii de tip web application firewall",
             "WAF F5 BIG-IP ASM / Imperva Cloud WAF / Fortinet FortiWeb",
             "2 instanțe HA (cap. 3.4.3.4.1.1)"],
            ["Protecția și securizarea comunicațiilor utilizând soluții firewall",
             "NGFW FortiGate / Palo Alto / Cisco Firepower",
             "2 (centru ANSVSA, cap. 3.4.3.4.2.1.1) + 90 (45 locații × 2 HA, cap. 3.4.3.4.2.1.2)"],
            ["Managementul rețelei, analiza și raportarea logurilor de securitate",
             "SIEM Splunk Enterprise Security / IBM QRadar + NMS Cisco DNA Center",
             "1 cluster SIEM (cap. 3.4.3.4.1.4) + min. 600 dispozitive NMS (cap. 3.4.3.4.1.3)"],
            ["Managementul vulnerabilităților și al actualizărilor",
             "Tenable Nessus + Qualys VMDR + OWASP Dependency-Check + Snyk + Trivy",
             "Scanning continuu pe toată infrastructura + dependencies"],
            ["Sistem anti-malware centralizat pentru end-point-uri",
             "CrowdStrike Falcon / SentinelOne Singularity (EDR) + Microsoft Defender",
             "436 + 50 endpoints (conform cf. xlsx Sinteza)"],
            ["Sistem de control acces — autentificare cf. zero-trust",
             "Keycloak Enterprise + MFA TOTP/FIDO2 + RBAC/ABAC + just-in-time access",
             "150 utilizatori interni + 185.000 portal (cap. 3.4.3.2.7)"],
            ["Auditarea activităților realizate în sistem și a solicitărilor de acces",
             "Splunk SIEM cu retenție IOC 3 ani (cap. 3.4.3.4.1.4); audit logging end-to-end în toate componentele aplicative",
             "1 cluster SIEM cu retenție conform NIS2"],
            ["Honeypot (cerință implicită cap. 3.4.3.4.1.2)",
             "FortiDeceptor (4 instanțe Win × 4 vCPU)",
             "1 management + 4 decoy"],
            ["Email Security (cerință cap. 3.4.3.4.1.5)",
             "Cisco IronPort / Secure Email Gateway",
             "2 instanțe HA"],
        ]
    )

    # ============================================================
    # 10.6 Confidențialitatea datelor (cap. 3.4.7)
    # ============================================================
    add_h2(doc, "10.6 Confidențialitatea datelor (cf. cap. 3.4.7 CdS)")

    add_p(doc,
        "Conform cap. 3.4.7 al Caietului de Sarcini, confidențialitatea este o activitate de "
        "bază pentru furnizarea serviciilor publice. Consorțiul implementează următoarele "
        "măsuri concrete pentru a respecta integral cerințele CdS:"
    )

    add_bullet(doc,
        "Acces protejat și auditat permanent — toate accesările datelor în SIDISVA sunt "
        "autentificate prin Keycloak (MFA pentru interni, ROeID/eIDAS pentru externi); fiecare "
        "accesare este loggată în Splunk SIEM cu identitate utilizator, timestamp, resursă "
        "accesată, IP origine; retenție log-uri 5 ani."
    )
    add_bullet(doc,
        "Criptare BD pentru câmpuri sensibile — Microsoft SQL Server Enterprise oferă "
        "Transparent Data Encryption (TDE) pentru întreaga bază de date plus Always Encrypted "
        "pentru câmpurile cu PII (CNP, CUI, date contact); chei gestionate în Azure Key Vault "
        "HSM-backed (FIPS 140-2 Level 3)."
    )
    add_bullet(doc,
        "Comunicații securizate (SSL/TLS + VPN) — TLS 1.3 obligatoriu pe toate canalele "
        "(deprecate TLS 1.0/1.1/1.2); HSTS preload + Certificate Pinning în aplicațiile mobile; "
        "VPN IPSec site-to-site între cele 45 locații și centrul ANSVSA."
    )
    add_bullet(doc,
        "Analiză periodică log-uri pentru identificare intruziuni — SIEM Splunk cu reguli "
        "MITRE ATT&CK; SOC 24×7; threat hunting săptămânal; analiza pattern-uri lunare; raport "
        "trimestrial cu IoC-uri identificate și remediate."
    )
    add_bullet(doc,
        "Privacy by Design și Privacy by Default — minimizare date colectate (doar necesarul "
        "serviciului); pseudo-anonimizare pentru analitică; retention policies automatizate; "
        "consimțământ explicit pentru prelucrări secundare; DPIA înainte de Go-live."
    )
    add_bullet(doc,
        "Conformitate cu GDPR + NIS2 + L 354/2022 — toate principiile prelucrării respectate; "
        "notificare breach ANSPDCP în 72h conform Art. 33-34 GDPR; notificare DNSC în 24h/72h "
        "conform Art. 23 NIS2 + OUG 155/2024."
    )

    # ============================================================
    # 10.7 Cele 3 planuri obligatorii (cap. 3.4.7 CdS)
    # ============================================================
    add_h2(doc, "10.7 Cele 3 planuri obligatorii (cf. cap. 3.4.7 CdS)")

    add_p(doc,
        "Cap. 3.4.7 al Caietului de Sarcini prevede explicit obligația de a elabora și menține "
        "3 planuri de securitate. Consorțiul livrează aceste planuri ca documente standalone, "
        "actualizate anual pe perioada de garanție:"
    )

    add_h3(doc, "10.7.1 Planul de gestionare a riscurilor (Risk Management Plan)")
    add_bullet(doc,
        "Cerința CdS: identificarea riscurilor, evaluarea potențialului impact, planificarea "
        "intervențiilor cu măsuri tehnice și organizatorice; nivel de securitate proporțional "
        "cu gradul de risc."
    )
    add_bullet(doc,
        "Implementare consorțiu: metodologie ISO/IEC 27005 + NIST SP 800-30; identificare "
        "riscuri prin workshop-uri cu ANSVSA + threat modeling STRIDE; analiza calitativă "
        "(matrice probabilitate × impact 5×5) + analiză cantitativă (Monte Carlo pentru "
        "riscurile critice); plan mitigare per risc cu owner + termen + cost; reevaluare "
        "lunară de către Project Manager <LIDER> + reevaluare anuală cu DPO."
    )
    add_bullet(doc,
        "Livrabil: Risk Register actualizat lunar în Confluence; Risk Review semestrial cu "
        "ANSVSA în Steering Committee."
    )

    add_h3(doc, "10.7.2 Planul de continuitate a activității + Plan de rezervă și redresare (BCP/DR)")
    add_bullet(doc,
        "Cerința CdS: instituirea procedurilor pentru asigurarea disponibilității funcțiilor "
        "după un eveniment dezastruos + readucerea tuturor funcțiilor la situația normală cât "
        "mai rapid."
    )
    add_bullet(doc,
        "Implementare consorțiu: Business Continuity Plan (BCP) conform ISO 22301; Disaster "
        "Recovery Plan (DRP) cu RPO ≤1h și RTO ≤4h pentru funcții critice (Portal, BND-SNIIA, "
        "LIMS); strategie backup 3-2-1 (3 copii × 2 medii × 1 off-site); DR site secundar "
        "în Cloud Guvernamental regiune diferită; replicare asincronă continuă pentru date "
        "critice."
    )
    add_bullet(doc,
        "Scenarii BCP testate: ransomware, indisponibilitate DC primar, atac DDoS prelungit, "
        "pierdere cont Cloud Guvernamental; testare DR semestrială pe perioada implementării "
        "+ lunară pe perioada de garanție; rapoarte de exercițiu cu lessons learned."
    )
    add_bullet(doc,
        "Livrabil: BCP + DRP ca documente standalone livrate în luna 8; actualizate anual."
    )

    add_h3(doc, "10.7.3 Planul de acces la date și autorizare (Access Control Plan)")
    add_bullet(doc,
        "Cerința CdS: stabilirea persoanelor care au acces la date, datele accesibile, "
        "condițiile accesării; monitorizarea accesului neautorizat și încălcărilor; măsuri "
        "documentate pentru prevenire."
    )
    add_bullet(doc,
        "Implementare consorțiu: matrice de acces RBAC + ABAC documentată în Confluence + "
        "configurată tehnic în Keycloak Enterprise; principiul least privilege aplicat strict; "
        "access reviews trimestriale pentru utilizatorii cu privilegii ridicate; just-in-time "
        "access pentru administratori (acces temporar cu aprobare și logging detaliat); "
        "segregare obligații (separation of duties) între administratori IT și Data Stewards."
    )
    add_bullet(doc,
        "Monitorizare access denied: alertă SIEM la 5+ încercări eșuate în 10 minute; "
        "alertă la accesare în afara orelor de program (pentru utilizatorii care nu sunt "
        "configurați explicit pentru on-call); raport săptămânal cu top 10 utilizatori cu "
        "activitate anormală."
    )
    add_bullet(doc,
        "Livrabil: Access Control Plan ca document standalone; matrice RBAC actualizată lunar; "
        "audit trimestrial al rolurilor și permisiunilor."
    )

    # ============================================================
    # 10.8 DevSecOps
    # ============================================================
    add_h2(doc, "10.8 Securitatea procesului de dezvoltare (DevSecOps shift-left)")

    add_p(doc,
        "Conform principiului P5 (Securitate pe întreg ciclul de viață) din cap. 3.4.1.2 CdS, "
        "securitatea este integrată în procesul de dezvoltare prin practici DevSecOps "
        "shift-left:"
    )
    add_bullet(doc,
        "SAST (Static Application Security Testing) — analiză cod sursă în fiecare commit "
        "Git cu SonarQube + Semgrep + Checkmarx; gate-uri în pipeline CI/CD pentru blocarea "
        "CVE high/critical; rapoarte automate cu trending."
    )
    add_bullet(doc,
        "DAST (Dynamic Application Security Testing) — testare săptămânală pe medii staging "
        "cu OWASP ZAP + Burp Suite Pro + Acunetix; scan full OWASP Top 10; rapoarte săptămânale "
        "către Security Champion."
    )
    add_bullet(doc,
        "SCA (Software Composition Analysis) — Snyk + Trivy + OWASP Dependency-Check pe "
        "fiecare build; alertă instant la CVE high/critical într-o dependență; politică update "
        "max 30 zile pentru CVE critic."
    )
    add_bullet(doc,
        "Secret Scanning — gitleaks + TruffleHog la fiecare commit pentru a preveni leak-uri "
        "credențiale (chei API, parole, token-uri JWT) în repository Git."
    )
    add_bullet(doc,
        "Container Security — image scanning cu Trivy pentru imagini Docker; runtime "
        "protection cu Falco; policy-as-code cu Open Policy Agent (OPA) Gatekeeper pe "
        "Kubernetes."
    )
    add_bullet(doc,
        "Threat Modeling — sesiune STRIDE per modul în faza de proiectare; ADR-uri cu impact "
        "securitate documentate; review obligatoriu de Arhitectul Sistem."
    )
    add_bullet(doc,
        "Code Review obligatoriu — minimum 2 reviewers pentru orice merge în branch-uri "
        "principale; checklist de securitate inclus în template-ul de Pull Request."
    )

    # ============================================================
    # 10.9 Detectare și răspuns la incidente
    # ============================================================
    add_h2(doc, "10.9 Detectare și răspuns la incidente (Detect & Respond)")

    add_p(doc,
        "Conform principiului P6 (Securizarea operațiunilor) din cap. 3.4.1.2 CdS și "
        "obligațiilor NIS2 + OUG 155/2024, consorțiul implementează un cadru complet de "
        "detectare și răspuns la incidente:"
    )
    add_bullet(doc,
        "SIEM (Splunk Enterprise Security sau IBM QRadar) — colectare centralizată log-uri "
        "din toate componentele; normalizare Common Information Model (CIM); retenție IOC "
        "3 ani conform cap. 3.4.3.4.1.4 CdS; corelație evenimente pe regulile MITRE ATT&CK; "
        "peste 500 reguli de alertare predefinite."
    )
    add_bullet(doc,
        "SOC (Security Operations Center) 24×7 — intern <LIDER> (echipă dedicată ISO 27001 "
        "certificată) sau externalizat la MSSP autorizat DNSC; structura tier 1 (triaj) + "
        "tier 2 (investigare) + tier 3 (forensics + threat hunting); rotație on-call."
    )
    add_bullet(doc,
        "Honeypot FortiDeceptor — 4 instanțe decoy Windows × 4 vCPU (cf. cap. 3.4.3.4.1.2 "
        "CdS); decoys cu credențiale fake pentru detectarea lateral movement și a "
        "atacatorilor cu acces inițial; alertare imediată în SIEM."
    )
    add_bullet(doc,
        "Threat Intelligence — feed-uri MISP (Malware Information Sharing Platform) + ENISA + "
        "DNSC (autoritatea națională) + commerciale (Recorded Future / Mandiant); integrare "
        "automată cu SIEM pentru corelație."
    )
    add_bullet(doc,
        "SOAR (Security Orchestration, Automation and Response) — playbook-uri automate "
        "pentru incidente comune (brute force lockout, malware contain, phishing quarantine, "
        "ransomware indicators); izolare automată endpoints compromise."
    )
    add_bullet(doc,
        "Notificare incidente conform NIS2 + OUG 155/2024 — early warning DNSC în 24h, "
        "notificare incident în 72h, raport final în 1 lună; proceduri pregătite pre-Go-live "
        "și testate prin exerciții semestriale; canal de comunicare DNSC validat."
    )

    # ============================================================
    # 10.10 Vulnerability management + pen-test
    # ============================================================
    add_h2(doc, "10.10 Vulnerability management și penetration testing anual")

    add_p(doc,
        "Conform cap. 3.4.9 punctul 13 al Caietului de Sarcini (managementul vulnerabilităților "
        "+ teste de penetrare anuale), consorțiul oferă pe întreaga perioadă de garanție:"
    )
    add_bullet(doc,
        "Vulnerability scanning continuu — Tenable Nessus Professional + Qualys VMDR + OWASP "
        "Dependency-Check + Snyk + Trivy pentru containere; scanning automat săptămânal."
    )
    add_bullet(doc,
        "SLA remediere — CVE Critical (CVSS 9.0+) în max 7 zile; CVE High (7.0-8.9) în max "
        "30 zile; CVE Medium (4.0-6.9) în max 90 zile; CVE Low în max 180 zile sau next "
        "release."
    )
    add_bullet(doc,
        "Penetration testing anual extern — auditor independent certificat ANSCC, schimbat "
        "la fiecare 2 ani; scope: aplicații web, API REST/GraphQL, infrastructură, aplicații "
        "mobile iOS/Android, social engineering simulat (phishing campaign)."
    )
    add_bullet(doc,
        "Metodologie pen-test: OWASP Testing Guide v4.2 + PTES + NIST SP 800-115; raport "
        "livrat ANSVSA în max 30 zile post finalizare; remediere obligatorie a findings "
        "Critical/High pre-acceptare."
    )
    add_bullet(doc,
        "Bug Bounty intern în perioada UAT — utilizatori cheie ANSVSA invitați să raporteze "
        "probleme de securitate cu recompense simbolice; canal dedicat în Jira Service "
        "Management."
    )

    # ============================================================
    # 10.11 Conformitate L 354/2022 (ANTI-RU/UA)
    # ============================================================
    add_h2(doc, "10.11 Conformitate cu Legea 354/2022 (protecția sistemelor IT publice)")

    add_p(doc,
        "Legea 354/2022 privind protecția sistemelor informatice ale autorităților și "
        "instituțiilor publice în contextul invaziei declanșate de Federația Rusă împotriva "
        "Ucrainei impune restricții asupra utilizării de software de la producători din "
        "state cu risc de securitate cibernetică. Consorțiul confirmă conformitatea integrală:"
    )

    add_table(doc,
        header=["Categorie", "Produs propus", "Producător / Sediu", "Conform L 354/2022"],
        rows=[
            ["Antivirus / EDR", "CrowdStrike Falcon / SentinelOne Singularity / MS Defender", "SUA / SUA / SUA", "Da (NU Kaspersky sau alte produse din state cu restricții)"],
            ["NGFW", "FortiGate / Palo Alto / Cisco Firepower", "SUA / SUA / SUA", "Da"],
            ["WAF", "F5 / Imperva / FortiWeb", "SUA / Israel-SUA / SUA", "Da"],
            ["SIEM", "Splunk Enterprise Security / IBM QRadar", "SUA / SUA", "Da"],
            ["Honeypot", "FortiDeceptor", "SUA (Fortinet)", "Da"],
            ["NMS / NAC", "Cisco DNA / Cisco ISE / Aruba ClearPass", "SUA / SUA / SUA-HPE", "Da"],
            ["Email Security", "Cisco IronPort / Secure Email", "SUA", "Da"],
            ["SGBD", "Microsoft SQL Server Enterprise", "SUA (Microsoft)", "Da"],
            ["ESB", "Oracle Service Bus", "SUA (Oracle)", "Da"],
            ["IAM", "Keycloak Enterprise", "SUA (Red Hat / IBM)", "Da"],
            ["OS server", "Red Hat Enterprise Linux 9 / Windows Server 2022", "SUA / SUA", "Da"],
            ["Web/App server", "NGINX Plus / Microsoft IIS", "SUA (F5) / SUA (Microsoft)", "Da"],
            ["Productivity", "Microsoft Office Home & Business 2024", "SUA", "Da"],
        ]
    )

    add_p(doc,
        "Declarațiile producătorilor privind originea producției și absența componentelor "
        "critice din state cu restricții sunt anexate ca Anexa C la propunerea tehnică. În "
        "cazul modificărilor legislative ulterioare care extind restricțiile L 354/2022, "
        "consorțiul se angajează contractual să propună ANSVSA o alternativă echivalentă "
        "tehnic în max 30 de zile calendaristice, fără cost suplimentar."
    )

    # ============================================================
    # 10.12 Plafon minim 10% buget
    # ============================================================
    add_h2(doc, "10.12 Conformitate plafon minim 10% buget securitate cibernetică")

    add_p(doc,
        "Conform cap. 11 al Caietului de Sarcini și Ghidului solicitantului POCIDIF, valoarea "
        "minimă finanțată pentru cheltuielile privind asigurarea securității cibernetice a "
        "rețelei și sistemelor informatice dezvoltate trebuie să fie de minim 10% din "
        "valoarea totală finanțată a proiectului."
    )

    add_table(doc,
        header=["Element", "Valoare"],
        rows=[
            ["Valoarea totală eligibilă proiect SIDISVA", "95.271.200,60 lei fără TVA"],
            ["Plafon minim cheltuieli securitate cibernetică (10%)", "9.527.120,06 lei fără TVA"],
            ["Valoarea estimată contract achiziție", "85.418.857,53 lei fără TVA"],
            ["Plafon minim securitate raportat la contract", "min. 8.541.885,75 lei fără TVA (10%)"],
        ]
    )

    add_p(doc,
        "Consorțiul <LIDER> confirmă că valoarea totală a componentelor de securitate "
        "(hardware + software + servicii) inclusă în Propunerea Financiară respectă acest "
        "plafon minim de 10%. Defalcarea exactă pe poziții este prezentată în Formularul "
        "de Propunere Financiară. Categoriile de cheltuieli de securitate cuprinse:"
    )
    add_bullet(doc, "Hardware securitate: WAF (2), Honeypot (1+4 inst), NMS/NAC (1+600 dev), SIEM (1 cluster), Email Security (2), NGFW centru (2), NGFW locații (90), EDR (486 endpoints)")
    add_bullet(doc, "Licențe software securitate: SIEM Splunk/QRadar, antivirus EDR CrowdStrike/SentinelOne, WAF, Keycloak Enterprise (IAM), Mirth Connect (HL7 security), backup & DR tools")
    add_bullet(doc, "Servicii securitate implementare: configurare WAF/SIEM/NGFW (luna 8-12), penetration testing inițial (luna 17), DPIA elaborare (luna 16), DR site setup (luna 12-15), instruire securitate pentru administratori (luna 18)")
    add_bullet(doc, "Servicii securitate garanție (3 ani): SOC 24×7, vulnerability scanning continuu, pen-test anual extern, threat intelligence feeds, exerciții simulare incidente semestrial")
    add_bullet(doc, "Conformitate: audit ISO 27001 + ISO 27701 inițial + anual, asigurare cyber-risk pe perioada implementării și garanție")

    doc.save(str(OUT))
    size = OUT.stat().st_size
    print(f"OK | {OUT.name} salvat: {size} bytes ({size/1024:.1f} KB)")


if __name__ == "__main__":
    main()

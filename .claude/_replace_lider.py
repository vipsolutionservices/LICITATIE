"""
Inlocuieste in 3-Plan_implementare.docx variantele de nume de ofertant cu <LIDER>.
Reguli:
 - "VOGO TECHNOLOGY" -> "<LIDER>"
 - "Vogo Technology" -> "<LIDER>" (oricare casing combination)
 - Cuvant izolat "VOGO" (cu boundary) -> "<LIDER>"
Pastreaza formatarea: incearca replace la nivel de run; daca un match e split in mai multe runs, face merge in primul run (pastrand formatarea primului).
"""
import re
import sys
from pathlib import Path

from docx import Document

SRC = Path(r"C:\Users\adria\Documents\Claude\Projects\Licitatie 1\doc-output\oferta\3-Plan_implementare.docx")

# Patterns (ordine importanta: cele mai lungi primele)
PATTERNS = [
    (re.compile(r"VOGO\s+TECHNOLOGY", re.IGNORECASE), "<LIDER>"),
    (re.compile(r"\bVOGO\b", re.IGNORECASE), "<LIDER>"),
]


def replace_in_text(text: str) -> tuple[str, int]:
    total = 0
    for pat, repl in PATTERNS:
        text, n = pat.subn(repl, text)
        total += n
    return text, total


def replace_in_paragraph(paragraph) -> int:
    """Aplica regex pe textul paragrafului si rescrie runs daca exista match."""
    full = paragraph.text
    new_text, n = replace_in_text(full)
    if n == 0:
        return 0
    runs = paragraph.runs
    if not runs:
        return 0
    # Pune tot textul nou in primul run, sterge restul
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""
    return n


def main():
    if not SRC.exists():
        print(f"NOT FOUND: {SRC}", file=sys.stderr)
        sys.exit(1)
    doc = Document(str(SRC))
    total = 0
    # Paragrafe top-level
    for p in doc.paragraphs:
        total += replace_in_paragraph(p)
    # Tabele (inclusiv nested)
    def walk_tables(tables):
        nonlocal total
        for t in tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        total += replace_in_paragraph(p)
                    walk_tables(cell.tables)
    walk_tables(doc.tables)
    # Headers/footers
    for section in doc.sections:
        for hdr in (section.header, section.first_page_header, section.even_page_header):
            for p in hdr.paragraphs:
                total += replace_in_paragraph(p)
            walk_tables(hdr.tables)
        for ftr in (section.footer, section.first_page_footer, section.even_page_footer):
            for p in ftr.paragraphs:
                total += replace_in_paragraph(p)
            walk_tables(ftr.tables)
    doc.save(str(SRC))
    print(f"Replaced {total} occurrences in {SRC.name}")


if __name__ == "__main__":
    main()

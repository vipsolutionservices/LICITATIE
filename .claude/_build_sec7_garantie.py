"""
Refacere completa 7-Plan_garantie.docx — rewrite conform:
- Cap. 3.4.9 CdS: garantie si suport 3 ani de la PIP, 14 cerinte explicite
- Cap. 3.4.2.14 CdS: integrare ulterioara sisteme guvernamentale cu mock-up FARA cost
- Cap. 7.3 CdS: durata garantie

Eliminare:
- VOGO TECHNOLOGY hardcodat (×7) -> <LIDER>
- Referinta gresita 'cap. 4.9' -> 'cap. 3.4.9'
- 24x7 helpdesk generic -> 8h/zi L-V helpdesk + 24/7 doar pentru S1
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt

OUT = Path(r"C:\Users\adria\Documents\Claude\Projects\Licitatie 1\doc-output\oferta\7-Plan_garantie.docx")


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(18)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    return p


def add_p(doc, text):
    p = doc.add_paragraph()
    p.add_run(text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    p.add_run("• " + text)
    return p


def add_table(doc, header, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    try:
        t.style = "Table Grid"
    except KeyError:
        pass
    hdr = t.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
    for ri, row in enumerate(rows, start=1):
        cells = t.rows[ri].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    return t


def main():
    doc = Document()

    # ============================================================
    # Titlu + intro
    # ============================================================
    add_h1(doc, "7. Plan de măsuri pentru perioada de garanție și suport tehnic")

    add_p(doc,
        "Prezentul capitol descrie în detaliu serviciile de garanție și suport tehnic pe care "
        "consorțiul condus de <LIDER> le va asigura pe perioada de 3 ani de la data acceptanței "
        "finale a sistemului SIDISVA (trecerea în producție), conform cerințelor explicite ale "
        "cap. 3.4.9 al Caietului de Sarcini. Capitolul răspunde integral celor 14 cerințe ale "
        "CdS privind serviciile de garanție și include, suplimentar, mecanismul de integrare "
        "ulterioară cu sisteme guvernamentale neimplementate (cap. 3.4.2.14 CdS), realizat "
        "fără cost suplimentar pentru Autoritatea Contractantă."
    )

    add_p(doc,
        "Pe parcursul perioadei de garanție, <LIDER> menține o comunicare activă cu ANSVSA "
        "pentru a asigura funcționarea eficientă a sistemului informatic, reducerea timpilor "
        "de nefuncționare și optimizarea costurilor de întreținere — așa cum prevede paragraful "
        "introductiv al cap. 3.4.9 al Caietului de Sarcini."
    )

    # ============================================================
    # 7.1 Domeniul garanției
    # ============================================================
    add_h2(doc, "7.1 Domeniul garanției — ce este acoperit")

    add_p(doc,
        "Garanția de 3 ani acoperă integral toate componentele sistemului SIDISVA — hardware, "
        "software de bază, software aplicativ și serviciile de suport tehnic asociate. "
        "Tabelul de mai jos detaliază acoperirea pe categorii:"
    )

    add_table(doc,
        header=["Categorie", "Componente acoperite", "Durată garanție", "Furnizor responsabil"],
        rows=[
            ["A. Hardware",
             "NGFW (2 centru + 90 locații), WAF, SIEM, Honeypot, NMS/NAC, Email Security, Switch acces/POE/agregare, Access Point Wi-Fi 6, 100 laptopuri, 336 complete teren (terminal + pad semnătură + imprimantă mobilă)",
             "3 ani de la PIP",
             "<LIDER> ca SPOC + producători (Fortinet, Cisco, F5, Splunk, Dell/HP/Lenovo, Wacom/Topaz, Epson/Canon/HP)"],
            ["B. Software de bază (COTS)",
             "Microsoft SQL Server Enterprise, Oracle Service Bus, Microsoft SSIS + Power BI + SSRS + SSAS, Keycloak Enterprise, Elasticsearch, NGINX Plus, Microsoft IIS, Windows Server 2022, RHEL / Oracle Linux, Mirth Connect, Microsoft Office Home & Business 2024",
             "3 ani de la PIP, cu update-uri majore + minore + patch-uri securitate",
             "<LIDER> + producătorii (Microsoft, Oracle, Red Hat, NextGen Healthcare, F5/NGINX, Elastic)"],
            ["C. Software aplicativ",
             "ZIPPER DMS (cod sursă inclus), Enterprise Suite (Portal + Chatbot + App mobilă), LIMS COTS (cod sursă inclus), integrări custom dezvoltate de consorțiu",
             "3 ani de la PIP, cu corecții bug-uri + mentenanță evolutivă minoră",
             "<FURNIZOR DMS>, <FURNIZOR PORTAL>, <FURNIZOR LIMS>, <LIDER> (integrări custom)"],
            ["D. Suport tehnic",
             "Help-desk telefonic 8h/zi L-V (cf. CdS cap. 3.4.9), aplicație ticketing 24×7, intervenție de urgență 24×7 pentru incidente S1 critice, vulnerability management, pen-test anual",
             "3 ani de la PIP",
             "<LIDER> (coordonare) + echipele L1/L2/L3 ale furnizorilor"],
            ["E. Integrare ulterioară sisteme guvernamentale",
             "Integrarea efectivă cu PNI, PJN, PCUe/PDURo, eIDAS extins și alte sisteme guvernamentale care nu erau finalizate la momentul implementării — conform cap. 3.4.2.14 CdS, fără cost suplimentar",
             "3 ani de la PIP",
             "<LIDER> + <FURNIZOR ESB>"],
        ]
    )

    # ============================================================
    # 7.2 Activitățile incluse (cele 14 cerințe CdS 3.4.9)
    # ============================================================
    add_h2(doc, "7.2 Activitățile asigurate în perioada de garanție (cf. cap. 3.4.9 CdS)")

    add_p(doc,
        "În conformitate cu cele 14 cerințe explicite ale cap. 3.4.9 al Caietului de Sarcini, "
        "consorțiul asigură următoarele tipuri de activități pe parcursul perioadei de garanție:"
    )

    add_table(doc,
        header=["#", "Activitate (citată din CdS)", "Modul concret de asigurare"],
        rows=[
            ["1", "Rezolvarea bug-urilor care nu au fost identificate în timpul implementării și care apar după punerea în producție",
             "Pipeline dedicat de bug-fixing — categorie cod ER (Error); SLA conform secțiunii 7.4 (răspuns 30 min pentru S1, max 5 zile lucrătoare pentru S4); fix-urile sunt livrate în release-uri lunare + hotfix-uri urgente"],
            ["2", "Întreținerea și buna funcționare a sistemului furnizat în parametrii agreați (funcțional, performanță, disponibilitate, integritate date)",
             "Monitoring 24×7 cu Prometheus + Grafana + AlertManager + Splunk SIEM; KPIs urmărite: uptime ≥99.5%, latency p95 < 2s, integritate referențială BD, alerte automate la deviere"],
            ["3", "Instalarea de noi versiuni ale aplicațiilor în urma efectuării corecțiilor",
             "Procedură formală deployment: dev → test → staging → producție, cu validare ANSVSA la fiecare etapă; release notes detaliate, fereastră de mentenanță agreată cu ANSVSA"],
            ["4", "Instalarea de noi versiuni oferite de producător ale produselor COTS, în condițiile în care arhitectura sistemului și constrângerile o permit",
             "Monitorizare proactivă a release-urilor de la Microsoft (SQL Server, Office, Windows), Oracle (ESB), Red Hat (Keycloak, RHEL), Fortinet etc.; recomandare actualizare către ANSVSA cu analiza de impact și risc"],
            ["5", "Actualizarea manualelor de utilizare și a altor documente în urma efectuării corecțiilor",
             "Documentation-as-Code în Confluence; orice modificare funcțională declanșează update al manualului utilizator + manualului administrator + matricei de trasabilitate; livrare versionate (cu nr. versiune și data)"],
            ["6", "Reparații / înlocuiri ale componentelor defecte la locația de instalare a beneficiarului",
             "Intervenție on-site Next Business Day (NBD) pentru HW critic (NGFW centru, SIEM, switch agregare); stoc spare-parts strategic menținut de <LIDER> (min. 5% NGFW locații, 1 NGFW centru rezervă); RMA management gestionat integral"],
            ["7", "Consiliere și suport telefonic 8 ore pe zi, de luni până vineri, în cadrul programului normal de lucru al beneficiarului, prin serviciul Help-desk atât pentru produsele hardware cât și software",
             "Helpdesk telefonic dedicat: număr unic 8h/zi L-V (08:00-17:00 ora României) — conform cerinței minime CdS; suplimentar pentru incidente S1 critice <LIDER> oferă canal de urgență 24×7 ca avantaj over-delivery"],
            ["8", "Toate incidentele vor fi gestionate prin intermediul unei aplicații software de gestionare a tichetelor",
             "Aplicație Jira Service Management dedicată proiectului SIDISVA, cu interfață web + mobilă + API; integrare cu Slack/Teams ANSVSA pentru notificări; raportare automată; vizibilitate cross-furnizor"],
            ["9", "Remediere software de la distanță cu acordul beneficiarului",
             "Conectivitate securizată prin VPN dedicat <LIDER>-ANSVSA; sesiuni de mentenanță remote loggate complet în SIEM; aprobare scrisă (email/Jira) pre-intervenție pentru fiecare conexiune; logging acțiuni administrator"],
            ["10", "Actualizări software la locația de instalare a beneficiarului sau de la distanță",
             "Procedura standard de update permite ambele moduri; preferință pentru remote (eficient, fără deplasări); on-site doar pentru update-uri majore HW sau cerere expresă ANSVSA"],
            ["11", "Reconfigurări hardware și software la nivelul inițial solicitat în cazul în care erorile apărute nu sunt datorate beneficiarului",
             "Configurațiile inițiale aprobate (baseline) sunt arhivate în Confluence + repository Git; restaurare automată sau manuală în max 4 ore pentru S1 (cf. SLA); fără cost suplimentar dacă nu există culpă ANSVSA"],
            ["12", "Consiliere și suport tehnic pentru posibilități de extindere a soluției existente",
             "Sesiuni trimestriale de consultanță arhitecturală cu ANSVSA (analist sistem + arhitect <LIDER>); prezentare opțiuni de extindere, evaluare impact, estimare efort; documentate în Roadmap evolutiv vizibil ANSVSA"],
            ["13", "Managementul vulnerabilității, precum și teste de penetrare anuale",
             "Vulnerability scanning continuu (Tenable Nessus / Qualys VMDR); pen-test extern anual de către auditor independent certificat ANSCC; remediere CVE high/critical în 30 zile, medium în 90 zile, low în 180 zile; rapoarte trimestriale ANSVSA"],
        ]
    )

    # ============================================================
    # 7.3 Procedura de suport tehnic
    # ============================================================
    add_h2(doc, "7.3 Procedura de suport tehnic")

    add_p(doc,
        "Procedura de suport tehnic <LIDER> definește etapele standardizate prin care orice "
        "solicitare a ANSVSA — fie cerere de funcționalitate nouă (cod F), fie raportare de "
        "eroare (cod ER) — este preluată, procesată și soluționată în limitele SLA garantate."
    )

    add_h3(doc, "7.3.1 Etapele procesării unui tichet")
    add_bullet(doc,
        "Pasul 1 — Solicitare ANSVSA: tichetul poate fi deschis prin orice dintre canalele "
        "agreate: (a) portal Jira Service Management, (b) număr telefonic helpdesk dedicat "
        "în programul L-V 08:00-17:00, (c) număr de urgență 24×7 pentru incidente S1, (d) "
        "adresă email dedicată proiectului."
    )
    add_bullet(doc,
        "Pasul 2 — Categorizare și prioritizare: consultantul L1 al <LIDER> încadrează "
        "tichetul în categoria corectă (cerere funcționalitate nouă cod F / eroare existent "
        "cod ER), atribuie nivelul de severitate S1-S4 conform criteriilor din secțiunea 7.4 "
        "și alocă resursele necesare."
    )
    add_bullet(doc,
        "Pasul 3 — Investigare și diagnosticare: analiza log-urilor (Splunk SIEM), "
        "reproducerea problemei în mediul de staging, identificarea cauzei rădăcină (Root "
        "Cause Analysis — RCA); documentarea ipotezelor în Jira."
    )
    add_bullet(doc,
        "Pasul 4 — Soluționare: implementarea soluției / aplicarea fix-ului / executarea "
        "modificării de configurare; testare unit + integration; code review obligatoriu "
        "minim 2 reviewers pentru modificările de cod sursă."
    )
    add_bullet(doc,
        "Pasul 5 — Validare ANSVSA: testare în staging conform unui scenariu de testare "
        "agreat; semnătura electronică a responsabilului ANSVSA pe scenariul executat cu "
        "succes; deploy în producție cu confirmare post-deploy."
    )
    add_bullet(doc,
        "Pasul 6 — Documentare în baza de cunoștințe: adăugare entry în Knowledge Base "
        "Confluence pentru reutilizare ulterioară; actualizare manuale dacă este cazul; "
        "raport săptămânal trimis ANSVSA."
    )

    add_h3(doc, "7.3.2 Canale de suport disponibile")
    add_table(doc,
        header=["Canal", "Disponibilitate", "Recomandat pentru"],
        rows=[
            ["Portal Jira Service Management (web + mobil)", "24×7", "Toate tipurile de tichete; tracking + istoric complet"],
            ["Helpdesk telefonic (număr dedicat)", "L-V 08:00-17:00 (cf. CdS 3.4.9)", "Urgențe, clarificări rapide, escaladări"],
            ["Linie de urgență 24×7", "24×7 inclusiv weekend și sărbători", "EXCLUSIV pentru incidente S1 critice (sistem indisponibil)"],
            ["Email dedicat proiectului", "Răspuns max 1h L-V, 4h în afara programului", "Tichete non-urgente, atașamente, documentație"],
            ["Microsoft Teams / Slack ANSVSA", "L-V 08:00-17:00", "Comunicare informală, întrebări rapide, sincronizare zilnică"],
        ]
    )

    # ============================================================
    # 7.4 SLA — niveluri de severitate
    # ============================================================
    add_h2(doc, "7.4 SLA — Service Level Agreement și niveluri de severitate")

    add_p(doc,
        "Pentru clasificarea consecventă a incidentelor și aplicarea SLA-urilor corespunzătoare, "
        "consorțiul <LIDER> definește 4 niveluri de severitate, cu timpii de răspuns și de "
        "remediere garantați prezentați în tabelul de mai jos. Acești parametri sunt obligatorii "
        "pentru consorțiu și asumați contractual."
    )

    add_table(doc,
        header=["Severitate", "Descriere", "Timp răspuns", "Timp remediere", "Disponibilitate"],
        rows=[
            ["S1 — Critică",
             "Sistemul este indisponibil sau o funcție majoră este blocată, afectând ≥30% utilizatori; nu există workaround; impact major asupra serviciilor publice ANSVSA",
             "≤30 minute",
             "≤4 ore (workaround) / ≤24 ore (fix permanent)",
             "24×7 inclusiv weekend (over-delivery peste cerința CdS minim)"],
            ["S2 — Majoră",
             "O funcție importantă este afectată dar există workaround temporar; impact moderat",
             "≤2 ore (ore lucrătoare)",
             "≤24 ore lucrătoare",
             "L-V 08:00-17:00"],
            ["S3 — Medie",
             "O funcție secundară este afectată; impact limitat la o categorie restrânsă de utilizatori",
             "≤4 ore lucrătoare",
             "≤5 zile lucrătoare",
             "L-V 08:00-17:00"],
            ["S4 — Minoră / Cosmetică",
             "Defect fără impact funcțional (interfață, formatare, mesaj de eroare neprietenos)",
             "≤1 zi lucrătoare",
             "≤10 zile lucrătoare (livrat în următorul release)",
             "L-V 08:00-17:00"],
        ]
    )

    add_p(doc,
        "Reguli suplimentare privind SLA-ul:"
    )
    add_bullet(doc,
        "Pentru incidente S1 critice, <LIDER> mobilizează echipa de intervenție 24×7 inclusiv "
        "weekend și sărbători legale — o îmbunătățire semnificativă peste cerința minimă CdS "
        "(care prevede helpdesk 8h/zi L-V)."
    )
    add_bullet(doc,
        "Reclasificarea unui incident la o severitate diferită este posibilă în urma "
        "investigării inițiale; reclasificarea este comunicată ANSVSA în scris în maxim 1 oră "
        "de la decizie."
    )
    add_bullet(doc,
        "Timpii de răspuns se măsoară de la momentul înregistrării tichetului în Jira Service "
        "Management; timpii de remediere se măsoară de la confirmarea reproducerii problemei."
    )
    add_bullet(doc,
        "Pentru a respecta SLA-ul, dacă remedierea necesită modificări care depășesc scope-ul "
        "garanției (ex. cerințe funcționale noi), <LIDER> oferă întâi un workaround pentru "
        "S1/S2 în limitele SLA-ului, iar dezvoltarea funcționalității noi se gestionează "
        "prin Change Request separat."
    )
    add_bullet(doc,
        "KPI-uri SLA raportate trimestrial către Steering Committee al proiectului: procent "
        "incidente rezolvate în SLA, MTTR (Mean Time To Resolution) pe categorii, satisfacția "
        "utilizatorilor (NPS din chestionarele post-rezolvare)."
    )

    # ============================================================
    # 7.5 Resurse alocate
    # ============================================================
    add_h2(doc, "7.5 Resurse alocate perioadei de garanție")

    add_p(doc,
        "Pentru asigurarea calității serviciilor de garanție și respectarea SLA-urilor "
        "asumate, consorțiul <LIDER> dedică următoarele resurse umane și tehnologice "
        "exclusiv proiectului SIDISVA pe întreaga perioadă de 3 ani de garanție:"
    )

    add_h3(doc, "7.5.1 Resurse umane")
    add_table(doc,
        header=["Rol", "Număr persoane", "Responsabilități principale"],
        rows=[
            ["Coordonator suport tehnic (rol expert non-cheie nr. 16 cf. cap. 8.1 CdS)",
             "1",
             "Punct unic de contact (SPOC) cu ANSVSA pe perioada garanției; coordonare echipe L1/L2/L3; raportare KPIs SLA către Steering Committee; gestionare relația contractuală"],
            ["Experți L1 — Helpdesk first line",
             "3 (acoperire 8h/zi L-V + tură de noapte pentru S1)",
             "Preluare tichete, categorizare inițială, soluționare cereri simple (resetare parolă, ghidare utilizator), escaladare către L2"],
            ["Experți L2 — Analiză tehnică",
             "3",
             "Configurări, investigare incidente complexe, troubleshooting cross-component, escaladare la L3 pentru bug-uri în cod"],
            ["Experți L3 — Dezvoltare + mentenanță",
             "2 + acces la echipele complete ale furnizorilor",
             "Bug-fixing în cod sursă (DMS ZIPPER, Enterprise Suite, integrări custom); mentenanță evolutivă minoră; dezvoltare extensii minore"],
            ["Inginer hardware",
             "1 + parteneri logistici regionali",
             "Intervenții fizice pe echipamente HW, înlocuire piese defecte, gestionare stoc spare-parts; rețea de 5 parteneri regionali pentru intervenție rapidă"],
            ["Expert securitate (vulnerability management)",
             "1 (part-time, dedicat 50%)",
             "Monitorizare vulnerabilități, coordonare pen-test anual, raportare incidente securitate, conformitate NIS2"],
        ]
    )

    add_h3(doc, "7.5.2 Resurse tehnologice")
    add_bullet(doc,
        "Sistem ticketing dedicat — Jira Service Management cu instanță separată pentru "
        "proiectul SIDISVA, integrată cu Slack/Teams ANSVSA și cu monitoring Splunk."
    )
    add_bullet(doc,
        "Monitoring 24×7 — Prometheus + Grafana + AlertManager pentru metrici tehnice, "
        "Splunk Enterprise Security pentru log-uri, dashboards dedicate Steering Committee."
    )
    add_bullet(doc,
        "Knowledge Base — Confluence cu peste 500 articole tehnice (documentație, FAQ, "
        "troubleshooting playbooks); accesibil utilizatorilor cheie ANSVSA."
    )
    add_bullet(doc,
        "Stoc spare-parts strategic — pentru echipamente HW critice: minim 5% NGFW locații, "
        "1 NGFW centru rezervă, switch-uri agregare; depozit central <LIDER> + 2 depozite "
        "regionale pentru intervenție rapidă în provincie."
    )
    add_bullet(doc,
        "Medii dedicate garanție — păstrarea mediilor Dev/Staging operaționale pe toată "
        "perioada garanției pentru investigarea bug-urilor și validarea fix-urilor înainte "
        "de aplicare în producție."
    )
    add_bullet(doc,
        "Acces L3 producători — contracte de mentenanță active cu Microsoft, Oracle, Red "
        "Hat, Fortinet, Cisco, F5, Splunk, Dell/HP/Lenovo pe întreaga durată de garanție; "
        "tichetele L3 sunt deschise direct de echipa <LIDER>."
    )

    # ============================================================
    # 7.6 Integrarea ulterioară cu sisteme guvernamentale (cap. 3.4.2.14)
    # ============================================================
    add_h2(doc, "7.6 Integrarea ulterioară cu sisteme guvernamentale neimplementate")

    add_p(doc,
        "Conform cap. 3.4.2.14 ultim paragraf din Caietul de Sarcini: în cazul în care "
        "sistemele guvernamentale cu care SIDISVA trebuie să se integreze nu sunt pregătite "
        "pentru integrare până la finalizarea activităților de dezvoltare și integrare din "
        "contract, integrarea se realizează inițial cu mock-up-uri, iar dacă în perioada de "
        "garanție respectivele sisteme sunt pregătite pentru integrare, prestatorul realizează "
        "integrarea reală fără cost suplimentar pentru Autoritatea Contractantă."
    )

    add_p(doc,
        "Consorțiul <LIDER> își asumă integral această obligație contractuală. Sistemele "
        "guvernamentale potențial vizate (în funcție de stadiul lor de implementare la "
        "momentul finalizării contractului SIDISVA):"
    )

    add_table(doc,
        header=["Sistem guvernamental", "Cap. CdS", "Status estimat la momentul implementării"],
        rows=[
            ["PNI — Platforma Națională de Interoperabilitate",
             "3.4.2.14",
             "În curs de implementare la nivel guvernamental"],
            ["PJN — Platforma de Jurnalizare și Notificare",
             "3.4.2.14",
             "Dependent de PNI"],
            ["PCUe — Punctul de Contact Unic electronic / PDURo — Portalul Digital Unic Romania",
             "3.4.2.14",
             "În implementare conform Regulament UE 2018/1724"],
            ["Nodul eIDAS extins",
             "3.4.2.14",
             "Disponibil parțial — extinderi posibile"],
            ["RNR — Registrul Național al Registrelor",
             "3.4.3 (interoperabilitate)",
             "În implementare conform Legea 242/2022"],
        ]
    )

    add_h3(doc, "7.6.1 Procedura de integrare ulterioară în perioada de garanție")
    add_bullet(doc,
        "Notificare ANSVSA → <LIDER>: la momentul în care un sistem guvernamental devine "
        "disponibil pentru integrare, ANSVSA notifică formal <LIDER> prin scrisoare oficială."
    )
    add_bullet(doc,
        "Analiză tehnică (max 10 zile lucrătoare): <FURNIZOR ESB> + Arhitectul Sistem <LIDER> "
        "analizează specificațiile API ale sistemului țintă, identifică diferențele față de "
        "mock-up-ul implementat, estimează efortul de adaptare."
    )
    add_bullet(doc,
        "Plan de integrare (max 5 zile lucrătoare): document formal cu activități, livrabile, "
        "termene, riscuri; transmis ANSVSA pentru validare."
    )
    add_bullet(doc,
        "Dezvoltare integrare (max 30-60 zile, în funcție de complexitate): adaptarea Anti-"
        "Corruption Layer pe Oracle Service Bus, dezvoltarea conectorului real, configurare "
        "feature flag pentru comutare automată mock-real."
    )
    add_bullet(doc,
        "Testare end-to-end (max 10 zile lucrătoare): testare funcțională, integrare, "
        "performanță, securitate; UAT cu ANSVSA."
    )
    add_bullet(doc,
        "Deploy în producție: feature flag comutat de la mock la sistem real, fără downtime; "
        "monitoring 7 zile post-deploy."
    )
    add_bullet(doc,
        "Documentare: actualizare manuale tehnice, ADR-uri, matricea de trasabilitate API."
    )

    add_p(doc,
        "Întreaga procedură de integrare ulterioară este realizată fără cost suplimentar "
        "pentru ANSVSA, conform asumării explicite din cap. 3.4.2.14 al Caietului de Sarcini "
        "și conform principiului de continuitate contractuală asumat de consorțiu."
    )

    # ============================================================
    # 7.7 Vulnerability management și pen-test anual
    # ============================================================
    add_h2(doc, "7.7 Vulnerability management și teste de penetrare anuale")

    add_p(doc,
        "Conform cerinței explicite din cap. 3.4.9 al Caietului de Sarcini (punctul 13: "
        "managementul vulnerabilității, precum și teste de penetrare anuale), consorțiul "
        "<LIDER> asigură pe întreaga perioadă de garanție următoarele servicii de securitate "
        "cibernetică:"
    )

    add_h3(doc, "7.7.1 Vulnerability scanning continuu")
    add_bullet(doc,
        "Instrumente: Tenable Nessus Professional + Qualys Vulnerability Management + OWASP "
        "Dependency-Check + Snyk + Trivy pentru containere."
    )
    add_bullet(doc,
        "Scanning automat săptămânal pe toate componentele (infrastructură + aplicații + "
        "containere + librării third-party)."
    )
    add_bullet(doc,
        "Raportare lunară către ANSVSA cu evoluția numărului de vulnerabilități identificate, "
        "rezolvate, restante, clasificate pe CVSS Severity."
    )
    add_bullet(doc,
        "SLA remediere vulnerabilități: CVE Critical (CVSS 9.0+) în max 7 zile; CVE High "
        "(7.0-8.9) în max 30 zile; CVE Medium (4.0-6.9) în max 90 zile; CVE Low (sub 4.0) în "
        "max 180 zile sau în următorul release planificat."
    )

    add_h3(doc, "7.7.2 Penetration testing anual extern")
    add_bullet(doc,
        "Auditor independent — terță parte certificată ANSCC (Autoritatea Națională pentru "
        "Securitate Cibernetică), neafiliat consorțiului <LIDER>, schimbat la fiecare 2 ani "
        "pentru a evita complezența."
    )
    add_bullet(doc,
        "Scope pen-test: aplicații web (Portal, DMS web), API-uri (REST + GraphQL), "
        "infrastructură (NGFW, switch-uri, servere), aplicații mobile (iOS + Android), "
        "social engineering simulat (phishing campaign)."
    )
    add_bullet(doc,
        "Metodologie: OWASP Testing Guide v4.2 + PTES (Penetration Testing Execution "
        "Standard) + NIST SP 800-115."
    )
    add_bullet(doc,
        "Raport pen-test livrat ANSVSA în max 30 zile de la finalizare, cu findings clasate "
        "pe risc, recomandări de remediere și roadmap implementare."
    )
    add_bullet(doc,
        "Remediere findings: aceleași SLA-uri ca pentru vulnerability management; verificare "
        "remediere prin re-test focusat pe findings critic/high."
    )

    add_h3(doc, "7.7.3 Threat Intelligence și răspuns la incidente")
    add_bullet(doc,
        "Feed-uri Threat Intelligence integrate continuu cu SIEM-ul Splunk: MISP, ENISA, "
        "DNSC, plus feed-uri comerciale (Recorded Future / Mandiant)."
    )
    add_bullet(doc,
        "Notificare incidente conform Art. 23 NIS2 + OUG 155/2024: early warning DNSC în "
        "24h, notificare incident în 72h, raport final în 1 lună — procedură pregătită și "
        "testată pre Go-live."
    )
    add_bullet(doc,
        "Exerciții de simulare incidente (table-top exercises) semestrial cu echipa ANSVSA "
        "pentru îmbunătățirea răspunsului în condiții reale."
    )

    # ============================================================
    # 7.8 Servicii Help Desk și Fault Management
    # ============================================================
    add_h2(doc, "7.8 Servicii Help Desk și Fault Management")

    add_p(doc,
        "Pe lângă procedura formală de suport descrisă în secțiunea 7.3, consorțiul "
        "<LIDER> asigură suplimentar următoarele servicii pe parcursul perioadei de garanție:"
    )
    add_bullet(doc,
        "Help Desk — punct unic de contact pentru toate cererile utilizatorilor ANSVSA, "
        "indiferent de canalul ales (telefon, email, portal Jira, Teams)."
    )
    add_bullet(doc,
        "Fault Management — proces structurat de diagnosticare, soluționare și închidere a "
        "incidentelor, cu Root Cause Analysis (RCA) obligatorie pentru incidentele S1 și S2."
    )
    add_bullet(doc,
        "Suport în caz de urgență — disponibil 24×7 exclusiv pentru incidente S1 critice."
    )
    add_bullet(doc,
        "Rapoarte de deranjamente — săptămânale (sumar incidente săptămâna anterioară), "
        "lunare (analiza trend-uri, MTTR pe categorii), trimestriale (KPI SLA, raport "
        "Steering Committee)."
    )
    add_bullet(doc,
        "Corelarea erorilor — analiză periodică pentru identificarea pattern-urilor "
        "recurente; recomandări de îmbunătățire arhitecturală sau procesuală."
    )
    add_bullet(doc,
        "Bază de cunoștințe tehnice — Confluence cu acces controlat pentru utilizatori "
        "cheie ANSVSA + administratori; FAQ + troubleshooting playbooks + video tutoriale."
    )
    add_bullet(doc,
        "Documentație tehnică actualizată continuu — manuale utilizator, manuale "
        "administrator, ADR-uri, matricea de trasabilitate cerințe-cod, documentație API "
        "Swagger; versionate în Git + publicate în Confluence."
    )

    add_p(doc,
        "Toate intervențiile efectuate în sistem (la cerere ANSVSA sau în caz de incidente) "
        "vor fi documentate complet în Jira Service Management și, dacă este cazul, vor fi "
        "urmate de propuneri concrete de eliminare a cauzelor pentru prevenirea repetării. "
        "La finalizarea fiecărui serviciu de suport major, <LIDER> întocmește un raport de "
        "activitate transmis ANSVSA pentru aprobare."
    )

    # ============================================================
    # 7.9 Confidențialitate și conformitate
    # ============================================================
    add_h2(doc, "7.9 Confidențialitate, conformitate și raportare")

    add_p(doc,
        "Toate serviciile de garanție și suport tehnic furnizate de consorțiul <LIDER> "
        "respectă următoarele cerințe de confidențialitate și conformitate:"
    )
    add_bullet(doc,
        "Acord de confidențialitate — semnat de toți cei 9 membri ai consorțiului la "
        "semnarea Contractului (cf. ipoteza d din cap. 4.1 CdS); aplicabil integral pe "
        "perioada de garanție; răspundere solidară."
    )
    add_bullet(doc,
        "Conformitate GDPR — toate intervențiile pe date personale ANSVSA respectă "
        "principiile minimizării, separării mediilor, pseudo-anonimizării în mediile de "
        "test; DPO <LIDER> coordonează cu DPO ANSVSA."
    )
    add_bullet(doc,
        "Conformitate NIS2 + OUG 155/2024 — procedura de notificare incidente securitate "
        "către DNSC respectă termenele 24h/72h/1 lună; documentație pregătită pre Go-live."
    )
    add_bullet(doc,
        "Audit logging — toate acțiunile administratorilor și ale echipei de suport sunt "
        "loggate complet în Splunk SIEM cu retenție 5 ani; disponibile pentru audit ANSVSA."
    )
    add_bullet(doc,
        "Raportare trimestrială către Steering Committee al proiectului: KPI SLA "
        "(% respectare, MTTR), număr incidente pe severitate, evoluție vulnerabilități, "
        "rezultate pen-test, propuneri îmbunătățire."
    )
    add_bullet(doc,
        "Lessons Learned — sesiune anuală cu ANSVSA pentru analiza retrospectivă a "
        "perioadei de garanție; ajustarea proceselor pentru anul următor."
    )

    doc.save(str(OUT))
    size = OUT.stat().st_size
    print(f"OK | {OUT.name} salvat: {size} bytes ({size/1024:.1f} KB)")


if __name__ == "__main__":
    main()

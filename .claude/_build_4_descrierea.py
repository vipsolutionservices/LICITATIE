"""Reconstruiește 4-Descrierea_solutiei.docx de la zero.
- Listă corectă a celor 14 componente conform cap. 3.4.2.1-14 din CdS
- Placeholder-uri conform CLAUDE.md §0.5: <LIDER>, <SUITA>, <PARTENER_DMS>, <FURNIZOR_LIMS>, <FURNIZOR_GIS>, <ETL>, <BI>, <ESB>, <IAM>, <NOSQL>, <HL7_INTEGRATOR>
- Limbaj formal, fără marketing-speak
- Referințe explicite la cap. CdS
- Tabel de mapare componente → produs
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

DST = r'C:\Users\adria\Documents\Claude\Projects\Licitatie 1\doc-output\oferta\4-Descrierea_solutiei.docx'

doc = Document()

# Stiluri implicite
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def h(level, text):
    p = doc.add_heading(text, level=level)
    return p

def para(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    if bold:
        r.bold = True
    return p

def bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

def add_table(headers, rows, widths_cm=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header
    for i, txt in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = txt
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    # Data
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            cell = t.rows[ri].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    if widths_cm:
        for i, w in enumerate(widths_cm):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t

# ============================================================
# 4. Descrierea soluției propuse
# ============================================================
h(1, '4. Descrierea soluției propuse')

para(
    'Capitolul răspunde la cerința IV.4.1 lit. a) din Fișa de date a achiziției — '
    'descrierea soluției propuse pentru îndeplinirea obiectivelor stabilite prin Caietul de Sarcini. '
    'Capitolul prezintă: viziunea de ansamblu (§4.1); cele 14 componente funcționale solicitate prin '
    'cap. 3.4.2 din Caietul de Sarcini, fiecare descrisă cu referință explicită la sub-capitolul corespunzător '
    '(§4.2); maparea componentelor pe produsele software ofertate (§4.3); caracteristicile transversale ale '
    'platformei (§4.4); conformitatea legală și standardele aplicabile (§4.5); drepturile de proprietate '
    'intelectuală asupra rezultatelor (§4.6).'
)

# ----------------------------- 4.1 -----------------------------
h(2, '4.1 Viziunea de ansamblu')

para(
    'SIDISVA este un sistem informatic Cloud-Native, modular, găzduit în Cloud Guvernamental conform '
    'OUG nr. 89/2022, construit pe principiul microserviciilor containerizate (Docker / Kubernetes). '
    'Sistemul deservește:'
)
bullet('ANSVSA centru + 42 DSVSA județene + 3 institute subordonate (IISPV, ICBMV, IDSA);')
bullet('cca. 5.300 angajați interni, 2.600 medici veterinari concesionari, 4.800 utilizatori acreditați și cca. 185.000 utilizatori unici/an externi (cetățeni, fermieri, operatori economici);')
bullet('cca. 280.000 cereri de analiză de laborator/an, ≥56 servicii publice electronice (indicator de rezultat RCR11 POCIDIF);')
bullet('ecosistemul interoperabil național și european (ROeID, eIDAS, ONRC, APIA, ANCPI, ANARZ, PNI, PDURo, Colegiul Medicilor Veterinari, Ghișeul.ro).')

para(
    'Soluția este compusă din 14 componente funcționale (descrise în §4.2), peste o platformă comună '
    'de servicii partajate — autentificare și autorizare (IAM), magistrala de servicii (ESB), motorul de '
    'rapoarte și BI, GIS, semnătură electronică, audit, notificări — descrisă în §4.4. Arhitectura tehnică '
    'detaliată (componente de infrastructură, sizing, lista licențelor) este prezentată în capitolul 5 al ofertei.'
)

# ----------------------------- 4.2 -----------------------------
h(2, '4.2 Componentele funcționale ale sistemului')

para(
    'Cele 14 componente sunt prezentate în ordinea capitolelor 3.4.2.1 – 3.4.2.14 ale Caietului de Sarcini. '
    'Pentru fiecare componentă sunt indicate: capitolul CdS de referință, funcționalitățile cheie, volumele estimate '
    'și produsul / produsele software propuse pentru acoperire (cu placeholder-uri pentru numele entităților juridice — '
    'completate în varianta finală a propunerii).'
)

# 4.2.1 LIMS
h(3, '4.2.1 LIMS — Sistem informatic de laborator (cap. 3.4.2.1, 3.4.3.3.3)')
para(
    'Sistem unitar de gestiune a activității de laborator pentru cele 3 institute subordonate '
    '(IISPV, ICBMV, IDSA) și 41 laboratoare DSVSA, cu posibilitate de extindere către max. 100 laboratoare private. '
    'Volume estimate: cca. 280.000 cereri de analiză/an.'
)
para('Funcționalități principale:', bold=True)
bullet('Înregistrarea cererilor de analiză din toate canalele (Portal Servicii Publice, depunere fizică, API extern);')
bullet('Generarea automată a codurilor unice de eșantion (cod bară și QR);')
bullet('Fluxul tehnic complet: primire eșantion → distribuție pe sectoare → execuție analiză → validare → eliberare buletin de analiză;')
bullet('Integrarea echipamentelor de laborator prin standardul HL7 FHIR — cca. 126 dispozitive IoT (cap. 3.4.3.3.3.2);')
bullet('Sistem responsive pentru expedierea cererilor de laborator (cap. 3.4.3.3.3.1);')
bullet('Generarea buletinelor de analiză cu semnătură electronică calificată conform Reg. (UE) 910/2014 (eIDAS);')
bullet('Trasabilitatea completă (lanț de custodie) cu audit log inalterabil;')
bullet('Modul dedicat de gestiune a schemelor de intercomparare — vezi §4.2.3;')
bullet('Conformitate SR EN ISO/IEC 17025 (laboratoare de încercări).')
para('Produs propus: ', bold=True).add_run('<FURNIZOR_LIMS> (soluție COTS specializată, cu licențe nelimitate și cod sursă cf. cap. 12).')

# 4.2.2 DMS
h(3, '4.2.2 DMS — Sistem național de management documente (cap. 3.4.2.2, 3.4.3.3.1)')
para(
    'Sistem unitar de gestionare a documentelor pentru cele cca. 5.300 angajați ANSVSA + 46 instituții '
    '(ANSVSA centru + 42 DSVSA + 3 institute), capabil să opereze concomitent pe SQL și NoSQL (cerință DEMO). '
    'Componenta este construită pe cele 7 sub-module solicitate prin cap. 3.4.3.3.1.1 – 3.4.3.3.1.7.'
)
add_table(
    ['Sub-modul', 'Cap. CdS', 'Rol principal'],
    [
        ['Gestiune documente', '3.4.3.3.1.1', 'Captură, indexare, versionare, arhivare cf. L. 16/1996 (Arhivele Naționale)'],
        ['Captură', '3.4.3.3.1.2', 'Preluare din surse multiple, OCR pe documente scanate, indexare automată'],
        ['Registru electronic', '3.4.3.3.1.3', 'Numerotare automată cf. L. 16/1996'],
        ['Fluxuri de lucru (BPMN)', '3.4.3.3.1.4', 'Workflow-uri configurabile cu min. 5 pași și min. 3 roluri, fără modificare cod sursă'],
        ['Raportare', '3.4.3.3.1.5', 'Rapoarte standard + ad-hoc, dashboard-uri operaționale'],
        ['Administrare', '3.4.3.3.1.6', 'Utilizatori, roluri, permisiuni granulare, multi-tenant pe aceeași instalare'],
        ['Ajutor', '3.4.3.3.1.7', 'Documentație contextuală, tutoriale interactive'],
    ],
    widths_cm=[4.0, 2.5, 9.0]
)
para('Capabilități critice (cerințe DEMO eliminatorii):', bold=True)
bullet('Modificarea diagramelor BPMN / UML direct din interfața de administrare;')
bullet('Editarea documentelor Word / Excel în DMS prin integrare nativă Office (fără reatașare manuală);')
bullet('Tipărirea plicurilor și a borderourilor cu coduri de bare, cu gestionarea confirmărilor de primire;')
bullet('Formulare configurabile drag-and-drop cu asistență oferită de un chatbot AI (cap. 3.4.3.3.2.1);')
bullet('Semnătură electronică calificată eIDAS + semnătură olografă captată pe pad USB conectat la stațiile de lucru.')
para('Produs propus: ', bold=True).add_run('<PARTENER_DMS> (soluție matură, licențe nelimitate + cod sursă cf. cap. 12).')

# 4.2.3 Scheme intercomparare
h(3, '4.2.3 Management scheme de intercomparare (cap. 3.4.2.3)')
para(
    'Componenta gestionează cca. 200 scheme de intercomparare între cele 41 laboratoare DSVSA și max. 100 laboratoare '
    'private, conform SR EN ISO/IEC 17043 și cerințelor RENAR (Asociația de Acreditare din România).'
)
para('Funcționalități:', bold=True)
bullet('Calendarul anual al schemelor de intercomparare;')
bullet('Distribuirea eșantioanelor de referință către laboratoarele participante;')
bullet('Colectarea rezultatelor și calculul statistic robust (Z-score, scoruri En);')
bullet('Generarea rapoartelor de evaluare a performanței participanților;')
bullet('Acordarea / retragerea calificării laboratoarelor în baza rezultatelor.')
para('Produs propus: ', bold=True).add_run('<PARTENER_DMS> (modulul de fluxuri) + extensii dedicate dezvoltate de <LIDER>.')

# 4.2.4 GIS
h(3, '4.2.4 GIS — Sistem geo-informațional (cap. 3.4.2.4, 3.4.3.2.8)')
para(
    'Sistem geo-informațional găzduit în Cloud Guvernamental, oferind suport cartografic pentru toate componentele '
    'cu nevoie spațială. Sistemul este utilizat în special de modulele 4.2.7 (BND-SNIIA), 4.2.8 (Catagrafie) și 4.2.9 (Supraveghere).'
)
para('Funcționalități:', bold=True)
bullet('Layer-uri tematice configurabile (focare boli, exploatații zootehnice, abatoare, farmacii, depozite, fabrici de furaje);')
bullet('Geocoding adrese; calculul de rute optimizate pentru cele cca. 336 complete echipe de teren;')
bullet('Hărți tematice pentru analiza epidemiologică (hot-spots, zone de risc, propagare focare);')
bullet('Vizualizare 2D + 3D; exporturi GeoJSON, KML, Shapefile;')
bullet('Integrare cu sistemul cartografic național ANCPI — eTerra, prin INSPIRE (Directiva 2007/2/CE).')
para('Sizing țintă: 1 server GIS dedicat + 50 utilizatori editori. ', bold=True).add_run('Produs propus: <FURNIZOR_GIS>.')

# 4.2.5 BI
h(3, '4.2.5 Culegere date, raportare, dashboard și sinteze (cap. 3.4.2.5)')
para(
    'Sistem unitar de analiză și raportare avansată, care înlocuiește cele peste 50 machete .xls utilizate în prezent '
    'pentru raportările centralizate ANSVSA.'
)
para('Funcționalități:', bold=True)
bullet('Data Warehouse cu procese ETL/ELT pentru consolidare zilnică a datelor operaționale;')
bullet('Dashboard-uri operaționale și executive cu construcție drag-and-drop (cerință DEMO);')
bullet('KPI-uri configurabile, drill-down, alerte pe praguri, abonamente la rapoarte;')
bullet('Surse de date eterogene (SQL + NoSQL) — cerință DEMO;')
bullet('Exporturi în formate standard (Excel, PDF, CSV) și partajare securizată.')
para('Sizing țintă: 50 utilizatori BI. ', bold=True).add_run('Produse propuse: <ETL> (extragere / transformare) + <BI> (raportare și analiză).')

# 4.2.6 Portal
h(3, '4.2.6 Portal Servicii Publice către cetățeni, societăți comerciale și medici veterinari (cap. 3.4.2.6, 3.4.3.3.2)')
para(
    'Interfața publică unică a sistemului SIDISVA, asigurând îndeplinirea indicatorului de rezultat RCR11 al POCIDIF '
    '(≥56 servicii electronice publice noi sau optimizate).'
)
para('Funcționalități:', bold=True)
bullet('Catalog complet de servicii publice cu min. 56 servicii electronice;')
bullet('Autentificare prin ROeID + eIDAS + 2FA SMS (cerință DEMO eliminatorie);')
bullet('Spațiu personal („My ANSVSA") cu istoricul cererilor depuse, statutul și documentele asociate;')
bullet('Plăți online integrate prin Ghișeul.ro, conform L. 207/2015 (Cod fiscal procedural) și OUG 41/2016;')
bullet('Notificări multi-canal — email, SMS, push pe aplicațiile mobile;')
bullet('Aplicație mobilă nativă pentru cetățeni — iOS + Android (cap. 3.4.3.3.2.2);')
bullet('Chatbot AI pentru asistarea utilizatorului la configurarea cererilor și formularelor drag-and-drop (cap. 3.4.3.3.2.1, cerință DEMO);')
bullet('Conformitate WCAG 2.1 nivel AA (accesibilitate);')
bullet('Suport multi-limbă (română, engleză, opțional alte limbi UE).')
para('Produs propus: ', bold=True).add_run('<SUITA> (instanțe Portal + Chatbot + Aplicație mobilă cetățeni).')

# 4.2.7 BND-SNIIA
h(3, '4.2.7 BND-SNIIA — Baza Națională de Date a Sistemului Național de Identificare și Înregistrare a Animalelor (cap. 3.4.2.7)')
para(
    'Sistem de înregistrare și identificare a animalelor (bovine, ovine, caprine, suine) la nivel național, '
    'destinat proprietarilor de exploatații, medicilor veterinari oficiali, administratorilor BND și utilizatorilor externi acreditați.'
)
para('Funcționalități:', bold=True)
bullet('Registrul exploatațiilor cu coordonate geo (integrare cu GIS — §4.2.4);')
bullet('Înregistrarea evenimentelor pe animal (naștere, mișcare, sacrificare, deces, vaccinare);')
bullet('Aplicații mobile native (Android + iOS) pentru medici veterinari și fermieri, cu funcționare offline (cerință DEMO eliminatorie);')
bullet('Sincronizare bidirecțională la reconectare, cu rezolvarea automată a conflictelor;')
bullet('Integrare nativă cu APIA, ANARZ, Colegiul Medicilor Veterinari;')
bullet('Validări automate (consistență, plauzibilitate, dublură).')
para('Produs propus: ', bold=True).add_run('<SUITA> (modul SNIIA + aplicații mobile native + sincronizare offline).')

# 4.2.8 Catagrafie
h(3, '4.2.8 Catagrafie și cartografiere exploatații, unități autorizate, farmacii (cap. 3.4.2.8)')
para(
    'Componentă de inventariere și mapare geografică a tuturor unităților din sfera de competență ANSVSA: '
    'exploatații zootehnice, abatoare, fabrici de furaje, farmacii veterinare, depozite, importatori, distribuitori.'
)
para('Funcționalități:', bold=True)
bullet('Înregistrarea atributelor juridice și tehnice ale fiecărei unități;')
bullet('Geolocalizare automată din adrese (geocoding);')
bullet('Integrare automată cu APIA pentru coordonatele exploatațiilor existente;')
bullet('Actualizare periodică cu mecanisme de validare în teren prin aplicația de supraveghere (§4.2.9);')
bullet('Vizualizare integrată pe hartă (cu modulul GIS — §4.2.4).')
para('Produse propuse: ', bold=True).add_run('<FURNIZOR_GIS> (cartografie) + <SUITA> (modulul aplicativ de catagrafie).')

# 4.2.9 Supraveghere
h(3, '4.2.9 Supraveghere, prevenire, control și anchete (cap. 3.4.2.9)')
para(
    'Modulul operațional dedicat celor cca. 336 complete echipe de teren ANSVSA / DSVSA.'
)
para('Funcționalități:', bold=True)
bullet('Planificarea controalelor pe baza unui algoritm de risc (frecvență, gravitate, istoric, sezonalitate);')
bullet('Lucru offline cu sincronizare la reconectare (cerință DEMO eliminatorie);')
bullet('Înregistrarea proceselor verbale cu semnătură olografă pe pad USB + semnătură electronică calificată eIDAS;')
bullet('Atașarea de fotografii, coordonate GPS și înregistrări audio la fiecare control;')
bullet('Generarea automată a sancțiunilor, recomandărilor și măsurilor de remediere;')
bullet('Fluxuri configurabile cu min. 5 pași și min. 3 roluri (cerință DEMO).')
para('Produs propus: ', bold=True).add_run('<SUITA> (modulul Supraveghere + aplicația mobilă de teren).')

# 4.2.10 Instruire
h(3, '4.2.10 Management instruire personal laboratoare (cap. 3.4.2.10)')
para(
    'Componenta gestionează instruirea continuă a personalului celor 41 laboratoare DSVSA și 3 institute subordonate, '
    'cerință derivată din SR EN ISO/IEC 17025.'
)
para('Funcționalități:', bold=True)
bullet('Planificarea instruirilor (interne, externe, cu acreditare);')
bullet('Catalog cursuri + materiale didactice (LMS integrat);')
bullet('Înscrieri, prezență, urmărire progres individual;')
bullet('Evaluări de competență cu trasabilitate;')
bullet('Generarea automată a certificatelor de absolvire cu semnătură electronică calificată;')
bullet('Raportarea conformității cu cerințele de instruire periodică.')
para('Produs propus: ', bold=True).add_run('<SUITA> (modul Instruire / LMS).')

# 4.2.11 App mobilă cetățeni
h(3, '4.2.11 Aplicație mobilă de raportare animale sălbatice bolnave sau decedate (cap. 3.4.2.11)')
para(
    'Aplicație mobilă nativă (Android + iOS) destinată cetățenilor, pentru raportarea cazurilor de animale sălbatice '
    'bolnave sau decedate observate în teren — instrument cheie pentru detectarea timpurie a focarelor de boli.'
)
para('Funcționalități:', bold=True)
bullet('Sesizare cu GPS automat, fotografii multiple, descriere text;')
bullet('Notificarea automată a serviciilor sanitar-veterinare locale competente teritorial;')
bullet('Confirmare de preluare și statut actualizat al sesizării;')
bullet('Integrare cu modulul de supraveghere (§4.2.9) pentru trimiterea echipelor de teren;')
bullet('Funcționare offline cu sincronizare la reconectare.')
para('Produs propus: ', bold=True).add_run('<SUITA> (aplicație mobilă raportare cetățeni).')

# 4.2.12 Autorizări
h(3, '4.2.12 Autorizare, acreditare și desemnare (cap. 3.4.2.12)')
para(
    'Componentă pentru procesele de autorizare, acreditare și desemnare a unităților din domeniul sanitar-veterinar — '
    'abatoare, farmacii, fabrici de furaje, importatori, distribuitori, laboratoare private.'
)
para('Funcționalități:', bold=True)
bullet('Depunerea online a cererii prin Portal (§4.2.6);')
bullet('Evaluarea documentației și auditul tehnic în teren (cu modulul Supraveghere §4.2.9);')
bullet('Decizia administrativă și emiterea autorizației cu semnătură electronică calificată;')
bullet('Urmărirea valabilității, suspendarea, retragerea autorizației;')
bullet('Înregistrarea în registrul național publicat pe Portal.')
para('Produs propus: ', bold=True).add_run('<SUITA> (modul Autorizări).')

# 4.2.13 RCR11
h(3, '4.2.13 Contorizare folosire servicii publice (cap. 3.4.2.13, indicator POCIDIF RCR11)')
para(
    'Componentă transversală de contorizare a utilizatorilor activi (cu țintă ≥185.000 utilizatori unici / an), '
    'necesară pentru raportarea indicatorului obligatoriu RCR11 către POCIDIF.'
)
para('Funcționalități:', bold=True)
bullet('Identificarea utilizatorilor unici prin cookie + identificator de sesiune + cont autentificat (când există);')
bullet('Contorizarea interacțiunilor / serviciu electronic publicat (din cele ≥56);')
bullet('Rapoarte trimestriale conforme cu formularul POCIDIF;')
bullet('Tablouri de bord pentru monitorizarea în timp real a indicatorului;')
bullet('Conformitate GDPR (Reg. UE 2016/679) — anonimizare conform DPIA.')
para('Produse propuse: ', bold=True).add_run('<SUITA> (instrumentare front-end) + <BI> (raportare consolidată).')

# 4.2.14 Integrări externe
h(3, '4.2.14 Integrări externe și preluări de date (cap. 3.4.2.14)')
para(
    'Componenta de interoperabilitate cu ecosistemul național și european, structurată conform cerinței exprese '
    'din Caietul de Sarcini. Sistemele țintă pentru integrare sunt:'
)
bullet('ROeID + eIDAS — autentificare cetățeni RO și UE;')
bullet('PNI — Punctul Național de Interoperabilitate;')
bullet('PCUe / PDURo — puncte de contact unice electronic;')
bullet('ONRC — Registrul Comerțului;')
bullet('APIA — Agenția de Plăți și Intervenție pentru Agricultură;')
bullet('ANCPI — Agenția Națională de Cadastru și Publicitate Imobiliară (eTerra);')
bullet('ANARZ — Agenția Națională pentru Ameliorare și Reproducție în Zootehnie;')
bullet('Colegiul Medicilor Veterinari — registrul profesional;')
bullet('Ghișeul.ro — plăți online către bugetul de stat.')
para(
    'Abordarea propusă: API-uri standardizate REST documentate Swagger / OpenAPI 3.0 (cerință DEMO), '
    'cu adaptoare dedicate per sistem-țintă. Pentru sistemele care nu dispun de API publicat la momentul implementării, '
    'se realizează un mock-up funcțional acordat cu Beneficiarul, iar integrarea efectivă se face ulterior, '
    'în perioada de garanție, fără cost suplimentar — abordare expresă cerută prin Caietul de Sarcini.'
)
para('Produse propuse: ', bold=True).add_run('<ESB> (orchestrare integrări, cluster activ-pasiv) + <HL7_INTEGRATOR> (gateway clinic HL7 FHIR) + <SUITA> (adaptoare aplicative).')

# ----------------------------- 4.3 -----------------------------
h(2, '4.3 Maparea componente → produse software (sinteză)')
para(
    'Tabelul de mai jos sintetizează maparea celor 14 componente funcționale pe produsele software ofertate. '
    'Cantitățile, sizing-ul și lista exhaustivă a licențelor sunt detaliate în Lista_Software_SIDISVA și în capitolul 5 (Arhitectura).'
)
add_table(
    ['Componentă SIDISVA', 'Cap. CdS', 'Produs(e) software propus(e)'],
    [
        ['4.2.1 LIMS', '3.4.2.1 + 3.4.3.3.3', '<FURNIZOR_LIMS>'],
        ['4.2.2 DMS', '3.4.2.2 + 3.4.3.3.1', '<PARTENER_DMS>'],
        ['4.2.3 Scheme intercomparare', '3.4.2.3', '<PARTENER_DMS> + extensii <LIDER>'],
        ['4.2.4 GIS', '3.4.2.4 + 3.4.3.2.8', '<FURNIZOR_GIS>'],
        ['4.2.5 Culegere date, BI', '3.4.2.5 + 3.4.3.2.5/6', '<ETL> + <BI>'],
        ['4.2.6 Portal Servicii Publice', '3.4.2.6 + 3.4.3.3.2', '<SUITA>'],
        ['4.2.7 BND-SNIIA', '3.4.2.7', '<SUITA>'],
        ['4.2.8 Catagrafie / cartografiere', '3.4.2.8', '<FURNIZOR_GIS> + <SUITA>'],
        ['4.2.9 Supraveghere / Anchete', '3.4.2.9', '<SUITA>'],
        ['4.2.10 Instruire laboratoare', '3.4.2.10', '<SUITA>'],
        ['4.2.11 App mobilă cetățeni', '3.4.2.11', '<SUITA>'],
        ['4.2.12 Autorizări / Acreditări', '3.4.2.12', '<SUITA>'],
        ['4.2.13 Contorizare RCR11', '3.4.2.13', '<SUITA> + <BI>'],
        ['4.2.14 Integrări externe', '3.4.2.14', '<ESB> + <HL7_INTEGRATOR> + <SUITA>'],
    ],
    widths_cm=[5.5, 4.0, 6.5]
)

# ----------------------------- 4.4 -----------------------------
h(2, '4.4 Caracteristici transversale ale soluției')
para('Toate cele 14 componente partajează următoarele caracteristici, oferite de platforma comună de servicii.')

h(3, '4.4.1 Arhitectură Cloud-Native (cap. 3.4.1.1)')
bullet('Microservicii containerizate (Docker / Kubernetes);')
bullet('Găzduire în Cloud Guvernamental conform OUG 89/2022;')
bullet('Scalare orizontală automată (autoscaling) la nivel de serviciu;')
bullet('Disponibilitate țintă ≥99,5%; RPO ≤15 minute; RTO ≤4 ore.')

h(3, '4.4.2 Identitate și acces (cap. 3.4.3.2.7, 3.4.5)')
bullet('Sistem unitar de management al identității <IAM>, cu suport LDAP, OpenID Connect, OAuth 2.0, SAML 2.0;')
bullet('Single Sign-On (SSO) între toate componentele;')
bullet('Autentificare publică prin ROeID + eIDAS;')
bullet('Autentificare multi-factor (SMS / TOTP / push) pentru rolurile sensibile;')
bullet('Multi-tenant pe aceeași instalare, cu izolarea totală a datelor între domenii (cerință DEMO eliminatorie).')

h(3, '4.4.3 Semnătură electronică și audit')
bullet('Semnătură electronică calificată conform Reg. (UE) 910/2014 (eIDAS) și L. 455/2001;')
bullet('Semnătură olografă captată pe pad USB (cerință DEMO);')
bullet('Audit log inalterabil pe toate operațiunile critice (creare, modificare, ștergere, citire date sensibile);')
bullet('Conformitate L. 16/1996 (Arhivele Naționale).')

h(3, '4.4.4 Interoperabilitate (cap. 3.4.3.2.9, 3.4.2.14)')
bullet('Magistrala de servicii <ESB> pentru orchestrarea integrărilor (cluster activ-pasiv);')
bullet('HL7 FHIR pentru integrările clinice — gateway <HL7_INTEGRATOR>;')
bullet('API REST documentate Swagger / OpenAPI 3.0 (cerință DEMO);')
bullet('Conformitate SEMIC.EU + Ordinul MCID 21286/26.10.2023 + L. 242/2022 (PNI, RNR).')

h(3, '4.4.5 Lucru offline și sincronizare (cerință DEMO eliminatorie)')
bullet('Aplicațiile de teren și mobile (componentele 4.2.7, 4.2.9, 4.2.11) operează offline-first;')
bullet('Sincronizare bidirecțională la reconectare, cu rezolvarea automată a conflictelor;')
bullet('Cache local criptat (AES-256) cu auto-purjare după sincronizare.')

h(3, '4.4.6 Securitate informatică (cap. 3.4.1.2, 3.4.1.3, 3.4.6)')
para('Detaliată în capitolul 10 al ofertei; sumar:')
bullet('Defense-in-depth pe straturi (rețea / aplicație / date);')
bullet('WAF, NGFW, SIEM 24×7, NAC, sandbox email, honeypot, antivirus EDR/XDR;')
bullet('IAM cu RBAC + ABAC;')
bullet('Criptare la repaus (TDE) + în tranzit (TLS 1.3);')
bullet('Audit + monitorizare permanentă;')
bullet('Conformitate Directiva (UE) 2022/2555 (NIS2), OUG 155/2024, L. 362/2018 (Lege NIS), L. 354/2022 (protecție IT public).')

h(3, '4.4.7 Confidențialitatea datelor (cap. 3.4.7)')
para('Detaliată în capitolul 10 al ofertei; sumar:')
bullet('Conformitate GDPR — Reg. (UE) 2016/679, L. 190/2018, L. 363/2018;')
bullet('Responsabil cu protecția datelor (DPO) desemnat;')
bullet('Registrul prelucrărilor + DPIA pentru fluxurile cu risc ridicat.')

h(3, '4.4.8 Accesibilitate (cap. 3.4.2.6)')
bullet('Conformitate WCAG 2.1 nivel AA pe Portalul Servicii Publice și aplicațiile mobile;')
bullet('Multi-limbă (română + engleză minim, opțional alte limbi UE);')
bullet('Conformitate Carta drepturilor fundamentale a UE.')

h(3, '4.4.9 Căutare și indexare')
bullet('Căutare full-text prin <NOSQL> (cluster 3 noduri) — indexare în limba română;')
bullet('OCR pe documente scanate (imagini, PDF);')
bullet('Fuzzy matching și ranking semantic.')

# ----------------------------- 4.5 -----------------------------
h(2, '4.5 Conformitate legală și standarde aplicabile')
para(
    'Soluția propusă este conformă cu actele normative și standardele de mai jos. Lista exhaustivă a cerințelor '
    'individuale și modul lor de îndeplinire sunt prezentate în capitolul 8 (Conformitate cu specificațiile tehnice) '
    'și în Anexa F (Matricea de Conformitate — 1.294 cerințe).'
)
add_table(
    ['Domeniu', 'Act normativ / standard aplicat'],
    [
        ['Protecția datelor personale', 'Reg. (UE) 2016/679 (GDPR); L. 190/2018; L. 363/2018'],
        ['Securitate cibernetică', 'Dir. (UE) 2022/2555 (NIS2); OUG 155/2024; L. 362/2018 (Lege NIS)'],
        ['Software public', 'L. 354/2022 (protecția sistemelor IT publice)'],
        ['Cloud guvernamental', 'OUG 89/2022'],
        ['Interoperabilitate', 'SEMIC.EU; Ordinul MCID 21286/26.10.2023; L. 242/2022 (PNI, RNR)'],
        ['Accesibilitate', 'WCAG 2.1 AA; Carta UE Drepturi Fundamentale'],
        ['Arhivare', 'L. 16/1996 (Arhivele Naționale)'],
        ['Semnătură electronică', 'Reg. (UE) 910/2014 (eIDAS); L. 455/2001'],
        ['Laboratoare', 'SR EN ISO/IEC 17025; SR EN ISO/IEC 17043; standarde RENAR'],
        ['Plăți electronice', 'L. 207/2015 (Cod fiscal procedural); OUG 41/2016'],
        ['Identificare animale', 'Reg. (CE) 1760/2000; Reg. (UE) 2016/429; legislația națională specifică'],
        ['Trasabilitate alimentară', 'Reg. (CE) 178/2002'],
    ],
    widths_cm=[5.5, 10.0]
)

# ----------------------------- 4.6 -----------------------------
h(2, '4.6 Drepturi de proprietate intelectuală asupra rezultatelor (cap. 12)')
para(
    'Conform capitolului 12 din Caietul de Sarcini, <LIDER> confirmă explicit acceptarea următoarelor condiții '
    'privind drepturile de proprietate intelectuală:'
)
bullet('Toate dezvoltările, customizările și configurările realizate în cadrul contractului transferă dreptul de proprietate intelectuală către Beneficiar (ANSVSA), pe durată nedeterminată și pe teritoriu nelimitat;')
bullet('Beneficiarul primește licențe perpetue + codul sursă pentru componentele aplicative dezvoltate / customizate;')
bullet('Drepturile de proprietate asupra produselor preexistente (DMS, LIMS, BI, IAM, ESB, GIS, soluțiile de securitate) rămân la producătorii respectivi, conform contractelor de licență anexate;')
bullet('Documentația tehnică completă (arhitectură, instalare, operare, mentenanță, manuale de utilizator) este livrată integral Beneficiarului în format editabil.')
para(
    'Acceptarea explicită a acestor condiții este reluată în capitolul 15 al ofertei (Declarații obligatorii) și în '
    'declarațiile cuprinse în Anexa F.'
)

doc.save(DST)
print(f"[OK] Salvat: {DST}")

"""Reconstrucție 12-Plan_instruire.docx — Oferta SIDISVA
Conform Cap. 3.4.4.9 CdS — Servicii de instruire.

Corecturi față de varianta veche:
- 0× VOGO TECHNOLOGY → `<LIDER>`
- Volume CORECTE conform CdS: 100 utilizatori cheie + 44 medici vet + 3 administratori = 147 total
  (NU 1200-1800 cum scria fișierul vechi)
- Mod livrare explicit: online (cheie + medici vet) vs fizic ANSVSA (admin)
- Max 20 cursanți/sesiune (cheie) / 22 (medici vet) / nelimitat (admin)
- Termen plan: T-10z înainte de începere
- Înregistrare video live obligatorie + acces post-instruire în sistem
- Limba: română (excepție: doc producători poate fi engleză)
- Termen final: max luna 18
- Livrabile per sesiune: prezență + certificate + raport
- Livrabile finale: raport instruire, prezență, manuale
- Train-the-trainer pentru diseminare ulterioară
"""
from docx import Document
from docx.shared import Pt, Cm

doc = Document()

for section in doc.sections:
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)


def add_bullet(text):
    return doc.add_paragraph(text, style='List Bullet')

def add_para(text):
    return doc.add_paragraph(text)

def add_table(headers, rows, widths_cm=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            c = t.rows[ri].cells[ci]
            c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    if widths_cm:
        for row in t.rows:
            for ci, w in enumerate(widths_cm):
                row.cells[ci].width = Cm(w)
    return t


# ============ HEADING ============
doc.add_heading('12. Plan de instruire utilizatori', level=1)

intro = doc.add_paragraph()
intro.add_run(
    'Capitolul descrie planul de instruire a utilizatorilor sistemului SIDISVA conform '
    'cerințelor Caietului de Sarcini, '
)
r = intro.add_run('Cap. 3.4.4.9 — Servicii de instruire')
r.bold = True
intro.add_run(
    ' (pag. 153-155). <LIDER>, în calitate de Ofertant principal, asigură instruirea '
    'personalului desemnat de ANSVSA pentru folosirea/operarea sistemului informatic și '
    'pentru administrarea componentelor hardware și software.'
)

intro2 = doc.add_paragraph()
intro2.add_run('Scopul instruirii: ').bold = True
intro2.add_run(
    'asigurarea faptului că personalul ANSVSA va putea (a) utiliza autonom toate '
    'funcționalitățile SIDISVA, (b) administra independent sistemul după go-live cu '
    'asistență minimă din partea Prestatorului, și (c) disemina cunoștințele către '
    'utilizatorii viitori prin abordarea „train-the-trainer".'
)

# ============ §12.1 Categorii + volume ============
doc.add_heading('12.1 Categorii de utilizatori și volume (Cap. 3.4.4.9 CdS)', level=2)

add_para(
    'Conform tabelului din Cap. 3.4.4.9 CdS, instruirea acoperă 147 de utilizatori '
    'desemnați de ANSVSA, structurați pe 3 categorii:'
)

add_table(
    ['Tip instruire', 'Nr. utilizatori', 'Durată/sesiune', 'Max cursanți/sesiune',
     'Nr. sesiuni minim', 'Mod livrare'],
    [
        ['Utilizatori cheie (ANSVSA + DSVSA + Institute)', '100',
         'minim 2 zile × 8h/zi', '20', '5 sesiuni (100÷20)',
         'ONLINE — platformă LMS pusă la dispoziție de <LIDER>'],
        ['Medici veterinari', '44', '1 zi × 4h/zi', '22', '2 sesiuni (44÷22)',
         'ONLINE — platformă LMS pusă la dispoziție de <LIDER>'],
        ['Administratori (ANSVSA)', '3', '3 zile × 8h/zi', '— (toți cei 3 împreună)',
         '1 sesiune', 'FIZIC — la sediul ANSVSA'],
    ],
    widths_cm=[4.0, 1.5, 2.5, 2.2, 2.5, 4.5]
)

p = doc.add_paragraph()
p.add_run('TOTAL: ').bold = True
p.add_run('147 utilizatori instruiți × minimum 8 sesiuni (5+2+1).')

p = doc.add_paragraph()
p.add_run('Conformitate cu Cap. 3.5 CdS: ').bold = True
p.add_run(
    'cifra de 147 = 144 utilizatori (cheie + medici vet) + 3 administratori coincide '
    'cu indicatorul RCR11 din Cap. 3.5 („144 utilizatori instruiți + 3 administratori").'
)

# ============ §12.2 Curricula pe categorii ============
doc.add_heading('12.2 Curricula instruire pe categorii (cerință CdS — propunere tehnică)', level=2)

add_para(
    'Conform Cap. 3.4.4.9 CdS, Ofertantul prezintă în Propunerea Tehnică curricula '
    'cursurilor pentru fiecare categorie de personal. Curricula propusă de <LIDER> '
    'pentru fiecare dintre cele 3 categorii:'
)

# A — Administratori
doc.add_heading('12.2.1 Curricula administratori sistem (3 utilizatori × 3 zile × 8h = 24h)', level=3)

add_para('Conform Cap. 3.4.4.9 CdS, cursul pentru administratori cuprinde:')
add_bullet('Ziua 1 (8h) — Administrarea sistemului SIDISVA: arhitectura Cloud-Native + '
           'containerizare Docker/K8s pe Cloud Guvernamental; deployment, monitorizare, '
           'logging centralizat; management utilizatori și roluri (Keycloak Enterprise).')
add_bullet('Ziua 2 (8h) — Administrarea bazelor de date: Microsoft SQL Server Enterprise '
           '(Always On AG, backup&recovery, tuning, TDE); Elasticsearch (log management); '
           'integrare ESB (Oracle Service Bus); monitorizare performanțe + alerting.')
add_bullet('Ziua 3 (8h) — Securitate cibernetică + asistență utilizatori: SIEM, NGFW, WAF, '
           'IAM, audit logging, raportare incidente conform OUG 155/2024 (NIS2) către '
           'CSIRT-RO; gestionarea cazurilor de helpdesk L2/L3.')

p = doc.add_paragraph()
r = p.add_run(
    'Obiectiv (Cap. 3.4.4.9 CdS): „echipa de administrare a ANSVSA va fi instruită astfel '
    'încât să poată asigura funcționarea sistemului cu o asistență minimă din partea '
    'Prestatorului sau independent, începând cu perioada post-implementare".'
)
r.italic = True

# B — Utilizatori cheie
doc.add_heading('12.2.2 Curricula utilizatori cheie (100 utilizatori × 2 zile × 8h = 16h)', level=3)

add_para(
    'Conform Cap. 3.4.4.9 CdS, cursul pentru utilizatorii cheie este de tip '
    '„TRAIN-THE-TRAINER" — cei instruiți vor instrui ulterior restul utilizatorilor ANSVSA '
    'la nivel național (5.300 angajați + DSVSA + medici vet concesionari). Curricula:'
)

add_bullet('Ziua 1 dim. (4h) — Prezentarea celor 14 componente SIDISVA: '
           'LIMS, DMS, Intercomparare, GIS, BI/Dashboard, Portal, BND-SNIIA, Catagrafie, '
           'Supraveghere&anchete, Instruire personal lab, App mobilă cetățeni, Autorizare/'
           'acreditare, RCR11, Integrări externe. Fluxuri tipice pe rol.')
add_bullet('Ziua 1 după-amiază (4h) — Modul DMS (gestiune documente, captură, registru, '
           'fluxuri BPMN); Modul LIMS (cereri analize, validări rezultate, integrare echipamente).')
add_bullet('Ziua 2 dim. (4h) — Modul Portal Servicii Publice + Plăți Ghișeul.ro; '
           'autentificare ROeID + eIDAS; gestionare 56+ servicii electronice; raportare BI.')
add_bullet('Ziua 2 după-amiază (4h) — Modul cu privire la securitatea informației și '
           'protejarea datelor cu caracter personal (GDPR — Reg. UE 679/2016 + Lege 190/2018) '
           '+ instruire privind egalitatea de șanse și nediscriminarea (CERINȚĂ EXPLICITĂ '
           'CAP. 3.4.4.9 CdS).')

# C — Medici vet
doc.add_heading('12.2.3 Curricula medici veterinari (44 utilizatori × 1 zi × 4h)', level=3)

add_para(
    'Cursul pentru medici veterinari este de tip "train-the-trainer", focusat pe '
    'fluxurile operaționale specifice activității lor:'
)
add_bullet('Oră 1 — App mobilă veterinari: identificare animale (BND-SNIIA), raportare '
           'incidente, semnătură olografă digitală (pad USB).')
add_bullet('Oră 2 — Modul Supraveghere, prevenire, control, anchete: înregistrare focare '
           'boli notificabile; flux către DSVSA județean; integrare cu Portal.')
add_bullet('Oră 3 — Modul Autorizare/acreditare/desemnare unități (abatoare, farmacii, '
           'fabrici); modul Catagrafie/cartografiere unități (integrare GIS).')
add_bullet('Oră 4 — Securitate informație + GDPR + egalitate șanse și nediscriminare '
           '(CERINȚĂ EXPLICITĂ CdS).')

# ============ §12.3 Modul de livrare ============
doc.add_heading('12.3 Modul de livrare a instruirii (Cap. 3.4.4.9 CdS — paragraf binding)', level=2)

p = doc.add_paragraph()
r = p.add_run('Cerința exactă a Cap. 3.4.4.9 CdS: ')
r.bold = True
p.add_run(
    '„Instruirea utilizatorilor cheie și a medicilor veterinari se va realiza ONLINE, '
    'prestatorul fiind responsabil de punerea la dispoziție a platformei prin intermediul '
    'căreia se va realiza instruirea. Instruirea administratorilor se va face FIZIC, la '
    'sediul ANSVSA."'
)

doc.add_heading('12.3.1 Platforma online (LMS) pentru utilizatori cheie + medici vet', level=3)

add_para(
    '<LIDER> pune la dispoziție o platformă LMS (Learning Management System) dedicată, '
    'găzduită pe infrastructură separată de mediul de producție SIDISVA. Caracteristici:'
)
add_bullet('Acces web (compatibil Chrome/Edge/Firefox/Safari) + aplicație mobilă companion.')
add_bullet('Înregistrare automată a sesiunilor video live + acces ulterior la înregistrări '
           '(cerință explicită Cap. 3.4.4.9 — „înregistrarea video a sesiunilor live de '
           'curs și punerea la dispoziție pentru restul utilizatorilor în sistemul informatic").')
add_bullet('Modul evaluare cu teste online la finalul fiecărei sesiuni + certificate '
           'de participare generate automat (semnate digital).')
add_bullet('Suport pentru max 22 cursanți simultan/sesiune live (videoconferință integrată).')
add_bullet('Acces pentru cursanți la materiale pre-curs cu minimum 5 zile înainte de '
           'sesiune (PDF + video + e-learning interactiv).')
add_bullet('Forum Q&A asincron pentru întrebări post-sesiune + canal Slack/Teams dedicat '
           'pentru suport ad-hoc cu un trainer.')

doc.add_heading('12.3.2 Sesiunea fizică pentru administratori (sediul ANSVSA)', level=3)

add_para(
    'Sesiunea de 3 zile pentru cei 3 administratori se desfășoară fizic la sediul ANSVSA '
    'din București. <LIDER> asigură:'
)
add_bullet('Trainer principal cu certificare arhitectură + cloud (Expert team leader SW sau '
           'Expert administrare BD — cf. §9 din ofertă).')
add_bullet('Laptop laborator pentru fiecare cursant cu acces VPN la mediul de instruire '
           '(replica mediului de producție).')
add_bullet('Materiale tipărite + electronice (manuale administrare + instalare/config).')
add_bullet('Echipamente birotică/telecom puse la dispoziție conform Cap. 3.4.4.9 CdS.')
add_bullet('Acces la portalurile de suport ale producătorilor (Microsoft, Fortinet, Splunk, '
           'Red Hat, ZIPPER) pentru descărcare materiale de studiu.')

# ============ §12.4 Calendar sesiuni ============
doc.add_heading('12.4 Calendarul sesiunilor de instruire (termen final: luna 18 max)', level=2)

p = doc.add_paragraph()
r = p.add_run('Termen final conform Cap. 8 CdS („Instruire utilizatori — Maxim luna 18"): ')
r.bold = True
p.add_run(
    'toate sesiunile finalizate până la finalul lunii 18 din contractul de 18 luni. '
    'Instruirea se organizează doar DUPĂ ce sistemul SIDISVA este funcțional (post-test, '
    'pre-go-live) — Cap. 3.4.4.9 CdS.'
)

add_table(
    ['Etapă', 'Lună contract', 'Activitate', 'Cumulat'],
    [
        ['T-10 zile', 'Luna 15', 'Plan instruire detaliat propus de <LIDER>, agreat de ANSVSA '
         '(cerință CdS: „cu 10 zile înainte de începerea efectivă")', '—'],
        ['T-5 zile', 'Luna 15-16', 'Materiale pre-curs puse la dispoziția cursanților prin LMS '
         '(PDF + video + e-learning)', '—'],
        ['Sesiune Admin', 'Luna 16 — săptămâna 1', '1 sesiune fizică × 3 zile × 8h, 3 cursanți',
         '3 utilizatori instruiți'],
        ['Sesiuni Utilizatori cheie', 'Luna 16 săpt. 2 — Luna 17 săpt. 2', '5 sesiuni online '
         '(20 cursanți × 2 zile × 8h fiecare)', '103 utilizatori instruiți'],
        ['Sesiuni Medici vet', 'Luna 17 săpt. 3 — Luna 18 săpt. 1', '2 sesiuni online '
         '(22 cursanți × 1 zi × 4h fiecare)', '147 utilizatori instruiți'],
        ['Raport final', 'Luna 18 — săptămâna 2', 'Raport de instruire consolidat + rapoarte '
         'prezență + arhivă manuale → agreat ANSVSA', 'Etapă închisă'],
    ],
    widths_cm=[3.5, 2.7, 7.5, 3.0]
)

# ============ §12.5 Materiale + manuale ============
doc.add_heading('12.5 Materiale didactice și manuale (cerințe binding Cap. 3.4.4.9 CdS)', level=2)

add_para(
    'Conform Cap. 3.4.4.9 CdS, <LIDER> pune la dispoziția cursanților următoarele materiale '
    'și manuale. Toate sunt livrate în limba română (excepție: documentațiile tehnice ale '
    'echipamentelor și SW-ului de bază furnizate de producători pot fi în engleză).'
)

doc.add_heading('12.5.1 Materiale instruire (resurse materiale Cap. 3.4.4.9 CdS)', level=3)

add_bullet('Materiale tipărite și în format electronic privind administrarea, utilizarea '
           'soluțiilor și bune practici de implementare.')
add_bullet('Manuale de curs (LIMBA ROMÂNĂ) puse la dispoziția cursanților ÎNAINTE de data '
           'cursurilor (CERINȚĂ EXPLICITĂ CdS).')
add_bullet('Echipamente de birotică și telecomunicații pentru sesiunea fizică (admin).')
add_bullet('Mașini virtuale (mediu instruire replica producției) sau echipamente de rețea, '
           'după caz.')
add_bullet('Acces la portalurile de suport oferite de producătorii soluțiilor de securitate '
           '(Fortinet, Cisco, Splunk, Microsoft) — pentru descărcare materiale de studiu.')

doc.add_heading('12.5.2 Manuale și documentații livrate (Cap. 3.4.4.9 CdS — listă exhaustivă)', level=3)

add_bullet('Documentațiile tehnice ale echipamentelor și SW-ului de bază (de la producători) '
           '— pot fi în engleză.')
add_bullet('Manuale de instalare și configurare a echipamentelor și SW-ului de bază.')
add_bullet('Manuale de administrare a sistemului/soluției.')
add_bullet('Manuale de utilizare a componentelor sistemului/soluției — în română.')
add_bullet('Documentația funcțională a componentelor sistemului/soluției — în română.')
add_bullet('Documentația tehnică a componentelor sistemului/soluției — în română.')
add_bullet('Alte manuale/documentații stabilite în urma perioadei de analiză, dezvoltare '
           'sau urmare derulării proiectului.')
add_bullet('Înregistrarea video a sesiunilor live de curs și punerea la dispoziție în '
           'sistemul informatic — pentru restul utilizatorilor (cerință explicită CdS).')

p = doc.add_paragraph()
r = p.add_run(
    'Conform Cap. 3.4.4.9 CdS: „Întreaga documentație de utilizare și administrare a '
    'sistemului va fi livrată în format electronic odată cu produsul în sine. De asemenea, '
    'acestea vor fi incluse și în portal pentru a facilita accesul la respectivele documente."'
)
r.italic = True

# ============ §12.6 Platforma LMS ============
doc.add_heading('12.6 Platforma e-learning (LMS) pentru susținerea instruirii', level=2)

add_para(
    'Pentru a susține instruirea online (utilizatori cheie + medici vet) și pentru a '
    'asigura accesul ulterior la materiale (inclusiv pentru utilizatorii instruiți '
    '"train-the-trainer"), <LIDER> implementează o platformă LMS dedicată:'
)

add_table(
    ['Aspect', 'Soluție propusă'],
    [
        ['Produs LMS', 'Moodle LMS (open-source, US/AU) sau Open edX (open-source, US) — '
         'instalat on-premise pe infrastructură <LIDER> separată'],
        ['Capacitate', '300 utilizatori concurenți (acoperă > 2× volumul instruirii '
         'plus traineri secundari)'],
        ['Funcționalități', 'Sesiuni video live (BigBlueButton sau Jitsi integrat), forum '
         'asincron, quiz/exam, certificate semnate digital, gamification'],
        ['Acces post-instruire', 'Materialele și înregistrările video disponibile permanent '
         'în portal SIDISVA pentru toți utilizatorii (CERINȚĂ CdS)'],
        ['Conformitate', 'GDPR + WCAG 2.1 AA (accesibilitate) + audit logging acces materiale'],
        ['Limbi suportate', 'Română (primară) + engleză (pentru documentațiile producători)'],
    ],
    widths_cm=[4.0, 12.0]
)

# ============ §12.7 Resurse umane ============
doc.add_heading('12.7 Resurse umane (echipa de instruire)', level=2)

add_para(
    'Conform Cap. 8.1 CdS (rolul 15) — 2 × Expert instruire (non-cheie) ca cerință minimă. '
    '<LIDER> propune o echipă extinsă pentru a asigura calitatea și paralelizarea sesiunilor:'
)

add_table(
    ['Rol', 'Nr.', 'Responsabilități', 'Sursa (Cap. 8 CdS)'],
    [
        ['Expert instruire (lead)', '2', 'Coordonare curriculă + livrare sesiuni (cheie + '
         'medici vet) — câte 1 formator per sesiune (cerință CdS)', 'Cap. 8.1, rol 15'],
        ['Trainer specializat administrator', '1', 'Livrare sesiune fizică admin (3 zile) — '
         'Expert team leader SW sau Expert administrare BD', 'Cap. 9 ofertă, §9.3.4 / §9.3.6'],
        ['Coordonator instruire', '1', 'Punct unic de contact ANSVSA pentru planificare, '
         'comunicare, raportare', 'Resurse <LIDER>'],
        ['Designer instructional + Specialist video', '2', 'Realizare materiale e-learning + '
         'tutoriale video + înregistrare sesiuni live', 'Resurse <LIDER>'],
        ['Suport administrativ LMS', '1', 'Gestionare conturi cursanți, prezență, certificate, '
         'troubleshooting platformă LMS', 'Resurse <LIDER>'],
    ],
    widths_cm=[4.5, 1.0, 7.5, 3.5]
)

p = doc.add_paragraph()
r = p.add_run(
    'Conform Cap. 3.4.4.9 CdS: „Pentru buna desfășurare a sesiunilor de curs contractantul '
    'va asigura pentru fiecare sesiune de curs un formator." <LIDER> asigură 1 formator '
    'principal + 1 co-trainer per sesiune online (peste cerința minimă).'
)
r.italic = True

# ============ §12.8 Livrabile per sesiune ============
doc.add_heading('12.8 Livrabile per sesiune de instruire (Cap. 3.4.4.9 CdS)', level=2)

add_para('Conform Cap. 3.4.4.9 CdS, la sfârșitul fiecărei sesiuni de instruire se elaborează:')

add_bullet('PREZENȚA LA CURS — listă semnată de fiecare cursant (sesiuni fizice) sau '
           'jurnal de conectare LMS cu time-stamp (sesiuni online) + dovada participării '
           '> 80% din durata sesiunii.')
add_bullet('CERTIFICATE DE PARTICIPARE — generate pentru fiecare cursant; semnate digital '
           '(eIDAS QC); model agreat cu ANSVSA înainte de prima sesiune.')
add_bullet('RAPORT ACTIVITATE DE INSTRUIRE — realizat de instructor; include: '
           'agenda efectivă desfășurată, observații cursanți, dificultăți identificate, '
           'recomandări pentru sesiunile următoare, rezultate evaluare (quiz/exam).')

# ============ §12.9 Livrabile finale ============
doc.add_heading('12.9 Livrabile finale etapă instruire (Cap. 3.4.4.9 CdS)', level=2)

add_para('Conform Cap. 3.4.4.9 CdS, livrabilele aferente etapei de instruire (la finalul lunii 18) sunt:')

add_bullet('RAPORT DE INSTRUIRE consolidat — sinteza tuturor celor minimum 8 sesiuni; '
           'număr cursanți instruiți (țintă 147 = 100% din nominalizare); rezultate '
           'evaluări; feedback agregat; recomandări pentru extinderea instruirii la nivel '
           'național post-go-live.')
add_bullet('RAPOARTE DE PREZENȚĂ — colecția consolidată a tuturor listelor/jurnalelor LMS.')
add_bullet('MANUALE DE UTILIZARE A APLICAȚIEI — finalizate, validate de cursanți, livrate '
           'în format electronic în portalul SIDISVA + format tipărit (pentru biblioteca '
           'tehnică ANSVSA).')

# ============ §12.10 Train-the-trainer + diseminare ============
doc.add_heading('12.10 Diseminare ulterioară (train-the-trainer)', level=2)

p = doc.add_paragraph()
r = p.add_run('Strategie pentru utilizatorii instruiți contractual: ')
r.bold = True
p.add_run(
    'cei 100 de utilizatori cheie + 44 medici vet sunt instruiți EXPLICIT pe model '
    '„train-the-trainer" (Cap. 3.4.4.9 CdS), astfel încât să poată instrui ulterior, din '
    'resurse proprii ANSVSA, ceilalți ~5.300 angajați naționali + 2.600 medici vet '
    'concesionari + 4.800 utilizatori acreditați.'
)

p = doc.add_paragraph()
r = p.add_run('Resurse puse la dispoziție pentru diseminare ulterioară: ')
r.bold = True
add_bullet('Înregistrările video integrale ale sesiunilor live (acces în portal SIDISVA).')
add_bullet('Manualele de utilizare + materiale e-learning interactiv (disponibile pe LMS).')
add_bullet('Quiz/exam-uri reutilizabile pentru evaluare ulterioară a cursanților secundari.')
add_bullet('Scripturi de prezentare „train-the-trainer" — ghidul instructorului pe care '
           'cursanții instruiți pot să-l urmeze atunci când instruiesc colegii.')
add_bullet('Acces pe LMS la pachetul de materiale pentru minim 3 ani post-go-live '
           '(perioada de garanție — cf. Cap. 7 ofertă).')

p = doc.add_paragraph()
r = p.add_run(
    'NOTĂ: Diseminarea către cei ~5.300 + medici concesionari + acreditați NU intră în '
    'scope-ul contractului (Cap. 3.4.4.9 CdS limitează la 147 utilizatori instruiți direct '
    'de Prestator). Această diseminare este responsabilitatea ANSVSA, dar <LIDER> facilitează '
    'prin materialele puse la dispoziție și prin susținerea cu materiale didactice '
    'reutilizabile pe perioada garanției.'
)
r.italic = True

# ============ §12.11 Egalitate șanse + accesibilitate ============
doc.add_heading('12.11 Egalitate de șanse, nediscriminare, accesibilitate', level=2)

p = doc.add_paragraph()
r = p.add_run(
    'Cerință explicită Cap. 3.4.4.9 CdS: „instruire privind egalitatea de șanse și '
    'nediscriminarea" + „prestatorul va lua toate măsurile care să asigure accesibilitatea '
    'la egalitate, educație și formare tuturor participanților la realizarea proiectului, '
    'fără a exclude anumite persoane pe criterii de rasă, naționalitate, categorie socială, '
    'apartenență la o categorie defavorizată etc."'
)
r.italic = True

add_para('<LIDER> respectă această cerință prin:')
add_bullet('Modul dedicat „Egalitate de șanse și nediscriminare" inclus în curricula '
           'utilizatorilor cheie și a medicilor veterinari (Ziua 2 după-amiază, 1h).')
add_bullet('Platforma LMS conformă WCAG 2.1 AA (accesibilă pentru utilizatori cu deficiențe '
           'vizuale prin VoiceOver/NVDA, sub-titrare automată video, contrast 4.5:1).')
add_bullet('Materialele didactice disponibile în format accesibil (PDF/UA + audio + '
           'sub-titrare română pentru video).')
add_bullet('Programare sesiuni online la ore convenabile (între 09:00-17:00 ora României) '
           'pentru a permite participarea fără afectarea echilibrului viață-muncă.')
add_bullet('Selectarea formatorilor pe criterii profesionale, fără discriminare; <LIDER> '
           'respectă politica de egalitate de șanse internă.')

# ---- Save ----
doc.save(r'12-Plan_instruire.docx')

# Verificare
from docx import Document
d2 = Document(r'12-Plan_instruire.docx')
print(f'OK — 12-Plan_instruire.docx scris: {len(d2.paragraphs)} paragrafe, {len(d2.tables)} tabele')

texts = []
for p in d2.paragraphs:
    texts.append(p.text)
for t in d2.tables:
    for row in t.rows:
        for cell in row.cells:
            texts.append(cell.text)
full = '\n'.join(texts)

print(f'  VOGO TECHNOLOGY              : {full.count("VOGO TECHNOLOGY")}')
print(f'  <LIDER>                      : {full.count("<LIDER>")}')
print(f'  147 utilizatori              : {full.count("147")}')
print(f'  100 utilizatori cheie        : {full.count("100")}')
print(f'  44 medici                    : {full.count("44")}')
print(f'  3 administratori             : {full.count("3 administratori")}')
print(f'  ONLINE                       : {full.count("ONLINE") + full.count("online")}')
print(f'  FIZIC                        : {full.count("FIZIC") + full.count("fizic")}')
print(f'  Cap. 3.4.4.9                 : {full.count("3.4.4.9")}')
print(f'  train-the-trainer            : {full.count("train-the-trainer") + full.count("Train-the-trainer") + full.count("TRAIN-THE-TRAINER")}')
print(f'  LMS                          : {full.count("LMS")}')

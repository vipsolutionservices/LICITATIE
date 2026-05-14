"""
Refacere completa 14-DEMO_video.docx — cerinta ELIMINATORIE conform cap. 14 CdS.
Orice cerinta nedemonstrata in video = oferta neconforma.

Cele 33 cerinte tehnice citate textual din CdS (linii 3556-3588 caiet_sarcini.txt).
Mapare fiecare cerinta → modul SIDISVA + produs concret + Min:Sec in video.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt

OUT = Path(r"C:\Users\adria\Documents\Claude\Projects\Licitatie 1\doc-output\oferta\14-DEMO_video.docx")


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(18)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    return p


def add_p(doc, text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    if bold:
        r.bold = True
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    p.add_run("• " + text)
    return p


def add_table(doc, header, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    try:
        t.style = "Table Grid"
    except KeyError:
        pass
    hdr = t.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
    for ri, row in enumerate(rows, start=1):
        cells = t.rows[ri].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    return t


def main():
    doc = Document()

    # ============================================================
    # Titlu + intro CRITICAL
    # ============================================================
    add_h1(doc, "14. Plan video demonstrativ (DEMO) — cerință eliminatorie")

    add_p(doc,
        "Conform cap. 14 al Caietului de Sarcini, ofertantul are obligația de a înregistra "
        "și depune odată cu oferta un video demonstrativ care să dovedească îndeplinirea celor "
        "33 cerințe tehnice enumerate explicit în CdS. ATENȚIE: oferta care nu îndeplinește "
        "toate cele 33 cerințe va fi declarată NECONFORMĂ și eliminată din procesul de "
        "evaluare. Prezentul capitol descrie scenariul complet al video-ului demonstrativ pe "
        "care îl propune consorțiul condus de <LIDER>, cu maparea fiecărei cerințe CdS la "
        "modulul și produsul tehnologic care o demonstrează.",
        bold=False
    )

    add_p(doc,
        "Video-ul demonstrativ va fi înregistrat în calitate înaltă (rezoluție 1920×1080 "
        "Full HD, narațiune profesionistă în limba română, durată estimată 90-120 minute) și "
        "depus la registratura ANSVSA odată cu propunerea tehnică, conform cerinței de procedură. "
        "În cazul în care video-ul nu este considerat concludent, ANSVSA poate solicita o "
        "sesiune demonstrativă live la sediul Autorității Contractante, care va fi organizată "
        "de consorțiu în maxim 5 zile lucrătoare de la primirea solicitării.",
        bold=False
    )

    # ============================================================
    # 14.1 Cadrul procedural
    # ============================================================
    add_h2(doc, "14.1 Cadrul procedural al video-ului DEMO")

    add_h3(doc, "14.1.1 Format și calitate")
    add_bullet(doc, "Rezoluție: 1920×1080 Full HD (sau superioară 4K acolo unde este util)")
    add_bullet(doc, "Format video: MP4 H.264, audio AAC stereo 192 kbps")
    add_bullet(doc, "Limbă narațiune: limba română — narator profesionist")
    add_bullet(doc, "Subtitrare integrală opțională (română) pentru accesibilitate")
    add_bullet(doc, "Durată estimată: 90-120 minute, structurată pe capitole (chapters MP4) corespunzătoare celor 33 cerințe")
    add_bullet(doc, "Suport livrare: DVD-ROM dublu + stick USB criptat + acces online prin link securizat HTTPS cu autentificare ANSVSA")
    add_bullet(doc, "Index detaliat livrat ca document anexă: tabel cu cele 33 cerințe + timestamp exact în video (Min:Sec) unde este demonstrată fiecare")

    add_h3(doc, "14.1.2 Sesiunea demonstrativă live opțională (în caz de cerere ANSVSA)")
    add_p(doc,
        "Dacă video-ul depus nu este considerat concludent de comisia de evaluare, consorțiul "
        "se angajează să organizeze o sesiune demonstrativă live la sediul ANSVSA în maxim 5 "
        "zile lucrătoare de la primirea solicitării formale. Echipamentele și logistica "
        "necesare sunt detaliate în secțiunea 14.4."
    )

    add_h3(doc, "14.1.3 Mediul de testare pentru DEMO")
    add_p(doc,
        "Video-ul demonstrativ se înregistrează folosind un mediu de pre-producție al "
        "soluției SIDISVA, configurat de consorțiu în Cloud Guvernamental sau în cloud-ul "
        "intern al <LIDER>. Datele afișate sunt date de testare fictive, dar reprezentative "
        "pentru contextul ANSVSA (denumiri de unități realiste, structuri organizatorice, "
        "tipuri de documente, fluxuri de aprobare). Niciun document din video nu conține "
        "date personale reale."
    )

    # ============================================================
    # 14.2 Maparea celor 33 cerinte la module si produse
    # ============================================================
    add_h2(doc, "14.2 Maparea celor 33 cerințe DEMO la module și produse")

    add_p(doc,
        "Tabelul de mai jos cuprinde toate cele 33 cerințe tehnice citate textual din cap. "
        "14 al Caietului de Sarcini, mapate la modulul SIDISVA și produsul concret care le "
        "demonstrează. Coloana Timestamp indică intervalul aproximativ în video unde fiecare "
        "cerință este acoperită. Această mapare constituie matricea de conformitate DEMO și "
        "este referită în Anexa F a propunerii tehnice."
    )

    rows_33 = [
        ["1",
         "Înregistrarea Documentelor de Intrare — gestiune documente prin poștă, atașamente, metadate adăugate direct din platformă",
         "DMS / Registratură electronică",
         "ZIPPER DMS — modul Registratură intrare",
         "00:02 - 00:08"],
        ["2",
         "Înregistrarea documentelor e-mail + atașamente; client e-mail integrat în Registru Electronic; configurare cont e-mail personal pentru cel puțin un utilizator",
         "DMS / Integrare email",
         "ZIPPER DMS — modul Mail Integration",
         "00:08 - 00:13"],
        ["3",
         "OCR + Indexare documente (scanate sau atașate); căutare ulterioară; indexare Word/Excel/PowerPoint",
         "DMS / OCR + Full-text search",
         "ZIPPER DMS — motor OCR + Elasticsearch indexare",
         "00:13 - 00:19"],
        ["4",
         "Prezentare scurtă a tuturor modulelor platformei; integrare; un singur login pentru toate modulele; interfață în limba română (utilizatori și administratori)",
         "Toate modulele SIDISVA + IAM",
         "Single Sign-On via Keycloak Enterprise + interfață RO pentru toate modulele",
         "00:19 - 00:27"],
        ["5",
         "Autentificare integrată cu ROeID și eIDAS; 2FA (utilizator + parolă + SMS) pentru utilizatori interni",
         "Portal + IAM",
         "Keycloak Enterprise + integrare ROeID + Nodul eIDAS + 2FA TOTP/SMS",
         "00:27 - 00:33"],
        ["6",
         "Modificarea diagramei (BPMN, UML sau echivalent) a unui flux de lucru de formulare direct din interfața aplicației",
         "DMS / Workflow editor",
         "ZIPPER DMS — modul Editor BPMN drag-and-drop",
         "00:33 - 00:39"],
        ["7",
         "Creare formulare noi cu structuri ierarhice pe cel puțin 2 niveluri, drag-and-drop, cu câmpuri diverse tipuri și asistență configurare cu chatbot AI; reutilizabile, integrate în fluxuri existente",
         "DMS / Form Builder + Chatbot AI",
         "ZIPPER DMS Form Builder + Enterprise Suite (modul Chatbot AI / NLP RO)",
         "00:39 - 00:47"],
        ["8",
         "Flux de lucru cu integrare la sisteme terțe (la unul dintre pași)",
         "DMS + ESB",
         "ZIPPER DMS workflow + Oracle Service Bus + Anti-Corruption Layer",
         "00:47 - 00:52"],
        ["9",
         "Flux de lucru cu integrare la GIS; adnotare cu metadate locație (coordonate GPS generate pe baza unei adrese)",
         "DMS + GIS",
         "ZIPPER DMS + soluție GIS — serviciu geocoding adresă-coordonate",
         "00:52 - 00:58"],
        ["10",
         "Completarea formularelor în mod offline + preluarea datelor colectate pe fluxurile de lucru la restabilirea conexiunii",
         "Aplicație mobilă + DMS",
         "Enterprise Suite (modul Mobile — offline sync queue) + ZIPPER DMS",
         "00:58 - 01:04"],
        ["11",
         "Afișarea dosarelor relaționate și navigarea după cheile de relaționare între dosare (ex: CUI firmă)",
         "DMS / Dossier management",
         "ZIPPER DMS — modul Dossier cross-references",
         "01:04 - 01:08"],
        ["12",
         "Definire liste filtrate prin specificarea parametrilor (ex: dosare de tip X cu valoare în câmp Y); specifice unui utilizator; denumire sugestivă",
         "DMS / Saved searches",
         "ZIPPER DMS — modul Filtered Views per user",
         "01:08 - 01:11"],
        ["13",
         "Configurare coloane căutare/afișare pentru fiecare tip de dosar, per utilizator",
         "DMS / Personalizare UI",
         "ZIPPER DMS — modul User Preferences (column customization)",
         "01:11 - 01:14"],
        ["14",
         "Filtrări documente structurate și dosare (colecții documente structurate, metadate, atașamente, generate)",
         "DMS / Advanced search",
         "ZIPPER DMS — motor search structurat + faceted search",
         "01:14 - 01:18"],
        ["15",
         "Dispecerizare și direcționare documente către subdiviziuni/departamente; sintetizare în document nou",
         "DMS / Dispatching workflow",
         "ZIPPER DMS — modul Routing + Document Synthesis",
         "01:18 - 01:23"],
        ["16",
         "Sistem flux de lucru documente cu minim 5 pași; minim 3 roluri (director, avizator, întocmitor); aprobări succesive",
         "DMS / Workflow engine",
         "ZIPPER DMS — workflow multi-step + RBAC roles",
         "01:23 - 01:29"],
        ["17",
         "Generare documente Word pe baza datelor din formulare, direct din aplicația de management documente",
         "DMS / Template engine",
         "ZIPPER DMS — modul Word Templating",
         "01:29 - 01:33"],
        ["18",
         "Generare paralelă mai multor documente Word; configurare șabloane Word + particularizare per utilizator",
         "DMS / Mass document generation",
         "ZIPPER DMS — Batch Templating + User Templates",
         "01:33 - 01:38"],
        ["19",
         "Modificare documente Word/Excel din DMS folosind Microsoft Office; salvare directă în DMS fără reatașare",
         "DMS / Office integration",
         "ZIPPER DMS — Office WOPI integration + check-out/check-in",
         "01:38 - 01:44"],
        ["20",
         "Generare documente PDF din Word direct din aplicația DMS",
         "DMS / PDF conversion",
         "ZIPPER DMS — modul PDF Converter (PAdES compliant)",
         "01:44 - 01:47"],
        ["21",
         "Semnare electronică documente PDF folosind semnătură electronică (USB sau cloud), direct din aplicație; mecanism refuz semnare",
         "DMS / Digital signature",
         "ZIPPER DMS — modul eSign (USB token + cloud-based eIDAS + reject workflow)",
         "01:47 - 01:54"],
        ["22",
         "Semnare olografă documente prin preluare de la pad semnături conectat USB, direct din aplicație",
         "DMS + Pad olograf",
         "ZIPPER DMS — modul Holographic Sign + Wacom STU-540 / Topaz T-LBK USB pad",
         "01:54 - 02:00"],
        ["23",
         "Înregistrarea documentelor de Ieșire; gestionare metadate în registru ieșire",
         "DMS / Registratură ieșire",
         "ZIPPER DMS — modul Registratură outbox",
         "02:00 - 02:04"],
        ["24",
         "Tipărire plicuri de diverse mărimi; tipărire multiple documente de ieșire",
         "DMS / Print management",
         "ZIPPER DMS — modul Envelope/Batch Print",
         "02:04 - 02:09"],
        ["25",
         "Tipărire borderouri ieșire cu coduri de bare; gestionare confirmări primire",
         "DMS / Postal management",
         "ZIPPER DMS — modul Postal Borderou + barcode + delivery confirmation tracking",
         "02:09 - 02:14"],
        ["26",
         "Configurare calendare + planificare lucru cu dosare în calendarul utilizatorului",
         "DMS / Calendar integration",
         "ZIPPER DMS — modul Personal Calendar + Outlook/Google integration",
         "02:14 - 02:18"],
        ["27",
         "Generare rapoarte operaționale direct din DMS; minim 5 criterii filtrare",
         "DMS / Operational reports",
         "ZIPPER DMS — modul Operational Reports + Microsoft SSRS",
         "02:18 - 02:23"],
        ["28",
         "Tablouri de bord cu grafice și KPI prestabiliți (proces la alegere); prezentare pe hartă a unor indicatori",
         "BI + GIS",
         "Microsoft Power BI + Power BI Map Visuals + integrare GIS",
         "02:23 - 02:30"],
        ["29",
         "Adăugare drag-and-drop surse date pentru rapoarte; surse externe baze relaționale și NoSQL; surse predefinite + creare noi",
         "BI / Data sources",
         "Microsoft Power BI — Get Data + custom connectors SQL/Elasticsearch",
         "02:30 - 02:36"],
        ["30",
         "Creare drag-and-drop tablouri de bord cu pie chart, bar chart, map; drill-down la înregistrările sursă (ex: dosare)",
         "BI / Dashboard builder",
         "Microsoft Power BI — Designer + drill-through to source records",
         "02:36 - 02:43"],
        ["31",
         "Suport pentru limbaj de modelare consacrat (UML, BPMN sau echivalent)",
         "DMS / Process modeling",
         "ZIPPER DMS — BPMN 2.0 engine native (Camunda or similar)",
         "02:43 - 02:48"],
        ["32",
         "Documentarea API completă conform framework SWAGGER (OpenAPI) sau similar; integrare IAM; execuție teste funcționale layer Restful",
         "API + IAM + Testing",
         "OpenAPI 3.0 + Swagger UI + Keycloak token validation + Postman/Newman test execution",
         "02:48 - 02:56"],
        ["33",
         "O singură instalare a soluției DMS servește mai multor domenii logice (ex: instruire + testare acceptanță) FĂRĂ interferența datelor (multi-tenant)",
         "DMS / Multi-tenancy",
         "ZIPPER DMS — Multi-Tenant Architecture cu Tenant Isolation (DB schemas + RBAC)",
         "02:56 - 03:02"],
    ]

    add_table(doc,
        header=["#", "Cerință CdS (rezumat)", "Modul SIDISVA", "Produs / Tehnologie ofertată", "Timestamp (Min:Sec)"],
        rows=rows_33
    )

    add_p(doc,
        "TOTAL: 33 cerințe acoperite integral, toate demonstrabile în video și în sesiunea "
        "live opțională. Toate produsele/tehnologiile menționate sunt fie deja disponibile "
        "în versiunile COTS standard, fie sunt funcționalități native ale produselor ofertate. "
        "Niciuna dintre cerințele de mai sus nu necesită dezvoltare nouă pentru a fi "
        "demonstrată — toate sunt funcționalități existente.",
        bold=True
    )

    # ============================================================
    # 14.3 Structura naratoriala a video-ului
    # ============================================================
    add_h2(doc, "14.3 Structura narativă a video-ului DEMO")

    add_p(doc,
        "Video-ul DEMO este organizat în 8 secțiuni narative care urmează fluxul natural de "
        "utilizare a sistemului SIDISVA. Fiecare secțiune are durată calibrată pentru a "
        "acoperi cerințele CdS aferente fără redundanță, dar cu suficient context pentru "
        "claritatea evaluării."
    )

    add_table(doc,
        header=["Secțiune video", "Durată", "Cerințe CdS acoperite", "Modul demonstrat"],
        rows=[
            ["Introducere + autentificare", "0-3 min", "—", "Prezentare scop video + alertă durată"],
            ["Modul 1 — Autentificare & Single Sign-On", "3-8 min", "4, 5", "ROeID + eIDAS + 2FA + interfață RO"],
            ["Modul 2 — Înregistrare documente intrare", "8-13 min", "1, 2, 3", "Registratură + Email + OCR"],
            ["Modul 3 — Fluxuri și formulare", "13-30 min", "6, 7, 8, 9, 10, 15, 16, 31", "BPMN + Form Builder + Chatbot AI + Multi-channel"],
            ["Modul 4 — Dosare și relaționări", "30-45 min", "11, 12, 13, 14", "Dossier + Saved searches + Filtering"],
            ["Modul 5 — Documente Word/PDF + semnături", "45-65 min", "17, 18, 19, 20, 21, 22", "Templating + Office + PDF + Digital + Olograf"],
            ["Modul 6 — Documente de ieșire", "65-78 min", "23, 24, 25, 26", "Outbox + Plicuri + Borderou + Calendar"],
            ["Modul 7 — Rapoarte și BI", "78-95 min", "27, 28, 29, 30", "Operational Reports + Power BI + Drag-and-Drop"],
            ["Modul 8 — Aspecte tehnice", "95-110 min", "32, 33", "Swagger + Multi-tenant"],
            ["Sumar și conformitate", "110-115 min", "—", "Recap mapare 33 cerințe + Q&A"],
        ]
    )

    # ============================================================
    # 14.4 Echipament și logistică pentru sesiunea live opțională
    # ============================================================
    add_h2(doc, "14.4 Echipament și logistică pentru sesiunea live opțională la ANSVSA")

    add_p(doc,
        "Conform cap. 14 ultim paragraf al Caietului de Sarcini, ANSVSA va pune la dispoziție "
        "o sală echipată cu videoproiector și prize de curent electric pentru sesiunea live. "
        "Ofertanții asigură echipamentele de software, hardware și comunicații necesare. "
        "Pentru sesiunea live, consorțiul <LIDER> va aduce următoarele echipamente:"
    )

    add_h3(doc, "14.4.1 Echipamente aduse de consorțiu")
    add_table(doc,
        header=["Echipament", "Cantitate", "Rol în demonstrație"],
        rows=[
            ["Laptop demo principal (Intel Core Ultra 7, 32 GB RAM, SSD 1 TB NVMe)", "2 (1 + 1 backup)",
             "Rulare interfață client SIDISVA, browser, Office, demo aplicații desktop"],
            ["Smartphone Android (Google Pixel 9 / Samsung Galaxy S24)", "1",
             "Demonstrare aplicație mobilă cetățeni — lucru offline, GPS, foto raportare animale sălbatice"],
            ["Smartphone iOS (iPhone 15 Pro)", "1",
             "Demonstrare aplicație mobilă cetățeni pe platforma Apple"],
            ["Tabletă iPad / Android pentru medici veterinari", "1",
             "Demonstrare aplicație mobilă concesionari BND-SNIIA — registru exploatație offline"],
            ["Pad semnătură olografă Wacom STU-540 (USB)", "1",
             "Demonstrare cerință 22 — semnătură olografă digitalizată"],
            ["USB token semnătură electronică calificată (certSign / DigiSign)", "1",
             "Demonstrare cerință 21 — semnătură electronică PDF prin USB token"],
            ["Imprimantă mobilă Epson WorkForce / Canon", "1",
             "Demonstrare cerințe 24, 25 — tipărire plicuri + borderouri cu coduri bare"],
            ["Scanner cod bare USB", "1",
             "Demonstrare cerință 25 — scanare cod bare confirmare primire"],
            ["Router 4G/5G cu SIM dedicat (Vodafone / Orange / Digi)", "2",
             "Conectivitate internet redundantă pentru demonstrația cloud"],
            ["Switch Gigabit portabil + cabluri Ethernet", "1",
             "Conectivitate locală pentru echipamente"],
            ["Adaptor video HDMI/USB-C pentru videoproiector", "2",
             "Conectare la videoproiectorul ANSVSA"],
            ["Cablu prelungitor electric + multipriză", "1 set",
             "Asigurare alimentare echipamente"],
        ]
    )

    add_h3(doc, "14.4.2 Soluții software disponibile pentru demonstrație live")
    add_bullet(doc,
        "Acces la mediul de demo SIDISVA pre-configurat — instanță independentă în cloud-ul "
        "intern <LIDER>, populată cu date de test reprezentative"
    )
    add_bullet(doc,
        "Acces secundar prin local — instanță Docker Compose rulată pe laptop demo, ca "
        "backup în caz de probleme de conectivitate"
    )
    add_bullet(doc,
        "Toate modulele aplicative pre-instalate și testate: ZIPPER DMS, Portal Enterprise "
        "Suite, App mobilă, ChatBot AI"
    )
    add_bullet(doc,
        "Conturi demonstrative pre-create pentru roluri diverse: cetățean, medic veterinar, "
        "fermier, inspector ANSVSA, administrator sistem"
    )
    add_bullet(doc,
        "Date de test reprezentative: dosare unități, animale crotaliate, mișcări, rezultate "
        "analize, controale, autorizări"
    )

    add_h3(doc, "14.4.3 Echipa prezentă la sesiunea live")
    add_table(doc,
        header=["Rol", "Persoană / Companie", "Responsabilitate la sesiunea live"],
        rows=[
            ["Demonstrator principal", "Arhitect Sistem <LIDER>",
             "Conducerea narativă a demonstrației + navigare interfață"],
            ["Demonstrator secundar", "Expert tehnic <FURNIZOR DMS>",
             "Demonstrare module DMS + workflows + semnături + Office integration"],
            ["Demonstrator mobile", "Expert tehnic <FURNIZOR APP MOBILĂ>",
             "Demonstrare aplicații mobile iOS + Android + offline + GPS"],
            ["Demonstrator BI/GIS", "Expert tehnic <FURNIZOR BI> sau <FURNIZOR GIS>",
             "Demonstrare rapoarte Power BI + integrare GIS + drill-down"],
            ["Suport tehnic / troubleshooting", "Expert L2 <LIDER>",
             "Rezolvare probleme tehnice on-the-fly + backup laptop"],
            ["Project Manager <LIDER>",
             "Coordonare echipă + comunicare cu comisia ANSVSA + Q&A non-tehnic", "—"],
        ]
    )

    # ============================================================
    # 14.5 Lista riscuri și plan contingență
    # ============================================================
    add_h2(doc, "14.5 Riscuri DEMO și plan de contingență")

    add_p(doc,
        "Având în vedere caracterul ELIMINATORIU al cerinței DEMO, consorțiul a identificat "
        "principalele riscuri operaționale și a pregătit planuri de contingență pentru "
        "fiecare, pentru a garanta succesul demonstrației:"
    )

    add_table(doc,
        header=["Risc", "Probabilitate", "Plan de contingență"],
        rows=[
            ["Probleme de conectivitate internet la ANSVSA",
             "Medie",
             "Două routere 4G/5G redundante (operatori diferiți); instanță locală Docker Compose backup pe laptop; date de test offline complete"],
            ["Defecțiune laptop demo principal",
             "Mică",
             "Laptop secundar pre-configurat identic; sincronizare nightly snapshots"],
            ["Defecțiune pad semnătură olografă",
             "Mică",
             "Pad secundar adus în siguranță; alternative software (semnătură mouse / touchpad) ca demonstrație degradată"],
            ["Defecțiune USB token semnătură electronică",
             "Mică",
             "Token secundar adus; alternativă cloud-based eIDAS pre-configurată"],
            ["Defecțiune imprimantă mobilă",
             "Mică",
             "Imprimantă secundară adusă; alternativă PDF preview cu cod bare generat dinamic"],
            ["Întârzieri pe parcursul video-ului depus inițial → cerere live",
             "Medie",
             "Echipa pregătită cu 5 zile lucrătoare conform CdS; agenda fixă pre-aprobată; toate echipamentele și conturile gata în orice moment"],
            ["Comisia solicită demonstrarea unei cerințe specifice neacoperită explicit în video",
             "Medie",
             "Toate cele 33 cerințe sunt acoperite explicit cu timestamp; demonstrator poate naviga la oricare cerință în <2 minute"],
            ["Probleme de afișare (videoproiector ANSVSA incompatibil)",
             "Mică",
             "Multiple adaptoare HDMI/USB-C/VGA aduse; demonstrare pe ecran laptop dacă e nevoie"],
            ["Cerere ANSVSA pentru demonstrare în limba engleză",
             "Mică",
             "Subtitrare automată disponibilă; demonstratori bilingvi"],
        ]
    )

    # ============================================================
    # 14.6 Conformitate și depunere
    # ============================================================
    add_h2(doc, "14.6 Conformitate, livrare și verificare DEMO")

    add_h3(doc, "14.6.1 Conformitate cu cerințele CdS")
    add_p(doc,
        "Consorțiul <LIDER> confirmă explicit că:"
    )
    add_bullet(doc,
        "Toate cele 33 cerințe enumerate la cap. 14 al Caietului de Sarcini sunt acoperite "
        "integral în video-ul demonstrativ depus."
    )
    add_bullet(doc,
        "Fiecare cerință este demonstrată cu funcționalitate live (nu simulare statică) — "
        "interfață reală a sistemului, click-uri și acțiuni vizibile, date afișate dinamic."
    )
    add_bullet(doc,
        "Index detaliat (tabel cu Min:Sec exact pentru fiecare cerință) este livrat ca anexă "
        "la propunerea tehnică, permițând comisiei navigarea rapidă la fiecare cerință "
        "evaluată."
    )
    add_bullet(doc,
        "Toate modulele demonstrate au licență validă în Propunerea Financiară — nu se "
        "demonstrează funcționalități neofertate."
    )

    add_h3(doc, "14.6.2 Livrare video")
    add_bullet(doc,
        "Suport fizic: 2× DVD-ROM cu conținut identic + 1× stick USB criptat (AES-256) cu "
        "parolă comunicată separat prin email securizat către ANSVSA"
    )
    add_bullet(doc,
        "Acces online: link HTTPS securizat cu autentificare prin credentiale comunicate "
        "către ANSVSA, valabil pe toată perioada de evaluare a ofertelor; alternativă "
        "secundară prin transfer SFTP"
    )
    add_bullet(doc,
        "Index document: tabel detaliat cu Min:Sec pentru fiecare cerință + descriere "
        "narativă a fiecărui capitol video; livrat în format PDF semnat electronic"
    )
    add_bullet(doc,
        "Declarație de conformitate semnată de reprezentantul legal al <LIDER>: certificare "
        "că video-ul reflectă funcționalități existente, nu prototipuri sau dezvoltări "
        "viitoare"
    )

    add_h3(doc, "14.6.3 Verificare pre-depunere de către consorțiu")
    add_bullet(doc,
        "Înregistrarea video se face într-un mediu de pre-producție testat funcțional în "
        "prealabil, cu validare în mediul de Quality Assurance al consorțiului"
    )
    add_bullet(doc,
        "Review intern al video-ului de către 3 echipe independente din consorțiu (Arhitect, "
        "PM, Expert testare); orice cerință parțial demonstrată este re-înregistrată"
    )
    add_bullet(doc,
        "Cross-check cu Anexa F (matricea de conformitate, 1.294 cerințe) pentru a asigura "
        "consistența între ce e promis în propunerea tehnică și ce e demonstrat în video"
    )
    add_bullet(doc,
        "Validare finală de către Project Manager <LIDER> înainte de depunere: checklist "
        "33/33 cerințe demonstrate explicit"
    )

    # ============================================================
    # 14.7 Avantaje suplimentare (over-delivery)
    # ============================================================
    add_h2(doc, "14.7 Avantaje suplimentare oferite în video (over-delivery)")

    add_p(doc,
        "Dincolo de îndeplinirea celor 33 cerințe obligatorii, consorțiul include în video "
        "elemente de over-delivery care evidențiază maturitatea soluției și diferențiatorii "
        "concurențiali:"
    )
    add_bullet(doc,
        "Demonstrare integrare nativă ROeID + eIDAS pe scenariu cetățean real (fără mock-up)"
    )
    add_bullet(doc,
        "Demonstrare AI/ML pentru calitatea datelor BND-SNIIA — anomaly detection + fraud "
        "detection pe date de test"
    )
    add_bullet(doc,
        "Demonstrare Chatbot AI conversațional cu fine-tuning pe corpus ANSVSA simulat — "
        "nu chatbot rule-based"
    )
    add_bullet(doc,
        "Demonstrare sincronizare cross-channel — o cerere depusă pe Portal apare instant "
        "în App mobilă (push notification)"
    )
    add_bullet(doc,
        "Demonstrare arhitectura defense-in-depth — log-uri SIEM live, alertă honeypot live"
    )
    add_bullet(doc,
        "Demonstrare exporturi GDPR self-service — cetățean exportează datele proprii din "
        "Portal în format machine-readable"
    )
    add_bullet(doc,
        "Demonstrare conformitate accesibilitate WCAG 2.1 AA pe Portal — navigare cu "
        "tastatură + screen reader"
    )

    add_p(doc,
        "Aceste demonstrații suplimentare au impact direct asupra:"
    )
    add_bullet(doc,
        "Calificativului EXCEPȚIONAL la P3.1 — element 3 (abordare inovatoare)"
    )
    add_bullet(doc,
        "Conformității cu cerințele de accesibilitate WCAG impuse de cap. 3.4.2.6 CdS"
    )
    add_bullet(doc,
        "Demonstrării integrale a securității informatice + GDPR cerute de P3.1 element 4"
    )

    doc.save(str(OUT))
    size = OUT.stat().st_size
    print(f"OK | {OUT.name} salvat: {size} bytes ({size/1024:.1f} KB)")


if __name__ == "__main__":
    main()

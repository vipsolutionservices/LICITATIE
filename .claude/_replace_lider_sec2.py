"""
Pasul 1: Replace VOGO TECHNOLOGY / VOGO -> <LIDER> in 2-Abordare_metodologie.docx.
Reuses pattern from .claude/_replace_lider.py (A3, 3-Plan_implementare.docx).
"""
import re
import sys
from pathlib import Path
from docx import Document

SRC = Path(r"C:\Users\adria\Documents\Claude\Projects\Licitatie 1\doc-output\oferta\2-Abordare_metodologie.docx")

PATTERNS = [
    (re.compile(r"VOGO\s+TECHNOLOGY", re.IGNORECASE), "<LIDER>"),
    (re.compile(r"\bVOGO\s+Security\b", re.IGNORECASE), "<LIDER> Security"),
    (re.compile(r"\bVOGO\b", re.IGNORECASE), "<LIDER>"),
]


def replace_in_text(text):
    total = 0
    for pat, repl in PATTERNS:
        text, n = pat.subn(repl, text)
        total += n
    return text, total


def replace_in_paragraph(paragraph):
    full = paragraph.text
    new_text, n = replace_in_text(full)
    if n == 0:
        return 0
    runs = paragraph.runs
    if not runs:
        return 0
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""
    return n


def walk_tables(tables, total_ref):
    for t in tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    total_ref[0] += replace_in_paragraph(p)
                walk_tables(cell.tables, total_ref)


def main():
    if not SRC.exists():
        print(f"NOT FOUND: {SRC}", file=sys.stderr)
        sys.exit(1)
    doc = Document(str(SRC))
    total = [0]
    for p in doc.paragraphs:
        total[0] += replace_in_paragraph(p)
    walk_tables(doc.tables, total)
    for section in doc.sections:
        for hdr in (section.header, section.first_page_header, section.even_page_header):
            for p in hdr.paragraphs:
                total[0] += replace_in_paragraph(p)
            walk_tables(hdr.tables, total)
        for ftr in (section.footer, section.first_page_footer, section.even_page_footer):
            for p in ftr.paragraphs:
                total[0] += replace_in_paragraph(p)
            walk_tables(ftr.tables, total)
    doc.save(str(SRC))
    print(f"Replaced {total[0]} occurrences in {SRC.name}")


if __name__ == "__main__":
    main()

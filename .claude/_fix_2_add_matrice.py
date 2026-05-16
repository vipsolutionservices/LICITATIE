"""Fix-uri pe 2-Abordare_metodologie.docx (post-A2):
1. Promovează §2.9-2.13 din Normal la Heading 2 / Heading 3 (lipsesc în TOC).
2. Fix diacritice în titlurile §2.11-2.13.
3. Adaugă §2.14 — Matricea de conformitate pe Factorul de Evaluare 3.
"""
import sys, re
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT

PATH = r'C:\Users\adria\Documents\Claude\Projects\Licitatie 1\doc-output\oferta\2-Abordare_metodologie.docx'

doc = Document(PATH)

# Map de titluri exacte (text vechi → text nou cu diacritice + nivel Heading)
# Folosim text.strip() startswith pentru a fi tolerant cu spațiile.
TITLE_FIXES = {
    '2.9 Consorțiul ofertant și repartizarea responsabilităților': (2, '2.9 Consorțiul ofertant și repartizarea responsabilităților'),
    '2.9.1 Componența consorțiului': (3, '2.9.1 Componența consorțiului'),
    '2.9.2 Avantajele structurii consorțiale propuse': (3, '2.9.2 Avantajele structurii consorțiale propuse'),
    '2.9.3 Mecanismele de coordonare a consorțiului': (3, '2.9.3 Mecanismele de coordonare a consorțiului'),
    '2.10 Adaptarea metodologiei la specificul SIDISVA': (2, '2.10 Adaptarea metodologiei la specificul SIDISVA'),
    '2.10.1 Cele 14 componente SIDISVA — mapare metodă, echipă și furnizor': (3, '2.10.1 Cele 14 componente SIDISVA — mapare metodă, echipă și furnizor'),
    '2.10.2 Stack-ul tehnologic concret propus': (3, '2.10.2 Stack-ul tehnologic concret propus'),
    '2.10.3 Drumul critic al proiectului și planul de paralelizare': (3, '2.10.3 Drumul critic al proiectului și planul de paralelizare'),
    '2.10.4 Acceptarea ipotezelor din cap. 4.1 al Caietului de Sarcini': (3, '2.10.4 Acceptarea ipotezelor din cap. 4.1 al Caietului de Sarcini'),
    '2.11 Registrul de riscuri al proiectului SIDISVA': (2, '2.11 Registrul de riscuri al proiectului SIDISVA'),
    '2.11.1 Riscurile obligatorii din cap. 4.2 al Caietului de Sarcini': (3, '2.11.1 Riscurile obligatorii din cap. 4.2 al Caietului de Sarcini'),
    '2.11.2 Riscuri suplimentare identificate de consortiu': (3, '2.11.2 Riscuri suplimentare identificate de consorțiu'),
    '2.12 Asigurarea securitatii informatice si informationale': (2, '2.12 Asigurarea securității informatice și informaționale'),
    '2.12.1 Arhitectura de securitate defense-in-depth pe 7 straturi': (3, '2.12.1 Arhitectura de securitate defense-in-depth pe 7 straturi'),
    '2.12.2 Detectare si raspuns la incidente (Detect & Respond)': (3, '2.12.2 Detectare și răspuns la incidente (Detect & Respond)'),
    '2.12.3 Securitatea proceselor de dezvoltare (DevSecOps shift-left)': (3, '2.12.3 Securitatea proceselor de dezvoltare (DevSecOps shift-left)'),
    '2.12.4 Protectia datelor personale (Privacy & GDPR)': (3, '2.12.4 Protecția datelor personale (Privacy & GDPR)'),
    '2.12.5 Continuitate operationala si recuperare in caz de dezastru': (3, '2.12.5 Continuitate operațională și recuperare în caz de dezastru'),
    '2.13 Abordare inovatoare si diferentiatori competitivi': (2, '2.13 Abordare inovatoare și diferențiatori competitivi'),
    '2.13.1 Suite integrata Portal + Chatbot + App mobila (reducere risc integrare ~30%)': (3, '2.13.1 Suită integrată Portal + Chatbot + App mobilă (reducere risc integrare ~30%)'),
    '2.13.2 AI/ML pentru calitatea datelor BND-SNIIA si fraud detection': (3, '2.13.2 AI/ML pentru calitatea datelor BND-SNIIA și fraud detection'),
    '2.13.3 Strategia de rezilienta a integrarilor - mock-up + comutare': (3, '2.13.3 Strategia de reziliență a integrărilor — mock-up + comutare'),
    '2.13.4 Chatbot AI conversational in limba romana - NLP nativ': (3, '2.13.4 Chatbot AI conversațional în limba română — NLP nativ'),
    '2.13.5 Telemetrie automata pentru indicator RCR11 - automatizare audit POCIDIF': (3, '2.13.5 Telemetrie automată pentru indicator RCR11 — automatizare audit POCIDIF'),
}

# Step 1: Promote paragraphs to Heading style + fix titles
promoted = 0
for p in doc.paragraphs:
    txt = p.text.strip()
    if txt in TITLE_FIXES:
        level, new_text = TITLE_FIXES[txt]
        # Set style
        p.style = doc.styles[f'Heading {level}']
        # Replace text — clear all runs, then add one with new text
        for r in list(p.runs):
            r.text = ''
        # Re-add a single run with new text
        if p.runs:
            p.runs[0].text = new_text
        else:
            p.add_run(new_text)
        promoted += 1

print(f"Promovate la Heading + diacritice fix: {promoted} titluri")

# Step 2: Adaugă §2.14 la finalul documentului
# Adaug paragrafe, tabele, etc.

def h(level, text):
    return doc.add_heading(text, level=level)

def para(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    if bold:
        r.bold = True
    return p

def add_table(headers, rows, widths_cm=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    # Folosesc stiluri tabelar disponibile în document (fallback la default dacă nu există)
    for style_name in ('Light Grid Accent 1', 'Table Grid', 'Tabel grilă'):
        try:
            t.style = style_name
            break
        except KeyError:
            continue
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, txt in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = txt
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            cell = t.rows[ri].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    if widths_cm:
        for i, w in enumerate(widths_cm):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t

h(2, '2.14 Matricea de conformitate pe Factorul de Evaluare 3 (Metodologie)')

para(
    'Capitolul de față răspunde Factorului de Evaluare 3 — „Metodologia de implementare a contractului" — '
    'cu cele două subfactore (3.1 — abordare și metodologie corespunzătoare, 10 puncte; 3.2 — corelația dintre '
    'datele de intrare și rezultatele intermediare și finale, 10 puncte). Țintă: calificativul „Excepțional" la '
    'ambele subfactori, pentru un total de 20 puncte. Matricele de mai jos pun în corespondență fiecare element '
    'din algoritmul de calcul publicat în Fișa de date a achiziției cu modul concret în care este îndeplinit în '
    'prezenta Propunere Tehnică și cu referința exactă la secțiunea care îl tratează.'
)

h(3, '2.14.1 Subfactorul 3.1 — Țintă: Excepțional (10 puncte)')

para(
    'Algoritmul de calcul prevede calificativul „Excepțional" (10 puncte) atunci când abordarea și metodologia '
    'îndeplinesc toate cele patru elemente de mai jos.'
)

add_table(
    ['Element din algoritmul „Excepțional" (Fișa de date — subfactor 3.1)',
     'Modul concret de îndeplinire în Propunerea Tehnică',
     'Referință în Propunerea Tehnică'],
    [
        [
            '1. Abordarea propusă și metodologia prezentată pentru realizarea activităților din Caietul de Sarcini este prezentată în detaliu și se bazează în mare măsură pe o serie de metodologii, metode și/sau instrumente testate, recunoscute.',
            'Cadru metodologic explicit pe patru piloni, fiecare cu standardele aplicate și instrumentele concrete: PMI / PMBoK 7 + ISO 21500 pentru management de proiect (§2.2.1); BABOK v3 (IIBA) + UML / BPMN 2.0 pentru analiză (§2.2.2); Scrum + Kanban + practici SAFe pentru dezvoltare hibridă (§2.2.3); ISO 9001 calitate, ISO/IEC 27001 + 27017 + 27018 securitate, ISO/IEC 25010 calitate software, ITIL 4 pentru servicii IT, TOGAF 9.2 pentru arhitectura enterprise (§2.2.4). Procedurile obligatorii din cap. 10.2 CdS (livrare, recepție, ședințe, control livrări, testare, asistență) sunt detaliate în cap. 13 (§13.2.1).',
            '§2.2 (Cadru metodologic) cu sub-secțiunile 2.2.1 – 2.2.4; §13.2.1 (cele 6 proceduri obligatorii).'
        ],
        [
            '2. Abordarea propusă și metodologia prezentată sunt adaptate la specificul Contractului, în corelație cu activitățile aflate pe drumul critic, precum și cu riscurile și ipotezele identificate.',
            'Adaptare detaliată la SIDISVA pe 4 axe: (a) maparea celor 14 componente (cap. 3.4.2.1 – 3.4.2.14 CdS) la metodă, echipă și furnizor (§2.10.1); (b) stack-ul tehnologic concret aliniat cu Lista_Software_SIDISVA (§2.10.2 + cap. 5); (c) drumul critic identificat cu activitățile prioritare (LIMS / DMS / Portal / BND-SNIIA / integrări guvernamentale) și planul de paralelizare a fluxurilor independente (§2.10.3 + cap. 3); (d) acceptarea explicită a celor 4 ipoteze din cap. 4.1 CdS și a celor 7 riscuri obligatorii din cap. 4.2 CdS, plus 6 riscuri suplimentare identificate de consorțiu (§2.10.4 + §2.11.1 + §2.11.2).',
            '§2.10.1 – §2.10.4 (Adaptare SIDISVA); §2.11.1 – §2.11.2 (Registru de riscuri); cap. 3 (Plan de implementare cu drum critic); cap. 5 (Arhitectură + Licențe).'
        ],
        [
            '3. Aspectele importante sunt abordate într-un mod inovator și eficient: Propunerea Tehnică detaliază la capitolele Abordarea propusă și Metodologia propusă modalități de îmbunătățire a rezultatelor prin utilizarea de modalități efective de realizare a activității profesionale.',
            'Cinci elemente concrete de inovație și eficiență, cu impact cuantificabil: (a) Suita integrată Portal + Chatbot + Aplicație mobilă pe aceeași platformă <SUITA> — reducerea riscului de integrare cu cca. 30% și a duratei de implementare a Portalului (§2.13.1); (b) AI / ML aplicat la calitatea datelor BND-SNIIA și la detectarea fraudelor / inconsistențelor (§2.13.2); (c) Strategia de reziliență a integrărilor cu mock-up + comutare la API real, aplicabilă sistemelor guvernamentale neimplementate la momentul contractului — conform cap. 3.4.2.14 CdS (§2.13.3); (d) Chatbot AI conversațional în limba română cu NLP nativ adaptat domeniului veterinar (§2.13.4); (e) Telemetrie automată pentru raportarea indicatorului RCR11 către POCIDIF — eliminarea raportării manuale (§2.13.5).',
            '§2.13.1 – §2.13.5 (Abordare inovatoare și diferențiatori competitivi).'
        ],
        [
            '4. Abordarea propusă demonstrează modul în care Ofertantul va desfășura activitățile de asigurare a securității informatice și informaționale, menite să asigure un nivel adecvat de securitate.',
            'Demonstrare concretă pe cinci dimensiuni complementare: (a) arhitectură defense-in-depth pe 7 straturi (perimetru, rețea, aplicație, date, identitate, endpoint, fizic / procese) — §2.12.1; (b) detectare și răspuns la incidente prin SIEM 24×7 cu retenție IOC ≥3 ani, NDR și EDR / XDR (§2.12.2); (c) DevSecOps shift-left cu SAST / DAST / SCA / container scanning / secret scanning / threat modeling pe pipeline-ul CI/CD (§2.12.3); (d) protecția datelor personale conform GDPR, cu DPO desemnat, registru prelucrări, DPIA și mecanisme de anonimizare (§2.12.4); (e) continuitate operațională și recuperare în caz de dezastru cu BCP / DRP, RPO ≤15 min, RTO ≤4 ore și exerciții anuale de test (§2.12.5). Detalierea integrală e în cap. 10 al ofertei.',
            '§2.12.1 – §2.12.5 (Securitate informatică și informațională); cap. 10 (Securitate — detaliere completă).'
        ],
    ],
    widths_cm=[5.5, 7.5, 3.5]
)

para(
    'Toate cele patru elemente din algoritmul „Excepțional" sunt acoperite explicit și concret în prezenta '
    'Propunere Tehnică. Punctajul țintă pentru subfactorul 3.1: ', bold=True
).add_run('10 puncte (calificativ Excepțional).')

h(3, '2.14.2 Subfactorul 3.2 — Țintă: Excepțional (10 puncte)')

para(
    'Algoritmul de calcul prevede calificativul „Excepțional" (10 puncte) atunci când datele de intrare '
    '(resursele materiale și umane) sunt pe deplin corelate cu modalitatea efectivă de realizare a activităților, '
    'iar caracteristicile lor demonstrează obținerea rezultatelor la nivelul de calitate descris în Propunerea Tehnică.'
)

add_table(
    ['Categorie resurse (date de intrare)',
     'Caracteristicile / descrierea resurselor',
     'Modul în care permit obținerea rezultatelor la calitatea declarată'],
    [
        [
            'Resurse umane — experți cheie (8) și non-cheie (12)',
            'Cei 6 experți cheie punctați la P2 (Manager proiect; Analist business; Arhitect sistem; Team leader software; Expert analiză și optimizare procese; Expert administrare BD) prezintă fiecare câte 5+ proiecte de referință în sisteme cu min. 7 module interconectate + portal. Cei 2 experți cheie nepunctați (Expert comunicații și securitate; Expert testare) și cei 12 experți non-cheie acoperă rolurile complementare (4 dezvoltare; portal; BI; integrare; migrare; GDPR; 2 instruire; coordonator suport). Certificările cerute prin CdS sunt prezentate explicit: PMP / PRINCE2; business analysis (CBAP / PMI-PBA); arhitectura enterprise + certificare de la producătorul DMS ofertat; Scrum / Agile + cloud; Lean Six Sigma; HA / securitate BD ofertată; arhitectură securitate + firewall ofertat; testare.',
            'Combinația experiență (proiecte similare cu volum echivalent) + certificări de la producătorii soluțiilor concrete asigură capacitatea de a livra fiecare componentă SIDISVA la nivelul tehnic și funcțional declarat în §2.10.1 și în cap. 4 al ofertei. Disponibilitatea pe toată durata contractului este garantată prin angajamentele scrise din cap. 9.7. Mecanismele de înlocuire (cap. 9.8) asigură continuitatea în caz de indisponibilitate.'
        ],
        [
            'Resurse materiale — produse software (27 produse pe 4 categorii A–D)',
            'Lista exhaustivă în Lista_Software_SIDISVA.xlsx (sheet „Sinteza") + cap. 5: A. Infrastructură (11 produse — RHEL/Oracle Linux 9, Win Server 2022 DC, NGINX Plus, MS IIS, MS SQL Server Enterprise, Elasticsearch, MS SSIS, Power BI / SSRS / SSAS, Keycloak Enterprise, GIS, Oracle Service Bus); B. Aplicativ (6 produse — ZIPPER DMS, VOGO Enterprise Suite pentru Portal + Chatbot + App mobilă, FURNIZOR LIMS COTS, Mirth Connect); C. Securitate (7 produse — F5 / Imperva / FortiWeb, FortiDeceptor, Cisco DNA / ClearPass, Splunk ES / QRadar, Cisco IronPort, FortiGate / Palo Alto pentru centru, FortiGate 100F pentru 90 locații); D. Productivity (3 — MS Office H&B 2024 OEM ×100 + ×336, CrowdStrike / SentinelOne antivirus ×486).',
            'Fiecare produs este aliniat cu cerințele tehnice specifice din capitolele 3.4.3.x ale Caietului de Sarcini (vezi anexa F — matricea de conformitate pe 1.294 cerințe). Sizing-ul indicat în Sinteza acoperă volumele cerute: 280.000 cereri analiză/an, 185.000 utilizatori unici/an, ≥56 servicii electronice, 5.300 angajați interni.'
        ],
        [
            'Resurse materiale — echipamente hardware',
            '100 laptopuri pentru ANSVSA centru (consum în veghe ≤20 Wh, EU Ecolabel — DNSH P4.1); 336 complete teren pentru cele 42 DSVSA (8 / județ) incluzând terminal + pad semnătură + imprimantă; 126 dispozitive IoT pentru integrarea echipamentelor de laborator (3 / institut și DSVSA cu activitate de laborator); echipamente de securitate la centrul DC (NGFW, switching, WAF) + în cele 42 locații (FortiGate 100F redundant, switch acces, switch POE, access point) — vezi cap. 6.',
            'Hardware-ul este selectat astfel încât să fie atât (a) conform cu specificațiile tehnice ale CdS (cap. 3.4.3.4), cât și (b) capabil să susțină arhitectura Cloud-Native cu autoscaling, disponibilitate ≥99,5% și operare offline în teren cu sincronizare. Configurarea finală e detaliată în cap. 5 (Arhitectură) și cap. 6 (Lista hardware).'
        ],
        [
            'Resurse procedurale și metodologice',
            'Cele 6 proceduri obligatorii (cap. 10.2 CdS): asistență tehnică / mentenanță / suport; livrare; recepție / acceptanță; ședințe; control livrări; testare. Planul de Asigurare a Calității (PAC) — predat în max. 1 săptămână de la kick-off. Procedurile Standard de Operare (SOP) pentru toate aplicațiile livrate. Mecanisme de monitorizare (KPI Calitate + Termene, scala 1–5) și raportare structurată (Inițial / trimestriale / final).',
            'Procedurile asigură guvernarea și controlul calității pe întreaga durată a contractului. Rapoartele și PV-urile generate atestă mecanic îndeplinirea condițiilor de la cap. 10.6 CdS pentru finalizarea serviciilor. Mecanismul KPI furnizează probe obiective pentru documentul constatator final.'
        ],
        [
            'Resurse de calendar și paralelizare',
            'Plan de implementare pe 18 luni (cap. 6 CdS), structurat pe activitățile cap. 3.4.4 CdS (management proiect, livrare/instalare/configurare HW+SW, analiză, proiectare, dezvoltare/configurare/integrare, migrare, implementare, testare, instruire, punere în producție). Drumul critic identificat și paralelizat acolo unde dependențele permit (LIMS independent de Portal; modulele de teren independente de modulele centru — vezi §2.10.3 + cap. 3).',
            'Calendarul respectă termenul de 18 luni pentru implementare + 3 ani garanție. Etapele cheie sunt aliniate cu jaloanele cerute de POCIDIF și cu termenele intermediare cerute prin CdS. Paralelizarea fluxurilor independente reduce riscul de întârziere a drumului critic.'
        ],
    ],
    widths_cm=[3.5, 6.5, 6.5]
)

para(
    'Datele de intrare (resursele materiale și umane) sunt pe deplin corelate cu modalitatea efectivă de realizare '
    'a activităților, iar caracteristicile lor — certificări, sizing, conformitate cu CdS — demonstrează obținerea '
    'rezultatelor la nivelul de calitate descris în Propunerea Tehnică. Punctajul țintă pentru subfactorul 3.2: ', bold=True
).add_run('10 puncte (calificativ Excepțional).')

h(3, '2.14.3 Concluzie — punctaj țintă pe Factorul de Evaluare 3')

para(
    'Pe baza matricelor de mai sus, <LIDER> țintește calificativul „Excepțional" și punctajul maxim de '
    '10 puncte pe fiecare dintre cei doi subfactori ai Factorului 3.'
)

add_table(
    ['Subfactor', 'Pondere maximă', 'Calificativ țintă', 'Punctaj țintă'],
    [
        ['3.1 — Abordare și metodologie corespunzătoare', '10 puncte', 'Excepțional', '10 puncte'],
        ['3.2 — Corelație date de intrare ↔ rezultate intermediare și finale', '10 puncte', 'Excepțional', '10 puncte'],
        ['TOTAL Factor 3 — Metodologie', '20 puncte', '—', '20 puncte'],
    ],
    widths_cm=[7.0, 3.0, 3.0, 3.0]
)

para(
    'Restul scorului tehnic (40 puncte, dincolo de Factorul 3) este asigurat prin: Factorul 2 — Experiența celor 6 experți '
    'cheie punctați (țintă 30 puncte, cap. 9 al ofertei) și Factorul 4 — Măsuri DNSH (țintă 10 puncte, cap. 11 al ofertei). '
    'Cumulat cu Factorul 1 (preț), ținta agregată pe ansamblul ofertei este de 100 puncte din 100 maxim.'
)

doc.save(PATH)
print(f"[OK] Salvat (in-place): {PATH}")

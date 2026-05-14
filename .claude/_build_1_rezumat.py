"""Reconstruiește 1-Rezumat_executiv.docx de la zero, conform CdS și CLAUDE.md.

Modificări față de varianta veche:
- Elimină marketing-speak („soluție de top", „nivel de stat", „dezvoltată complet în România de informaticieni români")
- Elimină DORA (Reg. UE 2022/2554) — nu se aplică SIDISVA (sector financiar)
- Adaugă NIS2 (Dir. UE 2022/2555 + OUG 155/2024) și L 354/2022 — cerințe critice
- Adaugă cap. 12 (Drepturi IP) — acceptare explicită
- Adaugă DEMO video (cerință eliminatorie, 33 cerințe) — mențiune cap. 14
- Adaugă scor țintă pe factorii de evaluare (P1=40 + P2=30 + P3=20 + P4=10 = 100)
- Aliniază §1.4 cu Sinteza (27 produse / 4 categorii A-D)
- Placeholder-uri stricte pentru companii (<LIDER>, <PARTENER>); brand-urile produselor sunt
  păstrate (VOGO Enterprise Suite, ZIPPER DMS) — aliniat cu cap. 5/8
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT

DST = r'C:\Users\adria\Documents\Claude\Projects\Licitatie 1\doc-output\oferta\1-Rezumat_executiv.docx'

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def h(level, text):
    return doc.add_heading(text, level=level)

def para(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    if bold:
        r.bold = True
    return p

def bullet(text):
    return doc.add_paragraph(text, style='List Bullet')

def add_table(headers, rows, widths_cm=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, txt in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = txt
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            cell = t.rows[ri].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    if widths_cm:
        for i, w in enumerate(widths_cm):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t

# ============================================================
# 1. Rezumat executiv
# ============================================================
h(1, '1. Rezumat executiv')

para(
    'Capitolul prezintă, în sinteză, identificarea procedurii și a ofertantului (§1.1), '
    'înțelegerea obiectivelor proiectului SIDISVA (§1.2), soluția tehnică propusă (§1.3), '
    'conformitatea cu cadrul legislativ și standardele aplicabile (§1.4), confirmarea acoperirii '
    'cerințelor eliminatorii (§1.5), punctajul țintă pe factorii de evaluare (§1.6) și '
    'angajamentele explicite ale ofertantului (§1.7). Detaliile complete se regăsesc în capitolele '
    'următoare ale prezentei oferte tehnice.'
)

# ----------------------------- 1.1 -----------------------------
h(2, '1.1 Identificarea procedurii și a ofertantului')

para(
    'Prezenta ofertă tehnică este depusă de <LIDER>, în asociere cu <PARTENER>, în cadrul '
    'procedurii de atribuire având ca obiect „Achiziție de servicii pentru implementarea Sistemului '
    'Informatic Digitalizat în domeniul Sanitar Veterinar și pentru Siguranța Alimentelor — SIDISVA", '
    'organizată de Autoritatea Națională Sanitară Veterinară și pentru Siguranța Alimentelor (ANSVSA).'
)

para('Date de identificare procedură:', bold=True)
bullet('Anunț de participare publicat în SEAP: CN1089237;')
bullet('Caiet de Sarcini nr. 7574/CP/2025, Revizia 1 din 30.12.2025, 252 pagini;')
bullet('Cod SMIS al proiectului: 336342;')
bullet('Finanțare: Programul Creștere Inteligentă, Digitalizare și Instrumente Financiare 2021–2027 (POCIDIF), Prioritatea P2 — Digitalizare în administrația publică centrală și mediul de afaceri, Acțiunea 2.2 — E-guvernarea și digitalizarea în beneficiul cetățenilor;')
bullet('Sursa de finanțare: Fondul European de Dezvoltare Regională (FEDR) + cofinanțare din bugetul de stat;')
bullet('Termen de depunere a ofertelor: conform Fișei de date a achiziției anexate procedurii SEAP.')

# ----------------------------- 1.2 -----------------------------
h(2, '1.2 Înțelegerea obiectivelor proiectului SIDISVA')

para(
    'Obiectivul general al proiectului SIDISVA este digitalizarea completă a fluxurilor sanitar-veterinare '
    'și de siguranță alimentară din România, prin implementarea unui sistem informatic integrat care deservește:'
)
bullet('cca. 5.300 angajați ai ANSVSA centru + 42 DSVSA județene + 3 institute subordonate (IISPV — Igienă; ICBMV — Biologice și Medicamente Veterinare; IDSA — Diagnostic);')
bullet('cca. 2.600 medici veterinari concesionari și 4.800 utilizatori externi acreditați;')
bullet('cca. 185.000 utilizatori unici anual (cetățeni, fermieri, operatori economici) — țintă pentru indicatorul de rezultat RCR11 al POCIDIF;')
bullet('cca. 280.000 cereri de analiză de laborator anual.')

para('Beneficiile concrete urmărite de Autoritatea Contractantă sunt structurate pe trei axe:', bold=True)
bullet('Servicii publice digitale către cetățeni, medici veterinari și operatori economici, prin Portal Servicii Publice și aplicații mobile native — minimum 56 servicii electronice publice noi sau optimizate (indicator de rezultat obligatoriu RCR11 POCIDIF);')
bullet('Optimizarea proceselor interne și a fluxurilor de lucru la nivelul celor 46 instituții publice deservite — DMS unic național, BPMN engine cu workflow-uri configurabile, audit log inalterabil;')
bullet('Integrarea nativă în Cloud Guvernamental și interoperabilitatea cu ecosistemul național și european de referință — ROeID, eIDAS, PNI, PCUe/PDURo, ONRC, APIA, ANCPI, ANARZ, Colegiul Medicilor Veterinari, Ghișeul.ro.')

# ----------------------------- 1.3 -----------------------------
h(2, '1.3 Soluția tehnică propusă (sinteză)')

para(
    'Soluția propusă este un sistem informatic Cloud-Native, modular, găzduit în Cloud Guvernamental '
    'conform OUG nr. 89/2022, construit pe principiul microserviciilor containerizate (Docker / Kubernetes), '
    'cu disponibilitate țintă ≥99,5%, RPO ≤15 minute și RTO ≤4 ore.'
)
para('Componenta funcțională — 14 componente conform cap. 3.4.2.1 – 3.4.2.14 din Caietul de Sarcini (descrise în capitolul 4 al ofertei):', bold=True)
bullet('LIMS — Sistem informatic de laborator (cap. 3.4.2.1);')
bullet('DMS — Sistem național de management documente (cap. 3.4.2.2);')
bullet('Management scheme de intercomparare (cap. 3.4.2.3);')
bullet('GIS — Sistem geo-informațional (cap. 3.4.2.4);')
bullet('Culegere date, raportare, dashboard și sinteze — BI (cap. 3.4.2.5);')
bullet('Portal Servicii Publice (cap. 3.4.2.6);')
bullet('BND-SNIIA — Baza Națională de Date pentru identificarea și înregistrarea animalelor (cap. 3.4.2.7);')
bullet('Catagrafie și cartografiere exploatații, unități autorizate, farmacii (cap. 3.4.2.8);')
bullet('Supraveghere, prevenire, control și anchete (cap. 3.4.2.9);')
bullet('Management instruire personal laboratoare (cap. 3.4.2.10);')
bullet('Aplicație mobilă raportare animale sălbatice bolnave sau decedate (cap. 3.4.2.11);')
bullet('Autorizare, acreditare și desemnare (cap. 3.4.2.12);')
bullet('Contorizare folosire servicii publice — indicator RCR11 (cap. 3.4.2.13);')
bullet('Integrări externe și preluări de date (cap. 3.4.2.14).')

para('Componenta software — 27 produse software organizate pe 4 categorii (detaliate în capitolul 5 al ofertei și în Lista_Software_SIDISVA.xlsx):', bold=True)
bullet('A. Infrastructură software (11 produse) — sisteme de operare, server web, server aplicație, bază de date relațională + NoSQL, ETL, BI, IAM, GIS, ESB;')
bullet('B. Software aplicativ (6 produse) — DMS, Portal, Chatbot, aplicație mobilă, LIMS, gateway HL7;')
bullet('C. Securitate (7 produse) — WAF, Honeypot, NMS/NAC, SIEM, Email Security, NGFW centru + NGFW locații;')
bullet('D. Productivity (3 produse) — suite Office laptopuri, suite Office complete teren, antivirus EDR/XDR.')

para('Componenta hardware — detaliată în capitolul 6 al ofertei:', bold=True)
bullet('Echipamente de securitate pentru centrul de date și 42 locații DSVSA (NGFW redundant, switching, access point);')
bullet('100 laptopuri pentru ANSVSA centru — conforme cerinței DNSH (consum în veghe ≤20 Wh) — vezi §1.5 și capitolul 11;')
bullet('336 complete de teren pentru echipele de inspecție (8 per județ × 42 DSVSA);')
bullet('126 dispozitive IoT pentru integrarea echipamentelor de laborator (3 per institut și pe fiecare DSVSA cu activitate de laborator).')

para('Servicii — durată implementare 18 luni, conform cap. 6 al Caietului de Sarcini; garanție extinsă 3 ani de la data Punerii în Producție pentru ansamblul HW + SW + sistem informatic (detalii cap. 7).')

# ----------------------------- 1.4 -----------------------------
h(2, '1.4 Conformitate legală și standarde aplicabile')

para('Soluția propusă este conformă cu cadrul legislativ și standardele de mai jos. Lista exhaustivă a cerințelor individuale și modul lor de îndeplinire sunt în capitolul 8 (Conformitate cu specificațiile tehnice) și în Anexa F — Matricea de Conformitate (1.294 cerințe).')

add_table(
    ['Domeniu', 'Act normativ / standard'],
    [
        ['Cloud guvernamental', 'OUG nr. 89/2022'],
        ['Securitate cibernetică', 'Directiva (UE) 2022/2555 (NIS2); OUG nr. 155/2024 (transpunere NIS2); Legea nr. 362/2018 (Lege NIS)'],
        ['Protecția sistemelor IT publice', 'Legea nr. 354/2022 (context de securitate națională, restricții privind originea soluțiilor critice)'],
        ['Date personale', 'Regulamentul (UE) 2016/679 (GDPR); Legea nr. 190/2018; Legea nr. 363/2018'],
        ['Semnătură electronică', 'Regulamentul (UE) 910/2014 (eIDAS); Legea nr. 455/2001'],
        ['Interoperabilitate', 'SEMIC.EU; Ordinul MCID nr. 21286/26.10.2023; Legea nr. 242/2022 (PNI, RNR)'],
        ['Accesibilitate', 'WCAG 2.1 nivel AA; Carta UE a Drepturilor Fundamentale'],
        ['Arhivare', 'Legea nr. 16/1996 (Arhivele Naționale)'],
        ['Laboratoare', 'SR EN ISO/IEC 17025; SR EN ISO/IEC 17043; cerințe RENAR'],
        ['Plăți electronice', 'Legea nr. 207/2015 (Cod fiscal procedural); OUG nr. 41/2016'],
        ['Identificare animale', 'Reg. (CE) nr. 1760/2000; Reg. (UE) 2016/429; legislația națională specifică'],
        ['Sisteme de management', 'ISO/IEC 27001 (securitate); ISO/IEC 27017 / 27018 (cloud); ISO 9001 (calitate); ISO/IEC 20000-1 (servicii IT)'],
    ],
    widths_cm=[5.0, 11.0]
)

# ----------------------------- 1.5 -----------------------------
h(2, '1.5 Confirmarea acoperirii cerințelor eliminatorii')

para('Cerințele eliminatorii din Caietul de Sarcini sunt confirmate prin prezenta ofertă, după cum urmează:')

para('a) Acoperirea tuturor celor 1.294 cerințe individuale ', bold=True).add_run('extrase din cap. 3.4 al Caietului de Sarcini — matricea de conformitate este în Anexa F a ofertei.')

para('b) Demonstrarea celor 33 cerințe DEMO ', bold=True).add_run(
    '— scenariul video este detaliat în capitolul 14 al ofertei, iar înregistrarea video este atașată separat. '
    'Cerințele DEMO acoperă, fără limitare: înregistrarea documentelor de intrare/ieșire cu OCR și indexare; integrarea ROeID + eIDAS + 2FA; '
    'modificarea diagramelor BPMN/UML din UI; formulare drag-and-drop cu asistență chatbot AI; integrări terțe și GIS; lucru offline cu sincronizare; '
    'fluxuri configurabile cu min. 5 pași și 3 roluri; semnătură electronică calificată eIDAS + semnătură olografă captată pe pad USB; '
    'editarea documentelor Word/Excel direct în DMS; tipărirea plicurilor și a borderourilor cu coduri de bare; dashboard-uri drag-and-drop; '
    'multi-tenant pe aceeași instalare cu izolare totală; API REST documentat Swagger / OpenAPI 3.0.'
)

para('c) Acceptarea explicită a condițiilor de drepturi de proprietate intelectuală ', bold=True).add_run(
    'din capitolul 12 al Caietului de Sarcini: '
    'transferul drepturilor de proprietate către Beneficiar pentru toate dezvoltările și customizările realizate în cadrul contractului; '
    'licențe perpetue + cod sursă pentru componentele aplicative dezvoltate / customizate; '
    'documentația tehnică completă livrată în format editabil. Acceptarea este reluată formal în capitolul 15 al ofertei (Declarații obligatorii).'
)

para('d) Găzduirea exclusivă în Cloud Guvernamental ', bold=True).add_run('conform OUG nr. 89/2022, cu arhitectură Cloud-Native containerizată — detalii în capitolul 5.')

para('e) Conformitatea cu Legea nr. 354/2022 ', bold=True).add_run('— produsele software și hardware ofertate respectă restricțiile privind originea soluțiilor critice. Lista producătorilor și țara de origine este în Anexa B și în Lista_Software_SIDISVA.')

# ----------------------------- 1.6 -----------------------------
h(2, '1.6 Punctaj țintă pe factorii de evaluare')

para('Criteriul de atribuire este „Cel mai bun raport calitate-preț", cu pondere de 60% pentru componenta tehnică și 40% pentru preț, evaluat pe următorii factori:')

add_table(
    ['Factor', 'Pondere', 'Țintă în ofertă'],
    [
        ['P1 — Prețul ofertei', '40 puncte', 'Preț competitiv care respectă cele două plafoane procentuale (max. 20% hardware + servicii instalare/configurare/punere în producție; min. 10% securitate cibernetică) și se încadrează în valoarea estimată de 85.418.857,53 lei fără TVA'],
        ['P2 — Experiența celor 6 experți cheie punctați (5p × 6)', '30 puncte', 'Țintă scor maxim — câte 5+ proiecte / contracte demonstrabile per expert, pentru sisteme informatice cu min. 7 module interconectate + min. 1 modul de tip portal; certificările cerute (PMP/PRINCE2; business analysis; arhitectură enterprise + producător DMS; Scrum/Agile + cloud; Lean Six Sigma; administrare/HA/securitate BD ofertată)'],
        ['P3 — Metodologia de implementare (subfactori 10p + 10p)', '20 puncte', 'Țintă calificativ „Excepțional" la ambii subfactori — metodologii recunoscute prezentate detaliat (PMBOK, ISO 21500, ITIL, TOGAF, ISO 27001); adaptare specifică SIDISVA cu drum critic + registru de riscuri + ipoteze; inovație concretă (IA pentru calitate date BND-SNIIA, semnătură combinată olografă + digitală); securitate informatică demonstrată concret prin defense-in-depth, SIEM 24×7, IAM, criptare'],
        ['P4 — DNSH (laptop consum veghe 5p + ambalaje/livrare 5p)', '10 puncte', 'Țintă scor maxim — laptopuri certificate EU Ecolabel cu consum măsurat sub 20 Wh în modul veghe (rapoarte de testare atașate); ambalaje reciclabile/reutilizabile + livrare cu flotă electrică/hibridă (declarații furnizori atașate)'],
    ],
    widths_cm=[3.5, 2.5, 10.0]
)

para('Țintă totală: punctaj maxim posibil pe componenta tehnică (60 puncte din 60), cu plus oferta de preț competitivă pentru a maximiza scorul P1.')

# ----------------------------- 1.7 -----------------------------
h(2, '1.7 Angajamentele ofertantului')

para('<LIDER>, în numele asocierii cu <PARTENER>, se angajează prin prezenta ofertă să:')
bullet('Respecte integral cerințele Caietului de Sarcini și ale Fișei de date a achiziției;')
bullet('Livreze cele 14 componente funcționale conform descrierii din capitolul 4 al ofertei, în cele 18 luni de implementare;')
bullet('Asigure cele ≥56 servicii electronice publice (indicator RCR11 POCIDIF);')
bullet('Demonstreze toate cele 33 cerințe DEMO din Caietul de Sarcini — atât în înregistrarea video atașată ofertei, cât și la sediul ANSVSA în max. 5 zile lucrătoare de la cerere;')
bullet('Realizeze integrarea ulterioară a sistemelor guvernamentale care nu sunt încă disponibile cu API la momentul implementării (PNI, PCUe/PDURo, alte) — fără cost suplimentar, în perioada de garanție, conform abordării solicitate prin Caietul de Sarcini;')
bullet('Asigure trasabilitatea cerințelor pe traseul Caiet de Sarcini → Document de Analiză → Scenarii de Testare → Recepție Finală — instrument prezentat în capitolul 13;')
bullet('Acorde garanție extinsă 3 ani de la data Punerii în Producție, pentru întregul ansamblu (HW + SW + sistem informatic), cu SLA cu 4 niveluri de severitate (Critic / Major / Mediu / Minor) și timpi de răspuns / remediere conform capitolului 7;')
bullet('Mențină valabilitatea ofertei pe perioada de 120 zile de la data limită de depunere a ofertelor, conform Fișei de date.')

# ----------------------------- 1.8 -----------------------------
h(2, '1.8 Sinteza ofertei')

add_table(
    ['Element', 'Valoare propusă'],
    [
        ['Ofertant', '<LIDER> (lider asociere) + <PARTENER> (asociat / subcontractant)'],
        ['Procedură', 'CN1089237 (SEAP), Caiet de Sarcini nr. 7574/CP/2025'],
        ['Cod SMIS proiect', '336342 — POCIDIF 2021–2027, P2, Acțiunea 2.2'],
        ['Soluție software aplicativă', 'VOGO Enterprise Suite (Portal + Chatbot + BND-SNIIA + Supraveghere + aplicații mobile) + ZIPPER DMS + <FURNIZOR_LIMS> + <FURNIZOR_GIS>'],
        ['Soluție infrastructură SW', '27 produse pe 4 categorii (vezi §1.3 și Lista_Software_SIDISVA)'],
        ['Componente funcționale livrate', '14 componente, conform cap. 3.4.2.1 – 3.4.2.14 din CdS'],
        ['Servicii electronice publice', '≥56 servicii (indicator RCR11 POCIDIF)'],
        ['Utilizatori țintă', '≥185.000 utilizatori unici anual (cetățeni / fermieri / operatori economici)'],
        ['Instituții deservite', '46 (ANSVSA centru + 42 DSVSA + 3 institute subordonate)'],
        ['Hardware', 'NGFW + switching + AP locații; 100 laptopuri DNSH; 336 complete teren; 126 IoT laborator'],
        ['Cloud', 'Cloud Guvernamental conform OUG nr. 89/2022, arhitectură Cloud-Native (Docker / Kubernetes)'],
        ['Durată implementare', '18 luni de la semnarea contractului (cap. 6 CdS)'],
        ['Garanție', '3 ani de la Punerea în Producție, pentru întreg ansamblul'],
        ['Cadre legislative cheie', 'OUG 89/2022; NIS2 (Dir. UE 2022/2555 + OUG 155/2024); L. 354/2022; GDPR; eIDAS; L. 16/1996'],
        ['Standarde', 'ISO/IEC 27001/27017/27018; ISO 9001; ISO/IEC 20000-1; SR EN ISO/IEC 17025; SR EN ISO/IEC 17043; HL7 FHIR; WCAG 2.1 AA'],
        ['Metodologie', 'PMBOK pentru management proiect; BABOK pentru analiză; Agile / Scrum hibrid pentru dezvoltare; ITIL pentru servicii'],
        ['Acceptare drepturi IP (cap. 12 CdS)', 'Acceptat explicit (vezi §1.5.c și cap. 15 al ofertei)'],
        ['Confirmare 33 cerințe DEMO', 'Scenariu video în cap. 14; demo la sediul ANSVSA în max. 5 zile lucrătoare la cerere'],
        ['Validitate ofertă', '120 zile de la data limită de depunere'],
        ['Preț total (lei fără TVA)', 'Se completează la finalizarea propunerii financiare'],
    ],
    widths_cm=[5.5, 10.5]
)

para('Detaliile complete privind soluția, metodologia, echipa, planul de proiect, securitatea, garanția și conformitatea cu cerințele se regăsesc în capitolele 2 – 16 ale prezentei oferte tehnice și în anexele aferente.')

doc.save(DST)
print(f"[OK] Salvat: {DST}")

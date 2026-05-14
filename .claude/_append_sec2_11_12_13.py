"""
Appenda sectiunile 2.11 (Riscuri 7+6), 2.12 (Securitate informatica), 2.13 (Inovatie)
la 2-Abordare_metodologie.docx - finalizare task A2.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt

SRC = Path(r"C:\Users\adria\Documents\Claude\Projects\Licitatie 1\doc-output\oferta\2-Abordare_metodologie.docx")


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    return p


def add_p(doc, text):
    p = doc.add_paragraph()
    p.add_run(text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    p.add_run("• " + text)
    return p


def add_table(doc, header, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    try:
        t.style = "Table Grid"
    except KeyError:
        pass
    hdr = t.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
    for ri, row in enumerate(rows, start=1):
        cells = t.rows[ri].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    return t


def main():
    if not SRC.exists():
        raise SystemExit(f"NOT FOUND: {SRC}")
    doc = Document(str(SRC))

    # ============================================================
    # 2.11 Registrul de riscuri al proiectului
    # ============================================================
    add_h2(doc, "2.11 Registrul de riscuri al proiectului SIDISVA")

    add_p(doc,
        "In completarea procedurii generale de management al riscurilor prezentata in sectiunea "
        "2.3, consortiul a intocmit registrul de riscuri specific proiectului SIDISVA. Registrul "
        "include integral cele 7 riscuri minime identificate de Autoritatea Contractanta in "
        "cap. 4.2 al Caietului de Sarcini, plus 6 riscuri suplimentare identificate de consortiu "
        "pe baza experientei proprii in implementari similare (sisteme guvernamentale cu "
        "integrari multiple). Fiecare risc are atribuit: probabilitate (scala 1-5), impact "
        "(scala 1-5), scor (P x I), strategia de tratare (evitare / mitigare / transfer / acceptare), "
        "planul de mitigare, planul de contingenta si proprietarul (Risk Owner)."
    )
    add_p(doc,
        "Registrul de riscuri este un document viu, actualizat lunar de Project Manager <LIDER> "
        "si raportat in fiecare Steering Committee. Reevaluarea probabilitatilor si impactelor "
        "se face la fiecare milestone major al proiectului (M6 - analiza finalizata, M12 - HW "
        "instalat, M15 - dezvoltare finalizata, M18 - Go-live)."
    )

    add_h3(doc, "2.11.1 Riscurile obligatorii din cap. 4.2 al Caietului de Sarcini")
    add_table(doc,
        header=["#", "Risc", "P x I", "Strategie", "Plan de mitigare", "Owner"],
        rows=[
            ["R1", "Lipsa unei comunicari eficiente intre echipa beneficiarului si consortiu",
             "3x4=12", "Mitigare",
             "Plan formal de comunicare; Steering Committee lunar + Technical Committee saptamanal; platforma Jira+Confluence cu vizibilitate cross-furnizor; SLA reciproc pentru livrarea informatiilor.",
             "PM <LIDER> + reprezentant ANSVSA"],
            ["R2", "Extinderea perioadei de achizitie din intarzieri/blocaje administrative",
             "2x3=6", "Acceptare + buffer",
             "Buffer de 1 luna prevazut in plan; suplimentare efort prin alocare resurse de rezerva in lunile critice; flexibilitate prin paralelizarea activitatilor ne-blocante.",
             "PM <LIDER>"],
            ["R3", "Schimbarea factorilor de decizie din ANSVSA pe parcursul implementarii",
             "3x3=9", "Mitigare",
             "Prezentarea proiectului tuturor decidentilor cheie din luna 1; documentatie de proiect actualizata permanent pentru onboarding rapid al unor noi decidenti; arhiva deciziilor in Confluence.",
             "<LIDER> + ANSVSA"],
            ["R4", "Intarzieri in luarea deciziilor (validare livrabile, aprobare CR-uri)",
             "4x3=12", "Mitigare",
             "SLA reciproc: 10 zile lucratoare validare rapoarte trimestriale, 5 zile pentru observatii la livrabile, 3 zile pentru aprobari CR mici (cap. 10.3 CdS); proceduri CCB formale; escalare catre Steering Committee.",
             "ANSVSA + PM <LIDER>"],
            ["R5", "Fluctuatii la nivelul echipelor de proiect (plecari experti cheie)",
             "3x4=12", "Mitigare + transfer",
             "Inlocuire experti cu CV similar conform art. 162 HG 395/2016 (cap. 8.1 CdS); retention bonus pentru echipa cheie; cross-training intre furnizori pentru roluri critice; arhitect <LIDER> face mentoring tehnic.",
             "<LIDER> + furnizori"],
            ["R6", "Date/informatii insuficiente de la Autoritatea Contractanta",
             "3x4=12", "Mitigare",
             "SLA livrare informatii (cf. ipoteza c, sec. 2.10.4); escalare formala pentru intarzieri; plan B prin modelare ipotetica validata ulterior cu ANSVSA; documentarea ipotezelor in Document de Analiza.",
             "ANSVSA + Analist Business <LIDER>"],
            ["R7", "Adaugare cerinte/activitati noi din modificari legislative (NIS2, L 242/2022, L 354/2022)",
             "3x3=9", "Mitigare prin Change Request",
             "Procedura formala CR cu CCB; buffer de 5% timp/buget contingency; monitorizare permanenta a actelor normative cu impact (DPO + Expert Securitate <LIDER>); arhitectura flexibila pentru ajustari legislative.",
             "PM <LIDER>"],
        ]
    )

    add_h3(doc, "2.11.2 Riscuri suplimentare identificate de consortiu")
    add_table(doc,
        header=["#", "Risc", "P x I", "Strategie", "Plan de mitigare", "Owner"],
        rows=[
            ["R8", "Indisponibilitate sisteme guvernamentale tinta (PNI, ROeID extins, eIDAS, PJN, PCUe) la momentul integrarii",
             "4x5=20", "Mitigare arhitecturala",
             "Strategia mock-up + comutare (vezi sec. 2.13.3): Anti-Corruption Layer pe Oracle Service Bus; mock-uri OpenAPI complete; feature flags pentru comutare automata mock-real; clauza cap. 3.4.2.14 CdS - integrare ulterioara in garantie fara cost suplimentar.",
             "<FURNIZOR ESB> + Arhitect Sistem <LIDER>"],
            ["R9", "Modificari API la sisteme partenere (APIA, ONRC, ANARZ, Colegiul Medicilor Veterinari)",
             "3x3=9", "Mitigare arhitecturala",
             "Anti-Corruption Layer pe Oracle Service Bus pentru izolarea SIDISVA; versionare API conforma SEMIC.EU; monitor proactiv pentru schimbari (subscribe la canalele oficiale); contracte API formale cu fiecare partener.",
             "<FURNIZOR ESB>"],
            ["R10", "Incidente de securitate cibernetica pe perioada implementarii (atacuri, ransomware, exfiltrare date)",
             "2x5=10", "Mitigare + transfer",
             "Defense-in-depth pe 7 straturi (sec. 2.12.1); DevSecOps cu SAST/DAST in CI/CD; SIEM Splunk/QRadar cu SOC 24x7; Honeypot FortiDeceptor pentru early warning; asigurare cyber-risk; plan continuitate operationala (BCP).",
             "<LIDER> Security + Expert Securitate"],
            ["R11", "Migrare incompleta date din vechiul SNIIA, LIMS, registre xls",
             "3x4=12", "Mitigare iterativa",
             "Migrare iterativa pe loturi; perioada paralel run de min. 30 zile; plan de rollback documentat; validare calitate date prin algoritmi AI/ML de detectie anomalii (sec. 2.13.2); audit migrare semnat de ANSVSA.",
             "<FURNIZOR LIMS> + <LIDER>"],
            ["R12", "Performanta sub asteptari la incarcare cu 185.000 utilizatori unici/an portal",
             "2x4=8", "Mitigare arhitecturala",
             "Arhitectura Cloud-Native cu auto-scaling Kubernetes; testare load + stress (sec. 2.5 - Testare); CDN pentru asset-uri statice; cache Redis pentru date frecvent accesate; SLO definit (p95 sub 2s).",
             "Arhitect Sistem <LIDER>"],
            ["R13", "Adoptie slaba utilizatori (5.300 angajati ANSVSA + 2.600 medici vet concesionari + 4.800 utilizatori acreditati)",
             "3x3=9", "Mitigare prin change management",
             "Plan instruire stratificat (Train-the-Trainer + sesiuni grupuri mici + e-learning); 144 utilizatori cheie + 3 administratori instruiti formal (cf. cap. 3.5 CdS); suport on-call in primele 3 luni post Go-live; chatbot AI pentru auto-asistenta; manuale video.",
             "<LIDER> + Expert Instruire"],
        ]
    )

    # ============================================================
    # 2.12 Asigurarea securitatii informatice si informationale
    # ============================================================
    add_h2(doc, "2.12 Asigurarea securitatii informatice si informationale")

    add_p(doc,
        "Securitatea cibernetica si protectia informatiilor sunt integrate by design pe intreaga "
        "arhitectura SIDISVA si pe intregul ciclu de viata al proiectului. Aceasta sectiune "
        "demonstreaza concret modul in care consortiul va indeplini cele 6 principii de securitate "
        "cibernetica enumerate la cap. 3.4.1.2 al Caietului de Sarcini (Conformitate, Optimizare "
        "costuri, Responsabilitati partajate, Protectia informatiilor, Securitate pe intreg "
        "ciclul de viata, Securizarea operatiunilor) si cum se conformeaza cu urmatoarele "
        "cadre normative obligatorii:"
    )
    add_bullet(doc,
        "Regulamentul (UE) 2016/679 GDPR + Legea 190/2018 + Legea 363/2018 (protectia datelor "
        "personale)"
    )
    add_bullet(doc,
        "Directiva (UE) 2022/2555 NIS2, transpusa in Romania prin OUG 155/2024 (securitate "
        "cibernetica pentru entitati esentiale si importante)"
    )
    add_bullet(doc,
        "Legea 354/2022 privind protectia sistemelor informatice ale autoritatilor publice "
        "in contextul invaziei declansate de Federatia Rusa impotriva Ucrainei"
    )
    add_bullet(doc,
        "Legea 362/2018 + OUG 119/2020 (nivel comun ridicat de securitate a retelelor si "
        "sistemelor informatice)"
    )
    add_bullet(doc,
        "ISO/IEC 27001:2022 (Sistem de Management al Securitatii Informatiei - SMSI), "
        "ISO/IEC 27701:2019 (Privacy Information Management - PIMS), ISO/IEC 27017 (cloud "
        "security), ISO/IEC 27018 (PII in cloud)"
    )
    add_bullet(doc,
        "NIST Cybersecurity Framework v2.0 (Identify / Protect / Detect / Respond / Recover / "
        "Govern)"
    )
    add_bullet(doc,
        "OWASP Top 10 + OWASP ASVS Level 3 (aplicatii web), OWASP MASVS Level 1 (aplicatii mobile)"
    )

    # 2.12.1 Defense in depth
    add_h3(doc, "2.12.1 Arhitectura de securitate defense-in-depth pe 7 straturi")
    add_p(doc,
        "Modelul defense-in-depth aplicat la SIDISVA distribuie mecanismele de protectie pe "
        "7 straturi independente, fiecare cu produsele si parametrii sai specifici (toate "
        "cantitatile conform cap. 3.4.3.4 al Caietului de Sarcini):"
    )
    add_table(doc,
        header=["Strat", "Mecanisme implementate", "Produse / Tehnologii"],
        rows=[
            ["1. Fizic",
             "Gazduire in Cloud Guvernamental (Azure RO) cu certificare ISO 27001 + ISAE 3402; control acces fizic conform politicilor Cloud Guvernamental; protectie environmental",
             "Cloud Guvernamental (Azure Stack RO)"],
            ["2. Retea perimetru",
             "Firewall redundant centru (2 instante, min. 10 VDOM); firewall in 45 locatii DSVSA (90 instante); micro-segmentare; VPN IPSec; protectie DDoS",
             "FortiGate / Palo Alto (centru), FortiGate 100F (locatii)"],
            ["3. Retea interna",
             "Segmentare VLAN per zona (DMZ, business, management, backup); NAC 802.1X pentru min. 600 dispozitive; NMS pentru detectie anomalii trafic",
             "Cisco DNA / ClearPass"],
            ["4. Aplicatie",
             "WAF cu protectie OWASP Top 10 (2 instante, NU per useri cf. cap. 3.4.3.4.1.1 CdS); Email Security cu filtrare + sandbox (2 instante); SAST + DAST automate in CI/CD",
             "F5 / Imperva / FortiWeb (WAF); Cisco IronPort (email)"],
            ["5. Identitate",
             "IAM Keycloak Enterprise (150 interni + 185.000 portal); federatie ROeID (cetateni RO) + eIDAS (UE); MFA obligatoriu pentru utilizatori interni; RBAC granular; audit logging end-to-end",
             "Keycloak Enterprise + ROeID + eIDAS"],
            ["6. Date",
             "Encryption at rest (TDE Microsoft SQL Server Enterprise, BitLocker, Azure Disk Encryption); encryption in transit (TLS 1.3 obligatoriu); Azure Key Vault HSM-backed pentru chei; data masking pentru medii test/dev",
             "MS SQL TDE, Azure Disk Encryption, Azure Key Vault HSM"],
            ["7. Endpoint",
             "EDR pentru 486 endpoints (436 + 50); MS Defender for Cloud; control USB; politici Group Policy; patch management automat",
             "CrowdStrike / SentinelOne EDR + MS Defender"],
        ]
    )

    # 2.12.2
    add_h3(doc, "2.12.2 Detectare si raspuns la incidente (Detect & Respond)")
    add_bullet(doc,
        "SIEM (Splunk Enterprise Security sau IBM QRadar) - colectare centralizata log-uri "
        "din toate componentele, normalizate Common Information Model; retentie Indicators of "
        "Compromise (IOC) pe 3 ani conform cap. 3.4.3.4.1.4 CdS; corelatie evenimente pe "
        "regulile MITRE ATT&CK."
    )
    add_bullet(doc,
        "Security Operations Center (SOC) 24x7 - fie intern <LIDER> (echipa dedicata ISO 27001 "
        "certificata), fie externalizat la un furnizor MSSP autorizat DNSC; tier 1 + tier 2 + "
        "tier 3 analisti; on-call rotation."
    )
    add_bullet(doc,
        "Honeypot FortiDeceptor - 4 instante Windows x 4 vCPU (cf. cap. 3.4.3.4.1.2 CdS); "
        "decoys cu credentiale fake pentru detectarea lateral movement si a atacatorilor cu "
        "acces initial."
    )
    add_bullet(doc,
        "Threat Intelligence - feed-uri MISP (Malware Information Sharing Platform), feed-uri "
        "ENISA, feed-uri DNSC (autoritatea nationala), feed-uri commerciale (Recorded Future / "
        "Mandiant); integrare cu SIEM pentru corelatie automata."
    )
    add_bullet(doc,
        "SOAR (Security Orchestration, Automation and Response) - playbook-uri automate "
        "pentru incidente comune (brute force, malware detected, phishing, ransomware "
        "indicators); izolare automata endpoints compromise; notificari catre SOC."
    )
    add_bullet(doc,
        "Notificare obligatorie incidente - conform Art. 23 NIS2 + OUG 155/2024: notificare "
        "early warning DNSC in 24h, notificare incident in 72h, raport final in 1 luna; "
        "proceduri pregatite si testate pre Go-live."
    )

    # 2.12.3 DevSecOps
    add_h3(doc, "2.12.3 Securitatea proceselor de dezvoltare (DevSecOps shift-left)")
    add_bullet(doc,
        "SAST (Static Application Security Testing) - analiza cod sursa in fiecare commit "
        "Git, gate-uri in pipeline CI/CD pentru blocarea CVE high/critical; instrumente: "
        "SonarQube, Semgrep, Checkmarx (pentru cod propriu) + scan-uri specifice limbajelor "
        "(ESLint security, Bandit pentru Python, gosec pentru Go etc.)."
    )
    add_bullet(doc,
        "DAST (Dynamic Application Security Testing) - testare saptamanala pe medii staging; "
        "instrumente: OWASP ZAP, Burp Suite Pro, Acunetix; scan-uri full OWASP Top 10."
    )
    add_bullet(doc,
        "Dependency Scanning (SCA - Software Composition Analysis) - Snyk + Trivy + OWASP "
        "Dependency-Check pe fiecare build; alert daca apare CVE high/critical intr-o "
        "dependenta; politici de actualizare cu maxim 30 zile."
    )
    add_bullet(doc,
        "Secret Scanning - gitleaks + TruffleHog ruleaza la fiecare commit pentru a preveni "
        "leak-uri de credentiale (chei API, parole, token-uri JWT) in repository."
    )
    add_bullet(doc,
        "Container Security - image scanning cu Trivy pentru toate imaginile Docker; runtime "
        "protection cu Falco pentru detectie comportament anomal in pod-uri; policy-as-code cu "
        "Open Policy Agent (OPA) Gatekeeper pe Kubernetes."
    )
    add_bullet(doc,
        "Penetration Testing - intern trimestrial (echipa <LIDER> Security cu certificari "
        "OSCP/CEH) + audit extern anual de catre o terta parte independenta certificata ANSCC; "
        "rezultate raportate catre Steering Committee + remediere obligatorie pre-acceptanta."
    )
    add_bullet(doc,
        "Bug Bounty intern pentru perioada UAT - utilizatori cheie ANSVSA invitati sa raporteze "
        "probleme de securitate identificate, cu recompense simbolice; canal dedicat in Jira "
        "Service Management."
    )

    # 2.12.4 Privacy / GDPR
    add_h3(doc, "2.12.4 Protectia datelor personale (Privacy & GDPR)")
    add_p(doc,
        "SIDISVA prelucreaza date personale ale cetatenilor romani (date de identificare via "
        "ROeID, date contact, exploatatii animale), date privind sanatatea publica (cu impact "
        "indirect - sanitar-veterinar), date ale profesionistilor (medici veterinari "
        "concesionari, angajati ANSVSA). Consortiul aplica urmatoarele masuri specifice:"
    )
    add_bullet(doc,
        "DPIA (Data Protection Impact Assessment) - elaborat inainte de Go-live conform Art. 35 "
        "GDPR; revizuit la fiecare modificare semnificativa a prelucrarii; coordonare cu DPO "
        "ANSVSA + DPO <LIDER>."
    )
    add_bullet(doc,
        "Privacy by Design + Privacy by Default - minimizare date colectate (doar ce e strict "
        "necesar serviciului); pseudo-anonimizare unde e tehnic posibil; retention policies "
        "automatizate cu stergere/anonimizare la expirare; consimtamant explicit pentru "
        "prelucrari secundare."
    )
    add_bullet(doc,
        "Exercitarea drepturilor persoanelor vizate (Art. 15-22 GDPR) - interfata self-service "
        "in Portal pentru: dreptul de acces (export date proprii), rectificare, stergere "
        "(drept de a fi uitat), portabilitate (export in format machine-readable), restrictionare, "
        "opozitie; ticket automat catre DPO pentru cereri complexe."
    )
    add_bullet(doc,
        "Registrul activitatilor de prelucrare (Art. 30 GDPR) - mentinut electronic in DMS "
        "(ZIPPER DMS), actualizat la fiecare modificare; export disponibil pentru ANSPDCP "
        "(Autoritatea Nationala de Supraveghere a Prelucrarii Datelor cu Caracter Personal) "
        "la cerere."
    )
    add_bullet(doc,
        "Notificare breach (Art. 33-34 GDPR) - procedura pregatita pentru notificare ANSPDCP "
        "in 72h si informare persoane vizate cand e cazul; testata in exercitii simulate."
    )
    add_bullet(doc,
        "Pregatit pentru Platforma de Jurnalizare-Notificare (PJN) - cand va fi disponibila "
        "la nivel guvernamental (cf. cap. 3.4.2.14 CdS), fiecare cetatean va putea fi "
        "notificat automat cand datele sale sunt accesate de un operator ANSVSA."
    )

    # 2.12.5 Continuitate
    add_h3(doc, "2.12.5 Continuitate operationala si recuperare in caz de dezastru")
    add_bullet(doc,
        "Strategy backup 3-2-1 - 3 copii ale datelor (1 productie + 2 backup-uri), 2 medii "
        "diferite (online + offline), 1 copie off-site (regiune diferita Cloud Guvernamental); "
        "backup automat zilnic + transactional log la 15 minute."
    )
    add_bullet(doc,
        "RPO (Recovery Point Objective) tinta: 1 ora - pierdere maxima admisibila de date "
        "in caz de incident."
    )
    add_bullet(doc,
        "RTO (Recovery Time Objective) tinta: 4 ore - timp maxim admisibil pentru restaurarea "
        "serviciului in caz de dezastru."
    )
    add_bullet(doc,
        "DR Site secundar - in Cloud Guvernamental, regiune geografica diferita; replicare "
        "asincrona continua pentru date critice; comutare automata DNS in caz de failure "
        "centru primar."
    )
    add_bullet(doc,
        "Teste DR - semestrial pe perioada implementarii contractului, lunar pe perioada "
        "garantiei; documentate prin rapoarte de exercitiu cu lessons learned."
    )
    add_bullet(doc,
        "Business Continuity Plan (BCP) - document standalone elaborat in luna 8, actualizat "
        "anual; include scenarii: ransomware, pierdere DC primar, indisponibilitate Cloud "
        "Guvernamental, atac DDoS prelungit; proceduri runbook pentru fiecare scenariu."
    )

    # ============================================================
    # 2.13 Abordare inovatoare si diferentiatori
    # ============================================================
    add_h2(doc, "2.13 Abordare inovatoare si diferentiatori competitivi")

    add_p(doc,
        "Pentru a livra valoare adaugata semnificativa peste cerintele minime ale Caietului "
        "de Sarcini - element esential pentru obtinerea calificativului EXCEPTIONAL la "
        "subfactorul 3.1 (criteriul abordare inovatoare si eficienta) - consortiul propune "
        "5 elemente concrete de inovatie si diferentiere. Fiecare element are impact direct "
        "verificabil asupra eficientei, calitatii sau riscului proiectului."
    )

    # 2.13.1
    add_h3(doc, "2.13.1 Suite integrata Portal + Chatbot + App mobila (reducere risc integrare ~30%)")
    add_p(doc,
        "Spre deosebire de abordarile tipice multi-furnizor pe componente conexe, in cadrul "
        "consortiului trei componente strans interdependente - Portal Servicii Publice, "
        "Chatbot AI conversational si Aplicatii mobile native pentru cetateni - sunt furnizate "
        "de acelasi membru (<FURNIZOR PORTAL> / <FURNIZOR CHATBOT> / <FURNIZOR APP MOBILA>), "
        "prin module ale aceleiasi platforme integrate Enterprise Suite. Avantajele concrete, "
        "masurabile:"
    )
    add_bullet(doc,
        "Single Sign-On nativ intre cele 3 canale - un singur cont cetatean folosit identic "
        "in Portal web, Chatbot si App mobila, fara re-autentificare la trecerea intre canale."
    )
    add_bullet(doc,
        "UX consistent garantat - acelasi design system, aceleasi terminologii, acelasi "
        "flow conversational; reduce timpul de instruire cetateni si rata de abandon."
    )
    add_bullet(doc,
        "Sincronizare in timp real a starii utilizatorului - o cerere depusa in Portal e "
        "vizibila imediat in App mobila; notificarea push la decizia ANSVSA ajunge prin canalul "
        "preferat al cetateanului."
    )
    add_bullet(doc,
        "Mentenanta unica - un singur producator pentru cele 3 module, un singur ciclu de "
        "release (eliminand decalajul de versiune intre componente), un singur SLA, un singur "
        "punct de escaladare in garantie."
    )
    add_bullet(doc,
        "Reducerea cu ~30% a riscului de incompatibilitate API - estimat pe baza experientei "
        "consortiului din implementari similare unde componentele eterogene au generat "
        "incidente de integrare pe perioada UAT."
    )

    # 2.13.2
    add_h3(doc, "2.13.2 AI/ML pentru calitatea datelor BND-SNIIA si fraud detection")
    add_p(doc,
        "Caietul de Sarcini solicita explicit, la cap. 3.4.2.7, implementarea de algoritmi "
        "ML/AI pentru monitorizarea calitatii datelor, alerte pentru fraude si alte anomalii, "
        "predictii legate de volumul de actiuni sanitar-veterinare. Consortiul propune "
        "implementarea concreta a 4 modele dedicate, antrenate pe datele istorice ANSVSA + SNIIA:"
    )
    add_bullet(doc,
        "Anomaly Detection - Isolation Forest + LSTM Autoencoders pentru depistarea "
        "automata a anomaliilor in declararile de evenimente BND-SNIIA: inregistrari de "
        "nastere imposibile temporal (ex: 2 fatari consecutive in 30 zile pentru aceeasi "
        "femela bovina), miscari de animale incompatibile geografic (transport peste 1000 km "
        "in 6 ore), schimbari de proprietar suspecte."
    )
    add_bullet(doc,
        "Prediction Volumes - XGBoost si Prophet pentru estimarea volumului de actiuni "
        "sanitar-veterinare asteptate per DSVSA judeteana / per luna, permitand alocarea "
        "anticipata a resurselor inspectie si a stocurilor de crotalii / vaccinuri / "
        "reagenti laborator."
    )
    add_bullet(doc,
        "Fraud Detection - Graph Neural Networks (GNN) pentru detectarea pattern-urilor de "
        "frauda: aceleasi animale crotaliate de mai multe ori (re-utilizare ilegala crotalii "
        "pierdute), transferuri ciclice intre exploatatii conectate (laundering provenienta), "
        "scheme de frauda subventii APIA."
    )
    add_bullet(doc,
        "Data Quality Scoring - sistem de scoring per inregistrare BND (0-100) bazat pe "
        "consistenta temporala, geografica, biologica; evidentiere automata pe harta GIS a "
        "zonelor cu calitate slaba a datelor (potential pentru control fizic / re-verificare); "
        "dashboard de calitate vizibil utilizatorilor ANSVSA."
    )
    add_p(doc,
        "Toate modelele AI/ML vor fi explicabile (XAI - Explainable AI, prin tehnici SHAP/LIME), "
        "respectand principiul transparentei algoritmice obligatoriu pentru sistemele "
        "guvernamentale (Regulamentul UE AI Act). Re-training-ul modelelor este programat "
        "trimestrial pe perioada garantiei."
    )

    # 2.13.3
    add_h3(doc, "2.13.3 Strategia de rezilienta a integrarilor - mock-up + comutare")
    add_p(doc,
        "Pentru integrarile cu sisteme guvernamentale in curs de implementare la nivel "
        "national (Platforma Nationala de Interoperabilitate - PNI, Platforma de "
        "Jurnalizare-Notificare - PJN, Punctul de Contact Unic electronic - PCUe / Portalul "
        "Digital Unic - PDURo), unde Caietul de Sarcini (cap. 3.4.2.14) permite explicit "
        "implementarea cu mock-up si finalizarea integrarii in perioada de garantie, "
        "consortiul aplica o strategie sofisticata in 4 pasi:"
    )
    add_bullet(doc,
        "Anti-Corruption Layer pe Oracle Service Bus - pentru fiecare sistem tert, un strat "
        "de adaptare izoleaza SIDISVA de modelele de date si API-urile specifice; modificarile "
        "viitoare ale sistemelor partenere afecteaza doar layer-ul, nu componentele SIDISVA."
    )
    add_bullet(doc,
        "Mock-uri OpenAPI complete - pentru fiecare sistem tinta, mock-uri implementate "
        "generate din specificatiile publicate sau previzionate; mock-urile raspund cu date "
        "realiste pentru testare si UAT; documentatie Swagger completa pentru fiecare API."
    )
    add_bullet(doc,
        "Switch automatic prin feature flags - comutare instantanee mock-real pentru fiecare "
        "integrare, fara modificare cod si fara downtime; configurabil per mediu (dev/test/prod); "
        "rollback rapid in caz de probleme."
    )
    add_bullet(doc,
        "Conformitate cu Ordinul MCID 21286/26.10.2023 (standarde interoperabilitate) - "
        "toate API-urile expuse de SIDISVA respecta cadrul SEMIC.EU + EIF (European "
        "Interoperability Framework); catalog API publicat in RNR (Registrul National al "
        "Registrelor) cand acesta va fi disponibil."
    )

    # 2.13.4
    add_h3(doc, "2.13.4 Chatbot AI conversational in limba romana - NLP nativ")
    add_p(doc,
        "Componenta Chatbot din Enterprise Suite implementeaza procesare a limbajului natural "
        "(NLP) nativ in limba romana, fine-tunata pe corpus-ul public ANSVSA (legislatie, "
        "ghiduri, FAQ-uri, comunicate). Capabilitati concrete:"
    )
    add_bullet(doc,
        "Intent Classification cu peste 90% acuratete pe cazurile uzuale - depunere cerere "
        "(analiza, autorizare, raportare animale salbatice), verificare status cerere "
        "anterior depusa, intrebari frecvente (durate procesare, costuri, documente necesare), "
        "ghidare procedurala pas-cu-pas."
    )
    add_bullet(doc,
        "Fallback warm transfer la operator uman - cand acuratetea modelului scade sub un "
        "prag configurabil (default 60%), conversatia e transferata automat unui operator "
        "ANSVSA cu istoric conversatie + context cetatean preincarcat (Salesforce Service "
        "Cloud / similar)."
    )
    add_bullet(doc,
        "Retrieval-Augmented Generation (RAG) - cunostintele ANSVSA (legislatie, ghiduri "
        "interne) sunt indexate vectorial in Elasticsearch + retrieved la fiecare interogare; "
        "elimina riscul de halucinatie caracteristic LLM-urilor pure."
    )
    add_bullet(doc,
        "Conformitate DEMO #7 + #6 - Chatbot-ul asista utilizatorul la configurarea "
        "formularelor drag-and-drop prin sugestii contextuale (DEMO #7); si permite "
        "modificarea diagramelor BPMN/UML prin comenzi conversationale (parte din DEMO #6)."
    )
    add_bullet(doc,
        "Multi-canal - acelasi chatbot expus pe Portal web (widget), pe App mobila "
        "(integrare nativa), pe WhatsApp Business (canal optional) si pe e-mail (raspunsuri "
        "automate pentru intrebari simple)."
    )

    # 2.13.5
    add_h3(doc, "2.13.5 Telemetrie automata pentru indicator RCR11 - automatizare audit POCIDIF")
    add_p(doc,
        "Indicatorul de rezultat obligatoriu POCIDIF RCR11 - Utilizatori de servicii, "
        "produse si procese digitale publice noi si optimizate, numar utilizatori anual - "
        "necesita, conform cap. 3.4.2.13 CdS, crearea unei modalitati facile de probare si "
        "contorizare a utilizatorilor. Consortiul propune o solutie automatizata end-to-end:"
    )
    add_bullet(doc,
        "Event streaming in timp real - fiecare accesare a unui serviciu electronic SIDISVA "
        "(56 servicii prevazute la cap. 3.5) genereaza automat un eveniment Apache Kafka, "
        "consumat de pipeline-ul Microsoft SSIS si incarcat in Data Warehouse Microsoft SQL "
        "Server Enterprise."
    )
    add_bullet(doc,
        "Dashboard auditor dedicat in Microsoft Power BI - vizualizare RCR11 pe diverse "
        "dimensiuni (perioada, tipuri servicii, demografice utilizatori, demografice zone "
        "geografice); autentificare via Keycloak Enterprise cu rol Auditor POCIDIF; "
        "export Excel/PDF pentru raportare oficiala."
    )
    add_bullet(doc,
        "Export automat lunar catre Agentia pentru Digitalizarea Romaniei (ADR) - in format "
        "standardizat conform Catalogului Serviciilor Publice (CSP) gestionat de ADR; "
        "elimina munca manuala de raportare lunara."
    )
    add_bullet(doc,
        "Anonimizare automata - pentru raportarea publica catre cetateni, datele sunt "
        "anonimizate prin k-anonymity (k>=5) si differential privacy; doar datele agregate "
        "sunt publicabile."
    )
    add_bullet(doc,
        "Trail de audit complet - pentru fiecare conturizare RCR11, log-uri SIEM cu "
        "timestamp + sursa evenimentului + tipul serviciului accesat; retentie 5 ani conform "
        "cerintei de sustenabilitate POCIDIF."
    )

    add_p(doc,
        "Aceste 5 elemente de inovatie au impact direct verificabil asupra:"
    )
    add_bullet(doc,
        "Punctaj P3.1 - element 3 EXCEPTIONAL abordare inovatoare si eficienta (10p)."
    )
    add_bullet(doc,
        "Conformitate DEMO video - cerintele 7 (chatbot AI configurare formulare), 8 "
        "(integrari terte), 9 (integrare GIS cu metadate locatie), 10 (lucru offline), "
        "11 (dosare relationate)."
    )
    add_bullet(doc,
        "Indicator de sustenabilitate POCIDIF - RCR11 raportare automatizata reduce riscul "
        "de auditat din partea Comisiei Europene in perioada de sustenabilitate."
    )

    doc.save(str(SRC))
    size = SRC.stat().st_size
    print(f"OK | {SRC.name} salvat: {size} bytes ({size/1024:.1f} KB)")


if __name__ == "__main__":
    main()

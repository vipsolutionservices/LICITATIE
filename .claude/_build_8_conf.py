"""Reconstrucție 8-Conformitate_specificatii.docx — Oferta SIDISVA
- Curățat VOGO TECHNOLOGY → <LIDER>
- Normalizat VOGO ENTERPRISE BUSINESS SUITE → VOGO Enterprise Suite
- Tabel sintetic extins de la 6 la 14 rânduri reprezentative (toate componentele SIDISVA)
- Aliniat cu Anexa F (6 coloane: Nr / Cap CDS / Cerință / Responsabil / Răspuns / Doc ref)
- Statistici breakdown cerințe pe capitol
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ---- Set page margins ----
for section in doc.sections:
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

# ---- Heading 1 ----
h1 = doc.add_heading('8. Conformitate cu specificațiile tehnice', level=1)

# ---- Intro paragraph ----
intro = doc.add_paragraph()
intro.add_run(
    'Capitolul răspunde la cerința IV.4.1 lit. e) din Fișa de date. '
    'Răspunsuri de tip „OK" / „100%" / „Soluția răspunde la cerință" NU sunt acceptate. '
    'Trimiterile doar la hyperlink-uri NU sunt acceptate. '
    '<LIDER>, în calitate de Ofertant principal, răspunde la fiecare cerință cu detalii concrete '
    'despre modul de îndeplinire, tehnologia folosită, capacitatea oferită și documentele '
    'justificative anexate. Versiunea integrală a matricei de conformitate este prezentată în '
    'Anexa F — Matricea de Conformitate.'
)

# ---- 8.1 Reguli formale ----
doc.add_heading('8.1 Reguli formale de redactare a răspunsurilor', level=2)

reguli = [
    'Fiecare cerință din Caietul de Sarcini primește un răspuns detaliat;',
    'Răspunsul indică concret modul în care cerința este sau va fi îndeplinită '
    '(tehnologie, configurație, capacitate, dovadă);',
    'Pentru fiecare răspuns sunt indicate: documentul anexă, pagina, paragraful relevant;',
    'Trimiterile către documente de la producători sunt însoțite de citarea textului concret '
    '(nu doar URL);',
    'Nu se admit răspunsuri prin simpla repetare a cerinței (cu schimbarea timpului verbal);',
    'Răspunsurile sunt asertive — formulări de tip „VOGO Enterprise Suite asigură…", '
    '„Soluția propusă de <LIDER> răspunde la cerință prin…".',
    'Fiecare răspuns este atribuit unui Responsabil din consorțiu (coloana „Responsabil" '
    'din Anexa F), pre-populat pe baza componentei tratate.',
]
for r in reguli:
    p = doc.add_paragraph(r, style='List Bullet')

# ---- 8.2 Structura Anexei F ----
doc.add_heading('8.2 Structura Anexei F — Matricea de Conformitate', level=2)

p82 = doc.add_paragraph()
p82.add_run(
    'Anexa F conține toate cele 1.294 cerințe extrase din Caietul de Sarcini, organizate în '
    '1.365 rânduri (incluzând 71 rânduri sub-header de secțiuni pentru navigare). '
    'Matricea este structurată pe 6 coloane:'
)

coloane = [
    ('Nr.', 'Numărul ordinal al cerinței (1 → 1.294).'),
    ('Cap. CDS', 'Capitolul din Caietul de Sarcini din care provine cerința '
                 '(ex. 3.4.3.2.4 — Soluție bază de date).'),
    ('Cerință (citat)', 'Textul integral al cerinței, copiat literal din Caietul de Sarcini.'),
    ('Responsabil', 'Membrul consorțiului care livrează cerința (pre-populat pe baza '
                    'componentei: <LIDER>, ZIPPER, VOGO, FURNIZOR LIMS etc.).'),
    ('Răspuns ofertant', 'Modul concret de îndeplinire — tehnologie, configurație, '
                         'capacitate, dovezi.'),
    ('Doc. referință', 'Documentul-suport: capitol din ofertă, anexă tehnică, fișa producător.'),
]
for nume, descr in coloane:
    p = doc.add_paragraph(style='List Bullet')
    r1 = p.add_run(f'{nume} — ')
    r1.bold = True
    p.add_run(descr)

# ---- 8.3 Tabel sintetic ----
doc.add_heading('8.3 Tabel sintetic de conformitate (extras reprezentativ)', level=2)

p83 = doc.add_paragraph()
p83.add_run(
    'Tabelul de mai jos prezintă un extras reprezentativ cu 14 cerințe cheie acoperind '
    'toate cele 14 componente SIDISVA. Versiunea integrală (toate cele 1.294 cerințe) este '
    'prezentată în Anexa F.'
)

# Tabel sintetic — 6 coloane aliniate cu Anexa F
header = ['Nr.', 'Cap. CDS', 'Cerință (citat scurt)', 'Responsabil',
          'Mod îndeplinire — răspuns sintetic', 'Doc. referință']

rows_data = [
    ('1', '3.4.1.1',
     'Sistem Cloud-Native găzduit în Cloud Guvernamental, compatibil containerizare',
     '<LIDER>',
     'Soluția este implementată ca microservicii containerizate (Docker) orchestrate '
     'prin Kubernetes, găzduită exclusiv în Cloud Guvernamental conform OUG 89/2022. '
     'Componentele aplicative respectă arhitectura Cloud-Native (stateless, '
     '12-Factor App, auto-scaling).',
     'Cap. 5.1, 5.2 ofertă; Anexa arhitectură'),

    ('2', '3.4.1.2',
     'Securitate cibernetică — NIS2, OUG 155/2024, Legea 354/2022, GDPR',
     '<LIDER> + FURNIZOR SECURITATE',
     'Conformitate completă cu Directiva (UE) 2022/2555 (NIS2) transpusă prin OUG 155/2024 '
     '(plan IR, SOC 24×7, raportare incidente CSIRT-RO); Lege 354/2022 — toți producătorii '
     'aleși au origine UE/SUA/IL (lista detaliată în Anexa B); GDPR — DPO desemnat, '
     'registru prelucrări, DPIA livrabilă; defense-in-depth pe 7 niveluri.',
     'Cap. 10 ofertă; Anexa H — analiză juridică'),

    ('3', '3.4.3.3.1',
     'DMS — performanță optimă pentru minim 5.301 utilizatori interni fără restricții',
     'ZIPPER',
     'ZIPPER DMS — soluție COTS cu licență perpetuă, număr nelimitat de utilizatori, '
     'cod sursă integral livrat conform cap. 12. Stress-testat la 7.500 useri concurenți. '
     'Arhitectură multi-nivel (web + aplicație + BD); client web + PWA + mobile.',
     'Cap. 4.2.2 ofertă; fișa tehnică ZIPPER'),

    ('4', '3.4.3.3.2',
     'Portal Servicii Publice — 56 servicii electronice; min 5.500 utilizatori interni + '
     'număr nelimitat utilizatori externi',
     'VOGO',
     'VOGO Enterprise Suite — Portal: catalog complet cu 56+ servicii electronice; '
     'autentificare ROeID + eIDAS + LDAP intern prin Keycloak; plăți online prin Ghișeul.ro; '
     'accesibilitate WCAG 2.1 AA; responsive (desktop/tablet/mobile); SSO cu DMS, LIMS, BND.',
     'Cap. 4.2.3 ofertă; Anexa wireframes'),

    ('5', '3.4.3.3.2.1',
     'Chatbot — modul extern bazat pe inteligență artificială; integrare LLAMA/OpenAI/Gemini; '
     'recunoaștere text și voce',
     'VOGO',
     'VOGO Enterprise Suite — Chatbot: motor NLP RO on-premise (fără dependență externă), '
     'cu modul LLM integrabil cu LLAMA, OpenAI și Gemini; canale chat web + WhatsApp + '
     'voce (TTS/STT); recunoaștere intent + entități; dialog management; asistă configurare '
     'formulare drag-and-drop.',
     'Cap. 4.2.4 ofertă; demonstrație video DEMO'),

    ('6', '3.4.3.3.2.2',
     'Aplicație mobilă cetățeni — dezvoltată nativ Android + iOS; publicare în magazine; '
     'raportare animale sălbatice + focare boli',
     'VOGO',
     'VOGO Enterprise Suite — Aplicație Mobilă: dezvoltare nativă Swift (iOS) + '
     'Kotlin (Android) + PWA companion; publicare TestFlight + App Store + Google Play; '
     'GPS, cameră foto cu watermark, lucru offline cu sincronizare la reconectare; '
     'integrare GIS central (WMS/WFS) + Keycloak (OAuth2 PKCE).',
     'Cap. 4.2.5 ofertă; manifest store + raport pentest MASVS L1'),

    ('7', '3.4.3.2.8',
     'Sistem GIS — gestionarea, analiza și vizualizarea datelor spațiale; hărți tematice '
     'focare boli + cartografiere unități',
     'FURNIZOR GIS',
     'Sistem GIS dedicat — server central + 50 editori; layere vector + raster; '
     'WMS/WFS publicat către componentele SIDISVA (Portal, App mobilă, BND-SNIIA); '
     'integrare APIA pentru coordonate exploatații; analize spațiale (buffer, intersecții, '
     'proximitate); editare directă pe hartă.',
     'Cap. 4.2.7 ofertă; fișă tehnică FURNIZOR GIS'),

    ('8', '3.4.3.2.6',
     'BI — 50 utilizatori; Data Warehouse + raportare + dashboards drag-and-drop',
     'FURNIZOR BI',
     'Microsoft Power BI + SSRS + SSAS — 50 useri named; surse date relaționale (MS SQL) + '
     'NoSQL (Elasticsearch); dashboards drag-and-drop; rapoarte programate + ad-hoc; '
     'înlocuiește 50+ machete .xls actuale; export PDF/Excel; integrare ETL prin SSIS.',
     'Cap. 4.2.8 ofertă; licențe Microsoft (BYOL Azure)'),

    ('9', '3.4.3.2.7',
     'IAM — autentificare unificată: utilizator/parolă, certificate X.509, Windows Native, '
     'SAML, multi-factor',
     '<LIDER>',
     'Keycloak Enterprise (open-source, US/Red Hat) — 150 utilizatori interni + 185.000 '
     'utilizatori portal/an; broker federat pentru ROeID + eIDAS + LDAP intern; OAuth2 / '
     'OIDC / SAML 2.0; MFA (TOTP, SMS, FIDO2); SSO pentru toate componentele SIDISVA.',
     'Cap. 5.3 ofertă; Anexa arhitectură IAM'),

    ('10', '3.4.3.4.1.4',
     'SIEM — colectare/stocare log-uri; corelare automată; alertare; instalare ca soluție SW',
     'FURNIZOR SECURITATE',
     'Splunk Enterprise Security sau IBM QRadar — cluster cu retention IOC 3 ani; '
     'stocare log-uri în NoSQL (Elasticsearch); >1000 rapoarte predefinite; SOC 24×7 '
     'cu playbook-uri IR; raportare incidente NIS2 către CSIRT-RO conform OUG 155/2024.',
     'Cap. 10 ofertă; fișa tehnică SIEM'),

    ('11', '3.4.3.4.1.1',
     'WAF — appliance virtual; HA; protecție OWASP Top 10; NU licențiat per useri',
     'FURNIZOR SECURITATE',
     'F5 Advanced WAF sau Imperva sau FortiWeb — appliance virtual deployabil pe '
     'VMware/Hyper-V/KVM; HA activ-pasiv 8 vCPU + 16GB RAM/instanță; throughput ≥3 Gbps HTTP; '
     '≥64 domenii admin; protecție OWASP Top 10 + bot mitigation + API security; licență '
     'pe throughput, nu pe utilizatori.',
     'Anexa B item 1; fișa tehnică producător'),

    ('12', '3.4.3.4.2.6',
     'Laptop — consum redus DNSH (max 20Wh în veghe); 100 buc pentru ANSVSA',
     '<LIDER>',
     'Laptop EnergyStar 8.0 + EU Ecolabel (sau echivalent — Blue Angel / Nordic Swan); '
     'consum măsurat în veghe < 20 Wh per raport testare anexat; CPU low-power '
     'din generație curentă; SSD; ecran low-power 14"; ambalaje reciclabile; livrare '
     'cu flotă electrică/hibridă conform 4.2 DNSH.',
     'Cap. 11 ofertă (DNSH); Anexa B item 12; raport test producător'),

    ('13', '3.4.3.4.4',
     'Integrare nativă ROeID + eIDAS + 2FA SMS (cerință DEMO eliminatorie)',
     '<LIDER> + VOGO',
     'Integrare nativă ROeID prin OAuth2 + SAML (broker Keycloak); eIDAS Connector cu '
     'noduri Romania + UE; 2FA SMS via SMSC dedicat (operatori RO); demonstrată live în '
     'DEMO video + DEMO sediu ANSVSA (cerință eliminatorie nr. 3 din lista de 33).',
     'Cap. 14 DEMO ofertă; Anexa integrări guvernamentale'),

    ('14', '3.4.3.3.4 + Cap. 12',
     'Drepturi IP — licențe perpetue + cod sursă integral; proprietate intelectuală '
     'către Beneficiar pentru dezvoltări/customizări',
     '<LIDER>',
     '<LIDER>, în nume propriu și al partenerilor consorțiului, ACCEPTĂ expres și fără '
     'rezerve condițiile cap. 12: toate dezvoltările/customizările trec în proprietatea '
     'ANSVSA; licențe perpetue + cod sursă pentru toate componentele aplicative; IP-ul '
     'produselor COTS preexistente (ZIPPER, VOGO Enterprise Suite, FURNIZOR LIMS) rămâne '
     'la producători, dar Beneficiarul primește licență perpetuă + cod sursă.',
     'Cap. 5.4 ofertă; Declarații IP anexate (cap. 15)'),

    ('…', 'restul',
     '[Continuat în Anexa F — toate cele 1.294 cerințe extrase]',
     'Conform pre-populare',
     'Răspuns detaliat per cerință',
     'Anexa F'),
]

table = doc.add_table(rows=1 + len(rows_data), cols=6)
table.style = 'Light Grid Accent 1'

# Header
for i, h in enumerate(header):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(9)

# Data rows
for ri, data in enumerate(rows_data, start=1):
    for ci, val in enumerate(data):
        cell = table.rows[ri].cells[ci]
        cell.text = val
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)

# Column widths (approximative; Word respects total)
widths = [Cm(0.8), Cm(2.0), Cm(4.5), Cm(2.5), Cm(6.0), Cm(2.7)]
for row in table.rows:
    for ci, w in enumerate(widths):
        row.cells[ci].width = w

# ---- 8.4 Statistici cerințe pe capitol ----
doc.add_heading('8.4 Distribuție cerințe Anexa F pe capitole', level=2)

p84 = doc.add_paragraph()
p84.add_run(
    'Cele 1.294 cerințe extrase din Caietul de Sarcini sunt repartizate astfel:'
)

stat_data = [
    ('3.4 — Cadru general serviciilor', '11'),
    ('3.4.1 — Cerințe nefuncționale', '42'),
    ('3.4.2 — Cerințe funcționale (servicii)', '118'),
    ('3.4.3 — Cerințe tehnice componente', '980'),
    ('3.4.4 — Echipamente hardware', '97'),
    ('3.4.5-3.4.9 — Cerințe diverse (raportare, instruire, suport)', '46'),
    ('TOTAL', '1.294'),
]

stat_table = doc.add_table(rows=1 + len(stat_data), cols=2)
stat_table.style = 'Light Grid Accent 1'

# Header
hdr = stat_table.rows[0].cells
hdr[0].text = 'Capitol / Secțiune'
hdr[1].text = 'Nr. cerințe'
for c in hdr:
    for p in c.paragraphs:
        for r in p.runs:
            r.bold = True

for ri, (cap, nr) in enumerate(stat_data, start=1):
    cells = stat_table.rows[ri].cells
    cells[0].text = cap
    cells[1].text = nr
    if cap == 'TOTAL':
        for c in cells:
            for p in c.paragraphs:
                for r in p.runs:
                    r.bold = True

# Column widths for stat table
for row in stat_table.rows:
    row.cells[0].width = Cm(12)
    row.cells[1].width = Cm(3)

# ---- 8.5 Trasabilitate ----
doc.add_heading('8.5 Trasabilitate cerințe → soluție', level=2)

p85 = doc.add_paragraph()
p85.add_run(
    'Trasabilitatea cerințelor pe traseul Caiet de Sarcini → Ofertă Tehnică → Document de '
    'Analiză → Scenarii de Testare → Recepție Finală este asigurată de <LIDER>, în calitatea '
    'sa de Ofertant principal, prin Matricea de Trasabilitate a Cerințelor (RTM) — instrument '
    'central al managementului calității proiectului. Această matrice permite ANSVSA să '
    'verifice oricând, în orice etapă a proiectului, modul în care fiecare cerință din '
    'Caietul de Sarcini este îndeplinită concret.'
)

p85b = doc.add_paragraph()
p85b.add_run(
    'RTM este menținut activ pe toată durata contractului (18 luni implementare + 36 luni '
    'garanție) și este livrabil obligatoriu la fiecare recepție parțială și la recepția '
    'finală. Forma electronică (export Excel + acces read-only în portalul de management '
    'al proiectului) este pusă la dispoziția responsabililor ANSVSA.'
)

# ---- Save ----
doc.save(r'8-Conformitate_specificatii.docx')
print('OK — 8-Conformitate_specificatii.docx scris.')

# Verificare
from docx import Document
d2 = Document(r'8-Conformitate_specificatii.docx')
print(f'Verificare: {len(d2.paragraphs)} paragrafe, {len(d2.tables)} tabele')
for ti, t in enumerate(d2.tables):
    print(f'  Tabel {ti}: {len(t.rows)} rânduri × {len(t.columns)} coloane')

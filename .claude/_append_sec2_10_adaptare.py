"""
Appendă secțiunea 2.10 "Adaptarea metodologiei la specificul SIDISVA"
la 2-Abordare_metodologie.docx.

Sub-secțiuni:
 - 2.10.1 Cele 14 componente — mapare metodă / echipă / furnizor (TABEL)
 - 2.10.2 Stack-ul tehnologic concret propus (TABEL)
 - 2.10.3 Drumul critic al proiectului (jaloane fixe cap. 6+7.3)
 - 2.10.4 Acceptarea celor 4 ipoteze din cap. 4.1
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt

SRC = Path(r"C:\Users\adria\Documents\Claude\Projects\Licitatie 1\doc-output\oferta\2-Abordare_metodologie.docx")


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    return p


def add_p(doc, text):
    p = doc.add_paragraph()
    p.add_run(text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    p.add_run("• " + text)
    return p


def add_table(doc, header, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    try:
        t.style = "Table Grid"
    except KeyError:
        pass
    hdr = t.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
    for ri, row in enumerate(rows, start=1):
        cells = t.rows[ri].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    return t


def main():
    if not SRC.exists():
        raise SystemExit(f"NOT FOUND: {SRC}")
    doc = Document(str(SRC))

    # ============ 2.10 Adaptarea metodologiei la specificul SIDISVA ============
    add_h2(doc, "2.10 Adaptarea metodologiei la specificul SIDISVA")

    add_p(doc,
        "SIDISVA este un proiect de complexitate ridicată care presupune dezvoltarea, configurarea "
        "și integrarea a 14 componente funcționale interconectate, în 18 luni calendaristice, "
        "pentru 46 instituții publice (ANSVSA + 42 DSVSA județene + 3 institute subordonate), "
        "circa 185.000 utilizatori unici anuali ai portalului și un buget eligibil de "
        "95.271.200,60 lei. Metodologia generală PMI/PMBoK + BABOK + Agile-Scrum SAFe expusă în "
        "secțiunile 2.2 anterioare a fost adaptată concret la următoarele specificități ale "
        "contractului, conform cerinței calificativului EXCEPȚIONAL la subfactorul 3.1 din "
        "metodologia de evaluare (cap. 13 din Caietul de Sarcini)."
    )

    # ---- 2.10.1 ----
    add_h3(doc, "2.10.1 Cele 14 componente SIDISVA — mapare metodă, echipă și furnizor")
    add_p(doc,
        "Pentru fiecare dintre cele 14 componente funcționale enumerate la cap. 3.4.2 din Caietul "
        "de Sarcini, consorțiul a definit explicit furnizorul responsabil, produsul/tehnologia "
        "propusă și strategia de implementare (configurare COTS vs. dezvoltare custom vs. mix). "
        "Această mapare este sursa adevărului pentru toate planurile operaționale subsidiare "
        "(WBS, plan resurse, plan de testare) și se reflectă identic în Anexa F — Matricea de "
        "conformitate (1.294 cerințe atomice mapate la responsabil)."
    )
    add_table(doc,
        header=["Componenta SIDISVA (cap. CdS)", "Furnizor", "Produs / Tehnologie", "Strategia de implementare"],
        rows=[
            ["1. LIMS — Laboratory Info Mgmt (3.4.2.1)", "<FURNIZOR LIMS>", "LIMS COTS + cod sursă",
             "Configurare COTS + dezvoltări specifice ANSVSA (HL7 prin Mirth Connect, integrare echipamente IoT)"],
            ["2. DMS — Document Mgmt (3.4.2.2)", "<FURNIZOR DMS>", "ZIPPER DMS + cod sursă",
             "Configurare nucleu + dezvoltări fluxuri ANSVSA (registre, fluxuri aprobare, integrare pad-uri semnătură olografă)"],
            ["3. Mgmt scheme intercomparare (3.4.2.3)", "<FURNIZOR DMS> + <FURNIZOR LIMS>", "DMS + LIMS configurate",
             "Configurare în DMS și LIMS, fără cod nou — folosit fluxuri standard"],
            ["4. GIS (3.4.2.4)", "<FURNIZOR GIS>", "Soluție GIS dedicată",
             "Configurare + layere ANSVSA + integrare APIA pentru coordonate exploatații"],
            ["5. Raportare/BI/Dashboard (3.4.2.5)", "<FURNIZOR BI> + <FURNIZOR ETL>", "MS Power BI + SSRS + SSAS / MS SSIS",
             "Configurare Data Warehouse + cube-uri SSAS + dashboards Power BI + ETL din >50 surse"],
            ["6. Portal Servicii Publice (3.4.2.6)", "<FURNIZOR PORTAL>", "Enterprise Suite (Portal)",
             "Configurare + dezvoltare 56 servicii electronice + integrare ROeID, eIDAS, ONRC, Ghiseul.ro"],
            ["7. BND-SNIIA + mobile vet + mobile fermieri (3.4.2.7)", "<LIDER> + <FURNIZOR APP MOBILĂ>", "Custom + Enterprise Suite (Mobile)",
             "Dezvoltare full BND-SNIIA + apps native iOS/Android (Swift/SwiftUI + Kotlin/Compose) + offline + AI/ML"],
            ["8. Catagrafie/cartografiere unități (3.4.2.8)", "<FURNIZOR DMS> + <FURNIZOR GIS>", "DMS + GIS",
             "Configurare DMS + layere GIS + import date xls existente"],
            ["9. Supraveghere/control/anchete (3.4.2.9)", "<FURNIZOR DMS>", "ZIPPER DMS + pad semnătură olografă",
             "Configurare DMS + integrare pad-uri USB + dosare unitate partajate"],
            ["10. Mgmt instruire personal laboratoare (3.4.2.10)", "<FURNIZOR DMS>", "ZIPPER DMS configurat",
             "Configurare DMS — calendar, înscrieri, feedback, materiale"],
            ["11. App cetățeni — raportare animale sălbatice (3.4.2.11)", "<FURNIZOR APP MOBILĂ>", "Enterprise Suite (Mobile)",
             "Native iOS + Android + GPS + foto + GIS + ROeID"],
            ["12. Autorizare/acreditare/desemnare (3.4.2.12)", "<FURNIZOR DMS> + <FURNIZOR PORTAL>", "DMS + Portal",
             "Configurare DMS pentru dosare + servicii Portal pentru depunere online + semnătură digitală"],
            ["13. Contorizare folosire servicii publice — RCR11 (3.4.2.13)", "<FURNIZOR PORTAL>", "Enterprise Suite (modul telemetrie)",
             "Configurare modul telemetrie + integrare Data Warehouse + dashboard auditor"],
            ["14. Integrări externe (3.4.2.14)", "<LIDER> + <FURNIZOR ESB>", "Custom + Oracle Service Bus",
             "API-uri pe Oracle Service Bus + mock-up-uri pentru sisteme guvernamentale neimplementate (PNI, PJN, PCUe)"],
        ]
    )
    add_p(doc,
        "Observație: componentele 3 (scheme intercomparare), 8 (catagrafie), 10 (instruire lab), "
        "12 (autorizare/acreditare) sunt explicit prevăzute în Caietul de Sarcini să fie "
        "implementate folosind soluția DMS — adică prin configurare, nu dezvoltare. Acest fapt "
        "reduce semnificativ efortul de dezvoltare net și riscul tehnic asociat și permite "
        "concentrarea resurselor pe componentele inovatoare (BND-SNIIA cu AI/ML, app mobilă, "
        "integrări guvernamentale)."
    )

    # ---- 2.10.2 ----
    add_h3(doc, "2.10.2 Stack-ul tehnologic concret propus")
    add_p(doc,
        "Pentru a elimina orice ambiguitate privind soluția tehnică ofertată — element esențial "
        "pentru calificativul EXCEPȚIONAL la subfactorul 3.1 — tabelul de mai jos prezintă "
        "produsele și tehnologiile propuse pe fiecare nivel arhitectural. Toate cantitățile și "
        "sizing-urile respectă cerințele exacte din cap. 3.4.3.2 + 3.4.3.4 al Caietului de "
        "Sarcini. Toate produsele COTS sunt licențiate BYOL (Bring Your Own License) prin "
        "Microsoft Azure în Cloud Guvernamental, conform cap. 3.4.3.1."
    )
    add_table(doc,
        header=["Strat arhitectural", "Componentă", "Produs / Tehnologie", "Sizing"],
        rows=[
            ["A. Infrastructură", "OS server Linux", "Red Hat Enterprise Linux 9 / Oracle Linux 9", "20-30 VM"],
            ["A. Infrastructură", "OS server Windows", "Windows Server 2022 Datacenter", "Min. 4 + 6-10 VM (inclusiv 4 honeypot)"],
            ["A. Infrastructură", "Web server", "NGINX Plus", "2 noduri × 16 cores, cluster activ-activ"],
            ["A. Infrastructură", "App server", "Microsoft IIS", "2 noduri × 16 cores"],
            ["A. Infrastructură", "SGBD principal", "Microsoft SQL Server Enterprise (TDE, cache fusion)", "2 noduri × 8 cores"],
            ["A. Infrastructură", "SGBD NoSQL", "Elasticsearch (cluster 3 noduri)", "Log-uri, retenție conform NIS2"],
            ["A. Infrastructură", "ETL", "Microsoft SSIS", "16 cores"],
            ["A. Infrastructură", "BI", "Microsoft Power BI + SSRS + SSAS", "50 utilizatori"],
            ["A. Infrastructură", "IAM", "Keycloak Enterprise", "150 utilizatori interni + 185.000 portal"],
            ["A. Infrastructură", "GIS", "Soluție GIS dedicată — <FURNIZOR GIS>", "1 server + 50 editori"],
            ["A. Infrastructură", "ESB", "Oracle Service Bus", "Cluster activ-pasiv, 16 cores/nod"],
            ["B. Software aplicativ", "DMS", "ZIPPER DMS", "Utilizatori nelimitați + cod sursă"],
            ["B. Software aplicativ", "Portal", "Enterprise Suite (modul Portal)", "Nelimitat + perpetuu"],
            ["B. Software aplicativ", "Chatbot AI / NLP RO", "Enterprise Suite (modul Chatbot)", "1 instanță"],
            ["B. Software aplicativ", "App mobilă", "Enterprise Suite (modul Mobile)", "iOS + Android + PWA, dev accounts Apple 99$/an + Google free"],
            ["B. Software aplicativ", "LIMS", "<FURNIZOR LIMS> COTS", "Utilizatori nelimitați + cod sursă + 3 ani L2/L3"],
            ["B. Software aplicativ", "Integrare HL7", "Mirth Connect (NextGen)", "1 instanță"],
            ["C. Securitate", "WAF", "F5 / Imperva / FortiWeb", "2 buc (OWASP Top 10)"],
            ["C. Securitate", "Honeypot", "FortiDeceptor", "1 + 4 instanțe Win × 4 vCPU"],
            ["C. Securitate", "NMS / NAC", "Cisco DNA / ClearPass", "Min. 600 dispozitive"],
            ["C. Securitate", "SIEM", "Splunk Enterprise Security / IBM QRadar", "1 cluster (IOC 3 ani)"],
            ["C. Securitate", "Email Security", "Cisco IronPort", "2 buc (filtrare + sandbox)"],
            ["C. Securitate", "NGFW centru", "FortiGate / Palo Alto", "2 buc, redundant, 10 VDOM"],
            ["C. Securitate", "NGFW locații", "FortiGate 100F", "90 buc (45 locații × 2)"],
            ["C. Securitate", "Antivirus EDR", "CrowdStrike / SentinelOne", "436 + 50 endpoints"],
            ["D. Productivity", "MS Office laptopuri", "Microsoft Office Home & Business 2024 OEM", "100 buc"],
            ["D. Productivity", "MS Office complete teren", "Microsoft Office Home & Business 2024 OEM", "336 buc (8 × 42 DSVSA)"],
        ]
    )
    add_p(doc,
        "Conformitate cu Legea 354/2022 privind protecția sistemelor informatice ale autorităților "
        "publice (în contextul invaziei declanșate de Federația Rusă împotriva Ucrainei): "
        "stack-ul propus include exclusiv producători din state cu jurisdicții compatibile "
        "(SUA, UE, Israel, Țările de Jos) — nicio componentă din state cu restricții impuse de "
        "L 354/2022. Antivirus-ul EDR (CrowdStrike / SentinelOne) respectă această cerință "
        "fundamentală, iar acest fapt va fi documentat explicit în Propunerea Tehnică, "
        "secțiunea 10 — Securitate informatică."
    )

    # ---- 2.10.3 ----
    add_h3(doc, "2.10.3 Drumul critic al proiectului și planul de paralelizare")
    add_p(doc,
        "Pe baza jaloanelor obligatorii din cap. 6 + 7.3 ale Caietului de Sarcini (Analiză — luna 6, "
        "Proiectare — luna 7, Livrare HW/SW + PIP — luna 12, Dezvoltare/configurare/integrare — "
        "luna 15, Migrare — luna 16, Implementare — luna 17, Testare + Instruire + Go-live — luna 18, "
        "fără a depăși 21.05.2027), drumul critic identificat pentru SIDISVA este următorul:"
    )
    add_table(doc,
        header=["Lună", "Activitate critică", "Livrabile principale", "Dependențe"],
        rows=[
            ["L1-L6", "Analiză sistem (toate cele 14 componente)",
             "Document de Analiză aprobat (per componentă), specificații funcționale, modele BPMN/UML",
             "Acces ANSVSA + utilizatori cheie pentru workshop-uri (cf. ipoteză c)"],
            ["L6-L7", "Proiectare sistem",
             "Document de Proiectare aprobat, modele de date logice + fizice, arhitectură de detaliu, ADR-uri",
             "Validare ANSVSA (max 10 zile)"],
            ["L4-L12", "Livrare/instalare/configurare HW + SW + PIP",
             "Medii Dev/Test/Stage/Prod configurate în Cloud Guvernamental, monitoring activ",
             "Activare cont Cloud Guvernamental Azure RO; aprovizionare HW endpoints"],
            ["L7-L15", "Dezvoltare/configurare/integrare (drumul critic principal)",
             "Cod sursă livrat (DMS configurat, LIMS configurat, BND-SNIIA dezvoltat, Portal/Chatbot/Mobile dezvoltate, GIS configurat, BI configurat, integrări API live)",
             "Drum critic real: BND-SNIIA + Integrări guvernamentale (vezi mai jos)"],
            ["L15-L16", "Migrare date",
             "Date migrate din vechiul SNIIA, LIMS, registre xls, BND existent — validare calitate prin AI/ML",
             "Disponibilitatea exporturilor de la sistemele actuale"],
            ["L16-L17", "Implementare (rollout pe medii)",
             "Toate componentele instalate și configurate în mediul de producție Cloud Guvernamental",
             "Migrare validată"],
            ["L17-L18", "Testare integrată + UAT",
             "Rapoarte testare (funcțional, integrare, performanță, securitate Pen-Test, DR), UAT acceptat",
             "PIP HW completă, dev finalizat"],
            ["L18", "Instruire utilizatori + Go-live",
             "144 utilizatori + 3 administratori instruiți; sistem live cu suport on-call",
             "UAT semnat fără obiecții"],
        ]
    )

    add_p(doc,
        "Componentele de pe drumul critic — și deci cu prioritate maximă în alocarea resurselor "
        "consorțiului — sunt:"
    )
    add_bullet(doc,
        "BND-SNIIA (cap. 3.4.2.7) — cea mai complexă componentă, integrează 6+ surse de date "
        "(SNIIA actual, APIA, ANARZ, Colegiul Medicilor Veterinari, exploatații, evenimente "
        "animale); 2 aplicații mobile native (medici vet + fermieri) cu lucru offline; algoritmi "
        "ML/AI obligatorii pentru calitate date + fraud detection (cf. cerinței explicite din "
        "cap. 3.4.2.7); aproximativ 400.000 fermieri + 2.600 medici vet concesionari beneficiari."
    )
    add_bullet(doc,
        "Integrările cu sisteme guvernamentale (cap. 3.4.2.14) — PSCID-ROeID, Nodul eIDAS, "
        "Platforma Națională de Interoperabilitate (PNI), PCUe/PDURo, Platforma de Jurnalizare "
        "și Notificare (PJN), ONRC, APIA, ANCPI, ANARZ, Colegiul Medicilor Veterinari, SNIIA, "
        "Ghiseul.ro. Unele dintre acestea sunt în curs de implementare la nivel guvernamental — "
        "strategia de mitigare prin mock-up este detaliată în secțiunea 2.13.3."
    )
    add_bullet(doc,
        "LIMS + migrarea datelor din sistemele de laborator actuale — implică integrare cu "
        "echipamente IoT, conformitate SR EN ISO/CEI 17043 (RENAR), suport HL7 prin Mirth Connect "
        "și un volum foarte mare (280.000 cereri analize/an, 33.000 utilizatori unici)."
    )

    add_p(doc,
        "Plan de paralelizare pentru încadrarea în 18 luni, în pofida volumului mare de "
        "dezvoltare aferent celor 14 componente:"
    )
    add_bullet(doc,
        "Analiza (L1-L6) este structurată pe 9 sub-echipe paralele, fiecare furnizor responsabil "
        "pentru analiza propriei componente, sincronizate săptămânal de Analyst Business "
        "<LIDER> în Technical Committee."
    )
    add_bullet(doc,
        "<FURNIZOR ESB> începe configurarea Oracle Service Bus din luna 4 (în paralel cu "
        "analiza), pentru ca până în luna 8 backbone-ul de integrare să fie disponibil și "
        "celelalte componente să poată dezvolta direct pe API-uri publicate."
    )
    add_bullet(doc,
        "Echipele DMS, Portal, Chatbot, GIS, BI rulează dezvoltare paralelă cu BND-SNIIA, "
        "sincronizate la milestone-uri majore (M6 = analiză finalizată, M12 = HW+PIP gata, "
        "M15 = dezvoltare finalizată)."
    )
    add_bullet(doc,
        "Aplicația mobilă (3 platforme: iOS nativ, Android nativ, PWA) — dezvoltare paralelă "
        "cu echipa Portal, începută în luna 4, conform planului de efort estimat la 830 om-zile "
        "(692 dezvoltare + 138 contingency 20%) — vezi planul de implementare detaliat din "
        "secțiunea 3 a propunerii tehnice."
    )
    add_bullet(doc,
        "Securitatea cibernetică (WAF, SIEM, NGFW, EDR, NAC, Honeypot, Email Security) — "
        "configurare începută din luna 8 (după ce mediile țintă sunt disponibile), cu validare "
        "în testarea Pen-Test din luna 18."
    )

    add_p(doc,
        "Riscul cel mai mare al drumului critic — disponibilitatea sistemelor guvernamentale "
        "țintă (PNI, eIDAS extins, PJN) la momentul integrării — este atenuat prin clauza "
        "explicită din cap. 3.4.2.14 al Caietului de Sarcini: dacă sistemele guvernamentale "
        "țintă nu sunt finalizate până la finalul perioadei de implementare, integrarea se "
        "realizează cu mock-up-uri OpenAPI și apoi cu sistemele reale în perioada de garanție, "
        "fără cost suplimentar pentru ANSVSA. Această clauză transformă riscul tehnologic extern "
        "într-un risc gestionabil contractual."
    )

    # ---- 2.10.4 ----
    add_h3(doc, "2.10.4 Acceptarea ipotezelor din cap. 4.1 al Caietului de Sarcini")
    add_p(doc,
        "<LIDER> și ceilalți membri ai consorțiului acceptă integral cele 4 ipoteze enunțate "
        "de Autoritatea Contractantă în cap. 4.1 al Caietului de Sarcini, cu următoarele "
        "precizări operaționale și măsuri de adaptare:"
    )
    add_table(doc,
        header=["Ipoteză (cap. 4.1 CdS)", "Poziția consorțiului", "Măsuri operaționale"],
        rows=[
            ["a) Conținutul serviciilor descris în Caietul de Sarcini reflectă forma și detaliile cunoscute Achizitorului la data realizării. Specificațiile de integrare cu sisteme terțe se referă la forma actuală a acestor sisteme.",
             "ACCEPTATĂ",
             "Mecanism formal de Change Request pentru orice modificare a sistemelor țintă identificată în faza de analiză; ADR-uri actualizate la fiecare modificare API a sistemelor partenere; anti-corruption layer pe Oracle Service Bus pentru izolarea SIDISVA de modificările sistemelor externe."],
            ["b) Nu se prevăd schimbări ale cadrului instituțional și legal care să afecteze major implementarea.",
             "ACCEPTATĂ",
             "Monitorizare continuă a actelor normative cu impact (NIS2 / OUG 155/2024, L 242/2022 PNI, L 354/2022, GDPR); raportare lunară către ANSVSA în Steering Committee a oricărei modificări legislative cu potențial impact; propunere CR pentru adaptare."],
            ["c) Toate informațiile relevante și disponibile la Achizitor vor fi puse la dispoziția Prestatorului într-un timp rezonabil.",
             "ACCEPTATĂ cu propunere de SLA reciproc",
             "Consorțiul propune ca în acordul de lucru să fie definit un SLA reciproc pentru livrarea informațiilor de către ANSVSA: maxim 5 zile lucrătoare pentru cereri standard, 10 zile lucrătoare pentru cereri complexe; orice depășire = trigger pentru extensie de termen, fără penalizare consorțiu."],
            ["d) Prestatorul va semna acord de confidențialitate și de garantare GDPR la semnarea Contractului.",
             "ACCEPTATĂ integral",
             "Consorțiul deține deja proceduri interne de confidențialitate și prelucrare date conforme ISO/IEC 27001:2022 și ISO/IEC 27701:2019 (PIMS); DPO desemnat pre-contractual; acordul de confidențialitate va fi semnat de toți cei 9 membri ai consorțiului, cu răspundere solidară."],
        ]
    )

    doc.save(str(SRC))
    size = SRC.stat().st_size
    print(f"OK | {SRC.name} salvat: {size} bytes ({size/1024:.1f} KB)")


if __name__ == "__main__":
    main()

"""
Genereaza matrice_conformitate_pe_capitole.docx (v4 — abordare surgicala).

Strategie: copiez EXACT structura din anexa_f_conformitate - old.docx
si fac doar 2 modificari minime, pastrand formatare/layout/coloane/section headers:

 1. Replace VOGO TECHNOLOGY / VOGO -> <LIDER> in TOT textul (paragrafe + tabel).
    Pastreaza formatarea prin manipulare la nivel de run.

 2. Pre-completare coloana 3 "Raspuns ofertant" cu produsul din Sinteza
    (MAP cap -> produs). Coloana 4 "Document referinta" ramane goala.

Rezultat: document IDENTIC cu sursa, dar cu raspunsuri pre-completate pe baza Sinteza.
"""
import re
import sys
import shutil
from pathlib import Path

sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn

ROOT = Path(r"C:\Users\adria\Documents\Claude\Projects\Licitatie 1\doc-output\oferta")
SRC = ROOT / "anexa_f_conformitate - old.docx"
OUT = ROOT / "matrice_conformitate_pe_capitole.docx"


# MAP cap CdS -> Produs din Sinteza, cu <LIDER> in loc de VOGO TECHNOLOGY.
# Pentru "VOGO Enterprise Suite" pastram numele de PRODUS asa cum apare in Sinteza,
# pentru ca este nume de produs nu de companie (acest fisier va trece prin replace
# VOGO->LIDER ulterior, deci scriu deja "<LIDER> Enterprise Suite" pentru consistenta).
MAP = {
    '3.4':         '<LIDER> (coordonare integrator)',
    '3.4.1.1':     '<LIDER> (arhitect sistem)',
    '3.4.1.2':     'Toate produsele C. Securitate (WAF / Honeypot / NAC / SIEM / Email / NGFW)',
    '3.4.1.3':     'Toate produsele C. Securitate + <LIDER> (conformitate legala)',
    '3.4.2.1':     'FURNIZOR LIMS COTS',
    '3.4.2.2':     'ZIPPER',
    '3.4.2.3':     'ZIPPER / <LIDER> Enterprise Suite',
    '3.4.2.5':     'Microsoft Power BI / SSRS / SSAS + Microsoft SSIS',
    '3.4.2.6':     '<LIDER> Enterprise Suite',
    '3.4.2.7':     '<LIDER> Enterprise Suite',
    '3.4.2.8':     'FURNIZOR GIS + <LIDER> Enterprise Suite',
    '3.4.2.9':     '<LIDER> Enterprise Suite',
    '3.4.2.10':    '<LIDER> Enterprise Suite',
    '3.4.2.11':    '<LIDER> Enterprise Suite (aplicatie mobila)',
    '3.4.2.12':    '<LIDER> Enterprise Suite',
    '3.4.2.14':    'Oracle Service Bus + Mirth Connect + <LIDER> Enterprise Suite',
    '3.4.3':       '<LIDER> (arhitect sistem)',
    '3.4.3.1':     'RHEL 9 / Oracle Linux 9 + Win Server 2022 DC (Cloud Guvernamental)',
    '3.4.3.2.1':   'NGINX Plus',
    '3.4.3.2.2':   'Microsoft IIS',
    '3.4.3.2.3':   'RHEL 9 / Oracle Linux 9 + Win Server 2022 DC',
    '3.4.3.2.4':   'Microsoft SQL Server Enterprise + Elasticsearch',
    '3.4.3.2.5':   'Microsoft SSIS',
    '3.4.3.2.6':   'Microsoft Power BI / SSRS / SSAS',
    '3.4.3.2.7':   'Keycloak Enterprise',
    '3.4.3.2.8':   'FURNIZOR GIS',
    '3.4.3.2.9':   'Oracle Service Bus',
    '3.4.3.3.1':   'ZIPPER',
    '3.4.3.3.1.1': 'ZIPPER',
    '3.4.3.3.1.2': 'ZIPPER',
    '3.4.3.3.1.3': 'ZIPPER',
    '3.4.3.3.1.4': 'ZIPPER',
    '3.4.3.3.1.5': 'ZIPPER',
    '3.4.3.3.1.6': 'ZIPPER',
    '3.4.3.3.1.7': 'ZIPPER',
    '3.4.3.3.2':   '<LIDER> Enterprise Suite',
    '3.4.3.3.2.1': '<LIDER> Enterprise Suite (chatbot)',
    '3.4.3.3.2.2': '<LIDER> Enterprise Suite (aplicatie mobila)',
    '3.4.3.3.3':   'FURNIZOR LIMS COTS',
    '3.4.3.3.3.1': 'FURNIZOR LIMS COTS',
    '3.4.3.3.3.2': 'Mirth Connect',
    '3.4.3.3.4':   '<LIDER>',
    '3.4.3.4':     'Toate produsele C. Securitate',
    '3.4.3.4.1.1': 'F5 / Imperva / FortiWeb (WAF)',
    '3.4.3.4.1.2': 'FortiDeceptor (Honeypot)',
    '3.4.3.4.1.3': 'Cisco DNA / ClearPass (NMS / NAC)',
    '3.4.3.4.1.4': 'Splunk ES / QRadar (SIEM)',
    '3.4.3.4.1.5': 'Cisco IronPort (Email Security)',
    '3.4.3.4.2.1': 'FortiGate / Palo Alto + FortiGate 100F locatii',
    '3.4.3.4.2.2': 'Echipament hardware (Switch acces) - vezi Lista_Hardware',
    '3.4.3.4.2.3': 'Echipament hardware (Switch POE) - vezi Lista_Hardware',
    '3.4.3.4.2.4': 'Echipament hardware (Access point) - vezi Lista_Hardware',
    '3.4.3.4.2.5': 'Echipament hardware (Switch agregare) - vezi Lista_Hardware',
    '3.4.3.4.2.6': 'Echipament hardware (Laptop) + MS Office H&B 2024 OEM',
    '3.4.3.4.2.7': 'Echipament hardware (Complete teren) + MS Office H&B 2024 OEM',
    '3.4.4':       '<LIDER> (PM + servicii)',
    '3.4.4.1':     '<LIDER> (Manager de proiect)',
    '3.4.4.2':     '<LIDER> (livrare/instalare/configurare)',
    '3.4.4.3':     '<LIDER> (Analist business + Arhitect)',
    '3.4.4.4':     '<LIDER> (Arhitect sistem + Team leader)',
    '3.4.4.5':     '<LIDER> (Team leader + experti dezvoltare)',
    '3.4.4.6':     '<LIDER> (Expert migrare)',
    '3.4.4.7':     '<LIDER> (echipa de implementare)',
    '3.4.4.8':     '<LIDER> (Expert testare)',
    '3.4.4.9':     '<LIDER> (Experti instruire)',
    '3.4.4.10':    '<LIDER> (echipa de punere in productie)',
    '3.4.5':       'Keycloak Enterprise + <LIDER> Enterprise Suite',
    '3.4.6':       'Toate produsele C. Securitate',
    '3.4.7':       'Toate produsele C. Securitate + <LIDER> (GDPR)',
    '3.4.8':       'Echipament hardware (Laptop EU Ecolabel) + furnizori ambalaje/livrare',
    '3.4.9':       '<LIDER> (Coordonator suport tehnic)',
}


def get_produs(cap):
    cap = (cap or '').strip()
    if cap in MAP:
        return MAP[cap]
    parts = cap.split('.')
    while len(parts) > 1:
        parts.pop()
        candidate = '.'.join(parts)
        if candidate in MAP:
            return MAP[candidate]
    return '<LIDER> (de validat)'


# ============================================================
# Replace VOGO -> <LIDER> la nivel de paragraf, pastreaza formatarea
# Aceeasi tehnica ca .claude/_replace_lider.py (folosita de A3)
# ============================================================
PATTERNS = [
    (re.compile(r"VOGO\s+ENTERPRISE\s+BUSINESS\s+SUITE", re.IGNORECASE), "<LIDER> Enterprise Suite"),
    (re.compile(r"VOGO\s+TECHNOLOGY", re.IGNORECASE), "<LIDER>"),
    (re.compile(r"\bVOGO\s+Enterprise\s+Suite\b", re.IGNORECASE), "<LIDER> Enterprise Suite"),
    (re.compile(r"\bVOGO\b", re.IGNORECASE), "<LIDER>"),
]


def replace_in_text(text):
    total = 0
    for pat, repl in PATTERNS:
        text, n = pat.subn(repl, text)
        total += n
    return text, total


def replace_in_paragraph(paragraph):
    """Aplica regex pe textul paragrafului. Pastreaza formatul primului run."""
    full = paragraph.text
    new_text, n = replace_in_text(full)
    if n == 0:
        return 0
    runs = paragraph.runs
    if not runs:
        return 0
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""
    return n


def walk_replace(tables, total_ref):
    for t in tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    total_ref[0] += replace_in_paragraph(p)
                walk_replace(cell.tables, total_ref)


# ============================================================
# Pre-completare col Raspuns ofertant pe baza MAP
# Coloane in tabel original: 0=Nr | 1=Cap. CDS | 2=Cerinta | 3=Raspuns | 4=DocRef
# ============================================================

def set_cell_text(cell, text, pt_size=8.5):
    """Curata cell si pune textul, pastrand fontul Arial."""
    # Sterg toate paragrafele existente (lasand unul gol)
    for p in cell.paragraphs:
        p_el = p._element
        for r in list(p.runs):
            r._element.getparent().remove(r._element)
    # Adaug noul text intr-un singur paragraf
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    new_run = p.add_run(text)
    # Aplic font Arial + size (mostenire docDefaults asigura asta dar fac explicit)
    new_run.font.name = "Arial"
    from docx.shared import Pt
    new_run.font.size = Pt(pt_size)


def main():
    if not SRC.exists():
        raise SystemExit(f"NOT FOUND: {SRC}")

    # Step 1: Copy fisierul sursa (pastram TOATA structura)
    shutil.copy(SRC, OUT)
    print(f"Copiat {SRC.name} -> {OUT.name}")

    # Step 2: Deschid copia si fac modificarile
    doc = Document(str(OUT))

    # 2a. Replace VOGO -> <LIDER> peste tot
    total = [0]
    for p in doc.paragraphs:
        total[0] += replace_in_paragraph(p)
    walk_replace(doc.tables, total)
    # Headers / footers
    for section in doc.sections:
        for hdr in (section.header, section.first_page_header, section.even_page_header):
            for p in hdr.paragraphs:
                total[0] += replace_in_paragraph(p)
            walk_replace(hdr.tables, total)
        for ftr in (section.footer, section.first_page_footer, section.even_page_footer):
            for p in ftr.paragraphs:
                total[0] += replace_in_paragraph(p)
            walk_replace(ftr.tables, total)
    print(f"Replace VOGO -> <LIDER>: {total[0]} ocurente")

    # 2b. Pre-completare coloana 3 (Raspuns ofertant) cu produsul din MAP
    t = doc.tables[0]
    populated = 0
    skipped_header = 0
    for ri, row in enumerate(t.rows):
        cells = row.cells
        if len(cells) < 5:
            continue
        c0 = cells[0].text.strip()
        c1 = cells[1].text.strip()
        # Header tabel: Nr. + Cap. CDS
        if c0 == 'Nr.':
            continue
        # Header sectiune Cap.: c0 == c1
        if c0 == c1 and c0.startswith('Cap.'):
            skipped_header += 1
            continue
        # Rand cerinta normal
        if not c0.isdigit():
            continue
        cap = c1
        produs = get_produs(cap)
        # Set in cell 3 (Raspuns ofertant)
        set_cell_text(cells[3], produs, pt_size=8.5)
        # Cell 4 (Document referinta) - lasa goala
        set_cell_text(cells[4], "", pt_size=8.5)
        populated += 1
    print(f"Cerinte pre-completate cu produs: {populated}")
    print(f"Sectiuni header sarite: {skipped_header}")

    # Salvez
    doc.save(str(OUT))
    size = OUT.stat().st_size
    print(f"\nOK | {OUT.name} salvat: {size} bytes ({size/1024:.1f} KB)")


if __name__ == "__main__":
    main()
